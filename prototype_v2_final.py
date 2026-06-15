from __future__ import annotations
import os
import re
import json
import time
import hashlib
import html
import io
import shutil
import subprocess
import sys
import threading
import warnings
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from datetime import datetime, timedelta
from urllib.parse import urljoin, urlparse, unquote
from collections import Counter, defaultdict
from typing import Optional, Iterable, Any
from email.utils import parsedate_to_datetime

try:
    import requests
except ImportError:
    requests = None

import pandas as pd
try:
    from bs4 import BeautifulSoup, Comment, Tag
except ImportError:
    BeautifulSoup = None
    Comment = None
    Tag = None

try:
    from bs4 import XMLParsedAsHTMLWarning
except Exception:
    XMLParsedAsHTMLWarning = None

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_BREAK
except ImportError:
    Document = None
    Pt = None
    Cm = None
    WD_BREAK = None

try:
    import matplotlib.pyplot as plt
except ImportError:
    plt = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from wordcloud import WordCloud
except ImportError:
    WordCloud = None


# =========================================================
# KONFIGURATION
# =========================================================

RUN_MODE = "crawl"
# "crawl" oder "report_from_existing_csv"

EXISTING_HITS_CSV = "zirp_berichte/zirp_seitenfunde_2026-04-15_16-57.csv"
ZIRP_BASE = "https://www.zirp.de"

# =========================================================
# STATIC EXCEL LINK (SOR/SCIP v2)
# =========================================================
# The crawler keeps the old structure, but the member universe is now loaded
# from the static Excel model whenever the file is available. This keeps the
# optimization story clean: Excel defines M and later P; the crawler enriches M
# with live evidence.
USE_STATIC_EXCEL_MEMBERS = os.getenv("USE_STATIC_EXCEL_MEMBERS", "1").strip().lower() not in {"0", "false", "no", "off"}
_DEFAULT_STATIC_ZIRP_EXCEL = (
    "zirp(3).xlsx"
    if (Path(__file__).resolve().parent / "zirp(3).xlsx").exists()
    else "sor_scip_full_static_zirp_dataset_updated.xlsx"
)
STATIC_ZIRP_EXCEL_PATH = os.getenv("STATIC_ZIRP_EXCEL_PATH", _DEFAULT_STATIC_ZIRP_EXCEL)
STATIC_ZIRP_MEMBER_SHEET = os.getenv("STATIC_ZIRP_MEMBER_SHEET", "Actor_Vector")
STATIC_ZIRP_SOURCE_PROFILE_SHEET = os.getenv("STATIC_ZIRP_SOURCE_PROFILE_SHEET", "Actor_Source_Profiles")
STATIC_ZIRP_KEYWORD_RADAR_SHEET = os.getenv("STATIC_ZIRP_KEYWORD_RADAR_SHEET", "InfoBox_Keyword_Radar")
STATIC_ZIRP_ACTOR_TYPE_RULES_SHEET = os.getenv("STATIC_ZIRP_ACTOR_TYPE_RULES_SHEET", "Actor_Type_Source_Rules")
STATIC_ZIRP_CRAWLER_SCORING_SHEET = os.getenv("STATIC_ZIRP_CRAWLER_SCORING_SHEET", "Crawler_Scoring_Config")

# Feedback-approved cached words are loaded at runtime and merged into an
# active relevance model. They change article scoring/classification, not the
# static member universe or SCIP constraints.
DYNAMIC_TERM_MEMORY_PATH = os.getenv("DYNAMIC_TERM_MEMORY_PATH", "data/feedback/dynamic_term_memory.json")
USE_DYNAMIC_RELEVANCE_TERMS = os.getenv("USE_DYNAMIC_RELEVANCE_TERMS", "1").strip().lower() not in {"0", "false", "no", "off"}

REPORT_TITLE = "ZIRP-Mitgliederbriefing"
REPORT_SUBTITLE = "Leiterinnen-Briefing zu aktuellen Entwicklungen bei ZIRP-Mitgliedern"

MAX_REPORT_EVENTS_PER_SECTION = 3
MAX_EVIDENCE_EVENTS_PER_SECTION = 10
MAX_PAGES_PER_MEMBER = 10
MAX_PAGES_FOR_BROAD_SITES = 6
MAX_SNIPPETS_PER_PAGE = 3
MAX_STRATEGIC_CONDENSATIONS = 3
CRAWL_WORKERS = 12
DISCOVERY_MIN_ARTICLE_LINK_SCORE = int(os.getenv("DISCOVERY_MIN_ARTICLE_LINK_SCORE", "5") or "5")
DISCOVERY_PROFILE_LINK_SCORE = int(os.getenv("DISCOVERY_PROFILE_LINK_SCORE", "3") or "3")
DISCOVERY_MIN_SOURCE_SCORE = int(os.getenv("DISCOVERY_MIN_SOURCE_SCORE", "24") or "24")
DISCOVERY_STRONG_SOURCE_SCORE = int(os.getenv("DISCOVERY_STRONG_SOURCE_SCORE", "38") or "38")
DISCOVERY_FALLBACK_PAGE_LIMIT = int(os.getenv("DISCOVERY_FALLBACK_PAGE_LIMIT", "3") or "3")

MIN_TEXT_LENGTH = 160
MIN_SNIPPET_WORDS = 6
MIN_EVENT_SCORE = 8
MIN_DECISION_SCORE = 4
MIN_ARTICLE_QUALITY_SCORE = 4

SLEEP_SECONDS = 0
REQUEST_TIMEOUT = 4
ARTICLE_TEXT_CACHE_ENABLED = os.getenv("ARTICLE_TEXT_CACHE_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}
ARTICLE_TEXT_CACHE_REFRESH_HOURS = int(os.getenv("ARTICLE_TEXT_CACHE_REFRESH_HOURS", "12") or "12")
ARTICLE_TEXT_CACHE_KEEP_DAYS = int(os.getenv("ARTICLE_TEXT_CACHE_KEEP_DAYS", "45") or "45")
DISCOVERY_CACHE_ENABLED = os.getenv("DISCOVERY_CACHE_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}
DISCOVERY_CACHE_REFRESH_HOURS = int(os.getenv("DISCOVERY_CACHE_REFRESH_HOURS", "6") or "6")
DISCOVERY_CACHE_KEEP_DAYS = int(os.getenv("DISCOVERY_CACHE_KEEP_DAYS", "14") or "14")
FORCE_CRAWL_REFRESH = os.getenv("FORCE_CRAWL_REFRESH", "0").strip().lower() in {"1", "true", "yes", "on"}
REPORT_LOOKBACK_DAYS = 14
PUBLIC_REPORT_DAYS = 7
REPORT_STYLE = "narrative"
INCLUDE_DAILY_CALENDAR = False
INCLUDE_EVIDENCE_SECTION = False
CREATE_WORDCLOUD = False
MAX_NARRATIVE_EVENTS = int(os.getenv("MAX_NARRATIVE_EVENTS", "16") or "16")
RUN_MEETING_OPTIMIZER_AFTER_REPORT = True
MEETING_OPTIMIZER_TIMEOUT = int(os.getenv("MEETING_OPTIMIZER_TIMEOUT", "360") or "360")

# =========================================================
# OPENAI / EDITORIAL ENRICHMENT
# =========================================================

USE_OPENAI_EDITORIAL = True
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "").strip()
OPENAI_TIMEOUT = 90
OPENAI_MAX_OUTPUT_TOKENS = 900
OPENAI_MAX_RETRIES = int(os.getenv("OPENAI_MAX_RETRIES", "0") or "0")
WEEKLY_REWRITE_MAX_TOKENS = int(os.getenv("WEEKLY_REWRITE_MAX_TOKENS", "900") or "900")
WEEKLY_DIRECT_MAX_TOKENS = int(os.getenv("WEEKLY_DIRECT_MAX_TOKENS", "420") or "420")
REPORT_LOCAL_CASCADE = os.getenv("REPORT_LOCAL_CASCADE", "1").strip().lower() not in {"0", "false", "no", "off"}
REPORT_DRAFT_MODEL = os.getenv("REPORT_DRAFT_MODEL", "qwen2.5:7b-instruct")
REPORT_POLISH_MODEL = os.getenv("REPORT_POLISH_MODEL", "qwen3:14b")
REPORT_DRAFT_MAX_TOKENS = int(os.getenv("REPORT_DRAFT_MAX_TOKENS", "760") or "760")
REPORT_POLISH_MAX_TOKENS = int(os.getenv("REPORT_POLISH_MAX_TOKENS", "620") or "620")
REPORT_OPENAI_MODEL = os.getenv("REPORT_OPENAI_MODEL", "gpt-4.1-mini")
REPORT_OPENAI_BASE_URL = os.getenv("REPORT_OPENAI_BASE_URL", "").strip()
REPORT_OPENAI_TIMEOUT = int(os.getenv("REPORT_OPENAI_TIMEOUT", "60") or "60")
SIGNAL_JUDGE_ENABLED = os.getenv("SIGNAL_JUDGE_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}
SIGNAL_JUDGE_MODEL = os.getenv("SIGNAL_JUDGE_MODEL", REPORT_OPENAI_MODEL)
SIGNAL_JUDGE_TOP_N = int(os.getenv("SIGNAL_JUDGE_TOP_N", "40") or "40")
SIGNAL_JUDGE_BATCH_SIZE = int(os.getenv("SIGNAL_JUDGE_BATCH_SIZE", "5") or "5")
SIGNAL_JUDGE_MAX_TOKENS = int(os.getenv("SIGNAL_JUDGE_MAX_TOKENS", "1400") or "1400")
EDITORIAL_STRICT_MODE = True

MODEL_DECIDES_RELEVANCE = False
MODEL_KEEP_ROLES = {"top_signal", "condensation", "watchlist"}

MAX_MODEL_TOP_SIGNALS = 3
MAX_MODEL_CONDENSATIONS = 3
MAX_MODEL_WATCHLIST = 3

# Proportionale, Groq-sichere Limits.
# Intern breit prÃ¼fen, Ã¶ffentlich eng kuratieren.
DYNAMIC_LIMITS_ENABLED = os.getenv("DYNAMIC_LIMITS_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}
SIGNAL_JUDGE_MIN = int(os.getenv("SIGNAL_JUDGE_MIN", "25") or "25")
SIGNAL_JUDGE_MAX = int(os.getenv("SIGNAL_JUDGE_MAX", "50") or "50")
SIGNAL_JUDGE_MEMBER_RATIO = float(os.getenv("SIGNAL_JUDGE_MEMBER_RATIO", "0.40") or "0.40")
NARRATIVE_EVENTS_MIN = int(os.getenv("NARRATIVE_EVENTS_MIN", "10") or "10")
NARRATIVE_EVENTS_MAX = int(os.getenv("NARRATIVE_EVENTS_MAX", "18") or "18")
NARRATIVE_EVENTS_JUDGE_RATIO = float(os.getenv("NARRATIVE_EVENTS_JUDGE_RATIO", "0.40") or "0.40")
PUBLIC_SCIP_CARDS_MIN = int(os.getenv("PUBLIC_SCIP_CARDS_MIN", "4") or "4")
PUBLIC_SCIP_CARDS_MAX = int(os.getenv("PUBLIC_SCIP_CARDS_MAX", "8") or "8")
PUBLIC_SCIP_SIGNAL_RATIO = float(os.getenv("PUBLIC_SCIP_SIGNAL_RATIO", "0.35") or "0.35")
SCIP_PROMPT_CARDS_MIN = int(os.getenv("SCIP_PROMPT_CARDS_MIN", "2") or "2")
SCIP_PROMPT_CARDS_MAX = int(os.getenv("SCIP_PROMPT_CARDS_MAX", "5") or "5")
SCIP_ARCHIVE_CARDS_MAX = int(os.getenv("SCIP_ARCHIVE_CARDS_MAX", "15") or "15")

DISCOVERY_SEED_PATHS = [
    "/news", "/news/", "/presse", "/presse/", "/aktuelles", "/aktuelles/",
    "/aktuell", "/aktuell/", "/meldungen", "/meldungen/", "/pressemitteilungen",
    "/pressemitteilungen/", "/newsroom", "/newsroom/", "/unternehmen/presse",
    "/de/news", "/de/presse", "/de/aktuelles", "/de/newsroom",
    "/events", "/veranstaltungen", "/termine", "/forschung", "/transfer",
    "/publikationen", "/studien", "/research", "/analysen", "/downloads",
    "/innovation", "/nachhaltigkeit", "/feed", "/rss",
]

ACTOR_TYPE_SEED_PATHS = {
    "academic_support": ["/aktuelles", "/news", "/forschung", "/transfer", "/veranstaltungen", "/presse"],
    "implementation_anchor": ["/presse", "/aktuelles", "/news", "/veranstaltungen", "/publikationen", "/downloads"],
    "profit_anchor": ["/presse", "/newsroom", "/news", "/innovation", "/nachhaltigkeit", "/events"],
    "context_actor": ["/aktuelles", "/news", "/presse", "/termine", "/veranstaltungen", "/publikationen"],
}

SOURCE_QUALITY_TERMS = [
    "pressemitteilung", "aktuelles", "news", "research", "studie", "studien",
    "publikation", "analyse", "veranstaltung", "event", "projekt", "innovation",
    "nachhaltigkeit", "forschung", "transfer", "förderung", "foerderung",
    "newsletter", "pressearchiv", "archiv", "position", "positionspapier",
    "jahresbericht", "geschäftsbericht", "geschaeftsbericht", "faktencheck",
    "gesundheitsbericht", "mittelstand", "esg", "green finance",
]

SOURCE_DISCOVERY_PATHS = [
    "/sitemap.xml", "/sitemap_index.xml", "/robots.txt", "/feed", "/rss",
    "/news/rss", "/aktuelles/rss", "/presse/rss",
]

PAGINATION_PATTERNS = [
    "?page={n}", "?p={n}", "?tx_news_pi1[currentPage]={n}",
    "?offset={offset}", "?start={offset}", "?skip={offset}", "/page/{n}/",
    "/seite/{n}/", "/archiv/page/{n}/",
]

SOURCE_DIFFICULTY_BONUS = {
    "low": 0,
    "medium": 2,
    "high": 3,
}

DOCUMENT_LINK_EXTENSIONS = (".pdf", ".docx")

DEFAULT_ACTOR_SOURCE_FAMILY_RULES = {
    "academic": {
        "priority_boost": 0,
        "seed_paths": ["/aktuelles", "/news", "/forschung", "/transfer", "/events", "/presse"],
        "source_quality_terms": ["forschung", "transfer", "projekt", "veranstaltung", "presse"],
        "pagination_patterns": [],
    },
    "company_industry": {
        "priority_boost": 20,
        "seed_paths": ["/presse", "/pressemitteilungen", "/newsroom", "/news-and-media", "/stories", "/media-stories", "/innovation", "/nachhaltigkeit", "/investor-relations", "/finanzberichte", "/publikationen", "/events"],
        "source_quality_terms": ["pressemitteilung", "newsroom", "innovation", "investition", "nachhaltigkeit", "produktion", "standort", "jahresbericht", "geschäftsbericht"],
        "pagination_patterns": [],
    },
    "bank_finance": {
        "priority_boost": 22,
        "seed_paths": ["/research", "/studien", "/analysen", "/publikationen", "/news", "/presse", "/mittelstand", "/esg", "/green-finance"],
        "source_quality_terms": ["research", "studie", "analyse", "mittelstand", "finanzierung", "förderung", "esg", "green finance"],
        "pagination_patterns": [],
    },
    "chamber_association": {
        "priority_boost": 25,
        "seed_paths": ["/aktuelles", "/presse", "/pressemitteilungen", "/termine", "/veranstaltungen", "/publikationen", "/downloads", "/archiv", "/pressearchiv", "/category", "/infodienst"],
        "source_quality_terms": ["pressemitteilung", "position", "positionspapier", "veranstaltung", "publikation", "archiv", "infodienst"],
        "pagination_patterns": ["/category/{n}/", "/archiv/page/{n}/"],
    },
    "health_social": {
        "priority_boost": 20,
        "seed_paths": ["/presse", "/aktuelles", "/pressemitteilungen", "/faktencheck", "/gesundheitsberichte", "/positionen", "/kampagnen", "/mediathek", "/publikationen"],
        "source_quality_terms": ["gesundheitsbericht", "faktencheck", "versorgung", "position", "kampagne", "publikation"],
        "pagination_patterns": [],
    },
    "energy_infrastructure": {
        "priority_boost": 25,
        "seed_paths": ["/presse", "/newsroom", "/aktuelles", "/projekte", "/baumassnahmen", "/innovation", "/technologie", "/netze", "/nachhaltigkeit", "/mediathek"],
        "source_quality_terms": ["projekt", "netze", "energiewende", "infrastruktur", "baumassnahme", "technologie", "nachhaltigkeit"],
        "pagination_patterns": [],
    },
    "media_culture": {
        "priority_boost": 12,
        "seed_paths": ["/presse", "/pressemitteilungen", "/news", "/saison", "/programm", "/publikationen", "/mediathek"],
        "source_quality_terms": ["pressemitteilung", "programm", "saison", "publikation", "mediathek"],
        "pagination_patterns": [],
    },
    "state_institution": {
        "priority_boost": 16,
        "seed_paths": ["/presse", "/aktuelles", "/themen", "/foerderung", "/förderung", "/programme", "/publikationen", "/downloads"],
        "source_quality_terms": ["förderung", "foerderung", "programm", "thema", "publikation", "pressemitteilung"],
        "pagination_patterns": [],
    },
}

GENERIC_STATIC_PROFILE_TERMS = {
    "it", "digital", "digitale", "digitaler", "digitales", "digitalisierung",
    "ki", "daten", "data", "software", "technologie", "technik", "land",
    "wirtschaft", "standort", "region", "wissen", "gesellschaftlich",
}

CONCRETE_DIGITAL_TERMS = {
    "prozessdigitalisierung", "digital health", "e-health", "bim", "plattform",
    "cloud", "erp", "automatisierung", "industrie 4.0", "digitaler zwilling",
    "cybersecurity", "e-government", "chatbot", "sensorik", "robotik",
    "datenplattform", "software-as-a-service", "ki-anwendung", "pilot",
}

HARD_NOISE_TERMS = [
    "kapitalmÃ¤rkte daily", "kapitalmÃ¤rkte daily", "marktkommentar",
    "trump droht iran", "usa: trump", "iran mit neuen attacken",
    "che-ranking", "bestnoten", "girls'day", "girlsday",
    "gewinnspiel", "podcast", "folge", "sendung", "ratgeber",
    "tipps", "ticket", "brustring", "business club",
]

EXCLUSION_TITLE_PATTERNS = [
    "wenn kinder", "impfreaktionen", "tipps", "ratgeber", "so gelingt",
    "was sie wissen", "was eltern", "familie", "rezepte", "podcast",
    "sendung", "folge", "tv", "audiothek", "download", "gewinnspiel",
]

EXCLUSION_URL_PATTERNS = [
    "/magazin/", "/ratgeber/", "/familie/", "/gesundheit/", "/podcast/",
    "/audiothek/", "/video/", "/download/", "/service/", "/faq/",
]

EXCLUSION_TEXT_PATTERNS = [
    "tipps fÃ¼r eltern", "kleine lÃ¼gen gehÃ¶ren", "welche beschwerden kÃ¶nnen auftreten",
    "gesundheitsratgeber", "verbraucherinnen und verbraucher",
    "alltagstipps", "so reagieren sie", "hilft gegen", "was tun bei",
]

HUB_PAGE_PATTERNS = [
    "news und pressemitteilungen", "aktuelles", "meldungen", "Ã¼bersicht",
    "alle meldungen", "presse", "newsroom", "messen & events",
]

LEADERSHIP_TERMS = [
    "strategie", "strategisch", "kooperation", "partnerschaft", "allianz",
    "netzwerk", "verbund", "plattform", "investition", "ausbau",
    "fÃ¶rderung", "pilot", "initiative", "programm", "transfer",
    "innovation", "digitalisierung", "industrie 4.0", "quantentechnologien",
    "6g", "fachkrÃ¤fte", "versorgung", "pflege", "gesundheitssystem",
    "standort", "produktion", "kapazitÃ¤t", "forschungspartner",
    "internationalisierung", "cluster", "infrastruktur", "transform",
    "klima", "energie", "resilienz", "arbeitsplÃ¤tze"
]

WEAK_SIGNAL_TERMS = [
    "girls'day", "girlsday", "stipendium", "antrittsvorlesung",
    "campus entdecken", "gemeinsam feiern", "bewerben", "willkommen",
]

STRONG_STRATEGIC_TERMS = [
    "quantentechnologien", "6g", "industrie 4.0", "mittelstand-digital",
    "gesundheitsversorgung", "koalitionsverhandlungen", "fachkrÃ¤fte",
    "kooperation", "partnerschaft", "allianz", "standort", "investition",
    "innovation", "digitalisierung", "transfer", "internationalisierung",
]

MACRO_COMMENTARY_TERMS = [
    "wachstumsprognosen", "prognosen", "aktien", "aktienmarkt",
    "kapitalmÃ¤rkte", "kapitalmarkt", "allzeithochs", "friedenshoffnung",
    "marktkommentar", "konjunktur", "daily", "einschÃ¤tzungen",
    "iwf", "bip", "inflations", "handelskonflikt", "makro"
]

INSTITUTIONAL_ACTION_TERMS = [
    "vereinbarung", "kooperation", "partnerschaft", "allianz", "gemeinsame",
    "netzwerk", "transfer", "ausbau", "strategie", "initiative",
    "programm", "internationalisierung", "projektpartner", "verbund",
    "forschungspartner", "cluster", "plattform", "konsortium"
]

REGIONAL_IMPLEMENTATION_TERMS = [
    "kommune", "kommunen", "kommunal", "rheinland-pfalz", "mainz", "trier",
    "kaiserslautern", "bingen", "wallmerod", "region", "land", "landes",
    "versorgung", "fachkrÃ¤fte", "klima", "umsetzung", "maÃŸnahme", "maÃŸnahmen",
    "strategie", "biodiversitÃ¤tsstrategie", "fÃ¶rderprogramm", "projektpartner"
]


# =========================================================
# VOLLSTÃ„NDIGE MITGLIEDERLISTE (best effort homepages)
# =========================================================

MEMBERS = [
    {"name": "Alexianer GmbH", "homepage": "https://www.alexianer.de/", "zirp_member_path": "", "seed_paths": []},
    {"name": "AOK Rheinland-Pfalz/Saarland", "homepage": "https://www.aok.de/rheinland-pfalz", "zirp_member_path": "", "seed_paths": []},
    {"name": "Architektenkammer Rheinland-Pfalz", "homepage": "https://www.diearchitekten.org", "zirp_member_path": "", "seed_paths": []},
    {"name": "Aturas GmbH", "homepage": "https://aturas.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "Barmherzige BrÃ¼der Trier gGmbH", "homepage": "https://www.bbtgruppe.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "BASF SE", "homepage": "https://www.basf.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "Bauern- und Winzerverband Rheinland-Pfalz SÃ¼d e.V.", "homepage": "https://www.bwv-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Bauwirtschaft Rheinland-Pfalz e.V.", "homepage": "https://www.bauwirtschaft-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Bernd Hummel Holding GmbH", "homepage": "https://www.bernd-hummel.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Bitburger Braugruppe GmbH", "homepage": "https://www.bitburger-braugruppe.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Boehringer Ingelheim Pharma GmbH & Co. KG", "homepage": "https://www.boehringer-ingelheim.com/de", "zirp_member_path": "", "seed_paths": []},
    {"name": "BRICKMAKERS AG", "homepage": "http://www.brickmakers.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Bundesagentur fÃ¼r Arbeit, Regionaldirektion Rheinland-Pfalz-Saarland", "homepage": "https://www.arbeitsagentur.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Caritasverband fÃ¼r die DiÃ¶zese Speyer e.V.", "homepage": "https://www.caritas-speyer.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Continental Automotive Technologies GmbH, Safety and Motion (SAM), Standort RheinbÃ¶llen", "homepage": "https://www.continental-corporation.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Debeka Versicherungsgruppe", "homepage": "https://www.debeka.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Deutsche Bundesbank, Hauptverwaltung in Rheinland-Pfalz und dem Saarland", "homepage": "https://www.bundesbank.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Deutsche Telekom AG", "homepage": "https://www.telekom.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Deutsche UniversitÃ¤t fÃ¼r Verwaltungswissenschaften Speyer", "homepage": "https://www.uni-speyer.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Deutscher Gewerkschaftsbund (DGB) Rheinland-Pfalz/Saarland", "homepage": "https://rlp-saarland.dgb.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Duale Hochschule Rheinland-Pfalz", "homepage": "https://www.dh-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "1. FSV Mainz 05 e.V.", "homepage": "https://www.mainz05.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Eckes-Granini Deutschland GmbH", "homepage": "https://www.eckes-granini.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Empolis Information Management GmbH", "homepage": "https://www.empolis.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "ENTEGA Plus GmbH", "homepage": "https://www.entega.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Entwicklungsagentur Rheinland-Pfalz e.V.", "homepage": "https://www.ea-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Ernst & Young GmbH", "homepage": "https://www.ey.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "Erster FC Kaiserslautern GmbH & Co. KGaA", "homepage": "https://fck.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Evangelische Kirche der Pfalz", "homepage": "https://www.evkirchepfalz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Fachingen Heil- und Mineralbrunnen GmbH", "homepage": "https://www.fachingen.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Gerolsteiner Brunnen GmbH & Co. KG", "homepage": "https://www.gerolsteiner.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "GLOBUS Markthallen Holding GmbH & Co. KG", "homepage": "https://www.globus.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Handwerkskammern Rheinland-Pfalz", "homepage": "https://handwerk-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Heberger Gruppe", "homepage": "https://www.heberger.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule fÃ¼r Wirtschaft und Gesellschaft Ludwigshafen", "homepage": "https://www.hwg-lu.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule Kaiserslautern", "homepage": "https://www.hs-kl.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule Koblenz", "homepage": "https://www.hs-koblenz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule Mainz", "homepage": "https://www.hs-mainz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule Trier", "homepage": "https://www.hochschule-trier.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Hochschule Worms", "homepage": "https://www.hs-worms.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "IKK SÃ¼dwest", "homepage": "https://www.ikk-suedwest.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "in.betrieb gGmbH", "homepage": "https://inbetrieb-mainz.de/de/", "zirp_member_path": "", "seed_paths": []},
    {"name": "Ingenieurkammer Rheinland-Pfalz", "homepage": "https://www.ing-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Investitions- und Strukturbank Rheinland-Pfalz (ISB)", "homepage": "https://www.isb.rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "IQIB - Institut fÃ¼r qualifizierende Innovationsforschung und -beratung GmbH", "homepage": "https://www.iqib.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "ITK Engineering GmbH", "homepage": "https://www.itk-engineering.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "JOBS FOR MOMS UG", "homepage": "https://jobsformoms.de/de-de/home", "zirp_member_path": "", "seed_paths": []},
    {"name": "Johannes Gutenberg-UniversitÃ¤t Mainz", "homepage": "https://www.uni-mainz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Joseph VÃ¶gele AG", "homepage": "https://www.voegele.info", "zirp_member_path": "", "seed_paths": []},
    {"name": "JUWI GmbH", "homepage": "https://www.juwi.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Karl GemÃ¼nden GmbH & Co. KG", "homepage": "https://www.gemuenden-bau.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "KassenÃ¤rztliche Vereinigung Rheinland-Pfalz", "homepage": "https://www.kv-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "KassenzahnÃ¤rztliche Vereinigung Rheinland-Pfalz", "homepage": "https://www.kzv-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Katholische Hochschule Mainz", "homepage": "https://www.kh-mz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "KPMG AG WirtschaftsprÃ¼fungsgesellschaft", "homepage": "https://kpmg.com/de/de/home", "zirp_member_path": "", "seed_paths": []},
    {"name": "KSB AG", "homepage": "https://www.ksb.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "KÃœBLER GmbH", "homepage": "https://www.kuebler-hallenheizungen.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "L Q M Marktforschung GmbH", "homepage": "", "zirp_member_path": "", "seed_paths": []},
    {"name": "LandesÃ¤rztekammer Rheinland-Pfalz", "homepage": "https://www.laek-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Landesbank Baden-WÃ¼rttemberg (LBBW)", "homepage": "https://www.lbbw.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "LandesPsychotherapeutenKammer Rheinland-Pfalz", "homepage": "https://www.lpk-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Landkreistag Rheinland-Pfalz", "homepage": "https://www.landkreistag.rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Leonardo PersonalKonzept GmbH", "homepage": "https://www.leonardo-personal.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Licharz GmbH", "homepage": "https://www.licharz.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "LOTTO Rheinland-Pfalz GmbH", "homepage": "https://www.lotto-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "LÃ–WEN ENTERTAINMENT GmbH", "homepage": "https://www.loewen.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Mainzer Stadtwerke AG", "homepage": "https://www.mainzer-stadtwerke.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Michelin Reifenwerke AG & Co. KGaA", "homepage": "https://www.michelin.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Ministerium fÃ¼r Wirtschaft, Verkehr, Landwirtschaft und Weinbau", "homepage": "https://mwvlw.rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "PFAFF Industriesysteme und Maschinen GmbH", "homepage": "https://www.pfaff-industrial.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "Pfalzwerke AG", "homepage": "https://www.pfalzwerke.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Provinzial Versicherung AG", "homepage": "http://www.provinzial.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Rheinland-PfÃ¤lzische Technische UniversitÃ¤t Kaiserslautern-Landau (RPTU)", "homepage": "https://regionalentwicklung-raumordnung.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "robotspaceship", "homepage": "https://www.robotspaceship.ai", "zirp_member_path": "", "seed_paths": []},
    {"name": "RPR1.", "homepage": "https://www.rpr1.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Schoenergie GmbH", "homepage": "http://www.schoenergie.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "SCHOTT AG", "homepage": "https://www.schott.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "Schuler Service GmbH & Co. KG", "homepage": "https://www.schuler-service.com", "zirp_member_path": "", "seed_paths": []},
    {"name": "SIMONA AG", "homepage": "https://www.simona.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "SmartFactory KL e.V.", "homepage": "https://smartfactory.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Sparkassenverband Rheinland-Pfalz", "homepage": "https://www.sv-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Staatskanzlei des Landes Rheinland-Pfalz", "homepage": "https://www.rlp.de/regierung/staatskanzlei", "zirp_member_path": "", "seed_paths": []},
    {"name": "StÃ¤dtetag Rheinland-Pfalz", "homepage": "https://www.staedtetag-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Steuerberaterkammer Rheinland-Pfalz", "homepage": "https://www.sbk-rlp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "SWR â€“ SÃ¼dwestrundfunk", "homepage": "https://www.swr.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Techniker Krankenkasse", "homepage": "https://www.tk.de/rp", "zirp_member_path": "", "seed_paths": []},
    {"name": "Technische Hochschule Bingen", "homepage": "https://www.th-bingen.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "thyssenkrupp Rasselstein GmbH", "homepage": "https://www.thyssenkrupp-steel.com/de/unternehmen/business-units/packaging-steel/", "zirp_member_path": "", "seed_paths": []},
    {"name": "Transdev GmbH", "homepage": "https://www.transdev.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "TÃœV Rheinland Berlin Brandenburg Pfalz e.V.", "homepage": "https://www.tuv.com/germany/de/", "zirp_member_path": "", "seed_paths": []},
    {"name": "UniversitÃ¤t Koblenz", "homepage": "http://www.uni-koblenz.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "UniversitÃ¤t Trier", "homepage": "https://www.uni-trier.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Vereinigte VR Bank Kur- und Rheinpfalz eG", "homepage": "https://www.vvrbank-krp.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "vero â€“ Verband der Bau- und Rohstoffindustrie", "homepage": "http://www.vero-baustoffe.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Villa Musica Rheinland-Pfalz", "homepage": "https://www.villamusica.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Vinzenz Pallotti University gGmbH", "homepage": "https://vp-uni.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Westenergie AG", "homepage": "https://www.westenergie.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "WHU â€“ Otto Beisheim School of Management", "homepage": "https://www.whu.edu", "zirp_member_path": "", "seed_paths": []},
    {"name": "Wilhelm Faber GmbH", "homepage": "https://www.wilhelm-faber.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "Zahnen Technik GmbH", "homepage": "https://zahnen-technik.de", "zirp_member_path": "", "seed_paths": []},
    {"name": "ZDF â€“ Zweites Deutsches Fernsehen", "homepage": "https://www.zdf.de", "zirp_member_path": "", "seed_paths": []},
]

# Keep the original hard-coded Python list as an explicit fallback snapshot.
# From this point on, MEMBERS is rebuilt from either Excel or this fallback.
FALLBACK_MEMBERS_RAW: list[dict[str, Any]] = [dict(member) for member in MEMBERS]


def clean_member_name(value: Any) -> str:
    """Normalize member names only for matching/de-duplication."""
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return "" if text.lower() in {"nan", "none"} else text


def member_key(value: Any) -> str:
    """Stable comparison key for patch/delete logic."""
    return clean_member_name(value).casefold()


def clean_excel_value(value: Any, default: str = "") -> str:
    """Return Excel cell content as a clean string without pandas NaN noise."""
    if value is None:
        return default
    try:
        if pd.isna(value):
            return default
    except Exception:
        pass
    text = str(value).strip()
    return default if text.lower() in {"nan", "none"} else text


def normalize_search_text(value: Any) -> str:
    text = str(value or "").lower()
    replacements = {
        "\u00c3\u00a4": "\u00e4",
        "\u00c3\u00b6": "\u00f6",
        "\u00c3\u00bc": "\u00fc",
        "\u00c3\u009f": "\u00df",
        "\u00e2\u20ac\u201c": "-",
        "\u00e2\u20ac\u201d": "-",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    text = text.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    return re.sub(r"\s+", " ", text).strip()


def safe_int(value: Any, default: int) -> int:
    try:
        if pd.isna(value):
            return default
    except Exception:
        pass
    try:
        return int(float(value))
    except Exception:
        return default


def with_member_ids(members: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Assign stable sequential IDs unless a row already has member_id."""
    enriched: list[dict[str, Any]] = []
    for idx, member in enumerate(members, start=1):
        row = dict(member)
        row["member_id"] = safe_int(row.get("member_id", idx), idx)
        enriched.append(row)
    return enriched


def _resolve_project_path(path_value: str | os.PathLike[str]) -> Path:
    path = Path(path_value)
    if path.is_absolute():
        return path
    return Path(__file__).resolve().parent / path


def excel_member_row_to_crawler_member(raw: dict[str, Any], fallback_id: int) -> Optional[dict[str, Any]]:
    """Convert one All_ZIRP_Members row into the crawler member schema.

    The crawler still only requires name/homepage, but preserving the static
    metadata keeps prototype_v2 aligned with the Excel-based SCIP model.
    """
    row = {str(k).strip(): v for k, v in raw.items()}
    name = clean_member_name(row.get("member_name") or row.get("mitglied") or row.get("name"))
    if not name:
        return None

    return {
        "member_id": safe_int(row.get("member_id"), fallback_id),
        "name": name,
        "homepage": clean_excel_value(row.get("homepage")),
        "zirp_member_path": "",
        "seed_paths": [],
        "actor_role": clean_excel_value(row.get("actor_role")),
        "cluster_count": row.get("cluster_count", ""),
        "max_cluster_relevance": row.get("max_cluster_relevance", ""),
        "broad_zirp_fields": clean_excel_value(row.get("broad_zirp_fields")),
        "top_static_clusters": clean_excel_value(row.get("top_static_clusters")),
    }


def load_members_from_static_excel(
    excel_path: str | os.PathLike[str] = STATIC_ZIRP_EXCEL_PATH,
    sheet_name: str = STATIC_ZIRP_MEMBER_SHEET,
) -> list[dict[str, Any]]:
    """Load member rows from the static SOR/SCIP Excel workbook.

    If the sheet contains the full member universe, it will become M.
    If it contains only update rows, it will be applied as a patch below.
    """
    path = _resolve_project_path(excel_path)
    print(f"Static Excel member path checked: {path.resolve()} exists={path.exists()}")
    if not USE_STATIC_EXCEL_MEMBERS or not path.exists():
        return []

    try:
        df = pd.read_excel(path, sheet_name=sheet_name)
    except Exception as exc:
        print(f"Static Excel members: konnte {path.name}/{sheet_name} nicht laden ({exc}); nutze Fallback-MEMBERS")
        return []

    members: list[dict[str, Any]] = []
    for idx, raw in df.iterrows():
        member = excel_member_row_to_crawler_member(dict(raw), fallback_id=idx + 1)
        if member:
            members.append(member)

    if members:
        print(f"Static Excel members: {len(members)} Zeilen aus {path.name}/{sheet_name} geladen")
    return members


def split_excel_terms(value: Any) -> list[str]:
    text = clean_excel_value(value)
    if not text:
        return []
    parts = re.split(r"[;,\n|]+", text)
    return [re.sub(r"\s+", " ", part).strip() for part in parts if re.sub(r"\s+", " ", part).strip()]


def source_priority_score(value: Any) -> int:
    text = clean_excel_value(value).lower()
    if "very high" in text or "sehr hoch" in text:
        return 100
    if "high" in text or "hoch" in text:
        return 80
    if "medium" in text or "mittel" in text:
        return 55
    if "low" in text or "niedrig" in text:
        return 30
    return 45


def load_actor_source_profiles(
    excel_path: str | os.PathLike[str] = STATIC_ZIRP_EXCEL_PATH,
    sheet_name: str = STATIC_ZIRP_SOURCE_PROFILE_SHEET,
) -> dict[str, list[dict[str, Any]]]:
    path = _resolve_project_path(excel_path)
    if not path.exists():
        return {}
    try:
        df = pd.read_excel(path, sheet_name=sheet_name)
    except Exception as exc:
        print(f"Actor source profiles: konnte {path.name}/{sheet_name} nicht laden ({exc})")
        return {}

    profiles: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for _, raw in df.iterrows():
        row = {str(k).strip(): v for k, v in dict(raw).items()}
        member = clean_member_name(row.get("Member") or row.get("member_name") or row.get("mitglied"))
        url = clean_excel_value(row.get("Source URL") or row.get("source_url") or row.get("url"))
        if not member or not url:
            continue
        profile = {
            "member": member,
            "url": url,
            "source_type": clean_excel_value(row.get("Source Type")),
            "priority": source_priority_score(row.get("Crawl Priority")),
            "difficulty": clean_excel_value(row.get("Crawler Difficulty")).lower(),
            "strategy": clean_excel_value(row.get("Recommended crawl strategy")),
            "follow_patterns": split_excel_terms(row.get("Follow patterns")),
            "extract_signals": split_excel_terms(row.get("Extract sections / signals")),
            "keyword_families": split_excel_terms(row.get("Keyword families")),
            "status": clean_excel_value(row.get("Status")),
        }
        profiles[member_key(member)].append(profile)

    for rows in profiles.values():
        rows.sort(key=lambda item: (-int(item.get("priority", 0)), item.get("url", "")))
    if profiles:
        print(f"Actor source profiles: {sum(len(v) for v in profiles.values())} Quellen aus {path.name}/{sheet_name} geladen")
    return dict(profiles)


def load_member_keyword_radar(
    excel_path: str | os.PathLike[str] = STATIC_ZIRP_EXCEL_PATH,
    sheet_name: str = STATIC_ZIRP_KEYWORD_RADAR_SHEET,
) -> dict[str, dict[str, Any]]:
    path = _resolve_project_path(excel_path)
    if not path.exists():
        return {}
    try:
        df = pd.read_excel(path, sheet_name=sheet_name)
    except Exception as exc:
        print(f"Keyword radar: konnte {path.name}/{sheet_name} nicht laden ({exc})")
        return {}

    radar: dict[str, dict[str, Any]] = {}
    for _, raw in df.iterrows():
        row = {str(k).strip(): v for k, v in dict(raw).items()}
        member = clean_member_name(row.get("member_name") or row.get("Member") or row.get("mitglied"))
        if not member:
            continue
        keywords = split_excel_terms(row.get("keywords_hit"))
        families = split_excel_terms(row.get("keyword_families_hit"))
        profile_text = clean_excel_value(row.get("searchable_profile_text"))
        strongest = clean_excel_value(row.get("strongest_static_signal"))
        all_terms = list(dict.fromkeys(keywords + families + split_excel_terms(strongest)))
        radar[member_key(member)] = {
            "member": member,
            "keywords": all_terms,
            "keyword_families": families,
            "profile_text": profile_text,
            "strongest_static_signal": strongest,
            "suggested_use": clean_excel_value(row.get("suggested_use")),
        }
    if radar:
        print(f"Keyword radar: {len(radar)} Mitglieder aus {path.name}/{sheet_name} geladen")
    return radar


def load_actor_type_source_rules(
    excel_path: str | os.PathLike[str] = STATIC_ZIRP_EXCEL_PATH,
    sheet_name: str = STATIC_ZIRP_ACTOR_TYPE_RULES_SHEET,
) -> dict[str, dict[str, Any]]:
    rules = {
        family: {
            "priority_boost": int(config.get("priority_boost", 0) or 0),
            "seed_paths": list(config.get("seed_paths", [])),
            "source_quality_terms": list(config.get("source_quality_terms", [])),
            "pagination_patterns": list(config.get("pagination_patterns", [])),
        }
        for family, config in DEFAULT_ACTOR_SOURCE_FAMILY_RULES.items()
    }
    path = _resolve_project_path(excel_path)
    if not path.exists():
        return rules
    try:
        df = pd.read_excel(path, sheet_name=sheet_name)
    except Exception:
        return rules

    loaded = 0
    for _, raw in df.iterrows():
        row = {str(k).strip(): v for k, v in dict(raw).items()}
        enabled = clean_excel_value(row.get("enabled") or row.get("Enabled") or "1").lower()
        if enabled in {"0", "false", "nein", "no", "off"}:
            continue
        family = clean_excel_value(row.get("actor_family") or row.get("Actor Family")).lower()
        if not family:
            continue
        base = rules.get(family, {"priority_boost": 0, "seed_paths": [], "source_quality_terms": [], "pagination_patterns": []})
        try:
            priority_boost = int(float(clean_excel_value(row.get("priority_boost") or row.get("Priority Boost") or base.get("priority_boost", 0))))
        except Exception:
            priority_boost = int(base.get("priority_boost", 0) or 0)
        seed_paths = split_excel_terms(row.get("seed_paths") or row.get("Seed Paths")) or list(base.get("seed_paths", []))
        source_terms = split_excel_terms(row.get("source_quality_terms") or row.get("Source Quality Terms")) or list(base.get("source_quality_terms", []))
        pagination_patterns = split_excel_terms(row.get("pagination_patterns") or row.get("Pagination Patterns")) or list(base.get("pagination_patterns", []))
        rules[family] = {
            "priority_boost": priority_boost,
            "seed_paths": seed_paths,
            "source_quality_terms": source_terms,
            "pagination_patterns": pagination_patterns,
        }
        loaded += 1
    if loaded:
        print(f"Actor type source rules: {loaded} Familien aus {path.name}/{sheet_name} geladen")
    return rules


def load_crawler_scoring_config(
    excel_path: str | os.PathLike[str] = STATIC_ZIRP_EXCEL_PATH,
    sheet_name: str = STATIC_ZIRP_CRAWLER_SCORING_SHEET,
) -> dict[str, Any]:
    path = _resolve_project_path(excel_path)
    if not path.exists():
        return {}
    try:
        df = pd.read_excel(path, sheet_name=sheet_name)
    except Exception:
        return {}

    config: dict[str, Any] = {}
    for _, raw in df.iterrows():
        row = {str(k).strip(): v for k, v in dict(raw).items()}
        key = clean_excel_value(row.get("key") or row.get("Key"))
        value = clean_excel_value(row.get("value") or row.get("Value"))
        if not key or value == "":
            continue
        config[key.strip().lower()] = value
    if config:
        print(f"Crawler scoring config: {len(config)} Regeln aus {path.name}/{sheet_name} geladen")
    return config


def is_full_member_universe(static_members: list[dict[str, Any]]) -> bool:
    """Heuristic: full ZIRP member sheet has many rows; update patch has only a few."""
    minimum_full_rows = int(os.getenv("STATIC_EXCEL_MIN_FULL_MEMBER_ROWS", "50") or "50")
    return len(static_members) >= minimum_full_rows


def apply_member_patch(
    fallback_members: list[dict[str, Any]],
    patch_members: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Apply the current ZIRP member cleanup to the fallback universe.

    This supports the uploaded workbook shape where All_ZIRP_Members contains
    only the two inserted rows instead of the full member universe.
    """
    delete_keys = {
        member_key(x)
        for x in os.getenv("STATIC_EXCEL_DELETE_MEMBERS", "Alexianer GmbH;Alexianer").split(";")
        if member_key(x)
    }

    patched: list[dict[str, Any]] = []
    for member in fallback_members:
        if member_key(member.get("name")) in delete_keys:
            continue
        patched.append({k: v for k, v in member.items() if k != "member_id"})

    seen = {member_key(member.get("name")) for member in patched}
    for member in patch_members:
        key = member_key(member.get("name"))
        if not key or key in seen:
            continue
        patched.append({k: v for k, v in member.items() if k != "member_id"})
        seen.add(key)

    return patched


def configured_deleted_member_keys() -> set[str]:
    return {
        member_key(x)
        for x in os.getenv("STATIC_EXCEL_DELETE_MEMBERS", "Alexianer GmbH;Alexianer").split(";")
        if member_key(x)
    }


def build_member_universe(
    *,
    fallback_members_raw: list[dict[str, Any]],
    static_members: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build crawler member universe M from Excel first, fallback second.

    Clean rule:
    - Full Excel member sheet => Excel defines M.
    - Small Excel update sheet => patch fallback M.
    - No Excel => fallback M.
    """
    if is_full_member_universe(static_members):
        members = [dict(member) for member in static_members]
        source = "full Excel universe"
    elif static_members:
        members = apply_member_patch(fallback_members_raw, static_members)
        source = "fallback universe + Excel patch"
    else:
        members = apply_member_patch(fallback_members_raw, [])
        source = "fallback hard-coded universe minus deleted members"

    delete_keys = configured_deleted_member_keys()
    if delete_keys:
        members = [member for member in members if member_key(member.get("name")) not in delete_keys]

    result = with_member_ids(members)
    names = {member_key(member.get("name")) for member in result}
    hfgg_present = any("hochschule für gesellschaftsgestaltung" in name for name in names)
    print(
        f"Member universe v2: {len(result)} Mitglieder ({source}); "
        f"Alexianer={'ja' if member_key('Alexianer GmbH') in names else 'nein'}; "
        f"HfM Mainz={'ja' if member_key('Hochschule für Musik Mainz') in names else 'nein'}; "
        f"HfGG={'ja' if hfgg_present else 'nein'}"
    )
    return result


_STATIC_EXCEL_MEMBERS = load_members_from_static_excel()
MEMBERS = build_member_universe(
    fallback_members_raw=FALLBACK_MEMBERS_RAW,
    static_members=_STATIC_EXCEL_MEMBERS,
)
ACTOR_SOURCE_PROFILES = load_actor_source_profiles()
MEMBER_KEYWORD_RADAR = load_member_keyword_radar()
ACTOR_TYPE_SOURCE_RULES = load_actor_type_source_rules()
CRAWLER_SCORING_CONFIG = load_crawler_scoring_config()


# =========================================================
# RELEVANZMODELL â€“ an ZIRP-PrioritÃ¤ten angelehnt
# =========================================================

STRATEGIC_PRIORITY_TERMS = {
    "transformation", "transformiert", "nachhaltigkeit", "international", "innovation",
    "innovationskraft", "wettbewerbsfÃ¤higkeit", "zukunftsfÃ¤higkeit", "wirtschaftsstandort",
    "fachkrÃ¤fte", "versorgung", "energie", "klima", "mobilitÃ¤t", "digitalisierung",
    "wissensallianz", "hochschule", "transfer", "kultur", "gesellschaft", "resilienz",
    "industriestandort", "rheinland-pfalz"
}

RLP_TERMS = {
    "rheinland-pfalz", "rlp", "pfalz", "rheinhessen", "eifel", "hunsrÃ¼ck", "mosel",
    "mainz", "koblenz", "trier", "worms", "kaiserslautern", "ludwigshafen", "neustadt",
    "speyer", "ingelheim", "wÃ¶rrstadt", "bingen", "vallendar", "frankenthal", "bitburg",
    "gerolstein", "kirn", "bad kreuznach", "pirmasens", "montabaur", "alzey", "rÃ¼lzheim"
}

NON_RLP_CONTEXT_TERMS = {
    "sachsen", "bayern", "baden-wÃ¼rttemberg", "stuttgart", "heilbronn", "freiburg",
    "konstanz", "dresden", "hamburg", "berlin", "mÃ¼nchen"
}

REPORT_SECTIONS = {
    "wirtschaftsentwicklung": {
        "title": "Wirtschaftslage und -entwicklung",
        "keywords": [
            "wachstum", "investition", "investitionen", "ausbau", "standort",
            "produktion", "innovation", "innovationen", "markt", "umsatz",
            "beschÃ¤ftigung", "arbeitsplÃ¤tze", "expansion", "strategie",
            "neuausrichtung", "transformationsprozess", "krise", "insolvenz",
            "auftrag", "bestellung", "betrieb", "anlage", "werk", "fabrik",
            "wertschÃ¶pfung", "fÃ¶rderung", "finanzierung", "wirtschaft"
        ],
    },
    "versorgung_gesundheit": {
        "title": "Versorgung und Gesundheitssystem",
        "keywords": [
            "gesundheitsversorgung", "versorgung", "krankenkasse", "pflege",
            "Ã¤rzte", "arzt", "gesundheit", "gesundheitssystem",
            "arzneimittel", "prÃ¤vention", "gkv", "patienten",
            "notfallversorgung", "therapie", "ambulant", "klinikum"
        ],
    },
    "fuehrungswechsel": {
        "title": "Führungswechsel",
        "keywords": [
            "geschÃ¤ftsfÃ¼hrer", "geschÃ¤ftsfÃ¼hrung", "vorstand", "ceo", "leitung",
            "ernennung", "nachfolge", "wechsel", "berufung", "prÃ¤sident",
            "vorsitz", "aufsichtsrat", "direktor", "dekan", "rektorin", "rektor",
            "neuer vorstand", "neue geschÃ¤ftsfÃ¼hrung", "personalie"
        ],
    },
    "soziales_engagement": {
        "title": "Soziales Engagement",
        "keywords": [
            "engagement", "spende", "spenden", "bildung", "fÃ¶rderung",
            "nachwuchs", "ausbildung", "verantwortung", "teilhabe",
            "gesellschaft", "nachhaltigkeit", "soziales", "gemeinwohl",
            "stiftung", "regional", "ehrenamt", "frÃ¼herkennung", "gesundheit",
            "vielfalt", "inklusion", "demokratie", "familie", "pflege"
        ],
    },
    "kooperationen": {
        "title": "Kooperationen mit anderen Unternehmen / Partnern",
        "keywords": [
            "kooperation", "kooperationen", "partnerschaft", "partner",
            "zusammenarbeit", "netzwerk", "verbund", "hochschule",
            "forschungseinrichtung", "gemeinsam", "projektpartner",
            "allianz", "cluster", "runde tisch", "konsortium",
            "verbÃ¤nde", "wissensallianz", "plattform"
        ],
    },
}

ZIRP_THEMENFELDER = {
    "wirtschaft": [
        "wirtschaft", "wachstum", "investition", "standort", "produktion",
        "umsatz", "beschÃ¤ftigung", "arbeitsplÃ¤tze", "markt", "auftrag",
        "bestellung", "wettbewerb", "industrie"
    ],
    "technologie": [
        "technologie", "digitalisierung", "ki", "kÃ¼nstliche intelligenz",
        "innovation", "industrie 4.0", "smart city", "e-health",
        "mobilitÃ¤t", "energiewende", "forschung", "entwicklung",
        "batterie", "zellfertigung", "webapp", "plattform", "app",
        "bim", "e-government", "automatisierung", "software"
    ],
    "nachhaltigkeit": [
        "nachhaltigkeit", "klimaneutralitÃ¤t", "umweltschutz", "co2",
        "ressourcen", "verantwortung", "zukunftsunternehmen",
        "Ã¶kologisch", "Ã¶konomisch", "photovoltaik", "solar",
        "energie", "batteriespeicher", "erneuerbar", "recycling"
    ],
    "kultur": [
        "kultur", "kreativwirtschaft", "innenstadt", "musik", "kunst",
        "publikation", "kulturregion", "kulturell", "identitÃ¤t"
    ],
    "wissen": [
        "wissen", "bildung", "weiterbildung", "hochschule",
        "wissenschaft", "kompetenzen", "lebenslanges lernen",
        "wissensraum", "forschung", "transfer"
    ],
    "gesellschaft": [
        "gesellschaft", "teilhabe", "gerechtigkeit", "zusammenhalt",
        "vielfalt", "gesundheit", "pflege", "wohnen", "familie",
        "partizipation", "medien", "soziales", "versorgung", "krankenkasse",
        "notfallversorgung", "arbeitsmarkt", "kommunen", "landesregierung"
    ],
}

POSITIVE_LINK_PATTERNS = [
    "presse", "news", "aktuell", "aktuelles", "mitteilung", "meldung",
    "berichte", "bericht", "publikation", "medien", "media", "newsroom",
    "insights", "unternehmen", "forschung", "projekt", "blog", "story"
]

NEGATIVE_LINK_PATTERNS = [
    "impressum", "datenschutz", "kontakt", "login", "karriere", "jobs",
    "bewerbung", "shop", "warenkorb", "newsletter", "abo", "suche",
    "search", "eventkalender", "veranstaltungskalender", "termine",
    "facebook.com", "instagram.com", "linkedin.com", "xing.com",
    "youtube.com", "mailto:", "tel:", ".jpg", ".jpeg", ".png", ".svg",
    ".gif", ".webp", ".zip"
]

BROAD_SITE_DOMAINS = {
    "swr.de",
    "zdf.de",
    "whu.edu",
}

BAD_TITLES = {
    "homepage", "startseite", "presse", "news", "aktuelles",
    "unternehmen", "service", "medien", "pressemitteilungen",
    "pressemitteilung", "aktuell", "Ã¼bersicht", "overview",
    "pressekontakt", "hauptnavigation", "aktuelle meldungen"
}

NOISE_PATTERNS = [
    "mehr dazu", "previous next", "zum inhalt springen", "home news",
    "newsletter", "teilen", "share", "pressekontakt", "copyright",
    "audiothek", "player", "stream", "video", "podcast", "sendung",
    "filmtipp", "tv-tipp", "episode", "folge", "tatort"
]

ROUTINE_PATTERNS = [
    "sommersemester", "wintersemester", "semesterstart",
    "einschreibung", "veranstaltungskalender", "termine",
    "uhr", "kurzfÃ¼hrung", "anmeldung"
]

GENERIC_ANCHORS = {
    "mehr", "mehr erfahren", "weiterlesen", "lesen", "mehr lesen",
    "details", "zum artikel", "zur meldung", "zur pressemitteilung"
}

GENERIC_CONTAINER_HINTS = [
    "header", "footer", "nav", "menu", "breadcrumb", "sidebar", "teaser",
    "related", "social", "share", "newsletter", "cookie", "consent", "hero"
]

HIDDEN_STYLE_PATTERNS = [
    "display:none", "visibility:hidden", "opacity:0", "font-size:0",
    "left:-9999", "text-indent:-9999", "transform:scale(0)",
    "clip-path:inset(100%)", "clip:rect(0,0,0,0)"
]

GERMAN_STOPWORDS = {
    "der", "die", "das", "und", "oder", "mit", "von", "im", "in", "am",
    "an", "zu", "zur", "zum", "den", "dem", "des", "ein", "eine", "einer",
    "eines", "einem", "einen", "ist", "sind", "war", "wurden", "wird",
    "werden", "auf", "fÃ¼r", "bei", "als", "auch", "aus", "durch", "Ã¼ber",
    "unter", "nach", "vor", "bis", "seit", "mehr", "neue", "neuen",
    "neuer", "neues", "dass", "sowie", "rund", "sehr", "bereits", "kann",
    "kÃ¶nnen", "dies", "diese", "dieser", "dieses", "sich", "nicht", "noch",
    "hier", "wir", "unser", "unsere", "ihr", "ihre", "seine",
    "heute", "morgen", "gestern"
}

EMERGENT_TERM_BLACKLIST = {
    "mainz", "koblenz", "trier", "worms", "kaiserslautern", "deutschland",
    "rheinland-pfalz", "unternehmen", "seite", "news", "presse",
    "aktuell", "aktuelle", "meldung", "meldungen", "lesen",
    "weiter", "zurÃ¼ck", "medien", "thema", "themen"
}


class ZIRPWeeklyAnalyzer:
    def __init__(self) -> None:
        if XMLParsedAsHTMLWarning is not None:
            warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
        self.thread_local = threading.local()
        self.session = None

        self.timeout = REQUEST_TIMEOUT
        self.sleep_seconds = SLEEP_SECONDS
        self.max_pages_per_member = MAX_PAGES_PER_MEMBER
        self.max_pages_for_broad_sites = MAX_PAGES_FOR_BROAD_SITES
        self.max_snippets_per_page = MAX_SNIPPETS_PER_PAGE
        self.min_text_length = MIN_TEXT_LENGTH
        self.min_event_score = MIN_EVENT_SCORE
        self.min_decision_score = MIN_DECISION_SCORE
        self.min_article_quality_score = MIN_ARTICLE_QUALITY_SCORE

        self.today = datetime.now()
        self.week_start = self.today - timedelta(days=7)

        self.script_dir = Path(__file__).resolve().parent
        self.output_dir = self.script_dir / "zirp_berichte"
        self.output_dir.mkdir(exist_ok=True)

        self.state_file = self.script_dir / "zirp_member_monitor_state.json"
        self.state = self.load_state()
        self.state_lock = threading.Lock()
        self.keyword_rows_lock = threading.Lock()
        self.article_text_cache = self.state.setdefault("_article_text_cache", {})
        self.discovery_cache = self.state.setdefault("_discovery_cache", {})
        self.signal_judgment_cache = self.state.setdefault("_signal_judgment_cache", {})
        print(
            f"Crawler cache file: {self.state_file.resolve()} "
            f"({len(self.discovery_cache)} discovery entries, {len(self.article_text_cache)} cached pages)"
        )

        self.member_homepages: list[dict[str, Any]] = []
        self.page_hits: list[dict[str, Any]] = []
        self.keyword_rows: list[dict[str, Any]] = []
        self.events: list[dict[str, Any]] = []
        self.source_trace_by_page: dict[tuple[str, str], dict[str, Any]] = {}
        self.new_actor_source_rows: list[dict[str, Any]] = []
        self.global_weighted_counter: Counter = Counter()
        self.crawl_workers = CRAWL_WORKERS

        self.use_openai_editorial = USE_OPENAI_EDITORIAL
        self.openai_model = OPENAI_MODEL
        self.openai_base_url = OPENAI_BASE_URL
        self.openai_timeout = OPENAI_TIMEOUT
        self.report_local_cascade = REPORT_LOCAL_CASCADE
        self.report_draft_model = REPORT_DRAFT_MODEL
        self.report_polish_model = REPORT_POLISH_MODEL
        self.report_openai_model = REPORT_OPENAI_MODEL
        self.report_openai_base_url = REPORT_OPENAI_BASE_URL
        self.report_openai_timeout = REPORT_OPENAI_TIMEOUT
        self.openai_client = self._build_openai_client()
        self.report_openai_client = self._build_report_openai_client()

        self.editorial_frames: dict[str, dict[str, Any]] = {}
        self.report_meta: dict[str, Any] = {}
        self.meeting_optimizer_started_at: float | None = None

        self.member_lookup = {
            m["name"]: {
                "member_id": m.get("member_id"),
                "homepage": m.get("homepage", ""),
                "seed_paths": m.get("seed_paths", []),
                "actor_role": m.get("actor_role", ""),
                "broad_zirp_fields": m.get("broad_zirp_fields", ""),
                "top_static_clusters": m.get("top_static_clusters", ""),
            }
            for m in MEMBERS
        }
        self.actor_source_profiles = ACTOR_SOURCE_PROFILES
        self.member_keyword_radar = MEMBER_KEYWORD_RADAR
        self.actor_type_source_rules = ACTOR_TYPE_SOURCE_RULES
        self.crawler_scoring_config = CRAWLER_SCORING_CONFIG

        self.active_relevance_model = self.build_active_relevance_model()
        self.active_strategic_priority_terms = self.active_relevance_model["strategic_priority_terms"]
        self.active_report_sections = self.active_relevance_model["report_sections"]
        self.active_zirp_themenfelder = self.active_relevance_model["zirp_themenfelder"]
        self.active_noise_terms = self.active_relevance_model["noise_terms"]

    def load_dynamic_filter_terms(self) -> dict[str, set[str]]:
        """Load feedback-approved cached words from dynamic_term_memory.json.

        Supported shapes:
        1. {"trusted_feedback": {"pilot": {...}}, "noise": {"podcast": {...}}}
        2. {"problem": ["fachkrÃ¤fte"], "action": ["pilot"], ...}
        """
        result = {
            "problem": set(),
            "action": set(),
            "asset": set(),
            "trusted_feedback": set(),
            "noise": set(),
            "wirtschaft": set(),
            "technologie": set(),
            "nachhaltigkeit": set(),
            "kultur": set(),
            "wissen": set(),
            "gesellschaft": set(),
        }
        if not USE_DYNAMIC_RELEVANCE_TERMS:
            return result

        path = _resolve_project_path(DYNAMIC_TERM_MEMORY_PATH)
        if not path.exists():
            return result

        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception as exc:
            print(f"Dynamic relevance terms: konnte {path.name} nicht laden ({exc})")
            return result

        def add_term(bucket: str, term: Any) -> None:
            cleaned = str(term or "").strip().lower()
            if cleaned and cleaned not in {"nan", "none"} and len(cleaned) >= 3:
                result.setdefault(bucket, set()).add(cleaned)

        for bucket, value in (data or {}).items():
            bucket_key = str(bucket).strip().lower()
            if bucket_key not in result:
                bucket_key = "trusted_feedback"

            if isinstance(value, list):
                for term in value:
                    add_term(bucket_key, term)
            elif isinstance(value, dict):
                for term, meta in value.items():
                    if isinstance(meta, dict):
                        score = float(meta.get("score", 0) or 0)
                        category = str(meta.get("category", bucket_key) or bucket_key).strip().lower()
                        if score <= -2 or category == "noise" or bucket_key == "noise":
                            add_term("noise", term)
                        elif score >= 2 or bucket_key == "trusted_feedback":
                            add_term("trusted_feedback", term)
                            if category in result and category not in {"trusted_feedback", "noise"}:
                                add_term(category, term)
                        elif category in result:
                            add_term(category, term)
                    else:
                        add_term(bucket_key, term)

        total_dynamic = sum(len(v) for v in result.values())
        if total_dynamic:
            print(f"Dynamic relevance terms: {total_dynamic} aktive Feedback-Begriffe geladen")
        return result

    def build_active_relevance_model(self) -> dict[str, Any]:
        """Merge static ZIRP dictionaries with feedback-approved dynamic words."""
        dynamic = self.load_dynamic_filter_terms()

        strategic_terms = set(STRATEGIC_PRIORITY_TERMS)
        strategic_terms |= dynamic.get("problem", set())
        strategic_terms |= dynamic.get("action", set())
        strategic_terms |= dynamic.get("asset", set())
        strategic_terms |= dynamic.get("trusted_feedback", set())

        report_sections = {
            key: {
                "title": value.get("title", key),
                "keywords": list(dict.fromkeys(list(value.get("keywords", []))))
            }
            for key, value in REPORT_SECTIONS.items()
        }
        report_sections["wirtschaftsentwicklung"]["keywords"].extend(sorted(dynamic.get("problem", set()) | dynamic.get("asset", set()) | dynamic.get("wirtschaft", set())))
        report_sections["versorgung_gesundheit"]["keywords"].extend(sorted(dynamic.get("gesellschaft", set()) | dynamic.get("problem", set())))
        report_sections["kooperationen"]["keywords"].extend(sorted(dynamic.get("action", set()) | dynamic.get("trusted_feedback", set())))
        report_sections["soziales_engagement"]["keywords"].extend(sorted(dynamic.get("gesellschaft", set()) | dynamic.get("kultur", set())))

        zirp_themenfelder = {key: list(value) for key, value in ZIRP_THEMENFELDER.items()}
        for field in ["wirtschaft", "technologie", "nachhaltigkeit", "kultur", "wissen", "gesellschaft"]:
            zirp_themenfelder.setdefault(field, [])
            zirp_themenfelder[field].extend(sorted(dynamic.get(field, set())))

        zirp_themenfelder["wirtschaft"].extend(sorted(dynamic.get("problem", set())))
        zirp_themenfelder["technologie"].extend(sorted(dynamic.get("asset", set()) | dynamic.get("action", set())))
        zirp_themenfelder["wissen"].extend(sorted(dynamic.get("action", set()) | dynamic.get("trusted_feedback", set())))
        zirp_themenfelder["gesellschaft"].extend(sorted(dynamic.get("problem", set())))

        # Deduplicate while preserving order.
        for cfg in report_sections.values():
            cfg["keywords"] = list(dict.fromkeys(str(x).strip().lower() for x in cfg["keywords"] if str(x).strip()))
        for field, terms in zirp_themenfelder.items():
            zirp_themenfelder[field] = list(dict.fromkeys(str(x).strip().lower() for x in terms if str(x).strip()))

        return {
            "strategic_priority_terms": strategic_terms,
            "report_sections": report_sections,
            "zirp_themenfelder": zirp_themenfelder,
            "noise_terms": set(HARD_NOISE_TERMS) | dynamic.get("noise", set()),
            "dynamic_terms": dynamic,
        }

    def strategic_terms_for_run(self) -> set[str]:
        return getattr(self, "active_strategic_priority_terms", STRATEGIC_PRIORITY_TERMS)

    def report_sections_for_run(self) -> dict[str, dict[str, Any]]:
        return getattr(self, "active_report_sections", REPORT_SECTIONS)

    def zirp_themenfelder_for_run(self) -> dict[str, list[str]]:
        return getattr(self, "active_zirp_themenfelder", ZIRP_THEMENFELDER)

    def noise_terms_for_run(self) -> set[str]:
        return getattr(self, "active_noise_terms", set(HARD_NOISE_TERMS))

    @staticmethod
    def bounded_int(value: Any, minimum: int, maximum: int) -> int:
        try:
            number = int(round(float(value)))
        except Exception:
            number = minimum
        return max(minimum, min(maximum, number))

    def checked_member_count(self) -> int:
        return len(self.member_homepages) or len(MEMBERS)

    def dynamic_signal_judge_top_n(self) -> int:
        if not DYNAMIC_LIMITS_ENABLED:
            return SIGNAL_JUDGE_TOP_N
        return self.bounded_int(
            self.checked_member_count() * SIGNAL_JUDGE_MEMBER_RATIO,
            SIGNAL_JUDGE_MIN,
            SIGNAL_JUDGE_MAX,
        )

    def dynamic_narrative_event_limit(self) -> int:
        if not DYNAMIC_LIMITS_ENABLED:
            return MAX_NARRATIVE_EVENTS
        return self.bounded_int(
            self.dynamic_signal_judge_top_n() * NARRATIVE_EVENTS_JUDGE_RATIO,
            NARRATIVE_EVENTS_MIN,
            NARRATIVE_EVENTS_MAX,
        )

    def dynamic_public_scip_card_limit(self, usable_signals: Optional[int] = None) -> int:
        if not DYNAMIC_LIMITS_ENABLED:
            return 6
        if usable_signals is None:
            try:
                usable_signals = len(self.weekly_events())
            except Exception:
                usable_signals = 0
        return self.bounded_int(
            max(int(usable_signals or 0), 1) * PUBLIC_SCIP_SIGNAL_RATIO,
            PUBLIC_SCIP_CARDS_MIN,
            PUBLIC_SCIP_CARDS_MAX,
        )

    def dynamic_scip_prompt_card_limit(self, usable_signals: Optional[int] = None) -> int:
        public_limit = self.dynamic_public_scip_card_limit(usable_signals)
        if not DYNAMIC_LIMITS_ENABLED:
            return min(3, public_limit)
        return self.bounded_int(
            public_limit * 0.5,
            SCIP_PROMPT_CARDS_MIN,
            min(SCIP_PROMPT_CARDS_MAX, public_limit),
        )

    @staticmethod
    def create_session() -> requests.Session:
        if requests is None:
            raise RuntimeError("requests package is required for crawling")
        session = requests.Session()
        session.headers.update({
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0 Safari/537.36"
            )
        })
        return session

    def current_session(self) -> requests.Session:
        session = getattr(self.thread_local, "session", None)
        if session is None:
            session = self.create_session()
            self.thread_local.session = session
        self.session = session
        return session

    def matching_member_config(self, member_name: str, config: dict[str, Any]) -> Any:
        direct = config.get(member_key(member_name))
        if direct:
            return direct
        target = member_key(member_name)
        target_tokens = {tok for tok in self.tokenize(target) if len(tok) >= 4}
        best_key = ""
        best_overlap = 0
        for key in config:
            if not key:
                continue
            if key in target or target in key:
                return config[key]
            key_tokens = {tok for tok in self.tokenize(key) if len(tok) >= 4}
            overlap = len(target_tokens & key_tokens)
            if overlap > best_overlap:
                best_key = key
                best_overlap = overlap
        if best_key and best_overlap >= 2:
            return config[best_key]
        return None

    def member_source_profiles_for(self, member_name: str) -> list[dict[str, Any]]:
        profiles = self.matching_member_config(member_name, self.actor_source_profiles)
        return list(profiles or [])

    def keyword_radar_for_member(self, member_name: str) -> dict[str, Any]:
        return dict(self.matching_member_config(member_name, self.member_keyword_radar) or {})

    def crawler_config_int(self, key: str, default: int) -> int:
        try:
            return int(float(str(self.crawler_scoring_config.get(key.lower(), default)).strip()))
        except Exception:
            return default

    def actor_source_family_for_member(self, member: dict[str, Any]) -> str:
        name = clean_member_name(member.get("name") or "")
        meta = self.member_lookup.get(name, {})
        role = clean_excel_value(member.get("actor_role") or meta.get("actor_role")).lower()
        combined = " ".join(
            [
                name,
                role,
                clean_excel_value(member.get("broad_zirp_fields") or meta.get("broad_zirp_fields")),
                clean_excel_value(member.get("top_static_clusters") or meta.get("top_static_clusters")),
            ]
        ).lower()

        if role == "academic_support" or any(x in combined for x in ["hochschule", "universität", "universitaet", "university", "rptu", "hfm ", "jgu "]):
            return "academic"
        if any(x in combined for x in ["bank", "sparkasse", "debeka", "versicherung", "isb", "lbbw", "bundesbank", "finanz"]):
            return "bank_finance"
        if any(x in combined for x in ["kammer", "verband", "handwerk", "bauwirtschaft", "bwv", "architekten", "ingenieur", "städtetag", "staedtetag", "landkreistag"]):
            return "chamber_association"
        if any(x in combined for x in ["aok", "ikk", "tk ", "kv ", "caritas", "bbt", "gesund", "ärzte", "aerzte", "zahn", "kranken", "pflege", "alexianer"]):
            return "health_social"
        if any(x in combined for x in ["stadtwerke", "pfalzwerke", "westenergie", "entega", "zahnen", "netz", "energie", "wasser", "infrastruktur"]):
            return "energy_infrastructure"
        if any(x in combined for x in ["swr", "zdf", "rpr", "villa musica", "kultur", "musik", "medien"]):
            return "media_culture"
        if any(x in combined for x in ["ministerium", "staatskanzlei", "bundesagentur", "stadt ", "landkreis", "kirche", "verwaltung"]):
            return "state_institution"
        return "company_industry"

    def static_member_fit_score(self, member_name: str, title: str, text: str) -> int:
        radar = self.keyword_radar_for_member(member_name)
        keywords = [
            normalize_keyword
            for normalize_keyword in (
                re.sub(r"\s+", " ", str(term or "").strip().lower())
                for term in radar.get("keywords", [])
            )
            if len(normalize_keyword) >= 3
        ]
        if not keywords:
            return 0
        haystack = normalize_search_text(f"{title}\n{text[:2600]}")
        has_concrete_digital = any(normalize_search_text(term) in haystack for term in CONCRETE_DIGITAL_TERMS)
        hits = []
        generic_hits = []
        for term in keywords[:80]:
            term_norm = normalize_search_text(term)
            if not term_norm or term_norm not in haystack:
                continue
            if term_norm in GENERIC_STATIC_PROFILE_TERMS:
                generic_hits.append(term)
                continue
            if term_norm in {"digitalisierung", "digital", "it", "ki", "daten"} and not has_concrete_digital:
                generic_hits.append(term)
                continue
            if len(term_norm) <= 3 and term_norm not in {"bim", "erp"}:
                generic_hits.append(term)
                continue
            hits.append(term)
        score = len(set(hits)) * 2
        if has_concrete_digital and generic_hits:
            score += min(len(set(generic_hits)), 2)
        elif generic_hits and hits:
            score += 1
        return min(score, 10)

    def actor_seed_paths_for_member(self, member: dict[str, Any]) -> list[str]:
        role = clean_excel_value(member.get("actor_role")).lower()
        paths = list(ACTOR_TYPE_SEED_PATHS.get(role, []))
        family = self.actor_source_family_for_member(member)
        family_rule = self.actor_type_source_rules.get(family, {})
        paths.extend(family_rule.get("seed_paths", []))
        fields = normalize_search_text(member.get("broad_zirp_fields", ""))
        if "technologie" in fields:
            paths.extend(["/innovation", "/digitalisierung", "/forschung"])
        if "nachhaltigkeit" in fields:
            paths.extend(["/nachhaltigkeit", "/klima", "/energie"])
        if "gesellschaft" in fields:
            paths.extend(["/veranstaltungen", "/projekte", "/aktuelles"])
        if "kultur" in fields:
            paths.extend(["/events", "/veranstaltungen", "/programm"])
        return self.deduplicate_list(paths)

    def _build_openai_client(self):
        if not self.use_openai_editorial:
            print("OpenAI editorial: deaktiviert (USE_OPENAI_EDITORIAL=False)")
            return None
        if OpenAI is None:
            print("OpenAI editorial: nicht aktiv (Python-Paket 'openai' ist nicht installiert)")
            return None

        self.load_local_env()
        self.openai_model = os.getenv("OPENAI_MODEL", self.openai_model).strip()
        self.openai_base_url = os.getenv("OPENAI_BASE_URL", self.openai_base_url).strip()
        self.report_draft_model = os.getenv("REPORT_DRAFT_MODEL", self.report_draft_model).strip()
        self.report_polish_model = os.getenv("REPORT_POLISH_MODEL", self.report_polish_model).strip()
        self.report_local_cascade = os.getenv(
            "REPORT_LOCAL_CASCADE",
            "1" if self.report_local_cascade else "0",
        ).strip().lower() not in {"0", "false", "no", "off"}
        api_key = os.getenv("OPENAI_API_KEY", "").strip()
        if self.openai_base_url and not api_key:
            api_key = "lm-studio"
        if not api_key:
            print("OpenAI editorial: nicht aktiv (OPENAI_API_KEY fehlt)")
            return None

        try:
            endpoint = self.openai_base_url or "OpenAI API"
            print(f"OpenAI editorial: aktiv ({self.openai_model}, endpoint={endpoint})")
            return OpenAI(
                api_key=api_key,
                base_url=self.openai_base_url or None,
                timeout=self.openai_timeout,
                max_retries=OPENAI_MAX_RETRIES,
            )
        except Exception as exc:
            print(f"OpenAI editorial: nicht aktiv (Client konnte nicht gebaut werden: {exc})")
            return None

    def _build_report_openai_client(self):
        if not self.use_openai_editorial:
            return None
        if OpenAI is None:
            return None

        self.load_local_env()
        self.report_openai_model = os.getenv("REPORT_OPENAI_MODEL", self.report_openai_model).strip()
        self.report_openai_base_url = os.getenv("REPORT_OPENAI_BASE_URL", self.report_openai_base_url).strip()
        api_key = os.getenv("REPORT_OPENAI_API_KEY", "").strip()
        if not api_key and not self.report_openai_base_url:
            api_key = os.getenv("OPENAI_API_KEY", "").strip()

        if not api_key or (not self.report_openai_base_url and not api_key.startswith("sk-")):
            print("Report editorial: Cloud-Modell nicht aktiv (REPORT_OPENAI_API_KEY fehlt)")
            return None

        try:
            endpoint = self.report_openai_base_url or "OpenAI API"
            print(f"Report editorial: aktiv ({self.report_openai_model}, endpoint={endpoint})")
            return OpenAI(
                api_key=api_key,
                base_url=self.report_openai_base_url or None,
                timeout=self.report_openai_timeout,
                max_retries=0,
            )
        except Exception as exc:
            print(f"Report editorial: nicht aktiv (Client konnte nicht gebaut werden: {exc})")
            return None

    def load_local_env(self) -> None:
        script_dir = Path(__file__).resolve().parent
        for env_path in [
            Path(".env"),
            script_dir / ".env",
            script_dir.parent / ".env",
            self.output_dir / ".env",
        ]:
            if not env_path.exists():
                continue
            try:
                for raw_line in env_path.read_text(encoding="utf-8").splitlines():
                    line = raw_line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    key, value = line.split("=", 1)
                    key = key.strip()
                    value = value.strip().strip('"').strip("'")
                    if key and key not in os.environ:
                        os.environ[key] = value
            except Exception:
                continue

    @staticmethod
    def _event_key(event: dict[str, Any]) -> str:
        return "|".join([
            str(event.get("mitglied", "")).strip().lower(),
            str(event.get("titel", "")).strip().lower(),
            str(event.get("datum", "")).strip(),
            re.sub(r"\s+", " ", str(event.get("snippet", ""))).strip().lower()[:180],
        ])

    def run(self) -> None:
        if RUN_MODE == "report_from_existing_csv":
            print("\n=== ZIRP-BERICHT AUS BESTEHENDER CSV ===\n")
            self.load_hits_from_existing_csv(EXISTING_HITS_CSV)
            self.rebuild_events_from_hits()
            self.deduplicate_events()
            self.apply_zirp_relevance_filter()
            self.apply_signal_judgment()
            self.create_outputs()
            print("\n=== FERTIG ===")
            print(f"Seitenfunde geladen: {len(self.page_hits)}")
            print(f"Ereignisse gebaut: {len(self.events)}")
            print(f"Berichtsrelevante Ereignisse: {len(self.weekly_events())}")
            return

        print("\n=== ZIRP-BERICHT CRAWL ===\n")
        for idx, member in enumerate(MEMBERS, start=1):
            name = member["name"]
            print(f"[{idx}/{len(MEMBERS)}] Mitglied: {name}")

            homepage = self.get_member_homepage(member)
            zirp_detail_url = ""
            if member.get("zirp_member_path"):
                zirp_detail_url = urljoin(ZIRP_BASE, member["zirp_member_path"])

            self.member_homepages.append({
                "mitglied": name,
                "member_id": member.get("member_id"),
                "zirp_detail_url": zirp_detail_url,
                "homepage": homepage or ""
            })

            if not homepage:
                print("   -> Keine externe Homepage gefunden")
                continue

            seed_paths = list(member.get("seed_paths", []) or []) + self.actor_seed_paths_for_member(member)
            pages, discovery_cache_status = self.collect_relevant_pages(
                homepage,
                seed_paths=seed_paths,
                member_name=name,
            )
            self.member_homepages[-1]["discovery_cache_status"] = discovery_cache_status
            self.member_homepages[-1]["actor_source_profile_count"] = len(self.member_source_profiles_for(name))
            print(f"   -> Relevante Artikelseiten: {len(pages)}")

            for hit, event in self.analyze_pages_parallel(name, homepage, pages):
                if hit:
                    self.page_hits.append(hit)
                if event:
                    self.events.append(event)

        discovery_cached = sum(1 for item in self.member_homepages if item.get("discovery_cache_status") == "cached_discovery")
        discovery_fresh = sum(1 for item in self.member_homepages if item.get("discovery_cache_status") == "fresh_discovery")
        if discovery_cached or discovery_fresh:
            print(f"Crawler discovery cache: {discovery_cached} Mitglieder aus Cache, {discovery_fresh} frisch entdeckt")

        cached_pages = sum(1 for hit in self.page_hits if hit.get("cache_status") == "cached_text")
        fresh_pages = sum(1 for hit in self.page_hits if hit.get("cache_status") == "fresh")
        if cached_pages or fresh_pages:
            refresh_note = " (FORCE_CRAWL_REFRESH aktiv)" if FORCE_CRAWL_REFRESH else ""
            print(f"Crawler text cache: {cached_pages} Seiten aus Cache rekonstruiert, {fresh_pages} frisch geladen{refresh_note}")

        self.deduplicate_events()
        self.apply_zirp_relevance_filter()
        self.apply_signal_judgment()
        self.create_outputs()
        self.update_actor_source_profile_workbook()
        self.update_performance_workbook()
        self.save_state()

    def analyze_pages_parallel(self, member_name: str, homepage: str, pages: list[str]):
        if len(pages) <= 1 or self.crawl_workers <= 1:
            for page_url in pages:
                yield self.analyze_page(member_name=member_name, homepage=homepage, page_url=page_url)
                if self.sleep_seconds:
                    time.sleep(self.sleep_seconds)
            return

        worker_count = min(self.crawl_workers, len(pages))
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            future_to_url = {
                executor.submit(self.analyze_page, member_name, homepage, page_url): page_url
                for page_url in pages
            }
            for future in as_completed(future_to_url):
                try:
                    yield future.result()
                except Exception as exc:
                    print(f"   -> Seite Ã¼bersprungen ({future_to_url[future]}): {exc}")

    def load_state(self) -> dict:
        if self.state_file.exists():
            try:
                return json.loads(self.state_file.read_text(encoding="utf-8"))
            except Exception:
                return {}
        return {}

    def save_state(self) -> None:
        self.prune_discovery_cache()
        self.prune_article_text_cache()
        self.state_file.write_text(
            json.dumps(self.state, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        print(
            f"Crawler cache saved: {len(self.discovery_cache)} discovery entries, "
            f"{len(self.article_text_cache)} pages, "
            f"{len(self.signal_judgment_cache)} signal judgments -> {self.state_file.resolve()}"
        )

    def record_discovered_actor_source(
        self,
        *,
        member_name: str,
        source_url: str,
        source_score: int,
        source_type: str = "Discovered source page",
    ) -> None:
        if not member_name or not source_url or source_score < 24:
            return
        existing = {
            self.normalize_url(str(profile.get("url", "")))
            for profile in self.member_source_profiles_for(member_name)
        }
        if self.normalize_url(source_url) in existing:
            return
        key = (member_key(member_name), self.normalize_url(source_url))
        if any((member_key(row.get("Member")), self.normalize_url(row.get("Source URL", ""))) == key for row in self.new_actor_source_rows):
            return
        self.new_actor_source_rows.append({
            "Member": member_name,
            "Source URL": self.normalize_url(source_url),
            "Source Type": source_type,
            "Crawl Priority": "High" if source_score >= 42 else "Medium",
            "Crawler Difficulty": "Medium",
            "Why it matters for SOR": f"Auto-discovered source-quality score {source_score}; contains dated strategic/source-page signals.",
            "Current crawler risk": "Auto-discovered; review after the next crawl.",
            "Recommended crawl strategy": "Crawl as ranked source first, follow dated article cards, publications and event links, then apply normal signal gates.",
            "Follow patterns": "",
            "Extract sections / signals": "",
            "Keyword families": "",
            "Status": "Discovered",
        })

    def update_actor_source_profile_workbook(self) -> None:
        if not self.new_actor_source_rows:
            return
        workbook_path = _resolve_project_path(STATIC_ZIRP_EXCEL_PATH)
        if not workbook_path.exists():
            print(f"Actor source profiles: Workbook nicht gefunden, keine Aktualisierung ({workbook_path})")
            return
        try:
            import openpyxl
        except Exception as exc:
            print(f"Actor source profiles: openpyxl fehlt, keine Workbook-Aktualisierung ({exc})")
            return
        try:
            wb = openpyxl.load_workbook(workbook_path)
            if STATIC_ZIRP_SOURCE_PROFILE_SHEET not in wb.sheetnames:
                ws = wb.create_sheet(STATIC_ZIRP_SOURCE_PROFILE_SHEET)
                headers = list(self.new_actor_source_rows[0].keys())
                ws.append(headers)
            else:
                ws = wb[STATIC_ZIRP_SOURCE_PROFILE_SHEET]
                headers = [cell.value for cell in ws[1] if cell.value]
                if not headers:
                    headers = list(self.new_actor_source_rows[0].keys())
                    ws.append(headers)

            existing_urls = {
                self.normalize_url(str(row[1] or ""))
                for row in ws.iter_rows(min_row=2, values_only=True)
                if len(row) >= 2 and row[1]
            }
            appended = 0
            for row in self.new_actor_source_rows:
                if self.normalize_url(row.get("Source URL", "")) in existing_urls:
                    continue
                ws.append([row.get(header, "") for header in headers])
                existing_urls.add(self.normalize_url(row.get("Source URL", "")))
                appended += 1
            if appended:
                wb.save(workbook_path)
                print(f"Actor source profiles: {appended} neue Quellen in {workbook_path.name}/{STATIC_ZIRP_SOURCE_PROFILE_SHEET} ergänzt")
        except Exception as exc:
            print(f"Actor source profiles: Workbook-Aktualisierung übersprungen ({exc})")

    def source_performance_rows(self) -> list[dict[str, Any]]:
        run_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        events_by_key: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
        for event in self.events:
            events_by_key[(member_key(event.get("mitglied")), self.normalize_url(str(event.get("url", ""))))].append(event)

        grouped: dict[tuple[str, str], dict[str, Any]] = {}
        for hit in self.page_hits:
            member_name = str(hit.get("mitglied", "") or "")
            page_url = self.normalize_url(str(hit.get("seite", "") or ""))
            if not member_name or not page_url:
                continue
            trace = self.source_trace_for_page(member_name, page_url)
            source_url = trace.get("source_url") or page_url
            key = (member_key(member_name), self.normalize_url(source_url))
            bucket = grouped.setdefault(key, {
                "run_date": run_date,
                "member_name": member_name,
                "actor_family": trace.get("actor_family", ""),
                "source_url": self.normalize_url(source_url),
                "source_type": trace.get("source_type", ""),
                "source_status": trace.get("source_status", ""),
                "source_priority": trace.get("source_priority", 0),
                "crawl_success": 0,
                "pages_crawled": 0,
                "articles_found": 0,
                "strong_signals_found": 0,
                "avg_signal_score": 0.0,
                "avg_decision_score": 0.0,
                "matched_other_members": 0,
                "top_signal_title": "",
                "last_content_date": "",
                "error_type": "",
                "next_action": "",
                "_scores": [],
                "_decision_scores": [],
                "_event_members": set(),
            })
            bucket["pages_crawled"] += 1
            bucket["crawl_success"] = 1
            if str(hit.get("page_type", "")) in {"article", "document"}:
                bucket["articles_found"] += 1
            hit_date = str(hit.get("datum", "") or "")
            if hit_date and (not bucket["last_content_date"] or hit_date > bucket["last_content_date"]):
                bucket["last_content_date"] = hit_date
            for event in events_by_key.get((member_key(member_name), page_url), []):
                score = float(event.get("score", 0) or 0)
                decision_score = float(event.get("decision_score", 0) or 0)
                bucket["_scores"].append(score)
                bucket["_decision_scores"].append(decision_score)
                bucket["_event_members"].add(str(event.get("mitglied", "") or ""))
                if decision_score >= 10:
                    bucket["strong_signals_found"] += 1
                if not bucket["top_signal_title"] or decision_score > float(bucket.get("_top_decision_score", -1)):
                    bucket["top_signal_title"] = str(event.get("titel", "") or "")
                    bucket["_top_decision_score"] = decision_score

        rows = []
        for bucket in grouped.values():
            scores = bucket.pop("_scores", [])
            decision_scores = bucket.pop("_decision_scores", [])
            event_members = bucket.pop("_event_members", set())
            bucket.pop("_top_decision_score", None)
            bucket["avg_signal_score"] = round(sum(scores) / len(scores), 2) if scores else 0
            bucket["avg_decision_score"] = round(sum(decision_scores) / len(decision_scores), 2) if decision_scores else 0
            bucket["matched_other_members"] = max(0, len(event_members) - 1)
            if bucket["strong_signals_found"] >= 2 or bucket["avg_decision_score"] >= 14:
                bucket["next_action"] = "boost / crawl again soon"
            elif bucket["pages_crawled"] >= 3 and bucket["strong_signals_found"] == 0:
                bucket["next_action"] = "downgrade / crawl rarely"
            elif bucket["articles_found"] == 0:
                bucket["next_action"] = "review source pattern"
            else:
                bucket["next_action"] = "keep monitoring"
            rows.append(bucket)
        rows.sort(key=lambda row: (-int(row.get("strong_signals_found", 0)), -float(row.get("avg_decision_score", 0)), row.get("member_name", "")))
        return rows

    def run_evaluation_rows(self) -> list[dict[str, Any]]:
        run_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        weekly_urls = {self.normalize_url(str(event.get("url", ""))) for event in self.weekly_events()}
        family_members = {member.get("name", ""): self.actor_source_family_for_member(member) for member in MEMBERS}
        buckets: dict[str, dict[str, Any]] = {}
        total_final_mentions = 0

        def bucket_for(family: str) -> dict[str, Any]:
            return buckets.setdefault(family, {
                "run_date": run_date,
                "actor_type": family,
                "members_checked": 0,
                "sources_crawled": set(),
                "pages_crawled": 0,
                "signals_found": 0,
                "strong_signals": 0,
                "average_score": 0.0,
                "average_decision_score": 0.0,
                "final_dashboard_mentions": 0,
                "share_of_output": 0.0,
                "_scores": [],
                "_decision_scores": [],
            })

        for member in MEMBERS:
            bucket_for(family_members.get(member.get("name", ""), "company_industry"))["members_checked"] += 1

        for hit in self.page_hits:
            member_name = str(hit.get("mitglied", "") or "")
            family = family_members.get(member_name)
            if not family:
                member_meta = dict(self.member_lookup.get(member_name, {}))
                member_meta["name"] = member_name
                family = self.actor_source_family_for_member(member_meta)
            bucket = bucket_for(family)
            trace = self.source_trace_for_page(member_name, str(hit.get("seite", "")))
            bucket["sources_crawled"].add(trace.get("source_url") or hit.get("seite", ""))
            bucket["pages_crawled"] += 1

        for event in self.events:
            member_name = str(event.get("mitglied", "") or "")
            family = family_members.get(member_name)
            if not family:
                member_meta = dict(self.member_lookup.get(member_name, {}))
                member_meta["name"] = member_name
                family = self.actor_source_family_for_member(member_meta)
            bucket = bucket_for(family)
            bucket["signals_found"] += 1
            score = float(event.get("score", 0) or 0)
            decision_score = float(event.get("decision_score", 0) or 0)
            bucket["_scores"].append(score)
            bucket["_decision_scores"].append(decision_score)
            if decision_score >= 10:
                bucket["strong_signals"] += 1
            if self.normalize_url(str(event.get("url", ""))) in weekly_urls:
                bucket["final_dashboard_mentions"] += 1
                total_final_mentions += 1

        rows = []
        for bucket in buckets.values():
            scores = bucket.pop("_scores", [])
            decision_scores = bucket.pop("_decision_scores", [])
            sources = bucket.pop("sources_crawled", set())
            bucket["sources_crawled"] = len(sources)
            bucket["average_score"] = round(sum(scores) / len(scores), 2) if scores else 0
            bucket["average_decision_score"] = round(sum(decision_scores) / len(decision_scores), 2) if decision_scores else 0
            bucket["share_of_output"] = round(bucket["final_dashboard_mentions"] / total_final_mentions, 4) if total_final_mentions else 0
            rows.append(bucket)
        rows.sort(key=lambda row: (-int(row.get("final_dashboard_mentions", 0)), row.get("actor_type", "")))
        return rows

    def update_performance_workbook(self) -> None:
        workbook_path = _resolve_project_path(STATIC_ZIRP_EXCEL_PATH)
        if not workbook_path.exists():
            print(f"Performance workbook: Workbook nicht gefunden ({workbook_path})")
            return
        try:
            import openpyxl
        except Exception as exc:
            print(f"Performance workbook: openpyxl fehlt ({exc})")
            return

        source_rows = self.source_performance_rows()
        evaluation_rows = self.run_evaluation_rows()
        try:
            wb = openpyxl.load_workbook(workbook_path)

            def replace_sheet(title: str, rows: list[dict[str, Any]], headers: list[str]) -> None:
                if title in wb.sheetnames:
                    ws = wb[title]
                    ws.delete_rows(1, ws.max_row)
                else:
                    ws = wb.create_sheet(title)
                ws.append(headers)
                for row in rows:
                    ws.append([row.get(header, "") for header in headers])
                ws.freeze_panes = "A2"
                for cell in ws[1]:
                    cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
                    cell.fill = openpyxl.styles.PatternFill("solid", fgColor="0F172A")
                for col_cells in ws.columns:
                    letter = col_cells[0].column_letter
                    max_len = max(len(str(cell.value or "")) for cell in col_cells[: min(len(col_cells), 40)])
                    ws.column_dimensions[letter].width = min(max(max_len + 2, 12), 60)

            replace_sheet(
                "Source_Performance_Log",
                source_rows,
                [
                    "run_date", "member_name", "actor_family", "source_url", "source_type",
                    "source_status", "source_priority", "crawl_success", "pages_crawled",
                    "articles_found", "strong_signals_found", "avg_signal_score",
                    "avg_decision_score", "matched_other_members", "top_signal_title",
                    "last_content_date", "error_type", "next_action",
                ],
            )
            replace_sheet(
                "Run_Evaluation",
                evaluation_rows,
                [
                    "run_date", "actor_type", "members_checked", "sources_crawled",
                    "pages_crawled", "signals_found", "strong_signals",
                    "average_score", "average_decision_score",
                    "final_dashboard_mentions", "share_of_output",
                ],
            )
            wb.save(workbook_path)
            print(
                f"Performance workbook: {len(source_rows)} Quellen und "
                f"{len(evaluation_rows)} Actor-Type-Zeilen aktualisiert"
            )
        except Exception as exc:
            print(f"Performance workbook: Aktualisierung fehlgeschlagen ({exc})")

    def cache_age_hours(self, iso_value: str) -> float:
        try:
            cached_at = datetime.fromisoformat(str(iso_value))
            return max(0.0, (self.today - cached_at).total_seconds() / 3600)
        except Exception:
            return 999999.0

    def discovery_cache_key(self, homepage: str, seed_paths: Optional[Iterable[str]]) -> str:
        seeds = "|".join(sorted(str(seed).strip() for seed in (seed_paths or []) if str(seed).strip()))
        discovery_version = (
            f"adaptive-v2|article={DISCOVERY_MIN_ARTICLE_LINK_SCORE}|"
            f"profile={DISCOVERY_PROFILE_LINK_SCORE}|source={DISCOVERY_MIN_SOURCE_SCORE}|"
            f"strong={DISCOVERY_STRONG_SOURCE_SCORE}|fallback={DISCOVERY_FALLBACK_PAGE_LIMIT}"
        )
        return self.make_hash(f"{self.normalize_url(homepage)}|{seeds}|{self.page_limit_for_homepage(homepage)}|{discovery_version}")

    def get_cached_discovered_pages(self, homepage: str, seed_paths: Optional[Iterable[str]]) -> Optional[list[str]]:
        if not DISCOVERY_CACHE_ENABLED or FORCE_CRAWL_REFRESH:
            return None
        key = self.discovery_cache_key(homepage, seed_paths)
        cached = self.discovery_cache.get(key)
        if not isinstance(cached, dict):
            return None
        if self.cache_age_hours(str(cached.get("cached_at", ""))) > DISCOVERY_CACHE_REFRESH_HOURS:
            return None
        pages = cached.get("pages")
        if not isinstance(pages, list):
            return None
        valid_pages = [str(page) for page in pages if self.is_valid_http_url(str(page))]
        return valid_pages or None

    def remember_discovered_pages(self, homepage: str, seed_paths: Optional[Iterable[str]], pages: list[str]) -> None:
        if not DISCOVERY_CACHE_ENABLED or not pages:
            return
        key = self.discovery_cache_key(homepage, seed_paths)
        with self.state_lock:
            self.discovery_cache[key] = {
                "cached_at": self.today.isoformat(timespec="seconds"),
                "homepage": self.normalize_url(homepage),
                "seed_paths": list(seed_paths or []),
                "pages": pages,
            }

    def prune_discovery_cache(self) -> None:
        if not DISCOVERY_CACHE_ENABLED:
            return
        max_hours = DISCOVERY_CACHE_KEEP_DAYS * 24
        for key, cached in list(self.discovery_cache.items()):
            if not isinstance(cached, dict):
                self.discovery_cache.pop(key, None)
                continue
            if self.cache_age_hours(str(cached.get("cached_at", ""))) > max_hours:
                self.discovery_cache.pop(key, None)

    def get_cached_article_payload(self, page_url: str) -> Optional[dict[str, Any]]:
        if not ARTICLE_TEXT_CACHE_ENABLED or FORCE_CRAWL_REFRESH:
            return None
        cached = self.article_text_cache.get(page_url)
        if not isinstance(cached, dict):
            return None
        if self.cache_age_hours(str(cached.get("cached_at", ""))) > ARTICLE_TEXT_CACHE_REFRESH_HOURS:
            return None
        required = {"title", "published_at", "article_quality_score", "text"}
        if not required.issubset(cached.keys()):
            return None
        payload = dict(cached)
        payload["cache_status"] = "cached_text"
        return payload

    def remember_article_text_payload(
        self,
        *,
        page_url: str,
        title: str,
        published_at: datetime,
        article_quality_score: int,
        text: str,
        content_hash: str,
    ) -> None:
        if not ARTICLE_TEXT_CACHE_ENABLED or not text:
            return
        with self.state_lock:
            self.article_text_cache[page_url] = {
                "cached_at": self.today.isoformat(timespec="seconds"),
                "title": title,
                "published_at": published_at.strftime("%Y-%m-%d"),
                "article_quality_score": article_quality_score,
                "text": text,
                "content_hash": content_hash,
            }

    def prune_article_text_cache(self) -> None:
        if not ARTICLE_TEXT_CACHE_ENABLED:
            return
        max_hours = ARTICLE_TEXT_CACHE_KEEP_DAYS * 24
        for url, cached in list(self.article_text_cache.items()):
            if not isinstance(cached, dict):
                self.article_text_cache.pop(url, None)
                continue
            if self.cache_age_hours(str(cached.get("cached_at", ""))) > max_hours:
                self.article_text_cache.pop(url, None)

    def load_hits_from_existing_csv(self, filepath: str) -> None:
        df = pd.read_csv(filepath)
        self.page_hits = df.to_dict(orient="records")

    def fetch_html(self, url: str) -> Optional[str]:
        try:
            response = self.current_session().get(url, timeout=self.timeout)
            response.raise_for_status()
            if "text/html" not in response.headers.get("Content-Type", ""):
                return None
            return response.text
        except Exception:
            return None

    def get_member_homepage(self, member: dict[str, Any]) -> Optional[str]:
        homepage = (member.get("homepage") or "").strip()
        if homepage:
            return self.normalize_url(homepage)

        zirp_member_path = (member.get("zirp_member_path") or "").strip()
        if zirp_member_path:
            return self.resolve_member_homepage(urljoin(ZIRP_BASE, zirp_member_path))
        return None

    def resolve_member_homepage(self, zirp_detail_url: str) -> Optional[str]:
        html = self.fetch_html(zirp_detail_url)
        if not html:
            return None

        soup = BeautifulSoup(html, "html.parser")
        external_links = []
        for a in soup.find_all("a", href=True):
            href = a["href"].strip()
            full = urljoin(zirp_detail_url, href)
            if not self.is_valid_http_url(full):
                continue
            if "zirp.de" in urlparse(full).netloc.lower():
                continue
            if any(x in full.lower() for x in ["linkedin.com", "facebook.com", "instagram.com", "youtube.com"]):
                continue
            external_links.append(self.normalize_url(full))

        return external_links[0] if external_links else None

    def profile_seed_urls(self, homepage: str, member_name: str) -> list[str]:
        urls: list[str] = []
        for profile in self.member_source_profiles_for(member_name):
            source_url = self.normalize_url(str(profile.get("url", "")))
            if self.is_valid_http_url(source_url):
                urls.append(source_url)
            for pattern in profile.get("follow_patterns", [])[:12]:
                pattern = str(pattern).strip().strip("/")
                if not pattern:
                    continue
                if pattern.startswith("http"):
                    urls.append(self.normalize_url(pattern))
                else:
                    urls.append(self.normalize_url(urljoin(homepage, pattern)))
        return self.deduplicate_list(urls)

    def sitemap_and_feed_urls(self, homepage: str) -> list[str]:
        urls: list[str] = []
        base = self.normalize_url(homepage).rstrip("/")
        for path in SOURCE_DISCOVERY_PATHS:
            discovery_url = self.normalize_url(urljoin(base + "/", path.lstrip("/")))
            body = self.fetch_text(discovery_url)
            if not body:
                continue
            if discovery_url.endswith("robots.txt"):
                for match in re.findall(r"(?im)^\s*Sitemap:\s*(\S+)", body):
                    urls.append(self.normalize_url(match))
                continue
            for match in re.findall(r"https?://[^\s<>'\"]+", body):
                url = self.normalize_url(match)
                if self.same_domain(homepage, url) and any(term in url.lower() for term in POSITIVE_LINK_PATTERNS + SOURCE_QUALITY_TERMS):
                    urls.append(url)
            for loc in re.findall(r"<loc>\s*([^<]+)\s*</loc>", body, flags=re.I):
                url = self.normalize_url(html.unescape(loc))
                if self.same_domain(homepage, url) and any(term in url.lower() for term in POSITIVE_LINK_PATTERNS + SOURCE_QUALITY_TERMS):
                    urls.append(url)
            for href in re.findall(r"<link>\s*([^<]+)\s*</link>", body, flags=re.I):
                url = self.normalize_url(html.unescape(href))
                if self.same_domain(homepage, url):
                    urls.append(url)
        return self.deduplicate_list(urls)[:40]

    def fetch_text(self, url: str) -> Optional[str]:
        try:
            response = self.current_session().get(url, timeout=self.timeout)
            response.raise_for_status()
            content_type = response.headers.get("Content-Type", "").lower()
            if not any(kind in content_type for kind in ["text/", "xml", "rss", "atom"]):
                return None
            return response.text
        except Exception:
            return None

    def is_document_url(self, url: str) -> bool:
        path = urlparse(url).path.lower()
        return path.endswith(DOCUMENT_LINK_EXTENSIONS)

    def document_title_from_url(self, url: str) -> str:
        name = unquote(Path(urlparse(url).path).name or "Dokument")
        name = re.sub(r"\.(pdf|docx)$", "", name, flags=re.I)
        name = re.sub(r"[-_]+", " ", name)
        return re.sub(r"\s+", " ", name).strip() or "Dokument"

    def extract_date_from_text(self, text: str) -> Optional[datetime]:
        patterns = [
            r"\b(\d{1,2})\.(\d{1,2})\.(20\d{2})\b",
            r"\b(20\d{2})-(\d{1,2})-(\d{1,2})\b",
            r"\b(20\d{2})/(\d{1,2})/(\d{1,2})\b",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if not match:
                continue
            try:
                groups = match.groups()
                if pattern.startswith(r"\b(\d"):
                    day, month, year = groups
                else:
                    year, month, day = groups
                return datetime(int(year), int(month), int(day))
            except Exception:
                continue
        year_match = re.search(r"\b(20(2[5-9]|3\d))\b", text)
        if year_match:
            try:
                return datetime(int(year_match.group(1)), 1, 1)
            except Exception:
                return None
        return None

    def fetch_binary(self, url: str) -> tuple[Optional[bytes], str]:
        try:
            response = self.current_session().get(url, timeout=max(self.timeout, 8))
            response.raise_for_status()
            content_type = response.headers.get("Content-Type", "").lower()
            if len(response.content or b"") > 8_000_000:
                return None, content_type
            return response.content, content_type
        except Exception:
            return None, ""

    def fetch_document_payload(self, url: str) -> Optional[dict[str, Any]]:
        data, content_type = self.fetch_binary(url)
        if not data:
            return None
        path = urlparse(url).path.lower()
        text = ""
        try:
            if path.endswith(".pdf") or "pdf" in content_type:
                try:
                    from pypdf import PdfReader
                except Exception:
                    return None
                reader = PdfReader(io.BytesIO(data))
                pages = []
                for page in reader.pages[:8]:
                    pages.append(page.extract_text() or "")
                text = "\n".join(pages)
            elif path.endswith(".docx") or "wordprocessingml" in content_type:
                if Document is None:
                    return None
                doc = Document(io.BytesIO(data))
                text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            else:
                return None
        except Exception:
            return None

        text = re.sub(r"\s+\n", "\n", text)
        text = re.sub(r"[ \t]+", " ", text).strip()
        if len(text) < self.min_text_length:
            return None
        title = self.document_title_from_url(url)
        published_at = self.extract_date_from_text(f"{title}\n{text[:4000]}") or self.today
        quality = min(max(len(text) // 400, 4), 18)
        if any(term in f"{url} {title}".lower() for term in ["studie", "bericht", "publikation", "position", "jahresbericht", "geschaeftsbericht", "geschäftsbericht"]):
            quality += self.crawler_config_int("pdf_docx_link_bonus", 8)
        return {
            "title": title,
            "published_at": published_at,
            "article_quality_score": quality,
            "text": text,
            "content_hash": self.make_hash(text),
        }

    def source_profile_for_url(self, member_name: str, url: str) -> dict[str, Any]:
        profiles = self.member_source_profiles_for(member_name)
        url_norm = self.normalize_url(url)
        for profile in profiles:
            profile_url = self.normalize_url(str(profile.get("url", "")))
            if profile_url and (url_norm == profile_url or url_norm.startswith(profile_url.rstrip("/") + "/")):
                return profile
        return {}

    def source_quality_score(self, source_url: str, soup: BeautifulSoup, profile: Optional[dict[str, Any]] = None, member_name: str = "") -> int:
        title = self.extract_title(soup, source_url).lower()
        text = self.extract_clean_text(soup)
        links = [
            self.normalize_url(urljoin(source_url, a.get("href", "")))
            for a in soup.find_all("a", href=True)
        ]
        score = 0
        if re.search(r"\b20(2[4-9]|3\d)\b", text):
            score += 16
        if re.search(r"\b\d{1,2}\.\d{1,2}\.\d{4}\b", text):
            score += 18
        if any(term in f"{title} {text[:2500]}".lower() for term in SOURCE_QUALITY_TERMS):
            score += 22
        if len([link for link in links if self.same_domain(source_url, link)]) >= 12:
            score += 10
        if any(link.lower().endswith((".pdf", ".docx")) for link in links):
            score += 12
        if any(term in title for term in SOURCE_QUALITY_TERMS):
            score += 10
        if member_name:
            member_meta = dict(self.member_lookup.get(member_name, {}))
            member_meta["name"] = member_name
            family = self.actor_source_family_for_member(member_meta)
            family_rule = self.actor_type_source_rules.get(family, {})
            if family != "academic":
                score += int(family_rule.get("priority_boost", 0) or 0)
            family_terms = [str(term).lower() for term in family_rule.get("source_quality_terms", [])]
            if family_terms and any(term in f"{source_url} {title} {text[:2500]}".lower() for term in family_terms):
                score += 12
        if len(text) < 500:
            score -= 18
        if any(term in source_url.lower() for term in ["datenschutz", "impressum", "kontakt", "login", "suche"]):
            score -= 30
        if profile:
            score += min(int(profile.get("priority", 0)) // 12, 9)
            score += SOURCE_DIFFICULTY_BONUS.get(str(profile.get("difficulty", "")).lower(), 0)
        return score

    def pagination_urls(self, source_url: str, member_name: str = "") -> list[str]:
        urls = []
        base = source_url.rstrip("/")
        patterns = list(PAGINATION_PATTERNS)
        if member_name:
            member_meta = dict(self.member_lookup.get(member_name, {}))
            member_meta["name"] = member_name
            family = self.actor_source_family_for_member(member_meta)
            patterns.extend(self.actor_type_source_rules.get(family, {}).get("pagination_patterns", []))
        for n in range(2, 6):
            for pattern in patterns:
                suffix = pattern.format(n=n, offset=(n - 1) * 10)
                if suffix.startswith("/"):
                    urls.append(self.normalize_url(base + suffix))
                elif "?" in source_url:
                    urls.append(self.normalize_url(source_url + "&" + suffix.lstrip("?")))
                else:
                    urls.append(self.normalize_url(source_url + suffix))
        return self.deduplicate_list(urls)

    def adaptive_page_limit(
        self,
        *,
        hard_limit: int,
        strong_candidate_count: int,
        strong_source_count: int,
        has_profile_sources: bool,
    ) -> int:
        fallback_page_limit = self.crawler_config_int("discovery_fallback_page_limit", DISCOVERY_FALLBACK_PAGE_LIMIT)
        if strong_candidate_count <= 0:
            return min(hard_limit, max(1, fallback_page_limit + min(strong_source_count, 2)))
        target = max(fallback_page_limit, strong_candidate_count)
        if has_profile_sources:
            target += 2
        if strong_source_count:
            target += min(strong_source_count, 3)
        return min(hard_limit, target)

    def collect_relevant_pages(
        self,
        homepage: str,
        seed_paths: Optional[Iterable[str]] = None,
        member_name: str = "",
    ) -> tuple[list[str], str]:
        profile_urls = self.profile_seed_urls(homepage, member_name) if member_name else []
        cache_seed_paths = list(seed_paths or []) + profile_urls
        cached_pages = self.get_cached_discovered_pages(homepage, cache_seed_paths)
        if cached_pages:
            return cached_pages, "cached_discovery"

        page_limit = self.page_limit_for_homepage(homepage)
        hub_pages = list(profile_urls)
        if seed_paths:
            for seed in seed_paths:
                seed_text = str(seed).strip()
                if seed_text.startswith("http"):
                    hub_pages.append(self.normalize_url(seed_text))
                else:
                    hub_pages.append(self.normalize_url(urljoin(homepage, seed_text)))
        for seed in DISCOVERY_SEED_PATHS:
            hub_pages.append(self.normalize_url(urljoin(homepage, seed)))
        hub_pages.append(self.normalize_url(homepage))
        hub_pages.extend(self.sitemap_and_feed_urls(homepage))

        candidates: dict[str, int] = {}
        ranked_sources: list[tuple[int, str]] = []

        for hub_url in self.deduplicate_list(hub_pages)[:32]:
            if not self.is_valid_http_url(hub_url):
                continue
            if not self.same_domain(homepage, hub_url):
                continue
            html = self.fetch_html(hub_url)
            if not html:
                continue

            soup = BeautifulSoup(html, "html.parser")
            self.remove_non_content_nodes(soup)
            profile = self.source_profile_for_url(member_name, hub_url) if member_name else {}
            source_score = self.source_quality_score(hub_url, soup, profile, member_name=member_name)
            if source_score >= self.crawler_config_int("discovery_min_source_score", DISCOVERY_MIN_SOURCE_SCORE):
                ranked_sources.append((source_score, hub_url))
                if member_name:
                    self.record_discovered_actor_source(
                        member_name=member_name,
                        source_url=hub_url,
                        source_score=source_score,
                    )

            for a in soup.find_all("a", href=True):
                href = a["href"].strip()
                full_url = self.normalize_url(urljoin(hub_url, href))
                if not self.is_valid_http_url(full_url):
                    continue
                if not self.same_domain(homepage, full_url):
                    continue

                anchor_text = a.get_text(" ", strip=True)
                score = self.article_link_score(full_url, anchor_text) + max(0, source_score // 12)
                minimum_score = (
                    self.crawler_config_int("discovery_profile_link_score", DISCOVERY_PROFILE_LINK_SCORE)
                    if profile
                    else self.crawler_config_int("discovery_min_article_link_score", DISCOVERY_MIN_ARTICLE_LINK_SCORE)
                )
                if score >= minimum_score:
                    candidates[full_url] = max(candidates.get(full_url, 0), score)
                    self.remember_source_trace(
                        member_name=member_name,
                        page_url=full_url,
                        source_url=hub_url,
                        source_score=source_score,
                        link_score=score,
                        profile=profile,
                    )

        if not candidates:
            source_pages = [url for _, url in sorted(ranked_sources, key=lambda x: (-x[0], x[1]))]
            fallback_limit = self.adaptive_page_limit(
                hard_limit=page_limit,
                strong_candidate_count=0,
                strong_source_count=len(source_pages),
                has_profile_sources=bool(profile_urls),
            )
            pages = self.deduplicate_list(source_pages + profile_urls + [self.normalize_url(homepage)])[:fallback_limit]
            for page_url in pages:
                self.remember_source_trace(
                    member_name=member_name,
                    page_url=page_url,
                    source_url=page_url,
                    source_score=0,
                    link_score=0,
                    profile=self.source_profile_for_url(member_name, page_url) if member_name else {},
                )
            self.remember_discovered_pages(homepage, cache_seed_paths, pages)
            return pages, "fresh_discovery"

        ranked = sorted(candidates.items(), key=lambda x: (-x[1], x[0]))
        source_pages = [url for _, url in sorted(ranked_sources, key=lambda x: (-x[0], x[1]))]
        expanded_sources = []
        for source_score, source_url in sorted(ranked_sources, key=lambda x: (-x[0], x[1]))[:3]:
            if source_score < self.crawler_config_int("discovery_strong_source_score", DISCOVERY_STRONG_SOURCE_SCORE):
                continue
            expanded_sources.extend(self.pagination_urls(source_url, member_name=member_name)[:8])
        adaptive_limit = self.adaptive_page_limit(
            hard_limit=page_limit,
            strong_candidate_count=len(ranked),
            strong_source_count=len([score for score, _ in ranked_sources if score >= self.crawler_config_int("discovery_strong_source_score", DISCOVERY_STRONG_SOURCE_SCORE)]),
            has_profile_sources=bool(profile_urls),
        )
        pages = self.deduplicate_list([url for url, _ in ranked] + source_pages + expanded_sources)[:adaptive_limit]
        for page_url in pages:
            if (member_key(member_name), self.normalize_url(page_url)) not in self.source_trace_by_page:
                self.remember_source_trace(
                    member_name=member_name,
                    page_url=page_url,
                    source_url=page_url,
                    source_score=0,
                    link_score=0,
                    profile=self.source_profile_for_url(member_name, page_url) if member_name else {},
                )
        self.remember_discovered_pages(homepage, cache_seed_paths, pages)
        return pages, "fresh_discovery"

    def page_limit_for_homepage(self, homepage: str) -> int:
        domain = urlparse(homepage).netloc.replace("www.", "").lower()
        if domain in BROAD_SITE_DOMAINS:
            return min(self.max_pages_per_member, self.max_pages_for_broad_sites)
        return self.max_pages_per_member

    def article_link_score(self, url: str, anchor_text: str) -> int:
        text = f"{url} {anchor_text}".lower()
        anchor_norm = re.sub(r"\s+", " ", anchor_text.strip().lower())

        if any(signal in text for signal in NEGATIVE_LINK_PATTERNS):
            return -10
        if any(signal in text for signal in self.noise_terms_for_run()):
            return -10
        if anchor_norm in GENERIC_ANCHORS:
            return -2
        if any(x in text for x in ["kontakt", "impressum", "datenschutz", "karriere", "jobs", "overview", "service"]):
            return -10

        score = 0
        if self.is_document_url(url):
            score += self.crawler_config_int("pdf_docx_link_bonus", 8)
            if any(p in text for p in ["studie", "bericht", "publikation", "positionspapier", "position", "jahresbericht", "geschäftsbericht", "geschaeftsbericht", "förderprogramm", "foerderprogramm"]):
                score += 4
        if re.search(r"/20\d{2}([/-]|$)", url.lower()):
            score += 6
        if re.search(r"\b\d{4}-\d{2}-\d{2}\b", url.lower()):
            score += 5
        if any(p in url.lower() for p in ["/news/", "/presse/", "/aktuelles/", "/meldung/", "/beitrag/", "/artikel/", "/story/", "/press-release/", "/media/"]):
            score += 4
        if any(p in url.lower() for p in ["/ratgeber/", "/podcast/", "/video/", "/veranstaltungen/", "/events/"]):
            score -= 5
        if any(p in text for p in POSITIVE_LINK_PATTERNS):
            score += 2
        if any(p in text for p in ["kooperation", "projekt", "investition", "strategie", "transfer", "digitalisierung", "versorgung", "fachkrÃ¤fte", "fachkraefte"]):
            score += 2
        wc = len(anchor_text.split())
        if 4 <= wc <= 20:
            score += 2
        elif wc > 30 or wc <= 1:
            score -= 2
        return score

    def remember_source_trace(
        self,
        *,
        member_name: str,
        page_url: str,
        source_url: str,
        source_score: int,
        link_score: int,
        profile: Optional[dict[str, Any]],
    ) -> None:
        if not member_name or not page_url:
            return
        member_meta = dict(self.member_lookup.get(member_name, {}))
        member_meta["name"] = member_name
        self.source_trace_by_page[(member_key(member_name), self.normalize_url(page_url))] = {
            "member": member_name,
            "page_url": self.normalize_url(page_url),
            "source_url": self.normalize_url(source_url),
            "source_score": int(source_score or 0),
            "link_score": int(link_score or 0),
            "source_type": str((profile or {}).get("source_type", "")) or "discovered",
            "source_priority": int((profile or {}).get("priority", 0) or 0),
            "source_status": str((profile or {}).get("status", "")),
            "actor_family": self.actor_source_family_for_member(member_meta),
        }

    def source_trace_for_page(self, member_name: str, page_url: str) -> dict[str, Any]:
        trace = self.source_trace_by_page.get((member_key(member_name), self.normalize_url(page_url)))
        if trace:
            return dict(trace)
        profile = self.source_profile_for_url(member_name, page_url)
        member_meta = dict(self.member_lookup.get(member_name, {}))
        member_meta["name"] = member_name
        return {
            "member": member_name,
            "page_url": self.normalize_url(page_url),
            "source_url": self.normalize_url(str(profile.get("url", "")) or page_url),
            "source_score": 0,
            "link_score": 0,
            "source_type": str(profile.get("source_type", "")) or "direct_or_cached",
            "source_priority": int(profile.get("priority", 0) or 0),
            "source_status": str(profile.get("status", "")),
            "actor_family": self.actor_source_family_for_member(member_meta),
        }

    def analyze_page(self, member_name: str, homepage: str, page_url: str):
        cache_status = "fresh"
        cached_payload = self.get_cached_article_payload(page_url)
        if cached_payload:
            title = str(cached_payload.get("title", "")).strip()
            try:
                published_at = datetime.strptime(str(cached_payload.get("published_at", "")), "%Y-%m-%d")
            except Exception:
                published_at = None
            article_quality_score = int(cached_payload.get("article_quality_score", 0) or 0)
            text = str(cached_payload.get("text", ""))
            content_hash = str(cached_payload.get("content_hash", "")) or self.make_hash(text)
            cache_status = "cached_text"
        else:
            soup = None
            if self.is_document_url(page_url):
                document_payload = self.fetch_document_payload(page_url)
                if not document_payload:
                    return None, None
                title = str(document_payload.get("title", "")).strip()
                published_at = document_payload.get("published_at")
                article_quality_score = int(document_payload.get("article_quality_score", 0) or 0)
                text = str(document_payload.get("text", ""))
                content_hash = str(document_payload.get("content_hash", "")) or self.make_hash(text)
            else:
                html = self.fetch_html(page_url)
                if not html:
                    return None, None

                soup = BeautifulSoup(html, "html.parser")
                title = self.extract_title(soup, page_url)
                published_at = self.extract_date(soup, "")

                article_root, article_quality_score = self.select_article_root(soup)
                text = self.extract_clean_text(article_root)
                content_hash = self.make_hash(text)

        if len(text) < self.min_text_length:
            return None, None

        if not published_at and not cached_payload and soup is not None:
            published_at = self.extract_date(soup, text)
        if not published_at or published_at > self.today:
            return None, None

        snippets = self.extract_snippets(text, title=title)
        main_snippet = snippets[0] if snippets else self.first_meaningful_paragraph(text)

        section_matches = self.detect_report_sections(text, title, main_snippet)
        if not any(section_matches.values()) and not (self.openai_client is not None and MODEL_DECIDES_RELEVANCE):
            return None, None

        themes = self.detect_zirp_themenfelder(f"{title}\n{text}")
        terms = self.find_relevant_terms(f"{title}\n{text}")
        direct_member_score = self.direct_member_signal_score(member_name, title, text, page_url)
        regional_score = self.regional_relevance_score(title, text, member_name)
        static_member_fit_score = self.static_member_fit_score(member_name, title, text)
        geo_penalty = self.geography_mismatch_penalty(title, text)
        page_type = self.classify_page_type(page_url, title, text, article_quality_score)

        with self.state_lock:
            is_updated = self.state.get(page_url) != content_hash
            self.state[page_url] = content_hash
        if cache_status == "cached_text":
            is_updated = False
        else:
            self.remember_article_text_payload(
                page_url=page_url,
                title=title,
                published_at=published_at,
                article_quality_score=article_quality_score,
                text=text,
                content_hash=content_hash,
            )

        member_meta = self.member_lookup.get(member_name, {})
        hit = {
            "member_id": member_meta.get("member_id"),
            "mitglied": member_name,
            "homepage": homepage,
            "seite": page_url,
            "titel": title,
            "datum": published_at.strftime("%Y-%m-%d"),
            "is_updated": is_updated,
            "page_type": page_type,
            "article_quality_score": article_quality_score,
            "direct_member_score": direct_member_score,
            "regional_score": regional_score,
            "static_member_fit_score": static_member_fit_score,
            "geo_penalty": geo_penalty,
            "cache_status": cache_status,
            "wirtschaftsentwicklung": ", ".join(sorted(section_matches["wirtschaftsentwicklung"])),
            "versorgung_gesundheit": ", ".join(sorted(section_matches["versorgung_gesundheit"])),
            "fuehrungswechsel": ", ".join(sorted(section_matches["fuehrungswechsel"])),
            "soziales_engagement": ", ".join(sorted(section_matches["soziales_engagement"])),
            "kooperationen": ", ".join(sorted(section_matches["kooperationen"])),
            "zirp_themenfelder": ", ".join(sorted(themes)),
            "snippets": " | ".join(snippets),
        }

        self.register_keyword_rows(
            member_name=member_name,
            page_url=page_url,
            published_at=published_at,
            terms=terms,
            section_matches=section_matches,
            themes=themes,
            snippets=snippets
        )

        if self.is_homepage_like(homepage, page_url):
            return hit, None
        if article_quality_score < self.min_article_quality_score:
            return hit, None
        if not self.has_valid_title_and_snippet(title, main_snippet):
            return hit, None
        if self.is_excludable_content(title, main_snippet, page_url, page_type):
            return hit, None
        if (
            not (self.openai_client is not None and MODEL_DECIDES_RELEVANCE)
            and not self.context_item_is_relevant(
            member_name=member_name,
            title=title,
            text=text,
            url=page_url,
            direct_member_score=direct_member_score,
            regional_score=regional_score,
            )
        ):
            return hit, None
        if (
            not (self.openai_client is not None and MODEL_DECIDES_RELEVANCE)
            and not self.is_strategically_relevant_event(title, main_snippet, page_url, direct_member_score, regional_score)
        ):
            return hit, None

        event = self.build_event(
            member_name=member_name,
            page_url=page_url,
            title=title,
            published_at=published_at,
            snippet=main_snippet,
            section_matches=section_matches,
            themes=themes,
            terms=terms,
            is_updated=is_updated,
            direct_member_score=direct_member_score,
            regional_score=regional_score,
            article_quality_score=article_quality_score,
            static_member_fit_score=static_member_fit_score,
        )

        event["score"] -= geo_penalty
        event["decision_score"] -= min(geo_penalty, 4)

        if self.openai_client is not None and MODEL_DECIDES_RELEVANCE:
            if not self.model_relevance_gate(event):
                return hit, None
        else:
            if not self.fallback_rule_relevance_gate(event):
                return hit, None

        return hit, event

    def rebuild_events_from_hits(self) -> None:
        self.events = []

        for hit in self.page_hits:
            try:
                published_at = datetime.strptime(str(hit["datum"]), "%Y-%m-%d")
            except Exception:
                continue

            if published_at > self.today:
                continue

            title = str(hit.get("titel", "")).strip()
            page_url = str(hit.get("seite", "")).strip()
            member_name = str(hit.get("mitglied", "")).strip()
            snippet_blob = str(hit.get("snippets", "")).strip()
            page_type = str(hit.get("page_type", "article")).strip().lower() or "article"
            article_quality_score = int(hit.get("article_quality_score", 0) or 0)
            direct_member_score = int(hit.get("direct_member_score", 0) or 0)
            regional_score = int(hit.get("regional_score", 0) or 0)
            static_member_fit_score = int(hit.get("static_member_fit_score", 0) or 0)
            geo_penalty = int(hit.get("geo_penalty", 0) or 0)

            snippets = [s.strip() for s in snippet_blob.split(" | ") if s.strip()]
            combined_snippet = " ".join(snippets[:2]).strip()
            main_snippet = combined_snippet if combined_snippet else (snippets[0] if snippets else "")

            section_matches = {
                "wirtschaftsentwicklung": self.csv_field_to_set(hit.get("wirtschaftsentwicklung", "")),
                "versorgung_gesundheit": self.csv_field_to_set(hit.get("versorgung_gesundheit", "")),
                "fuehrungswechsel": self.csv_field_to_set(hit.get("fuehrungswechsel", "")),
                "soziales_engagement": self.csv_field_to_set(hit.get("soziales_engagement", "")),
                "kooperationen": self.csv_field_to_set(hit.get("kooperationen", "")),
            }
            themes = self.csv_field_to_set(hit.get("zirp_themenfelder", ""))
            terms = self.find_relevant_terms(f"{title}\n{main_snippet}")

            self.register_keyword_rows(
                member_name=member_name,
                page_url=page_url,
                published_at=published_at,
                terms=terms,
                section_matches=section_matches,
                themes=themes,
                snippets=snippets[:2]
            )

            if page_type != "article":
                continue
            if article_quality_score < self.min_article_quality_score:
                continue
            if not title or len(title) < 5:
                continue
            if title.strip().lower() in {"homepage", "startseite"}:
                continue
            if len(main_snippet.split()) < 8:
                continue
            if self.is_excludable_content(title, main_snippet, page_url, page_type):
                continue

            event = self.build_event(
                member_name=member_name,
                page_url=page_url,
                title=title,
                published_at=published_at,
                snippet=main_snippet,
                section_matches=section_matches,
                themes=themes,
                terms=terms,
                is_updated=self.parse_bool(hit.get("is_updated", True)),
                direct_member_score=direct_member_score,
                regional_score=regional_score,
                article_quality_score=article_quality_score,
                static_member_fit_score=static_member_fit_score,
            )
            event["score"] -= geo_penalty
            event["decision_score"] -= min(geo_penalty, 4)

            if self.openai_client is not None and MODEL_DECIDES_RELEVANCE:
                if not self.model_relevance_gate(event):
                    continue
            else:
                if not self.fallback_rule_relevance_gate(event):
                    continue

            self.events.append(event)

    def extract_title(self, soup: BeautifulSoup, fallback_url: str) -> str:
        selectors = [
            ("meta", {"property": "og:title"}, "content"),
            ("meta", {"name": "twitter:title"}, "content"),
            ("meta", {"name": "title"}, "content"),
        ]
        for tag_name, attrs, attr_name in selectors:
            node = soup.find(tag_name, attrs=attrs)
            if node and node.get(attr_name, "").strip():
                return self.clean_title(node.get(attr_name, ""))

        for candidate in [soup.find("h1"), soup.find("title")]:
            if candidate is not None:
                value = candidate.get_text(" ", strip=True)
                if value:
                    return self.clean_title(value)
        return fallback_url

    def clean_title(self, title: str) -> str:
        title = re.sub(r"\s+", " ", title).strip()
        title = re.sub(r"\s+[|Â·Â»\-â€“]\s+.*$", "", title).strip()
        return title[:220]

    def remove_non_content_nodes(self, root: Tag) -> None:
        if root is None:
            return

        for comment in list(root.find_all(string=lambda s: isinstance(s, Comment))):
            comment.extract()

        for tag in list(root.find_all([
            "script", "style", "noscript", "svg", "form", "nav",
            "footer", "header", "aside", "template", "iframe"
        ])):
            try:
                tag.decompose()
            except Exception:
                continue

        for tag in list(root.find_all(True)):
            if tag is None or not isinstance(tag, Tag):
                continue

            try:
                tag_id = tag.get("id") or ""
                tag_classes = tag.get("class") or []
                if not isinstance(tag_classes, list):
                    tag_classes = [str(tag_classes)]
                tag_role = tag.get("role") or ""
                tag_aria = tag.get("aria-label") or ""

                attrs = " ".join([
                    str(tag_id),
                    " ".join(str(x) for x in tag_classes),
                    str(tag_role),
                    str(tag_aria),
                ]).lower()

                style = re.sub(r"\s+", "", (tag.get("style") or "").lower())
                hidden_attr = tag.has_attr("hidden") or tag.get("aria-hidden") == "true"

                if hidden_attr or any(p in style for p in HIDDEN_STYLE_PATTERNS):
                    tag.decompose()
                    continue

                if any(hint in attrs for hint in GENERIC_CONTAINER_HINTS) and tag.name in {"div", "section", "aside"}:
                    if self.text_len(tag.get_text(" ", strip=True)) < 120:
                        tag.decompose()
            except Exception:
                continue

    def select_article_root(self, soup: BeautifulSoup) -> tuple[Tag, int]:
        soup = BeautifulSoup(str(soup), "html.parser")
        self.remove_non_content_nodes(soup)

        candidates: list[tuple[int, Tag]] = []
        selectors = [
            "article",
            "main",
            "[role='main']",
            ".news-detail",
            ".presse-detail",
            ".article",
            ".post",
            ".entry-content",
            ".content",
            ".content-wrapper",
            ".main-content",
        ]

        for selector in selectors:
            for node in soup.select(selector):
                score = self.content_root_score(node)
                candidates.append((score + 5, node))

        for node in soup.find_all(["article", "main", "section", "div"], limit=250):
            score = self.content_root_score(node)
            if score >= 6:
                candidates.append((score, node))

        if not candidates:
            return soup, 0

        best_score, best_node = max(candidates, key=lambda x: x[0])
        return best_node, best_score

    def content_root_score(self, node: Tag) -> int:
        text_blocks = []
        p_count = 0
        li_count = 0
        heading_count = 0
        link_text_chars = 0
        total_text_chars = 0
        date_hits = 0

        for el in node.find_all(["h1", "h2", "h3", "p", "li", "time"]):
            txt = re.sub(r"\s+", " ", el.get_text(" ", strip=True)).strip()
            if not txt:
                continue
            total_text_chars += len(txt)
            if el.name == "p":
                p_count += 1
                text_blocks.append(txt)
            elif el.name == "li":
                li_count += 1
            elif el.name in {"h1", "h2", "h3"}:
                heading_count += 1
            elif el.name == "time":
                date_hits += 1

        for a in node.find_all("a"):
            link_text_chars += len(a.get_text(" ", strip=True))

        body_text_chars = sum(len(t) for t in text_blocks)
        avg_p_len = body_text_chars / p_count if p_count else 0
        link_density = (link_text_chars / total_text_chars) if total_text_chars else 1.0

        score = 0
        score += min(body_text_chars // 220, 10)
        score += min(p_count, 8)
        score += 2 if avg_p_len >= 90 else 0
        score += 2 if heading_count >= 1 else 0
        score += 1 if date_hits >= 1 else 0
        score -= 3 if li_count > p_count * 2 and li_count >= 10 else 0
        score -= 4 if link_density > 0.55 else 0
        score -= 2 if body_text_chars < 250 else 0
        return int(score)

    def extract_clean_text(self, root: Tag) -> str:
        self.remove_non_content_nodes(root)

        blocks: list[str] = []
        for el in root.find_all(["h1", "h2", "h3", "p", "blockquote"]):
            text = re.sub(r"\s+", " ", el.get_text(" ", strip=True)).strip()
            if not text:
                continue
            if self.is_boilerplate_text(text):
                continue
            if el.name == "p" and len(text.split()) < 8:
                continue
            if el.name in {"h1", "h2", "h3"} and len(text.split()) < 3:
                continue
            blocks.append(text)

        text = "\n".join(self.deduplicate_list(blocks))
        text = self.strip_invisible_unicode(text)
        return re.sub(r"\n{3,}", "\n\n", text).strip()

    def strip_invisible_unicode(self, text: str) -> str:
        invisible_chars = [
            "\u200b", "\u200c", "\u200d", "\ufeff",
            "\u202a", "\u202b", "\u202c", "\u202d", "\u202e",
            "\u2066", "\u2067", "\u2068", "\u2069",
        ]
        for ch in invisible_chars:
            text = text.replace(ch, "")
        return text

    def is_boilerplate_text(self, text: str) -> bool:
        lowered = text.lower()
        if len(lowered) < 25:
            return True
        if any(p in lowered for p in NOISE_PATTERNS):
            return True
        if re.fullmatch(r"[A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ0-9\s\-â€“|/:,\.]+", lowered) and len(lowered.split()) <= 5:
            return True
        if lowered in BAD_TITLES:
            return True
        return False

    def extract_date(self, soup: BeautifulSoup, text: str) -> Optional[datetime]:
        candidates = []

        meta_selectors = [
            ("meta", {"property": "article:published_time"}, "content"),
            ("meta", {"name": "article:published_time"}, "content"),
            ("meta", {"name": "publish-date"}, "content"),
            ("meta", {"name": "date"}, "content"),
            ("meta", {"property": "og:updated_time"}, "content"),
            ("meta", {"property": "article:modified_time"}, "content"),
            ("meta", {"name": "DC.date"}, "content"),
            ("meta", {"name": "datePublished"}, "content"),
            ("meta", {"itemprop": "datePublished"}, "content"),
        ]
        for tag_name, attrs, attr in meta_selectors:
            tag = soup.find(tag_name, attrs=attrs)
            if tag:
                content = tag.get(attr, "").strip()
                if content:
                    candidates.append(content)

        for time_tag in soup.find_all("time"):
            dt = time_tag.get("datetime", "").strip()
            txt = time_tag.get_text(" ", strip=True)
            if dt:
                candidates.append(dt)
            if txt:
                candidates.append(txt)

        for candidate in candidates:
            parsed = self.try_parse_date(candidate)
            if parsed:
                return parsed

        text_window = text[:3500]
        patterns = [
            r"\b(\d{4}-\d{2}-\d{2})\b",
            r"\b(\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(?:Z|[+-]\d{2}:?\d{2})?)\b",
            r"\b(\d{2}\.\d{2}\.\d{4})\b",
            r"\b(\d{1,2}\.\s*(?:Januar|Februar|MÃ¤rz|Maerz|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s*\d{4})\b",
        ]
        for pattern in patterns:
            match = re.search(pattern, text_window, flags=re.IGNORECASE)
            if match:
                parsed = self.try_parse_date(match.group(1))
                if parsed:
                    return parsed

        return None

    def try_parse_date(self, raw: str) -> Optional[datetime]:
        raw = raw.strip()
        date_formats = [
            "%Y-%m-%d",
            "%Y-%m-%dT%H:%M:%S",
            "%Y-%m-%dT%H:%M:%S%z",
            "%Y-%m-%dT%H:%M:%SZ",
            "%Y-%m-%dT%H:%M:%S.%fZ",
            "%d.%m.%Y",
        ]
        for fmt in date_formats:
            try:
                dt = datetime.strptime(raw, fmt)
                return dt.replace(tzinfo=None)
            except Exception:
                pass

        try:
            dt = parsedate_to_datetime(raw)
            return dt.replace(tzinfo=None)
        except Exception:
            pass

        german_months = {
            "januar": 1, "februar": 2, "mÃ¤rz": 3, "maerz": 3,
            "april": 4, "mai": 5, "juni": 6, "juli": 7,
            "august": 8, "september": 9, "oktober": 10,
            "november": 11, "dezember": 12
        }
        m = re.search(
            r"(\d{1,2})\.\s*(Januar|Februar|MÃ¤rz|Maerz|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s*(\d{4})",
            raw,
            flags=re.IGNORECASE
        )
        if m:
            try:
                return datetime(
                    int(m.group(3)),
                    german_months[m.group(2).lower()],
                    int(m.group(1))
                )
            except Exception:
                return None
        return None

    def extract_snippets(self, text: str, title: str = "") -> list[str]:
        blocks = [b.strip() for b in text.split("\n") if b.strip()]
        scored: list[tuple[int, str]] = []

        strategy_terms = [
            "strategie", "umsetzung", "projekt", "ausbau", "investition",
            "kooperation", "partnerschaft", "fÃ¶rderung", "weiterbildung",
            "innovation", "digitalisierung", "nachhaltigkeit", "standort",
            "personal", "ernennung", "leitung", "gesellschaft", "gesundheit",
            "produktion", "versorgung", "fachkrÃ¤fte"
        ]

        title_terms = {t for t in self.tokenize(title) if len(t) > 4}

        for idx, block in enumerate(blocks):
            word_count = len(block.split())
            if word_count < MIN_SNIPPET_WORDS or word_count > 90:
                continue

            lowered = block.lower()
            if any(p in lowered for p in NOISE_PATTERNS):
                continue
            if any(p in lowered for p in ROUTINE_PATTERNS):
                continue

            score = 0
            score += sum(1 for term in strategy_terms if term in lowered) * 3
            score += 3 if idx < 3 else 0
            score += 2 if re.search(r"\b\d{4}\b", lowered) else 0
            score += 2 if len(block) >= 120 else 0
            score += sum(1 for term in title_terms if term in lowered)
            score -= 2 if lowered == title.lower().strip() else 0

            if score > 0:
                scored.append((score, block))

        scored.sort(key=lambda x: (-x[0], len(x[1])))
        selected = [block for _, block in scored[:self.max_snippets_per_page]]
        return self.deduplicate_list(selected)

    def acronym_explanation_context(self, title: str, snippets: list[str]) -> str:
        combined = "\n".join([str(title)] + [str(snippet) for snippet in snippets])
        acronym_pattern = r"\b[A-ZÃ„Ã–Ãœ][A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ]*[A-ZÃ„Ã–Ãœ][A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ]*\b"
        names = []
        for match in re.finditer(acronym_pattern, combined):
            token = match.group(0).strip()
            if len(token) < 4 or token.upper() in {"ZIRP", "GMBH", "MINT"}:
                continue
            names.append(token)
        contexts = []
        for token in self.deduplicate_list(names)[:5]:
            pattern = re.compile(rf"([^.?!\n]{{0,180}}\b{re.escape(token)}\b[^.?!\n]{{0,220}}[.?!]?)")
            for match in pattern.finditer(combined):
                sentence = re.sub(r"\s+", " ", match.group(1)).strip()
                if len(sentence.split()) >= 8:
                    contexts.append(f"{token}: {sentence}")
                    break
        return " | ".join(self.deduplicate_list(contexts)[:4])

    def first_meaningful_paragraph(self, text: str) -> str:
        for block in text.split("\n"):
            block = block.strip()
            if len(block.split()) >= 12 and not self.is_boilerplate_text(block):
                return block
        return ""

    def tokenize(self, text: str) -> list[str]:
        tokens = re.findall(r"[A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]{3,}", text.lower())
        return [t for t in tokens if t not in GERMAN_STOPWORDS and len(t) >= 3]

    def find_relevant_terms(self, text: str) -> Counter:
        lowered = text.lower()
        found = Counter()

        for cfg in self.report_sections_for_run().values():
            for term in cfg["keywords"]:
                hits = re.findall(rf"\b{re.escape(term.lower())}\b", lowered)
                if hits:
                    found[term.lower()] += len(hits)

        for terms in self.zirp_themenfelder_for_run().values():
            for term in terms:
                hits = re.findall(rf"\b{re.escape(term.lower())}\b", lowered)
                if hits:
                    found[term.lower()] += len(hits)

        for token in self.tokenize(text):
            if any(root in token for root in [
                "nachhalt", "kooper", "innov", "bildung", "forsch",
                "gesund", "kultur", "digital", "energie", "arbeits",
                "pflege", "kommunal", "fÃ¼hrung", "standort"
            ]):
                found[token] += 1

        return found

    def detect_report_sections(self, text: str, title: str, snippet: str) -> dict[str, set]:
        lowered = f"{title}\n{snippet}\n{text}".lower()
        report_sections = self.report_sections_for_run()
        result = {k: set() for k in report_sections.keys()}

        for section_key, cfg in report_sections.items():
            for term in cfg["keywords"]:
                if re.search(rf"\b{re.escape(term.lower())}\b", lowered):
                    result[section_key].add(term)

        return result

    def detect_zirp_themenfelder(self, text: str) -> set[str]:
        lowered = text.lower()
        themes = set()
        for field, terms in self.zirp_themenfelder_for_run().items():
            for term in terms:
                if re.search(rf"\b{re.escape(term.lower())}\b", lowered):
                    themes.add(field)
                    break
        return themes

    def infer_main_section(self, section_matches: dict[str, set], text: str) -> str:
        scores = {section: len(terms) for section, terms in section_matches.items()}
        best_section = max(scores, key=scores.get)

        if scores[best_section] == 0:
            lowered = text.lower()
            if any(x in lowered for x in ["gesundheit", "versorgung", "krankenkasse", "arzt", "pflege"]):
                return "versorgung_gesundheit"
            if any(x in lowered for x in ["geschÃ¤ftsfÃ¼hrer", "vorstand", "leitung", "ernennung"]):
                return "fuehrungswechsel"
            if any(x in lowered for x in ["kooperation", "partnerschaft", "gemeinsam", "netzwerk"]):
                return "kooperationen"
            if any(x in lowered for x in ["spende", "engagement", "gesellschaft", "gesundheit"]):
                return "soziales_engagement"
            return "wirtschaftsentwicklung"

        return best_section

    def classify_page_type(self, url: str, title: str, text: str, article_quality_score: int) -> str:
        if self.is_document_url(url):
            return "document"
        combined = f"{url} {title}".lower()
        title_l = title.lower().strip()
        paragraphs = [b for b in text.split("\n") if len(b.split()) >= 8]

        negative_patterns = [
            "terminkalender", "kalender", "alle meldungen", "alle kategorien",
            "startseite", "Ã¼bersicht", "overview", "search", "suche", "login",
            "willkommen", "online-magazin", "fÃ¼r presse", "fÃ¼r presse und medien",
            "news und pressemitteilungen", "messen & events", "veranstaltungen",
            "campusleben", "service", "faq"
        ]
        if any(p in combined for p in negative_patterns):
            return "listing"
        if title_l in BAD_TITLES:
            return "listing"
        if any(p in title_l for p in HUB_PAGE_PATTERNS):
            return "listing"
        if article_quality_score < self.min_article_quality_score:
            return "listing"
        if len(paragraphs) < 3:
            return "listing"
        return "article"

    def has_valid_title_and_snippet(self, title: str, snippet: str) -> bool:
        title_clean = re.sub(r"\s+", " ", str(title)).strip()
        snippet_clean = re.sub(r"\s+", " ", str(snippet)).strip()
        title_norm = re.sub(r"\s+", " ", title_clean.lower()).strip(" -â€“")

        if not title_clean or len(title_clean) < 8:
            return False
        if re.fullmatch(r"\d{4}", title_clean):
            return False
        if title_norm in BAD_TITLES:
            return False
        if any(x in title_norm for x in ["homepage", "presse-service", "Ã¶ffentlichkeitsarbeit", "aktuelle meldungen"]):
            return False
        if len(snippet_clean.split()) < MIN_SNIPPET_WORDS:
            return False
        if title_norm == snippet_clean.lower().strip():
            return False

        return True

    def is_routine_event(self, title: str, snippet: str) -> bool:
        combined = f"{title} {snippet}".lower()
        return any(p in combined for p in ROUTINE_PATTERNS)

    def is_excludable_content(self, title: str, snippet: str, url: str, page_type: str) -> bool:
        combined = f"{title} {snippet} {url}".lower()

        if page_type not in {"article", "document"}:
            return True
        if self.is_hard_noise_item(title, snippet, url):
            return True
        if any(p in combined for p in EXCLUSION_TITLE_PATTERNS):
            return True
        if any(p in combined for p in EXCLUSION_URL_PATTERNS):
            return True
        if any(p in combined for p in EXCLUSION_TEXT_PATTERNS):
            return True
        if any(p in combined for p in HUB_PAGE_PATTERNS):
            return True
        if self.is_routine_event(title, snippet):
            return True
        if len(title.split()) < 3:
            return True

        return False

    def is_hard_noise_item(self, title: str, snippet: str, url: str = "") -> bool:
        combined = f"{title} {snippet} {url}".lower()
        if any(term in combined for term in self.noise_terms_for_run()):
            if not any(term in combined for term in ["rheinland-pfalz", "rlp", "kooperation", "projekt", "strategie", "investition", "versorgung", "transfer"]):
                return True
        return False

    def direct_member_signal_score(self, member_name: str, title: str, text: str, url: str) -> int:
        haystack = f"{title}\n{text[:1800]}\n{url}".lower()
        member_tokens = [t for t in self.tokenize(member_name) if len(t) >= 4]
        score = 0
        for token in member_tokens[:5]:
            if re.search(rf"\b{re.escape(token)}\b", haystack):
                score += 2

        member_homepage = self.member_lookup.get(member_name, {}).get("homepage", "")
        if member_homepage and self.same_domain(url, member_homepage):
            score += 2
        return score

    def regional_relevance_score(self, title: str, text: str, member_name: str) -> int:
        haystack = f"{title}\n{text[:2200]}\n{member_name}".lower()
        score = 0
        for term in RLP_TERMS:
            if term in haystack:
                score += 1
        return min(score, 4)

    def context_item_is_relevant(
        self,
        member_name: str,
        title: str,
        text: str,
        url: str,
        direct_member_score: int,
        regional_score: int,
    ) -> bool:
        combined = f"{title}\n{text[:2500]}\n{url}\n{member_name}".lower()

        if direct_member_score >= 3:
            return True

        member_tokens = [t for t in self.tokenize(member_name) if len(t) >= 4]
        token_hits = sum(
            1 for tok in member_tokens[:5]
            if re.search(rf"\b{re.escape(tok)}\b", combined)
        )
        if token_hits >= 2:
            return True

        if regional_score >= 2:
            return True

        strategic_hits = sum(1 for t in self.strategic_terms_for_run() if t in combined)
        rlp_hits = sum(1 for t in RLP_TERMS if t in combined)

        if rlp_hits >= 1 and strategic_hits >= 2:
            return True
        if direct_member_score >= 2 and strategic_hits >= 2:
            return True
        if direct_member_score >= 2 and any(
            term in combined
            for term in [
                "projekt", "kooperation", "transfer", "strategie", "digitalisierung",
                "fachkraefte", "fachkrÃ¤fte", "innovation", "studiengang", "programm",
            ]
        ):
            return True

        return False

    def geography_mismatch_penalty(self, title: str, text: str) -> int:
        combined = f"{title} {text[:1800]}".lower()
        penalty = sum(1 for term in NON_RLP_CONTEXT_TERMS if term in combined)
        rlp_hits = sum(1 for term in RLP_TERMS if term in combined)

        if rlp_hits >= 2:
            return 0

        return min(penalty * 2, 6)

    def is_homepage_like(self, homepage: str, page_url: str) -> bool:
        if not homepage or not page_url:
            return False
        return self.normalize_url(homepage) == self.normalize_url(page_url)

    def is_strategically_relevant_event(
        self,
        title: str,
        snippet: str,
        url: str,
        direct_member_score: int,
        regional_score: int,
    ) -> bool:
        combined = f"{title} {snippet} {url}".lower()

        if any(p in combined for p in NOISE_PATTERNS):
            return False
        if self.is_routine_event(title, snippet):
            return False
        if len(title.split()) < 3:
            return False

        weak_context = direct_member_score <= 1 and regional_score == 0
        entertainment_terms = ["musical", "konzert", "premiere", "festival", "sendung", "tv", "podcast", "folge"]
        if weak_context and any(t in combined for t in entertainment_terms):
            return False

        return True

    def leadership_relevance_gate(
        self,
        title: str,
        snippet: str,
        main_section: str,
        themes: set[str],
        direct_member_score: int,
        regional_score: int,
        article_quality_score: int,
    ) -> bool:
        text = f"{title}\n{snippet}".lower()

        leadership_hits = sum(1 for t in LEADERSHIP_TERMS if t in text)
        strong_hits = sum(1 for t in STRONG_STRATEGIC_TERMS if t in text)
        weak_hits = sum(1 for t in WEAK_SIGNAL_TERMS if t in text)
        macro_hits = sum(1 for t in MACRO_COMMENTARY_TERMS if t in text)
        institutional_hits = sum(1 for t in INSTITUTIONAL_ACTION_TERMS if t in text)
        regional_impl_hits = sum(1 for t in REGIONAL_IMPLEMENTATION_TERMS if t in text)

        structural_section = main_section in {
            "wirtschaftsentwicklung",
            "versorgung_gesundheit",
            "kooperationen",
            "fuehrungswechsel",
        }

        if weak_hits >= 1 and leadership_hits == 0 and strong_hits == 0:
            return False

        if article_quality_score < self.min_article_quality_score:
            return False

        if macro_hits >= 2 and institutional_hits == 0 and regional_impl_hits == 0 and regional_score == 0:
            return False

        if direct_member_score >= 4 and structural_section and leadership_hits >= 1:
            return True

        if regional_score >= 2 and (leadership_hits >= 2 or strong_hits >= 1):
            return True

        if main_section == "kooperationen" and (leadership_hits >= 1 or strong_hits >= 1 or institutional_hits >= 1):
            return True

        if main_section == "versorgung_gesundheit" and (
            "gesundheitsversorgung" in text or
            "versorgung" in text or
            "pflege" in text or
            "fachkrÃ¤fte" in text or
            "koalitionsverhandlungen" in text
        ):
            return True

        if strong_hits >= 2:
            return True

        if structural_section and len(themes) >= 2 and leadership_hits >= 2:
            return True

        return False

    def model_relevance_gate(self, event: dict[str, Any]) -> bool:
        """
        Let OpenAI make the final relevance decision after basic article hygiene.
        The rule-based gates remain as a fallback when the API is unavailable.
        """
        if not MODEL_DECIDES_RELEVANCE or self.openai_client is None:
            return False

        frame = self.enrich_event_editorially(event)
        self.editorial_frames[self._event_key(event)] = frame

        role = str(frame.get("role_in_briefing", "")).strip()
        confidence = str(frame.get("confidence", "")).strip()
        is_context_only = bool(frame.get("is_context_only", False))
        model_used = bool(frame.get("model_used", False))

        event["model_role_in_briefing"] = role
        event["model_confidence"] = confidence
        event["model_is_context_only"] = is_context_only
        event["model_used"] = model_used

        if role == "ignore":
            return False

        return role in MODEL_KEEP_ROLES

    def fallback_rule_relevance_gate(self, event: dict[str, Any]) -> bool:
        title = str(event.get("titel", "")).strip().lower()
        snippet = str(event.get("snippet", "")).strip()
        if title in BAD_TITLES or len(title) < 5:
            return False
        if len(snippet.split()) < MIN_SNIPPET_WORDS:
            return False
        if event.get("article_quality_score", 0) < self.min_article_quality_score:
            return False
        if self.is_hard_noise_item(title, snippet, str(event.get("url", ""))):
            return False
        if self.is_macro_context_without_member_action(event):
            return False
        return True

    def evidence_role_for_event(self, event: dict[str, Any]) -> str:
        if self.is_macro_context_without_member_action(event) or self.is_hard_noise_item(
            str(event.get("titel", "")),
            str(event.get("snippet", "")),
            str(event.get("url", "")),
        ):
            return "exclude"
        if self.is_low_value_narrative_event(event):
            return "watchlist"
        decision_score = int(event.get("decision_score", 0) or 0)
        direct_score = int(event.get("direct_member_score", 0) or 0)
        regional_score = int(event.get("regional_score", 0) or 0)
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        has_action = any(term in text for term in INSTITUTIONAL_ACTION_TERMS)
        has_strategy = any(term in text for term in STRONG_STRATEGIC_TERMS)
        if decision_score >= 14 and (direct_score >= 3 or regional_score >= 1) and (has_action or has_strategy):
            return "main_signal"
        if decision_score >= 8 and (direct_score >= 2 or regional_score >= 1 or has_action or has_strategy):
            return "supporting_context"
        return "watchlist"

    def is_macro_context_without_member_action(self, event: dict[str, Any]) -> bool:
        text = f"{event.get('titel', '')} {event.get('snippet', '')} {event.get('url', '')}".lower()
        macro_hits = sum(1 for term in MACRO_COMMENTARY_TERMS if term in text)
        hard_macro = any(term in text for term in ["trump", "iran", "usa:", "kapitalmÃ¤rkte daily", "kapitalmÃ¤rkte daily"])
        institutional_hits = sum(1 for term in INSTITUTIONAL_ACTION_TERMS if term in text)
        regional_hits = sum(1 for term in REGIONAL_IMPLEMENTATION_TERMS if term in text)
        if not (macro_hits or hard_macro):
            return False
        return (
            institutional_hits == 0
            and regional_hits == 0
            and int(event.get("regional_score", 0) or 0) == 0
            and int(event.get("direct_member_score", 0) or 0) < 5
        )

    def is_report_worthy_event(self, event: dict[str, Any]) -> bool:
        if self.is_macro_context_without_member_action(event):
            return False
        if self.is_low_value_narrative_event(event):
            return False
        decision_score = int(event.get("decision_score", 0) or 0)
        direct_score = int(event.get("direct_member_score", 0) or 0)
        regional_score = int(event.get("regional_score", 0) or 0)
        section = str(event.get("hauptsektion", ""))
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        institutional = any(term in text for term in INSTITUTIONAL_ACTION_TERMS)
        strategic = any(term in text for term in STRONG_STRATEGIC_TERMS)
        evidence_role = str(event.get("evidence_role", "")).strip()
        if evidence_role in {"main_signal", "supporting_context"}:
            return True
        if decision_score >= 14 and (direct_score >= 3 or regional_score >= 1):
            return True
        if section == "kooperationen" and (institutional or strategic) and decision_score >= 10:
            return True
        if section in {"versorgung_gesundheit", "wirtschaftsentwicklung"} and strategic and decision_score >= 10:
            return True
        return False

    def build_event(
        self,
        member_name: str,
        page_url: str,
        title: str,
        published_at: datetime,
        snippet: str,
        section_matches: dict[str, set],
        themes: set[str],
        terms: Counter,
        is_updated: bool,
        direct_member_score: int,
        regional_score: int,
        article_quality_score: int,
        static_member_fit_score: int = 0,
    ) -> dict[str, Any]:
        full_text = f"{title}\n{snippet}"
        main_section = self.infer_main_section(section_matches, full_text)
        recency_weight = self.event_recency_weight(published_at)

        keyword_strength = min(sum(len(v) for v in section_matches.values()) + min(sum(terms.values()), 10), 18)
        theme_bonus = min(len(themes) * 2, 6)
        update_bonus = 2 if is_updated else 0
        snippet_bonus = 3 if snippet else 0
        article_bonus = min(max(article_quality_score - self.min_article_quality_score, 0), 4)
        direct_bonus = min(direct_member_score, 6)
        regional_bonus = min(regional_score, 3)
        static_fit_bonus = min(static_member_fit_score, 10)
        strategic_bonus = sum(1 for t in self.strategic_terms_for_run() if t in full_text.lower())
        strategic_bonus = min(strategic_bonus, 4)

        score = (
            keyword_strength
            + theme_bonus
            + (recency_weight * 3)
            + update_bonus
            + snippet_bonus
            + article_bonus
            + direct_bonus
            + regional_bonus
            + static_fit_bonus
            + strategic_bonus
        )
        decision_score = self.decision_relevance_score(
            title=title,
            snippet=snippet,
            main_section=main_section,
            themes=themes,
            direct_member_score=direct_member_score,
            regional_score=regional_score,
            article_quality_score=article_quality_score,
            static_member_fit_score=static_member_fit_score,
        )

        member_meta = self.member_lookup.get(member_name, {})
        event = {
            "member_id": member_meta.get("member_id"),
            "mitglied": member_name,
            "titel": title,
            "datum": published_at.strftime("%Y-%m-%d"),
            "datum_obj": published_at,
            "url": page_url,
            "hauptsektion": main_section,
            "sektionssignale": {k: sorted(v) for k, v in section_matches.items()},
            "themenfelder": sorted(themes),
            "snippet": snippet,
            "score": score,
            "decision_score": decision_score,
            "recency_weight": recency_weight,
            "is_updated": is_updated,
            "direct_member_score": direct_member_score,
            "regional_score": regional_score,
            "static_member_fit_score": static_member_fit_score,
            "article_quality_score": article_quality_score,
        }
        event["evidence_role"] = self.evidence_role_for_event(event)
        return event

    def event_recency_weight(self, published_at: datetime) -> int:
        days = (self.today - published_at).days
        if days <= 2:
            return 4
        if days <= 7:
            return 3
        if days <= 31:
            return 2
        return 1

    def decision_relevance_score(
        self,
        title: str,
        snippet: str,
        main_section: str,
        themes: set[str],
        direct_member_score: int,
        regional_score: int,
        article_quality_score: int,
        static_member_fit_score: int = 0,
    ) -> int:
        text = f"{title}\n{snippet}".lower()

        implementation_terms = [
            "strategie", "umsetzung", "pilot", "programm", "einfÃ¼hrung",
            "ausbau", "fÃ¶rderung", "kooperation", "partnerschaft",
            "weiterbildung", "innovation", "projekt", "initiative",
            "investition", "erÃ¶ffnung", "ausgrÃ¼ndung"
        ]
        consequence_terms = [
            "standort", "wirtschaft", "fachkrÃ¤fte", "versorgung",
            "nachhaltigkeit", "gesellschaft", "gesundheit", "produktion",
            "arbeitsplÃ¤tze", "transform", "wettbewerb"
        ]
        decision_terms = [
            "neu", "neue", "ernannt", "berufen", "startet", "erweitert",
            "investiert", "kooperiert", "fÃ¶rdert", "verleiht", "erÃ¶ffnet",
            "beschlieÃŸt", "baut", "nimmt", "grÃ¼ndet"
        ]

        score = 0
        score += sum(1 for t in implementation_terms if t in text) * 2
        score += sum(1 for t in consequence_terms if t in text) * 2
        score += sum(1 for t in decision_terms if t in text)
        score += len(themes)
        score += min(direct_member_score, 4)
        score += min(regional_score, 2)
        score += min(static_member_fit_score, 6)
        score += 2 if article_quality_score >= 14 else 0
        score += min(sum(1 for t in self.strategic_terms_for_run() if t in text), 3)

        if main_section == "wirtschaftsentwicklung":
            score += 3
        elif main_section == "versorgung_gesundheit":
            score += 3
        elif main_section == "kooperationen":
            score += 3
        elif main_section == "fuehrungswechsel":
            score += 2
        elif main_section == "soziales_engagement":
            score += 1

        if len(snippet.split()) >= 14:
            score += 2

        return score

    def register_keyword_rows(
        self,
        member_name: str,
        page_url: str,
        published_at: datetime,
        terms: Counter,
        section_matches: dict[str, set],
        themes: set[str],
        snippets: list[str]
    ) -> None:
        section_term_lookup = {
            key: {t.lower() for t in vals}
            for key, vals in section_matches.items()
        }

        rows = []
        for term, count in terms.items():
            rows.append({
                "mitglied": member_name,
                "member_id": self.member_lookup.get(member_name, {}).get("member_id"),
                "seite": page_url,
                "datum": published_at.strftime("%Y-%m-%d"),
                "begriff": term,
                "count": count,
                "wirtschaftsentwicklung": int(term in section_term_lookup["wirtschaftsentwicklung"]),
                "versorgung_gesundheit": int(term in section_term_lookup["versorgung_gesundheit"]),
                "fuehrungswechsel": int(term in section_term_lookup["fuehrungswechsel"]),
                "soziales_engagement": int(term in section_term_lookup["soziales_engagement"]),
                "kooperationen": int(term in section_term_lookup["kooperationen"]),
                "zirp_themenfelder": ", ".join(sorted(themes)),
                "snippets": " | ".join(snippets),
            })
        if rows:
            with self.keyword_rows_lock:
                self.keyword_rows.extend(rows)

    def csv_field_to_set(self, value) -> set[str]:
        if value is None or pd.isna(value):
            return set()
        text = str(value).strip()
        return {x.strip() for x in text.split(",") if x.strip()} if text else set()

    @staticmethod
    def parse_bool(value) -> bool:
        if isinstance(value, bool):
            return value
        if value is None:
            return False
        text = str(value).strip().lower()
        return text in {"1", "true", "ja", "yes", "y"}

    def deduplicate_events(self) -> None:
        unique = {}
        for event in self.events:
            key = (
                event["mitglied"].strip().lower(),
                event["titel"].strip().lower(),
                event["datum"],
                re.sub(r"\s+", " ", str(event["snippet"])).strip().lower()[:180],
            )
            if key not in unique or event["decision_score"] > unique[key]["decision_score"]:
                unique[key] = event
        self.events = list(unique.values())

    def zirp_relevance_label(self, event: dict[str, Any]) -> tuple[str, str]:
        title = str(event.get("titel", "") or "")
        snippet = str(event.get("snippet", "") or "")
        url = str(event.get("url", "") or "")
        section = str(event.get("hauptsektion", "") or "")
        topics = str(event.get("themenfelder", "") or "")

        haystack = f"{title} {snippet} {url} {section} {topics}".lower()

        score = float(event.get("score", 0) or 0)
        decision_score = float(event.get("decision_score", 0) or 0)

        zirp_core_terms = [
            "rheinland-pfalz", "rlp", "standort", "wirtschaftsstandort", "region",
            "regional", "kommune", "kommunal", "transformation", "zukunft",
            "innovation", "digitalisierung", "ki", "k?nstliche intelligenz",
            "nachhaltigkeit", "energie", "klima", "fachkr?fte", "ausbildung",
            "versorgung", "gesundheit", "pflege", "wissenschaft", "hochschule",
            "transfer", "kooperation", "partnerschaft", "netzwerk", "dialog",
            "kultur", "gesellschaft", "teilhabe", "infrastruktur", "investition",
        ]

        weak_pr_terms = [
            "gewinnspiel", "rabatt", "angebot", "aktion", "promotion", "podcast",
            "folge", "sendung", "ticket", "ranking", "bestnoten", "auszeichnung",
            "sponsoring", "infoabend", "download", "ratgeber", "tipps",
            "veranstaltungskalender", "girls'day", "girlsday",
        ]

        action_terms = [
            "kooperation", "partnerschaft", "vereinbarung", "investition", "ausbau",
            "strategie", "programm", "projekt", "initiative", "f?rderung",
            "plattform", "runde tisch", "modellprojekt", "pilot", "netzwerk",
            "transfer", "konzession", "standortentscheidung", "batteriezellfertigung",
            "infrastruktur",
        ]

        convening_terms = [
            "gemeinsam", "partner", "netzwerk", "dialog", "runde tisch",
            "plattform", "austausch", "kooperation", "kommunen", "wirtschaft",
            "wissenschaft", "politik", "kultur", "gesellschaft", "hochschule",
            "verwaltung", "kammer", "verband",
        ]

        core_hits = sum(1 for term in zirp_core_terms if term in haystack)
        weak_hits = sum(1 for term in weak_pr_terms if term in haystack)
        action_hits = sum(1 for term in action_terms if term in haystack)
        convening_hits = sum(1 for term in convening_terms if term in haystack)

        if weak_hits >= 1 and core_hits < 2 and action_hits == 0:
            return "drop", "Schwaches PR- oder Routinethema ohne ausreichenden ZIRP-Bezug."

        if action_hits >= 1 and core_hits >= 2:
            return "keep", "Konkretes Zukunfts-, Standort- oder Transformationssignal mit ZIRP-Relevanz."

        if convening_hits >= 2 and core_hits >= 2:
            return "keep", "Geeignet als Anlass f?r Vernetzung, Austausch oder Transfer."

        if decision_score >= 8 and core_hits >= 2:
            return "keep", "Hohe strategische Relevanz f?r ZIRP-Themenfelder."

        if core_hits >= 2 or decision_score >= 5 or (score >= self.min_event_score and core_hits >= 1):
            return "watch", "Beobachtenswertes Signal, aber noch kein starker Handlungsanlass."

        return "drop", "Kein ausreichender Bezug zu ZIRP-Auftrag, Vernetzung oder Rheinland-Pfalz-Transformation."

    def apply_zirp_relevance_filter(self) -> None:
        filtered_events = []

        for event in self.events:
            label, reason = self.zirp_relevance_label(event)
            event["zirp_label"] = label
            event["zirp_reason"] = reason
            event["evidence_role"] = self.evidence_role_for_event(event)

            if label in {"keep", "watch"} and event["evidence_role"] != "exclude":
                filtered_events.append(event)

        before = len(self.events)
        after = len(filtered_events)
        roles = Counter(event.get("evidence_role", "unknown") for event in filtered_events)
        role_text = ", ".join(f"{role}: {count}" for role, count in sorted(roles.items()))
        print(f"ZIRP relevance filter: {before} Ereignisse gepr?ft, {after} behalten, {before - after} entfernt ({role_text})")

        self.events = filtered_events

    def signal_judgment_key(self, event: dict[str, Any]) -> str:
        payload = "|".join([
            str(event.get("url", "")),
            str(event.get("titel", "")),
            str(event.get("datum", "")),
            str(event.get("snippet", ""))[:240],
        ])
        return self.make_hash(payload)

    def apply_signal_judgment(self) -> None:
        if not SIGNAL_JUDGE_ENABLED:
            return
        if self.report_openai_client is None:
            print("Signal judge: Ã¼bersprungen (kein Groq/Report-Client aktiv)")
            return
        if not self.events:
            return

        signal_judge_limit = self.dynamic_signal_judge_top_n()
        candidates = self.rank_events_for_narrative([
            event for event in self.events
            if not self.is_macro_context_without_member_action(event)
        ])[:signal_judge_limit]

        uncached = []
        for event in candidates:
            key = self.signal_judgment_key(event)
            cached = self.signal_judgment_cache.get(key)
            if isinstance(cached, dict):
                self.apply_signal_judgment_to_event(event, cached, source="cache")
            else:
                uncached.append((key, event))

        if uncached:
            for start in range(0, len(uncached), SIGNAL_JUDGE_BATCH_SIZE):
                batch = uncached[start:start + SIGNAL_JUDGE_BATCH_SIZE]
                judgments = self.request_signal_judgment_batch([event for _, event in batch])
                for key, event in batch:
                    judgment = judgments.get(self.signal_judge_event_id(event))
                    if not isinstance(judgment, dict) or not judgment:
                        continue
                    self.signal_judgment_cache[key] = judgment
                    self.apply_signal_judgment_to_event(event, judgment, source="groq")

        before = len(self.events)
        kept = []
        for event in self.events:
            use = str(event.get("llm_recommended_use", "")).strip()
            if use == "drop":
                continue
            kept.append(event)
        self.events = kept

        source_counts = Counter(event.get("llm_judge_source", "none") for event in self.events)
        use_counts = Counter(event.get("llm_recommended_use", "unjudged") for event in self.events)
        print(
            "Signal judge: "
            f"{len(candidates)} Kandidaten geprueft, {before - len(kept)} entfernt, "
            f"Quellen={dict(source_counts)}, Nutzung={dict(use_counts)}"
        )

    def signal_judge_event_id(self, event: dict[str, Any]) -> str:
        return self.make_hash(f"{event.get('mitglied','')}|{event.get('titel','')}|{event.get('url','')}")[:12]

    def request_signal_judgment_batch(self, events: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
        items = []
        for event in events:
            items.append({
                "id": self.signal_judge_event_id(event),
                "member": event.get("mitglied", ""),
                "title": event.get("titel", ""),
                "date": event.get("datum", ""),
                "section": event.get("hauptsektion", ""),
                "evidence_role": event.get("evidence_role", ""),
                "snippet": self.clean_narrative_snippet(event.get("snippet", ""), max_chars=220),
                "rule_scores": {
                    "decision_score": event.get("decision_score", 0),
                    "score": event.get("score", 0),
                    "direct_member_score": event.get("direct_member_score", 0),
                    "regional_score": event.get("regional_score", 0),
                    "article_quality_score": event.get("article_quality_score", 0),
                },
            })

        system_prompt = (
            "Du bist ein strenger Signal-Redakteur fÃ¼r ein ZIRP-Leitungsbriefing in Rheinland-Pfalz. "
            "Bewerte kompakte Crawler-Funde. Bevorzuge konkrete Entscheidungen, Programme, Kooperationen, "
            "Investitionen, Transfer, Versorgung, Resilienz, Standort- oder Transformationsbezug. "
            "Stufe PR, Ausstellungen, Rankings, Girls-Day/Nachwuchs-PR, Sport-/Kulturkommunikation und allgemeine Marktkommentare niedrig ein. "
            "Gib ausschliesslich valides JSON zurueck."
        )
        user_prompt = (
            "Bewerte jedes Ereignis. Schema:\n"
            "{\"judgments\":[{\"id\":\"...\",\"strategic_score\":0-10,\"concreteness_score\":0-10,"
            "\"zirp_relevance\":0-10,\"recommended_use\":\"briefing|convening|watch|drop\","
            "\"reason\":\"kurz\"}]}\n\n"
            f"EREIGNISSE:\n{json.dumps(items, ensure_ascii=False, indent=2)}"
        )

        started = time.perf_counter()
        try:
            judge_model = os.getenv("SIGNAL_JUDGE_MODEL", self.report_openai_model or SIGNAL_JUDGE_MODEL).strip()
            response = self.report_openai_client.chat.completions.create(
                model=judge_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.1,
                max_tokens=SIGNAL_JUDGE_MAX_TOKENS,
                timeout=self.report_openai_timeout,
            )
            content = response.choices[0].message.content if response.choices else ""
            parsed = self.parse_json_object(content) or {}
            judgments = parsed.get("judgments", [])
            result = {}
            if isinstance(judgments, list):
                for item in judgments:
                    if isinstance(item, dict) and item.get("id"):
                        result[str(item["id"])] = item
            print(
                f"Signal judge batch: {len(events)} Ereignisse, {len(result)} Urteile, "
                f"duration={time.perf_counter() - started:.1f}s"
            )
            return result
        except Exception as exc:
            print(f"Signal judge batch fehlgeschlagen after {time.perf_counter() - started:.1f}s: {exc}")
            return {}

    def apply_signal_judgment_to_event(self, event: dict[str, Any], judgment: dict[str, Any], source: str) -> None:
        def bounded_int(value: Any) -> int:
            try:
                return max(0, min(10, int(value)))
            except Exception:
                return 0

        recommended = str(judgment.get("recommended_use", "watch")).strip().lower()
        if recommended not in {"briefing", "convening", "watch", "drop"}:
            recommended = "watch"
        strategic = bounded_int(judgment.get("strategic_score", 0))
        concrete = bounded_int(judgment.get("concreteness_score", 0))
        relevance = bounded_int(judgment.get("zirp_relevance", 0))

        event["llm_judge_source"] = source
        event["llm_strategic_score"] = strategic
        event["llm_concreteness_score"] = concrete
        event["llm_zirp_relevance"] = relevance
        event["llm_recommended_use"] = recommended
        event["llm_judge_reason"] = str(judgment.get("reason", "")).strip()

        event["decision_score"] = int(event.get("decision_score", 0) or 0) + strategic + concrete + relevance
        if recommended == "briefing":
            event["evidence_role"] = "main_signal"
        elif recommended == "convening":
            event["evidence_role"] = "supporting_context"
        elif recommended == "watch":
            event["evidence_role"] = "watchlist"

    def compact_signal_text(self, value: Any, max_chars: int = 240) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if len(text) <= max_chars:
            return text
        return text[: max_chars - 3].rstrip(" ,.;:-") + "..."

    def opportunity_text_for_event(self, event: dict[str, Any], frame: Optional[dict[str, Any]] = None) -> str:
        parts = [
            event.get("titel", ""),
            event.get("snippet", ""),
            event.get("hauptsektion", ""),
            " ".join(event.get("themenfelder", []) or []),
            event.get("zirp_reason", ""),
            event.get("llm_judge_reason", ""),
        ]
        if frame:
            parts.extend([
                frame.get("dominant_angle", ""),
                frame.get("world_relevance", ""),
                frame.get("zirp_relevance", ""),
                frame.get("recommended_next_step", ""),
            ])
        return re.sub(r"\s+", " ", " ".join(str(part or "") for part in parts)).strip()

    def extract_named_actors_for_signal(self, text: str) -> str:
        candidates = []
        patterns = [
            r"\b(?:UniversitÃ¤t|Hochschule|TH|RPTU|IQIB|DLR|DKKV|THW|Kreis|Stadt|Ministerium|Landkreis|Kammer|Verband|Initiative|Stiftung|Institut|Zentrum)\s+[A-ZÃ„Ã–Ãœ][\wÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]*(?:\s+[A-ZÃ„Ã–Ãœ][\wÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]*){0,5}",
            r"\b[A-ZÃ„Ã–Ãœ][a-zÃ¤Ã¶Ã¼ÃŸ]+(?:\s+[A-ZÃ„Ã–Ãœ][a-zÃ¤Ã¶Ã¼ÃŸ]+){1,3}\b",
        ]
        for pattern in patterns:
            for match in re.findall(pattern, text):
                cleaned = self.compact_signal_text(match, 80)
                if len(cleaned) >= 5 and cleaned.lower() not in {"rheinland pfalz"}:
                    candidates.append(cleaned)
        return ", ".join(self.deduplicate_list(candidates)[:8])

    def extract_concrete_assets_for_signal(self, text: str) -> str:
        assets = []
        asset_patterns = [
            r"\b[\wÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]*(?:Cloud|Radar|App|WebApp|Plattform|Tool|Karte|Einbaukarte|Dashboard|Digitaler Zwilling|FrÃ¼hwarnsystem|Programm|Projekt|Pilot|Modellprojekt|Studiengang|Roundtable|InnovationTalk|Demonstration|Demo|Messe|Zentrum|Netzwerk)[\wÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]*\b",
            r"\b[A-ZÃ„Ã–Ãœ]{2,}[A-ZÃ„Ã–Ãœ0-9\-]*\b",
        ]
        for pattern in asset_patterns:
            for match in re.findall(pattern, text):
                cleaned = self.compact_signal_text(match, 90)
                if len(cleaned) > 2 and cleaned.lower() not in {"der", "die", "das"}:
                    assets.append(cleaned)
        return ", ".join(self.deduplicate_list(assets)[:8])

    def extract_next_window_for_signal(self, text: str) -> str:
        date_patterns = [
            r"\b\d{1,2}\.\s*(?:Januar|Februar|MÃ¤rz|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s+\d{4}\b",
            r"\b\d{1,2}\.\d{1,2}\.\d{2,4}\b",
            r"\b(?:Januar|Februar|MÃ¤rz|April|Mai|Juni|Juli|August|September|Oktober|November|Dezember)\s+\d{4}\b",
            r"\b(?:FrÃ¼hjahr|Sommer|Herbst|Winter)\s+\d{4}\b",
            r"\bab\s+\d{4}\b",
            r"\bbis\s+\d{4}\b",
        ]
        matches = []
        for pattern in date_patterns:
            matches.extend(re.findall(pattern, text, flags=re.IGNORECASE))
        return ", ".join(self.deduplicate_list(self.compact_signal_text(match, 70) for match in matches)[:5])

    def infer_mechanism_for_signal(self, event: dict[str, Any], text: str) -> str:
        lowered = text.lower()
        checks = [
            ("digitales Werkzeug / Infrastruktur", ["tool", "plattform", "cloud", "app", "digital", "ki", "software", "dashboard", "zwilling"]),
            ("Pilot / Demonstrator / Umsetzung", ["pilot", "modellprojekt", "demonstration", "demo", "umsetzung", "praxis", "anlage", "system"]),
            ("Kooperation / Netzwerk / Transfer", ["kooperation", "partnerschaft", "netzwerk", "gemeinsam", "transfer", "verbund", "runde tisch"]),
            ("Qualifizierung / FachkrÃ¤fte / Bildung", ["studiengang", "weiterbildung", "ausbildung", "fachkrÃ¤fte", "qualifizierung", "hochschule"]),
            ("Regulierung / Verwaltung / Politik", ["verordnung", "gesetz", "ministerium", "kommune", "verwaltung", "land", "politik"]),
            ("Versorgung / Gesundheit / soziale Verantwortung", ["versorgung", "gesundheit", "pflege", "prÃ¤vention", "organspende", "patient"]),
            ("Standort / Investition / Wirtschaft", ["standort", "investition", "wirtschaft", "produktion", "markt", "arbeitsplÃ¤tze"]),
        ]
        found = [label for label, terms in checks if any(term in lowered for term in terms)]
        if found:
            return " + ".join(found[:3])
        section = str(event.get("hauptsektion", "") or "")
        return REPORT_SECTIONS.get(section, {}).get("title", section or "Beobachtung")

    def build_signal_opportunity_fields(self, event: dict[str, Any], frame: Optional[dict[str, Any]] = None) -> dict[str, str]:
        text = self.opportunity_text_for_event(event, frame)
        title = self.compact_signal_text(event.get("titel", ""), 140)
        snippet = self.compact_signal_text(event.get("snippet", ""), 260)
        member = self.compact_signal_text(event.get("mitglied", ""), 90)
        themes = ", ".join(event.get("themenfelder", []) or [])
        mechanism = self.infer_mechanism_for_signal(event, text)
        named_actors = self.extract_named_actors_for_signal(text)
        concrete_assets = self.extract_concrete_assets_for_signal(text)
        next_window = self.extract_next_window_for_signal(text)
        why = (
            self.compact_signal_text(event.get("llm_judge_reason", ""), 240)
            or self.compact_signal_text(event.get("zirp_reason", ""), 240)
            or self.compact_signal_text(frame.get("zirp_relevance", "") if frame else "", 240)
        )
        implementation_angle = self.compact_signal_text(
            frame.get("recommended_next_step", "") if frame else "",
            180,
        )
        if not implementation_angle:
            if concrete_assets:
                implementation_angle = "PrÃ¼fen, ob das genannte Tool, Projekt oder Format als Transferbeispiel fÃ¼r weitere Mitglieder taugt."
            elif "Kooperation" in mechanism or "Transfer" in mechanism:
                implementation_angle = "Sondieren, welche Akteure fÃ¼r einen kleinen Austausch oder Transfercheck fehlen."
            else:
                implementation_angle = "Als Beobachtung speichern und bei zweitem Signal im selben Themenfeld neu bewerten."

        scip_reasoning = self.compact_signal_text(
            f"{member} liefert mit '{title}' Evidenz fÃ¼r {mechanism}. "
            f"SCIP sollte dieses Signal mit Mitgliedern matchen, die komplementÃ¤re Rollen bei Umsetzung, Praxis, Finanzierung, Verwaltung oder Transfer haben.",
            320,
        )

        opportunity_strength = max(0, min(100, int(
            min(float(event.get("decision_score", 0) or 0) * 3, 45)
            + min(float(event.get("score", 0) or 0) * 1.5, 25)
            + (10 if concrete_assets else 0)
            + (10 if named_actors else 0)
            + (10 if next_window else 0)
        )))
        recommended_zirp_use = "beobachten"
        if opportunity_strength >= 78 and ("Pilot" in mechanism or "Kooperation" in mechanism or concrete_assets):
            recommended_zirp_use = "fÃ¼r direktes SondierungsgesprÃ¤ch nutzen"
        elif opportunity_strength >= 62:
            recommended_zirp_use = "fÃ¼r SCIP-Matching nutzen"
        elif opportunity_strength >= 45:
            recommended_zirp_use = "fÃ¼r Briefing nutzen"
        elif opportunity_strength < 25:
            recommended_zirp_use = "ignorieren"
        briefing_sentence = self.compact_signal_text(
            f"{title} ist fÃ¼r ZIRP vor allem als {mechanism.lower()} relevant, weil daraus {implementation_angle[:1].lower() + implementation_angle[1:]}",
            260,
        )

        return {
            "one_sentence_context": snippet or title,
            "why_it_matters": why,
            "mechanism": mechanism,
            "named_actors": named_actors,
            "concrete_assets": concrete_assets,
            "next_date_or_window": next_window,
            "implementation_angle": implementation_angle,
            "scip_reasoning": scip_reasoning,
            "opportunity_strength": str(opportunity_strength),
            "recommended_zirp_use": recommended_zirp_use,
            "briefing_sentence": briefing_sentence,
            "opportunity_keywords": self.compact_signal_text(", ".join(part for part in [themes, mechanism, concrete_assets] if part), 280),
        }

    def weekly_events_for_main_analysis(self) -> list[dict[str, Any]]:
        events = self.weekly_events()

        main_events = [event for event in events if event.get("evidence_role") == "main_signal"]
        context_events = [event for event in events if event.get("evidence_role") == "supporting_context"]
        watch_events = [event for event in events if event.get("evidence_role") == "watchlist"]

        def sort_key(event: dict[str, Any]) -> tuple[float, float]:
            return (
                float(event.get("decision_score", 0) or 0),
                float(event.get("score", 0) or 0),
            )

        main_events = sorted(main_events, key=sort_key, reverse=True)
        context_events = sorted(context_events, key=sort_key, reverse=True)
        watch_events = sorted(watch_events, key=sort_key, reverse=True)

        narrative_limit = self.dynamic_narrative_event_limit()
        selected = main_events[:6] + context_events[:5] + watch_events[:2]
        if selected:
            return selected[:narrative_limit]

        retained_events = [
            event for event in self.events
            if str(event.get("evidence_role", "")) in {"main_signal", "supporting_context", "watchlist"}
            and not self.is_macro_context_without_member_action(event)
        ]
        retained_events = sorted(retained_events, key=sort_key, reverse=True)
        return retained_events[:self.dynamic_narrative_event_limit()]

    def weekly_events(self) -> list[dict[str, Any]]:
        def select_events(days_back: int) -> list[dict[str, Any]]:
            today_start = self.today.replace(hour=0, minute=0, second=0, microsecond=0)
            cutoff = today_start - timedelta(days=days_back - 1)
            selected = []
            for event in self.events:
                if event["datum_obj"] < cutoff:
                    continue
                if (
                    not (self.openai_client is not None and MODEL_DECIDES_RELEVANCE)
                    and event["score"] < self.min_event_score
                    and str(event.get("evidence_role", "")) not in {"main_signal", "supporting_context", "watchlist"}
                ):
                    continue
                if len(str(event["snippet"]).split()) < MIN_SNIPPET_WORDS:
                    continue
                selected.append(event)

            member_counts = Counter(e["mitglied"] for e in selected)

            def adjusted_rank(e: dict[str, Any]):
                member_penalty = max(0, member_counts[e["mitglied"]] - 1) * 2
                return (
                    -(e["direct_member_score"] * 3 + e["regional_score"] * 2 - member_penalty),
                    -e["decision_score"],
                    -e["score"],
                    -e["datum_obj"].timestamp()
                )

            selected.sort(key=adjusted_rank)
            return selected

        events = select_events(REPORT_LOOKBACK_DAYS)
        if events:
            return events

        return []

    def build_section_profiles(self) -> list[dict[str, Any]]:
        weekly_events = self.weekly_events()
        profiles = []

        report_sections = self.report_sections_for_run()
        for section_key, cfg in report_sections.items():
            section_events = [e for e in weekly_events if e["hauptsektion"] == section_key]
            if not section_events:
                continue

            member_names = sorted({e["mitglied"] for e in section_events})
            top_events = sorted(
                section_events,
                key=lambda x: (
                    -self.briefing_rank_score(x, weekly_events),
                    -x["decision_score"],
                    -x["score"],
                    -x["datum_obj"].timestamp()
                )
            )[:MAX_REPORT_EVENTS_PER_SECTION]

            profiles.append({
                "section_key": section_key,
                "title": cfg["title"],
                "event_count": len(section_events),
                "member_count": len(member_names),
                "member_names": member_names,
                "top_events": top_events,
            })

        order = list(self.report_sections_for_run().keys())
        profiles.sort(key=lambda x: order.index(x["section_key"]))
        return profiles

    def build_evidence_profiles(self) -> dict[str, list[dict[str, Any]]]:
        weekly_events = self.weekly_events()
        result = {}

        for section_key in self.report_sections_for_run().keys():
            items = [e for e in weekly_events if e["hauptsektion"] == section_key]
            if self.openai_client is not None and MODEL_DECIDES_RELEVANCE:
                items = [
                    e for e in items
                    if self.editorial_frame_for(e).get("role_in_briefing") in MODEL_KEEP_ROLES
                ]
            else:
                items = [
                    e for e in items
                    if self.fallback_rule_relevance_gate(e)
                ]
            items.sort(
                key=lambda x: (
                    -self.briefing_rank_score(x, weekly_events),
                    -x["decision_score"],
                    -x["score"],
                    -x["datum_obj"].timestamp()
                )
            )
            result[section_key] = items[:MAX_EVIDENCE_EVENTS_PER_SECTION]

        return result

    def detect_emergent_week_topics(self, min_count: int = 2, top_n: int = 6) -> list[dict[str, Any]]:
        weekly_events = self.weekly_events()
        if not weekly_events:
            return []

        token_counter = Counter()
        member_spread: dict[str, set[str]] = defaultdict(set)
        event_spread: dict[str, set[tuple[str, str, str]]] = defaultdict(set)

        for event in weekly_events:
            text = f"{event.get('titel', '')} {event.get('snippet', '')}"
            tokens = self.tokenize(text)
            seen_in_event = set()

            for token in tokens:
                if token in GERMAN_STOPWORDS:
                    continue
                if token in EMERGENT_TERM_BLACKLIST:
                    continue
                if len(token) < 6:
                    continue
                if token.isdigit():
                    continue
                if token not in seen_in_event:
                    token_counter[token] += 1
                    member_spread[token].add(event["mitglied"])
                    event_spread[token].add((event["mitglied"], event["titel"], event["datum"]))
                    seen_in_event.add(token)

        emergent = []
        for token, count in token_counter.items():
            spread = len(member_spread[token])
            distinct_events = len(event_spread[token])
            score = count + (spread * 2)
            if count >= min_count and distinct_events >= 2:
                emergent.append({
                    "term": token,
                    "count": count,
                    "member_spread": spread,
                    "distinct_events": distinct_events,
                    "score": score,
                })

        emergent.sort(key=lambda x: (-x["score"], -x["count"], x["term"]))
        return emergent[:top_n]

    def classify_signal_scope(self, event: dict[str, Any]) -> str:
        if int(event.get("direct_member_score", 0)) >= 3:
            return "direktes Mitgliedssignal"
        return "Kontextsignal aus dem Mitgliedsumfeld"

    def classify_signal_strength(self, event: dict[str, Any], weekly_events: list[dict[str, Any]]) -> str:
        section = event.get("hauptsektion")
        same_section = [e for e in weekly_events if e.get("hauptsektion") == section]
        member_count = len({e["mitglied"] for e in same_section})
        decision_score = int(event.get("decision_score", 0))
        direct_member_score = int(event.get("direct_member_score", 0))
        regional_score = int(event.get("regional_score", 0))

        if len(same_section) >= 4 and member_count >= 3 and decision_score >= 12 and regional_score >= 2:
            return "strategischer Trend"
        if len(same_section) >= 2 and member_count >= 2 and decision_score >= 9:
            return "wiederholtes Muster"
        if decision_score >= 10 and direct_member_score >= 3:
            return "frÃ¼hes Signal"
        return "isolierter Fall"

    def build_what_happened(self, event: dict[str, Any]) -> str:
        title = str(event.get("titel", "")).strip()
        snippet = re.sub(r"\s+", " ", str(event.get("snippet", "")).strip())

        if len(snippet) > 220:
            snippet = snippet[:217].rstrip() + "..."

        if snippet:
            return f"Bei **{event['mitglied']}** fÃ¤llt die Meldung **â€ž{title}â€œ** auf. {snippet}"
        return f"Bei **{event['mitglied']}** fÃ¤llt die Meldung **â€ž{title}â€œ** auf."

    def build_why_it_matters(self, event: dict[str, Any], strength: str) -> str:
        section = event.get("hauptsektion")
        themes = set(event.get("themenfelder", []))
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()

        if section == "wirtschaftsentwicklung":
            if "standort" in text or "werk" in text or "investition" in text:
                return "Relevant, weil sich daran Investitions-, Auslastungs- oder Anpassungsdruck in wirtschaftlich wichtigen Standorten ablesen lÃ¤sst."
            if "fachkrÃ¤fte" in text or "personal" in text:
                return "Relevant, weil sich daran zeigt, wie stark PersonalengpÃ¤sse inzwischen in zentrale Versorgungs- und Wirtschaftsbereiche hineinwirken."
            if any(t in text for t in MACRO_COMMENTARY_TERMS):
                return "Relevant vor allem als wirtschaftlicher Kontext, weniger als eigenstÃ¤ndiges institutionelles Leitsignal."
            return "Relevant, weil das Signal auf wirtschaftliche Anpassung, Standortdruck oder Investitionsbedarf verweist."

        if section == "versorgung_gesundheit":
            return "Relevant, weil das Signal Fragen der Versorgung, sozialen Infrastruktur oder institutionellen TragfÃ¤higkeit in Rheinland-Pfalz berÃ¼hrt."

        if section == "kooperationen":
            return "Relevant, weil hier sichtbar wird, wo neue Kooperationsmodelle, institutionelle Antworten oder gemeinsame ProblemlÃ¶sungen entstehen."

        if section == "fuehrungswechsel":
            return "Relevant, weil Personalwechsel hÃ¤ufig auf verÃ¤nderte strategische PrioritÃ¤ten, neue Ansprechpartner oder Verschiebungen im institutionellen Kurs hindeuten."

        if section == "soziales_engagement":
            return "Relevant, weil sich daran ablesen lÃ¤sst, wo gesellschaftliche Verantwortung, regionale Verankerung oder sozialpolitische AnschlussfÃ¤higkeit sichtbar werden."

        if "gesundheit" in themes or "gesellschaft" in themes:
            return "Relevant, weil das Signal Fragen der Versorgung, sozialen Infrastruktur oder institutionellen TragfÃ¤higkeit berÃ¼hrt."

        return f"Relevant, weil es als **{strength}** Ã¼ber den Einzelfall hinaus beobachtungswÃ¼rdig ist."

    def build_zirp_relevance(self, event: dict[str, Any]) -> str:
        section = event.get("hauptsektion")
        themes = set(event.get("themenfelder", []))
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()

        if section == "versorgung_gesundheit" or any(
                x in text for x in ["fachkrÃ¤fte", "Ã¤rzte", "hausarzt", "pflege", "versorgung"]):
            return "FÃ¼r die ZIRP relevant, weil sich hier FachkrÃ¤ftesicherung, Daseinsvorsorge und regionale ZukunftsfÃ¤higkeit verbinden."

        if "standort" in text or "werk" in text or "investition" in text or "auslastung" in text:
            return "FÃ¼r die ZIRP relevant, weil das Signal an Diskussionen Ã¼ber WettbewerbsfÃ¤higkeit, Transformation und StandortstabilitÃ¤t anschlussfÃ¤hig ist."

        if "kommune" in text or "stadt" in text or "praxis" in text:
            return "FÃ¼r die ZIRP relevant, weil hier kommunale InnovationsfÃ¤higkeit und neue Organisationsantworten auf VersorgungsengpÃ¤sse sichtbar werden."

        if section == "kooperationen":
            return "FÃ¼r die ZIRP relevant, wenn daraus neue GesprÃ¤chsanlÃ¤sse, LernrÃ¤ume oder Formate zwischen Wirtschaft, Wissenschaft, Politik und Gesellschaft entstehen."

        if section == "soziales_engagement":
            return "FÃ¼r die ZIRP relevant, wenn sich daraus Themen fÃ¼r gesellschaftliche Wirkung, regionale Verantwortung oder anschlussfÃ¤hige ErzÃ¤hlungen ergeben."

        if "wirtschaft" in themes:
            return "FÃ¼r die ZIRP relevant, weil das Signal wirtschaftliche Strukturfragen mit programmatischen Zukunftsthemen verbindet."

        return "FÃ¼r die ZIRP relevant, wenn sich das Signal in den kommenden Wochen verdichtet und in GesprÃ¤che, Formate oder Themenplanung Ã¼bersetzt werden kann."

    def suggest_next_step(self, event: dict[str, Any], strength: str, scope: str) -> str:
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()

        if any(t in text for t in MACRO_COMMENTARY_TERMS):
            return "nur als Kontextsignal beobachten"
        if strength == "strategischer Trend":
            return "GesprÃ¤chspunkt vorbereiten"
        if "fachkrÃ¤fte" in text or "hausarzt" in text or "pflege" in text or "versorgung" in text:
            return "als mÃ¶gliches ZIRPzoom-Thema prÃ¼fen"
        if "kooperation" in text or "partnerschaft" in text:
            return "relevanten Akteur fÃ¼r AnschlussgesprÃ¤ch prÃ¼fen"
        if scope == "Kontextsignal aus dem Mitgliedsumfeld":
            return "weiter beobachten"
        if strength == "wiederholtes Muster":
            return "Talking Points vorbereiten"
        return "vorerst beobachten"

    def build_source_line(self, event: dict[str, Any]) -> str:
        return (
            f"**Quelle:** {event['mitglied']} | {event['datum']} | "
            f"{event['titel']} | {event['url']}"
        )

    def strategic_pattern_score(self, event: dict[str, Any], weekly_events: list[dict[str, Any]]) -> int:
        section = event.get("hauptsektion")
        themes = set(event.get("themenfelder", []))
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()

        section_peers = [e for e in weekly_events if e.get("hauptsektion") == section]
        member_spread = len({e["mitglied"] for e in section_peers})
        section_volume = len(section_peers)

        score = 0
        score += min(event.get("decision_score", 0), 20)
        score += min(event.get("regional_score", 0) * 2, 6)
        score += min(len(themes) * 2, 6)
        score += min(member_spread, 4)
        score += 2 if section_volume >= 3 else 0
        score += sum(1 for t in STRONG_STRATEGIC_TERMS if t in text)
        score += min(sum(1 for t in INSTITUTIONAL_ACTION_TERMS if t in text), 4)
        score += 2 if event.get("hauptsektion") in {"kooperationen", "versorgung_gesundheit"} else 0
        score += 1 if event.get("is_updated") else 0
        score -= min(sum(1 for t in MACRO_COMMENTARY_TERMS if t in text), 4)
        score -= 2 if any(t in text for t in WEAK_SIGNAL_TERMS) else 0
        return score

    def briefing_rank_score(self, event: dict[str, Any], weekly_events: list[dict[str, Any]]) -> int:
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()

        strong_hits = sum(1 for t in STRONG_STRATEGIC_TERMS if t in text)
        institutional_hits = sum(1 for t in INSTITUTIONAL_ACTION_TERMS if t in text)
        regional_impl_hits = sum(1 for t in REGIONAL_IMPLEMENTATION_TERMS if t in text)
        macro_hits = sum(1 for t in MACRO_COMMENTARY_TERMS if t in text)

        score = 0
        score += event.get("decision_score", 0) * 3
        score += event.get("regional_score", 0) * 3
        score += event.get("direct_member_score", 0) * 2
        score += event.get("recency_weight", 0) * 2
        score += len(event.get("themenfelder", [])) * 2
        score += strong_hits * 2
        score += institutional_hits * 3
        score += regional_impl_hits * 2
        score += 4 if event.get("hauptsektion") in {"kooperationen", "versorgung_gesundheit",
                                                    "wirtschaftsentwicklung"} else 0
        score -= macro_hits * 4
        score -= 4 if any(t in text for t in WEAK_SIGNAL_TERMS) else 0

        if macro_hits >= 2 and institutional_hits == 0 and regional_impl_hits == 0:
            score -= 10

        return score

    def watchlist_rank_score(self, event: dict[str, Any], weekly_events: list[dict[str, Any]]) -> int:
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        macro_hits = sum(1 for t in MACRO_COMMENTARY_TERMS if t in text)
        institutional_hits = sum(1 for t in INSTITUTIONAL_ACTION_TERMS if t in text)

        score = 0
        score += event.get("decision_score", 0) * 2
        score += event.get("recency_weight", 0) * 2
        score += len(event.get("themenfelder", []))
        score += 2 if event.get("is_updated") else 0
        score += 2 if any(t in text for t in STRONG_STRATEGIC_TERMS) else 0
        score += min(institutional_hits, 3)
        score -= min(macro_hits, 3)
        score -= 3 if event.get("hauptsektion") == "soziales_engagement" and not any(
            x in text for x in ["fachkrÃ¤fte", "versorgung", "kooperation", "transfer", "innovation"]
        ) else 0
        return score

    def select_top_signals(self, weekly_events: list[dict[str, Any]], n: int = 3) -> list[dict[str, Any]]:
        ranked = sorted(
            weekly_events,
            key=lambda x: (
                -self.briefing_rank_score(x, weekly_events),
                -x["decision_score"],
                -x["score"],
                -x["datum_obj"].timestamp()
            )
        )

        selected = []
        used_sections = set()
        used_members = set()

        for ev in ranked:
            sec = ev.get("hauptsektion")
            member = ev.get("mitglied")
            if member in used_members:
                continue
            if sec in used_sections and len(selected) < 2:
                continue
            selected.append(ev)
            used_sections.add(sec)
            used_members.add(member)
            if len(selected) >= n:
                break

        return selected[:n]

    def select_strategic_condensations(
            self,
            weekly_events: list[dict[str, Any]],
            top_signals: list[dict[str, Any]],
            n: int = 3
    ) -> list[dict[str, Any]]:
        remaining = [e for e in weekly_events if e not in top_signals]
        ranked = sorted(
            remaining,
            key=lambda x: (
                -self.strategic_pattern_score(x, weekly_events),
                -x["decision_score"],
                -x["score"],
                -x["datum_obj"].timestamp()
            )
        )

        selected = []
        used_sections = set()
        for ev in ranked:
            sec = ev.get("hauptsektion")
            if sec in used_sections and len(selected) < 2:
                continue
            selected.append(ev)
            used_sections.add(sec)
            if len(selected) >= n:
                break

        return selected[:n]

    def select_watchlist(
            self,
            weekly_events: list[dict[str, Any]],
            excluded_events: list[dict[str, Any]],
            n: int = 3
    ) -> list[dict[str, Any]]:
        remaining = [e for e in weekly_events if e not in excluded_events]
        ranked = sorted(
            remaining,
            key=lambda x: (
                -self.watchlist_rank_score(x, weekly_events),
                -x["datum_obj"].timestamp()
            )
        )

        selected = []
        used_members = {e["mitglied"] for e in excluded_events}
        for ev in ranked:
            if ev["mitglied"] in used_members:
                continue
            selected.append(ev)
            used_members.add(ev["mitglied"])
            if len(selected) >= n:
                break

        return selected[:n]

    def build_weekly_thesis(self, top_signals: list[dict[str, Any]]) -> str:
        text = " ".join(
            f"{e.get('titel', '')} {e.get('snippet', '')}" for e in top_signals
        ).lower()

        coop = any(x in text for x in ["kooperation", "partnerschaft", "allianz", "netzwerk", "transfer"])
        health = any(x in text for x in ["gesundheit", "versorgung", "pflege", "krankenkasse", "gesundheitssystem"])
        tech = any(
            x in text for x in ["quantentechnologien", "6g", "industrie 4.0", "digitalisierung", "innovation"])
        labor = any(x in text for x in ["fachkrÃ¤fte", "personal", "arbeitskrÃ¤fte"])
        regional = any(x in text for x in ["rheinland-pfalz", "mainz", "trier", "kaiserslautern", "kommun"])
        climate = any(x in text for x in ["klima", "biodiversitÃ¤t", "umweltschutz", "nachhaltigkeit"])

        if tech and coop:
            return (
                "Die Woche wird vor allem von **Kooperations- und Innovationssignalen** geprÃ¤gt. "
                "AuffÃ¤llig ist, dass technologische Zukunftsthemen nicht isoliert auftreten, sondern in "
                "**Partnerschaften, Transferbeziehungen und institutionellen AnschlussmÃ¶glichkeiten** sichtbar werden. "
                "FÃ¼r die ZIRP ist das relevant, weil daraus GesprÃ¤chs- und Formatpotenziale an der Schnittstelle von Wirtschaft, Wissenschaft und regionaler Entwicklung entstehen."
            )

        if health and labor:
            return (
                "Die Woche verdichtet sich vor allem in Richtung **Versorgung, FachkrÃ¤ftedruck und institutioneller ReaktionsfÃ¤higkeit**. "
                "Die relevanten Signale deuten darauf hin, dass soziale Infrastruktur und ZukunftsfÃ¤higkeit zunehmend gemeinsam betrachtet werden mÃ¼ssen. "
                "FÃ¼r die ZIRP ist das vor allem als Themen- und GesprÃ¤chsraum anschlussfÃ¤hig."
            )

        if regional and coop:
            return (
                "Die ausgewÃ¤hlten Signale sprechen weniger fÃ¼r eine breite Lageverschiebung als fÃ¼r eine **zunehmende regionale Verdichtung strategischer Kooperationen**. "
                "FÃ¼r die ZIRP ist dies relevant, weil daraus Anschlussstellen fÃ¼r Vernetzung, Austauschformate und gemeinsame Problemwahrnehmung entstehen."
            )

        if regional and climate:
            return (
                "Die Woche zeigt vor allem **regionale Umsetzungs- und Koordinationssignale** an der Schnittstelle von Klima, Organisation und institutioneller Zusammenarbeit. "
                "FÃ¼r die ZIRP ist das relevant, weil daraus konkrete GesprÃ¤chsanlÃ¤sse zu kommunaler und regionaler ZukunftsfÃ¤higkeit entstehen."
            )

        return (
            "Die Woche liefert einige belastbare Hinweise auf **strategische Anpassung, Kooperationsdynamik und sektorÃ¼bergreifende Reaktionsbedarfe**. "
            "FÃ¼r die ZIRP ergibt sich daraus vor allem GesprÃ¤chs-, Beobachtungs- und Strukturierungspotenzial."
        )

    def editorial_frame_schema(self) -> dict[str, Any]:
        return {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "event_type": {
                    "type": "string",
                    "enum": [
                        "institutional_cooperation",
                        "regional_implementation",
                        "investment_or_expansion",
                        "leadership_change",
                        "healthcare_or_supply_pressure",
                        "research_transfer",
                        "social_engagement",
                        "macro_context",
                        "other"
                    ]
                },
                "dominant_angle": {"type": "string"},
                "world_relevance": {"type": "string"},
                "zirp_relevance": {"type": "string"},
                "recommended_next_step": {"type": "string"},
                "confidence": {
                    "type": "string",
                    "enum": ["low", "medium", "high"]
                },
                "is_context_only": {"type": "boolean"},
                "role_in_briefing": {
                    "type": "string",
                    "enum": ["top_signal", "condensation", "watchlist", "ignore"]
                },
                "why_now": {"type": "string"},
                "tone": {
                    "type": "string",
                    "enum": ["confident", "careful", "observational"]
                },
            },
            "required": [
                "event_type",
                "dominant_angle",
                "world_relevance",
                "zirp_relevance",
                "recommended_next_step",
                "confidence",
                "is_context_only",
                "role_in_briefing",
                "why_now",
                "tone",
            ],
        }

    def weekly_meta_schema(self) -> dict[str, Any]:
        return {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "headline": {"type": "string"},
                "meta_line": {"type": "string"},
                "narrative_analysis": {"type": "string"},
                "weekly_thesis": {"type": "string"},
                "one_line_conclusion": {"type": "string"},
                "strategic_summary": {"type": "string"},
            },
            "required": [
                "headline",
                "meta_line",
                "narrative_analysis",
                "weekly_thesis",
                "one_line_conclusion",
                "strategic_summary",
            ],
        }

    def event_payload_for_model(self, event: dict[str, Any]) -> dict[str, Any]:
        return {
            "member_id": event.get("member_id"),
            "member_name": event.get("mitglied"),
            "title": event.get("titel"),
            "date": event.get("datum"),
            "url": event.get("url"),
            "main_section": event.get("hauptsektion"),
            "themes": event.get("themenfelder", []),
            "snippet": event.get("snippet", ""),
            "score": event.get("score", 0),
            "decision_score": event.get("decision_score", 0),
            "recency_weight": event.get("recency_weight", 0),
            "direct_member_score": event.get("direct_member_score", 0),
            "regional_score": event.get("regional_score", 0),
            "article_quality_score": event.get("article_quality_score", 0),
            "is_updated": event.get("is_updated", False),
            "signal_scope": self.classify_signal_scope(event),
            "signal_strength": self.classify_signal_strength(event, self.weekly_events()),
        }

    def build_editorial_system_prompt(self) -> str:
        return (
            "Du bist die zentrale Relevanzentscheidung fÃ¼r ein ZIRP-Leiterinnen-Briefing. "
            "Die regelbasierte Pipeline hat nur technische Vorarbeit geleistet: Artikelqualitaet, Datum, Titel, Snippet und grobe Themenzuordnung. "
            "Du entscheidest nun, ob das Ereignis ins Briefing gehoert. "
            "Nutze role_in_briefing strikt: top_signal fÃ¼r starke Leitungsrelevanz, condensation fÃ¼r verdichtende Muster, watchlist fÃ¼r beobachtenswerte schwaechere Signale, ignore fÃ¼r irrelevante oder zu duenne Funde. "
            "Ueberhoehe schwache Signale nicht, aber verwerfe interessante Signale nicht nur deshalb, weil einzelne Schlagworte fehlen. "
            "Bevorzuge konkrete Umsetzungs-, Kooperations-, Standort-, Versorgungs-, Fuehrungs-, Transfer- und regionale Zukunftsrelevanz. "
            "Makro-Kommentare ohne institutionelle Handlungsebene sind meist watchlist oder ignore. "
            "Gib ausschliesslich gueltiges JSON nach Schema zurueck."
        )

    def build_event_user_prompt(self, payload: dict[str, Any]) -> str:
        return (
            "Bewerte das folgende Ereignis fÃ¼r ein strategisches Wochenbriefing in Rheinland-Pfalz.\n\n"
            f"{json.dumps(payload, ensure_ascii=False, indent=2)}\n\n"
            "Ordne das Ereignis inhaltlich ein, entscheide seine redaktionelle Rolle "
            "und bleibe eng an den vorhandenen Signalen."
        )

    def request_structured_json(
            self,
            *,
            schema_name: str,
            schema: dict[str, Any],
            system_prompt: str,
            user_prompt: str,
    ) -> Optional[dict[str, Any]]:
        if self.openai_client is None:
            return None

        for _ in range(max(1, OPENAI_MAX_RETRIES + 1)):
            try:
                if self.openai_base_url:
                    result = self.request_local_structured_json(
                        schema_name=schema_name,
                        schema=schema,
                        system_prompt=system_prompt,
                        user_prompt=user_prompt,
                    )
                    if result:
                        return result
                    continue

                response = self.openai_client.responses.create(
                    model=self.openai_model,
                    input=[
                        {
                            "role": "system",
                            "content": [{"type": "input_text", "text": system_prompt}],
                        },
                        {
                            "role": "user",
                            "content": [{"type": "input_text", "text": user_prompt}],
                        },
                    ],
                    text={
                        "format": {
                            "type": "json_schema",
                            "name": schema_name,
                            "schema": schema,
                            "strict": EDITORIAL_STRICT_MODE,
                        }
                    },
                )

                raw_text = getattr(response, "output_text", "") or ""
                if raw_text:
                    return json.loads(raw_text)

                raw = response.model_dump() if hasattr(response, "model_dump") else None
                if raw:
                    extracted = self.extract_json_text_from_response(raw)
                    if extracted:
                        return json.loads(extracted)
            except Exception as exc:
                print(f"OpenAI editorial: strukturierte Anfrage fehlgeschlagen ({schema_name}): {exc}")
                continue

        return None

    def request_local_structured_json(
            self,
            *,
            schema_name: str,
            schema: dict[str, Any],
            system_prompt: str,
            user_prompt: str,
    ) -> Optional[dict[str, Any]]:
        started = time.perf_counter()
        prompt_chars = len(system_prompt) + len(user_prompt) + len(json.dumps(schema, ensure_ascii=False))
        print(
            f"LLM profile start: local_json schema={schema_name} model={self.openai_model} "
            f"prompt_chars={prompt_chars} max_tokens={OPENAI_MAX_OUTPUT_TOKENS} timeout={self.openai_timeout}s"
        )
        try:
            messages = [
                {
                    "role": "system",
                    "content": (
                        f"{system_prompt}\n\n"
                        "Gib ausschliesslich ein valides JSON-Objekt zurueck. "
                        "Kein Markdown, keine Erklaerung, keine Code-Fences. "
                        f"Das JSON-Objekt muss diesem Schema namens {schema_name} entsprechen: "
                        f"{json.dumps(schema, ensure_ascii=False)}"
                    ),
                },
                {"role": "user", "content": user_prompt},
            ]
            response = self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=messages,
                temperature=0.2,
                max_tokens=OPENAI_MAX_OUTPUT_TOKENS,
                timeout=self.openai_timeout,
            )
            content = response.choices[0].message.content if response.choices else ""
            duration = time.perf_counter() - started
            print(
                f"LLM profile end: local_json schema={schema_name} duration={duration:.1f}s "
                f"output_chars={len(content or '')} success={bool(content)}"
            )
            return self.parse_json_object(content)
        except Exception as exc:
            duration = time.perf_counter() - started
            print(
                f"OpenAI editorial: lokale JSON-Anfrage fehlgeschlagen ({schema_name}) "
                f"after {duration:.1f}s: {exc}"
            )
            return None

    def request_local_plain_text(
        self,
        *,
        system_prompt: str,
        user_prompt: str,
        call_name: str = "plain_text",
        max_tokens: Optional[int] = None,
        model: Optional[str] = None,
    ) -> Optional[str]:
        started = time.perf_counter()
        output_tokens = int(max_tokens or OPENAI_MAX_OUTPUT_TOKENS)
        request_model = (model or self.openai_model).strip()
        if self.openai_base_url and "qwen3" in request_model.lower():
            user_prompt = "/no_think\n" + user_prompt
            system_prompt = (
                system_prompt
                + " Antworte direkt ohne sichtbaren Denkprozess, ohne <think>-Block und ohne Vorrede."
            )
        prompt_chars = len(system_prompt) + len(user_prompt)
        approx_prompt_tokens = max(1, prompt_chars // 4)
        print(
            f"LLM profile start: {call_name} model={request_model} "
            f"prompt_chars={prompt_chars} approx_prompt_tokens={approx_prompt_tokens} "
            f"max_tokens={output_tokens} timeout={self.openai_timeout}s retries={OPENAI_MAX_RETRIES}"
        )
        try:
            request_kwargs = {
                "model": request_model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0.55,
                "max_tokens": output_tokens,
                "timeout": self.openai_timeout,
            }
            if self.openai_base_url:
                request_kwargs["extra_body"] = {
                    "think": False,
                    "options": {
                        "num_predict": output_tokens,
                        "temperature": 0.55,
                        "top_p": 0.9,
                        "repeat_penalty": 1.12,
                        "num_ctx": 4096,
                    }
                }
            response = self.openai_client.chat.completions.create(**request_kwargs)
            text = response.choices[0].message.content if response.choices else ""
            cleaned = self.clean_generated_narrative(text)
            duration = time.perf_counter() - started
            words = len(re.findall(r"\b\w+\b", cleaned, flags=re.UNICODE))
            print(
                f"LLM profile end: {call_name} duration={duration:.1f}s "
                f"output_chars={len(text or '')} cleaned_chars={len(cleaned or '')} words={words} success={bool(cleaned)}"
            )
            return cleaned or None
        except Exception as exc:
            duration = time.perf_counter() - started
            print(f"OpenAI editorial: lokale Textanalyse fehlgeschlagen ({call_name}) after {duration:.1f}s: {exc}")
            return None

    def request_report_plain_text(
        self,
        *,
        system_prompt: str,
        user_prompt: str,
        call_name: str = "report_text",
        max_tokens: Optional[int] = None,
    ) -> Optional[str]:
        if self.report_openai_client is None:
            return None

        started = time.perf_counter()
        output_tokens = int(max_tokens or WEEKLY_REWRITE_MAX_TOKENS)
        prompt_chars = len(system_prompt) + len(user_prompt)
        approx_prompt_tokens = max(1, prompt_chars // 4)
        print(
            f"LLM profile start: {call_name} model={self.report_openai_model} "
            f"prompt_chars={prompt_chars} approx_prompt_tokens={approx_prompt_tokens} "
            f"max_tokens={output_tokens} timeout={self.report_openai_timeout}s retries=0"
        )
        try:
            response = self.report_openai_client.chat.completions.create(
                model=self.report_openai_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.35,
                max_tokens=output_tokens,
                timeout=self.report_openai_timeout,
            )
            text = response.choices[0].message.content if response.choices else ""
            cleaned = self.clean_generated_narrative(text)
            duration = time.perf_counter() - started
            words = len(re.findall(r"\b\w+\b", cleaned, flags=re.UNICODE))
            print(
                f"LLM profile end: {call_name} duration={duration:.1f}s "
                f"output_chars={len(text or '')} cleaned_chars={len(cleaned or '')} words={words} success={bool(cleaned)}"
            )
            return cleaned or None
        except Exception as exc:
            duration = time.perf_counter() - started
            print(f"Report editorial: Cloud-Textanalyse fehlgeschlagen ({call_name}) after {duration:.1f}s: {exc}")
            return None

    def clean_generated_narrative(self, text: str) -> str:
        cleaned = str(text or "").strip()
        cleaned = re.sub(r"^```(?:text|markdown)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
        cleaned = re.sub(r"(?m)^\s*#{1,6}\s*", "", cleaned)
        cleaned = re.sub(r"(?m)^\s*\*\*\s*$", "", cleaned)
        cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"(?m)^\s*[-*]\s+(DO|WATCH|MAYBE|DROP):", r"\1:", cleaned)
        replacements = {
            "Pruef": "PrÃ¼f",
            "pruef": "prÃ¼f",
            "fÃ¼r": "fÃ¼r",
            "waere": "wÃ¤re",
            "kÃ¶nnte": "kÃ¶nnte",
            "muessen": "mÃ¼ssen",
            "Ã¼ber": "Ã¼ber",
            "schaerf": "schÃ¤rf",
            "oeffentlich": "Ã¶ffentlich",
            "GesprÃ¤ch": "GesprÃ¤ch",
            "SondierungsgesprÃ¤ch": "SondierungsgesprÃ¤ch",
        }
        for old, new in replacements.items():
            cleaned = cleaned.replace(old, new)
        cleaned = self.repair_mojibake(cleaned)
        cleaned = re.sub(r"(?im)^\s*Wozu:\s*$", "", cleaned)
        cleaned = re.sub(r"(?im)^\s*Zusammenfassung\s*$", "", cleaned)
        cleaned = re.sub(r"(?im)^\s*[-â€“â€”]{3,}\s*$", "", cleaned)
        cleaned = re.sub(r"(?im)^\s*Die Analyse beendet sich hiermit\.?\s*$", "", cleaned)
        cleaned = re.sub(r"Signalsignaten|Signalien", "Signale", cleaned)
        cleaned = re.sub(r"Ankeuren", "Akteuren", cleaned)
        cleaned = re.sub(r"(?i)Chancen maximieren und Risiken minimieren", "konkrete Anschlusschancen prÃ¼fen", cleaned)
        cleaned = re.sub(r"(?i)immer auf dem neuesten Stand sein", "die Entwicklung weiter verfolgen", cleaned)
        cleaned = re.sub(r"(?i)die richtigen Schritte unternehmen", "nÃ¤chste Schritte ableiten", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def parse_json_object(self, text: str) -> Optional[dict[str, Any]]:
        if not text:
            return None
        cleaned = self.extract_first_json_object(text)
        if not cleaned:
            return None
        try:
            value = json.loads(cleaned)
            return value if isinstance(value, dict) else None
        except Exception:
            return None

    def extract_first_json_object(self, text: str) -> str:
        cleaned = str(text).strip()
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()

        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return ""
        return cleaned[start:end + 1]

    def parse_json_object_lenient(self, text: str) -> Optional[dict[str, Any]]:
        try:
            value = json.loads(self.extract_first_json_object(text))
            return value if isinstance(value, dict) else None
        except Exception:
            return None

    def extract_json_text_from_response(self, raw: dict[str, Any]) -> Optional[str]:
        try:
            output = raw.get("output", [])
            for item in output:
                for content in item.get("content", []):
                    if content.get("type") in {"output_text", "text"} and content.get("text"):
                        return content["text"]
        except Exception:
            return None
        return None

    def build_rule_based_editorial_frame(self, event: dict[str, Any]) -> dict[str, Any]:
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        section = event.get("hauptsektion", "")
        strength = self.classify_signal_strength(event, self.weekly_events())

        event_type = "other"
        if section == "kooperationen":
            event_type = "institutional_cooperation"
        elif section == "versorgung_gesundheit":
            event_type = "healthcare_or_supply_pressure"
        elif section == "fuehrungswechsel":
            event_type = "leadership_change"
        elif any(x in text for x in ["investition", "ausbau", "standort", "werk", "produktion"]):
            event_type = "investment_or_expansion"
        elif any(x in text for x in ["transfer", "forschung", "hochschule", "wissenschaft"]):
            event_type = "research_transfer"
        elif any(x in text for x in MACRO_COMMENTARY_TERMS):
            event_type = "macro_context"

        is_context_only = event_type == "macro_context"
        role = "watchlist" if is_context_only else "top_signal"
        if strength == "isolierter Fall":
            role = "watchlist"
        elif strength == "wiederholtes Muster":
            role = "condensation"
        elif strength == "strategischer Trend":
            role = "top_signal"

        return {
            "event_type": event_type,
            "dominant_angle": self.build_why_it_matters(event, strength),
            "world_relevance": self.build_why_it_matters(event, strength),
            "zirp_relevance": self.build_zirp_relevance(event),
            "recommended_next_step": self.suggest_next_step(
                event,
                strength,
                self.classify_signal_scope(event),
            ),
            "confidence": "high" if event.get("decision_score", 0) >= 12 else "medium",
            "is_context_only": is_context_only,
            "role_in_briefing": role,
            "why_now": "zeitnahes Signal innerhalb des aktuellen Beobachtungsfensters",
            "tone": "careful" if is_context_only else "confident",
        }

    def enrich_event_editorially(self, event: dict[str, Any]) -> dict[str, Any]:
        fallback = self.build_rule_based_editorial_frame(event)

        payload = self.event_payload_for_model(event)
        result = self.request_structured_json(
            schema_name="zirp_editorial_frame",
            schema=self.editorial_frame_schema(),
            system_prompt=self.build_editorial_system_prompt(),
            user_prompt=self.build_event_user_prompt(payload),
        )

        if not result:
            fallback["role_in_briefing"] = "watchlist"
            fallback["confidence"] = "low"
            fallback["tone"] = "observational"
            fallback["model_used"] = False
            return fallback

        merged = dict(fallback)
        merged.update(result)
        merged["model_used"] = True
        return merged

    def enrich_weekly_events_editorially(self, weekly_events: list[dict[str, Any]]) -> None:
        self.editorial_frames = {}
        for event in weekly_events:
            key = self._event_key(event)
            self.editorial_frames[key] = self.enrich_event_editorially(event)

    def editorial_frame_for(self, event: dict[str, Any]) -> dict[str, Any]:
        return self.editorial_frames.get(
            self._event_key(event),
            self.build_rule_based_editorial_frame(event),
        )

    def scip_convening_radar_for_prompt(self) -> dict[str, Any]:
        radar = self.latest_meeting_recommendations()
        if not radar:
            return {"count": 0, "cards": []}

        cards = []
        prompt_card_limit = self.dynamic_scip_prompt_card_limit()
        for item in radar.get("cards", [])[:prompt_card_limit]:
            cards.append({
                "members": self.clean_dashboard_value(item.get("members", "")),
                "score": self.clean_dashboard_value(item.get("score", "")),
                "empfehlung": self.clean_dashboard_value(item.get("empfehlung", "")),
                "theme": self.clean_dashboard_value(item.get("theme", "")),
                "convening_typology": self.clean_dashboard_value(item.get("convening_typology", "")),
                "convening_stage": self.clean_dashboard_value(item.get("convening_stage", "")),
                "purpose_maturity": self.clean_dashboard_value(item.get("purpose_maturity", "")),
                "next_engagement_move": self.clean_dashboard_value(item.get("next_engagement_move", "")),
                "expected_outputs": self.clean_dashboard_value(item.get("expected_outputs", "")),
                "do_not_ask_for": self.clean_dashboard_value(item.get("do_not_ask_for", "")),
                "stage_gate_status": self.clean_dashboard_value(item.get("stage_gate_status", "")),
                "stage_gate_reason": self.clean_dashboard_value(item.get("stage_gate_reason", "")),
                "format": self.clean_dashboard_value(item.get("format", "")),
                "why": self.clean_dashboard_value(item.get("why", "")),
                "shared_tension": self.clean_dashboard_value(item.get("shared_tension", "")),
                "next_best_action": self.clean_dashboard_value(item.get("next_best_action", "")),
                "recent_signals": self.compact_signal_items(item.get("recent_signals", "")),
            })

        return {
            "count": radar.get("count", 0),
            "created_at": radar.get("created_at", ""),
            "pattern_file": radar.get("pattern_file", ""),
            "cards": cards,
        }

    def scip_cards_for_briefing(self, limit: int = 5) -> list[dict[str, Any]]:
        radar = self.scip_convening_radar_for_prompt()
        cards = radar.get("cards", [])
        if not isinstance(cards, list):
            return []
        return cards[:max(0, int(limit))]

    def scip_card_members_text(self, card: dict[str, Any]) -> str:
        members = self.clean_dashboard_value(card.get("members", ""))
        members = members.replace("|", " x ").replace("  ", " ")
        return members or "eine kuratierte Akteurskonstellation"

    def scip_card_theme_text(self, card: dict[str, Any]) -> str:
        return (
            self.clean_dashboard_value(card.get("theme", ""))
            or self.clean_dashboard_value(card.get("shared_tension", ""))
            or "eine konkrete Umsetzungsfrage"
        )

    def scip_card_next_step_text(self, card: dict[str, Any]) -> str:
        raw = (
            self.clean_dashboard_value(card.get("next_best_action", ""))
            or self.clean_dashboard_value(card.get("next_engagement_move", ""))
            or self.clean_dashboard_value(card.get("format", ""))
        )
        key = re.sub(r"\s+", "_", raw.strip().lower().replace("-", "_"))
        if key in {"brief_or_validate", "watch_or_reframe"}:
            return "zunächst eine kurze Validierung mit einem konkreten Ansprechpartner durchführen"
        if "sounding" in key or "sond" in key or "bilateral" in key:
            return "ein begrenztes Sondierungsgespräch mit klarer Prüfungsfrage vorbereiten"
        if "workshop" in key:
            return "einen kleinen Validierungsworkshop mit Entscheidungsfrage vorbereiten"
        return raw or "eine begrenzte Validierung der Rollen, Bedarfe und nächsten Schritte durchführen"

    def scip_anchor_sentence(self, card: dict[str, Any]) -> str:
        members = self.scip_card_members_text(card)
        theme = self.scip_card_theme_text(card)
        next_step = self.scip_card_next_step_text(card)
        return (
            f"Die SCIP-Kuratierung setzt hier bei {members} an: "
            f"Die Konstellation bündelt Rollen für {theme} und sollte {next_step}."
        )

    def strategic_opportunity_briefing_prompt(self, payload: dict[str, Any]) -> str:
        return f"""
Du schreibst einen klaren, nüchternen und strategischen Projekttext für eine ZIRP-Projektleitung.

Aus SCIP-Matchings, News-Signalen und Kontextinformationen soll ein Briefing entstehen, das erklärt:
warum das Thema relevant ist, wo der konkrete Engpass liegt, welche vorhandenen Potenziale genutzt werden können,
welche realistische Rolle ZIRP übernehmen kann und welcher nächste pragmatische Umsetzungsschritt sinnvoll wäre.

Der Text muss aus zwei gleich starken Hälften gebaut sein:
1. Kontext und Crawler-Funde erklären das reale Wochenbild, die Themenlage und die vorhandene Ausgangsbasis.
2. SCIP-kuratierte Karten erklären die Akteurskonstellationen, Rollen, Engpässe und nächsten Umsetzungsschritte.

Beide Ebenen tragen jeweils ungefähr 50 Prozent der Geschichte. Schreibe also weder einen generischen Kontextaufsatz
noch eine mechanische Liste von SCIP-Karten. Nutze die kuratierten Karten als konkretes Rückgrat für Akteure,
Rollenklärung und nächsten Schritt; nutze die Crawler-Funde als Evidenz und Kontext, warum diese Linien jetzt relevant sind.

TON UND STIL
Schreibe klar, direkt und nüchtern. Keine Floskeln. Keine übertriebene Beratersprache.
Keine abstrakten Begriffe ohne Erklärung. Jeder Absatz trägt einen konkreten Gedanken.
Der Text soll strategisch, aber natürlich klingen.

AUSGABEFORMAT
Gib ausschließlich das finale Briefing aus.

Format:
# [Titel als deutsche strategische Frage]

[Absatz 1: klare Leitthese. Beginne mit einer direkten Einordnung und formuliere, unter welcher Bedingung das Thema erfolgreich bearbeitet werden kann.]

[Absatz 2: konkreter Engpass. Benenne nüchtern, woran Umsetzung derzeit scheitert oder wo der zentrale Handlungsbedarf liegt.]

[Absatz 3: systemische Einordnung und vorhandene Ausgangsbasis. Erkläre kurz, warum einfache Einzelmaßnahmen nicht reichen, und zeige vorhandene Kompetenzen, Akteure, Initiativen, Bedarfe oder Ressourcen.]

[Absatz 4: SCIP-kuratierte Akteurslinien. Greife mindestens zwei konkrete SCIP-Karten auf und erkläre, welche Rollen die Akteure in einer prüfbaren Umsetzungskonstellation übernehmen könnten.]

[Absatz 5: Rolle von ZIRP und nächster Schritt. Formuliere eine realistische Rolle und schlage einen begrenzten Umsetzungsschritt vor, der aus Kontext plus SCIP-Karten folgt: idealerweise ein 90-Tage-Testfeld, eine Pilotphase oder ein Validierungsprozess.]

[Letzter Satz: klare, merkfähige Verdichtung.]

Keine Bulletpoints. Keine Zwischenüberschriften nach dem Titel. Keine Methodenerklärung. Kein Quellen- oder Evidenzdump.
Zielumfang: etwa 500 bis 700 Wörter, wenn die Evidenz dafür reicht. Bei schwacher Evidenz lieber kürzer und präzise bleiben.

TITELREGEL
Der Titel muss eine Frage sein und die eigentliche Engstelle oder Projektchance der Woche beschreiben.
Digitalisierung darf nur dann die Klammer bilden, wenn mehrere konkrete Signale digitale Anwendung, Prozessumsetzung, Dateninfrastruktur oder digitale Verwaltungspraxis belegen.
Wenn die Signale eher Versorgung, Fachkräfte, Bau, Kultur, Finanzierung, Nachhaltigkeit, Führung, Kooperation oder Umsetzung betreffen, muss genau dieses Muster im Titel stehen.

Gute Titelrichtungen:
- Können regionale Kooperationen die Umsetzungslücke schließen?
- Wird Wissenstransfer zum Beschleuniger konkreter Pilotvorhaben?
- Welche Akteurskonstellation kann aus regionalen Kompetenzen Umsetzung machen?
- Wird Fachkräftequalifizierung zur gemeinsamen Umsetzungsfrage?
- Welche Rolle kann ZIRP bei der Validierung eines Pilotfelds übernehmen?

FORMULIERUNGSREGELN
Schreibe direkte Sätze.
Vermeide Satzmuster wie: "Es fehlt nicht an ..., sondern ..."
Besser: "Die Region verfügt bereits über ... Der Engpass liegt in ..."
Vermeide unklare Wörter wie "Ansätze", wenn nicht erklärt wird, was gemeint ist.
Nutze konkrete Begriffe: Initiativen, Kooperationen, Kompetenzen, Bedarfe, Pilotvorhaben, Rollen, Nutzen, Ressourcen, Umsetzung.
Jeder Absatz hat maximal 3 bis 4 Sätze.
Der Text bleibt kompakt, aber nicht stichpunktartig.

INHALTLICHE LOGIK
Die Leitthese beantwortet direkt, warum das Thema relevant ist.
Der Engpass benennt eine konkrete Umsetzungs-, Koordinations-, Kapazitäts-, Rollen- oder Validierungsfrage.
Die systemische Einordnung erklärt, warum ein einzelner Akteur das Thema nicht allein lösen kann.
Die vorhandene Ausgangsbasis nennt konkrete Akteure oder Ressourcen aus den übergebenen Signalen und SCIP-Karten.
Der strategische Hebel beschreibt, welche Bündelung, Priorisierung, Rollenklärung oder Pilotierung Wirkung erzeugen kann, und muss an mindestens einer SCIP-Karte sichtbar werden.
Die Rolle von ZIRP ist realistisch: ZIRP löst das Problem nicht allein, sondern bringt Akteure, Bedarfe und Umsetzung in eine entscheidungsfähige Konstellation.
Der nächste Schritt ist begrenzt: ein kuratiertes Sondierungsgespräch, eine 90-Tage-Pilotphase oder ein Validierungsprozess mit klarer Leitfrage.

VERBOTENE FORMULIERUNGEN
Vermeide unbedingt:
- Die Analyse zeigt
- Es wurden X Signale identifiziert
- Besonders auffällig ist
- Actor Anchors
- Warum für ZIRP
- Opportunity Line
- Nächster Schritt:
- Das Signal deutet darauf hin
- Diese Meldung zeigt
- strategisch sichtbar wird
- belastbare Signale herausgefiltert
- ZIRP kann als Plattform für Austausch und Kooperation dienen
- Beide Akteure bringen ihre Expertise ein
- Die digitale Transformation ist ein zentrales Thema
- allgemeiner Austausch
- Networking
- Vernetzung ist wichtig

QUALITÄTSSTANDARD
Der Text ist gut, wenn eine Projektleitung nach dem Lesen drei Dinge klar versteht:
Warum ist das Thema relevant?
Wo liegt der konkrete Umsetzungsengpass?
Was kann ZIRP realistisch als nächstes tun?

Prüfe vor der Ausgabe:
Ist der Titel eine Frage?
Ist die Leitthese konkret?
Wird ein Engpass benannt?
Werden vorhandene Potenziale aus den Daten genutzt?
Tragen Kontext/Crawler-Funde und SCIP-kuratierte Karten jeweils sichtbar zur Geschichte bei?
Werden mindestens zwei SCIP-Karten als Akteurs- oder Umsetzungslinien genutzt?
Ist die ZIRP-Rolle realistisch und umsetzungsnah?
Endet der Text mit einem konkreten begrenzten Schritt?

INPUT:
{json.dumps(payload, ensure_ascii=False, indent=2)}
"""

    def build_weekly_meta_prompt(self, events: list[dict[str, Any]]) -> str:
        grouped = self.classified_signal_groups(events)
        prompt_payload = {
            "wochenfunde": grouped,
            "scip_convening_radar": self.scip_convening_radar_for_prompt(),
        }

        return (
            "Du bist Redakteur:in fÃ¼r eine strategische Wochenanalyse der ZIRP. "
            "Gib valides JSON zurueck, aber das Feld narrative_analysis muss wie ein fertiger deutscher Redaktionstext klingen. "
            "Crawler-Funde und SCIP-Convening-Radar tragen die Analyse gleichgewichtig. "
            "Die Crawler-Funde erklaeren Kontext, Evidenz und Wochenbild; die SCIP-Karten liefern die kuratierten Akteurslinien, Rollen und naechsten Schritte.\n\n"
            "In narrative_analysis steht keine Beratungsfolie, kein Action Plan und keine Liste von SCIP Cards. "
            "Schreibe keine englischen Ueberschriften wie High-Priority, Why This Matters, Next Best Action, Outcome Goal oder Key Takeaways. "
            "Schreibe keine nummerierte Prioritaetenliste und keine Markdown-Trennlinien. "
            "Der fertige Text darf die Labels DO:, WATCH:, MAYBE:, Wozu: oder Zusammenfassung nicht enthalten.\n\n"
            "Aufbau von narrative_analysis:\n"
            "Titel als Frage. Danach Fliesstext mit klaren Absaetzen: Leitthese, konkreter Engpass, systemische Einordnung, vorhandene Ausgangsbasis, strategischer Hebel, Rolle von ZIRP und naechster begrenzter Schritt.\n"
            "Mindestens zwei SCIP-Karten muessen sichtbar als Akteurs- oder Umsetzungslinien verarbeitet werden. "
            "Wenn eine Karte nur schwache Evidenz hat, formuliere sie als Validierungslinie, nicht als gesicherte Projektbehauptung.\n\n"
            "Stil: strategisch, nÃ¼chtern, redaktionell, menschlich. Nicht werblich, nicht Ã¼berzogen, nicht mechanisch. "
            "Verwende konkrete Begriffe wie Initiativen, Kompetenzen, Bedarfe, Pilotvorhaben, Rollen, Nutzen, Ressourcen und Umsetzung. "
            "Schreibe normale deutsche Umlaute. Arbeite ausschliesslich mit dem Input; erfinde keine Fakten, Akteure, Termine oder Kausalitaeten. "
            "Verwende niemals das Wort Signalien; korrekt ist Signale. "
            "Zentrale Projektnamen und Akronyme muessen bei erster Nennung kurz erklÃ¤rt werden, sofern explanation_context oder summary genug Kontext liefern. "
            "Wenn der Kontext fehlt, nicht so tun, als sei das Akronym erklÃ¤rt. "
            "Unklare Begriffe wie Bio-Treppe nur verwenden, wenn sie aus dem Input eindeutig erklÃ¤rt sind.\n\n"
            "Wichtige Leitfragen: Warum ist das Thema relevant? Wo liegt der konkrete Umsetzungsengpass? Was kann ZIRP realistisch als naechstes tun?\n\n"
            f"INPUT:\n{json.dumps(prompt_payload, ensure_ascii=False, indent=2)}"
        )

    def build_plain_narrative_prompt(self, events: list[dict[str, Any]]) -> str:
        grouped = self.compact_signal_groups_for_prompt(self.classified_signal_groups(events))
        prompt_payload = {
            "wochenfunde": grouped,
            "scip_convening_radar": self.scip_convening_radar_for_prompt(),
        }

        return (
            "Du bist Redakteur:in fÃ¼r eine strategische Wochenanalyse der ZIRP. "
            "Du schreibst einen fertigen deutschen Redaktionstext, keinen Plan, keine Auswertung der Aufgabe und keine Methodennotiz. "
            "Crawler-Funde und SCIP-Convening-Radar tragen den Text jeweils zur Haelfte. "
            "Die Crawler-Funde liefern Kontext und Evidenz; die SCIP-Karten liefern Akteurslogik, Rollen und naechste Umsetzungsschritte.\n\n"
            "Arbeitsweise vor dem Schreiben:\n"
            "Bestimme zuerst die tragende Entwicklung aus den Crawler-Funden: Welche Themen, Mitglieder und regionalen Fragen treten insgesamt hervor? "
            "Bestimme danach die zwei bis drei wichtigsten SCIP-Karten: Welche Akteurskonstellationen koennen daraus eine pruefbare Umsetzungslinie machen? "
            "Verwebe beide Ebenen gleichgewichtig. Bei schwacher Evidenz bleibt die SCIP-Karte eine Validierungslinie, aber sie verschwindet nicht aus der Analyse.\n\n"
            "So muss der fertige Text aussehen:\n"
            "Titel als Frage. Danach fuenf kurze, gut lesbare Absaetze: "
            "1. klare Leitthese; "
            "2. konkreter Engpass; "
            "3. systemische Einordnung und vorhandene Ausgangsbasis; "
            "4. strategischer Hebel; "
            "5. realistische Rolle von ZIRP und ein begrenzter naechster Schritt. "
            "Die Bloecke sollen wie ein internes Leitungsmemo klingen: menschlich, konkret, mit sauberer Urteilskraft. "
            "Schreibe nicht Ã¼ber Scores, Prioritaetsstufen, High-Priority/Mid-Priority, Outcome Goals oder Next Best Actions. "
            "Keine englischen ZwischenÃ¼berschriften. Keine Markdown-Trennlinien. Keine nummerierte Liste. Keine Bullet-Liste. Kein 'Below is'. Kein 'Key Takeaways'.\n\n"
            "Der letzte Absatz soll einen plausiblen naechsten Schritt enthalten, aber nicht mit einem sichtbaren Label beginnen. "
            "Keine Labels wie DO, WATCH oder MAYBE.\n\n"
            "Stil:\n"
            "Strategisch, nÃ¼chtern, redaktionell, aber menschlich. Der Text soll gelesen klingen, nicht generiert. "
            "Keine Beratungsfolie, keine To-do-Liste, keine mechanische Zusammenfassung der Convening-Karten. "
            "Arbeite ausschliesslich mit dem Ã¼bergebenen Input. Erfinde keine Fakten, keine Akteure, keine Termine und keine Kausalitaeten. "
            "Schreibe nicht aus der Wir-Perspektive, sondern aus Sicht von ZIRP. "
            "Verwende niemals das Wort Signalien; korrekt ist Signale. "
            "Wenn ein Projektname oder Akronym zentral ist, erklaere es bei der ersten Nennung mit einem kurzen Halbsatz aus explanation_context oder summary. "
            "Bewerte aus Sicht der ZIRP: relevant ist nicht PR, sondern ob daraus ein Zukunftsthema, ein Standortsignal, ein Transferanlass oder eine Vernetzungsfrage fÃ¼r Rheinland-Pfalz entsteht. "
            "Laenge: 300 bis 650 WÃ¶rter. Der Text muss vollstaendig enden, keine abgeschnittenen Saetze.\n\n"
            f"Input:\n{json.dumps(prompt_payload, ensure_ascii=False, indent=2)}"
        )

    def build_draft_rewrite_prompt(self, *, draft: str, events: list[dict[str, Any]]) -> str:
        selected = self.select_events_for_narrative(events)
        main_events = [e for e in selected if e.get("evidence_role") == "main_signal"][:4]
        supporting_events = [e for e in selected if e.get("evidence_role") == "supporting_context"][:4]
        watch_events = [e for e in selected if e.get("evidence_role") == "watchlist"][:2]
        used_ids = {id(e) for e in main_events + supporting_events + watch_events}
        if len(main_events) < 3:
            for event in selected:
                if id(event) in used_ids:
                    continue
                main_events.append(event)
                used_ids.add(id(event))
                if len(main_events) >= 3:
                    break

        def fact_item(event: dict[str, Any]) -> dict[str, Any]:
            role = event.get("evidence_role", self.evidence_role_for_event(event))
            return {
                "member": self.clean_dashboard_value(event.get("mitglied", "")),
                "title": self.clean_dashboard_value(event.get("titel", "")),
                "date": self.clean_dashboard_value(event.get("datum", "")),
                "evidence_role": role,
                "use_rule": (
                    "Rueckgrat der Analyse" if role == "main_signal"
                    else "nur zur Verdichtung" if role == "supporting_context"
                    else "nur als Beobachtung, nicht als Hauptthese"
                ),
                "section": REPORT_SECTIONS.get(event.get("hauptsektion", ""), {}).get("title", event.get("hauptsektion", "")),
                "snippet": self.clean_narrative_snippet(event.get("snippet", ""), max_chars=170),
                "why_relevant": self.clean_dashboard_value(event.get("llm_judge_reason", "")) or self.signal_relevance_reason(
                    event,
                    "hoch" if int(event.get("regional_score", 0) or 0) >= 2 else "mittel",
                    "hoch" if int(event.get("decision_score", 0) or 0) >= 18 else "mittel",
                ),
                "explanation_context": self.acronym_explanation_context(
                    str(event.get("titel", "")),
                    [str(event.get("snippet", ""))]
                ),
            }

        scip = self.scip_convening_radar_for_prompt()
        scip["cards"] = scip.get("cards", [])[:self.dynamic_scip_prompt_card_limit(len(events))]
        payload = {
            "regelbasierter_entwurf": self.clean_narrative_snippet(draft, max_chars=1600),
            "gliederung_der_fakten": {
                "main_signal": [fact_item(event) for event in main_events],
                "supporting_context": [fact_item(event) for event in supporting_events],
                "watchlist": [fact_item(event) for event in watch_events],
            },
            "gewichtung": "Kontext/Crawler-Funde und SCIP-kuratierte Karten tragen jeweils etwa 50 Prozent der Geschichte.",
            "scip_kuratierte_karten": scip,
        }
        return self.strategic_opportunity_briefing_prompt(payload)

    def build_fast_direct_report_prompt(self, events: list[dict[str, Any]]) -> str:
        facts = []
        for event in self.select_events_for_narrative(events)[:5]:
            facts.append({
                "member": event.get("mitglied", ""),
                "title": event.get("titel", ""),
                "role": event.get("evidence_role", self.evidence_role_for_event(event)),
                "snippet": self.clean_narrative_snippet(event.get("snippet", ""), max_chars=120),
            })
        return self.strategic_opportunity_briefing_prompt({
            "gewichtung": "Kontext/Crawler-Funde und SCIP-kuratierte Karten tragen jeweils etwa 50 Prozent der Geschichte.",
            "fakten": facts,
            "scip_kuratierte_karten": self.scip_convening_radar_for_prompt(),
        })

    def build_polish_report_prompt(self, *, draft: str, events: list[dict[str, Any]]) -> str:
        facts = []
        for event in self.select_events_for_narrative(events)[:4]:
            facts.append({
                "member": event.get("mitglied", ""),
                "title": event.get("titel", ""),
                "role": event.get("evidence_role", self.evidence_role_for_event(event)),
                "snippet": self.clean_narrative_snippet(event.get("snippet", ""), max_chars=90),
            })
        payload = {
            "draft_to_polish": self.clean_narrative_snippet(draft, max_chars=2600),
            "fact_check_anchor": facts,
            "gewichtung": "Beim Polieren die SCIP-kuratierten Karten nicht herauswaschen: Kontext und SCIP tragen jeweils etwa 50 Prozent der Geschichte.",
            "scip_kuratierte_karten": self.scip_convening_radar_for_prompt(),
        }
        return self.strategic_opportunity_briefing_prompt(payload)

    def compact_signal_groups_for_prompt(self, groups: dict[str, list[dict[str, Any]]]) -> dict[str, list[dict[str, Any]]]:
        limits = {
            "main_briefing_signals": 5,
            "convening_input_signals": 4,
            "watchlist_signals": 2,
            "appendix_signals": 1,
            "dropped_signals": 0,
        }
        compact: dict[str, list[dict[str, Any]]] = {}
        for key, items in groups.items():
            compact[key] = []
            for item in items[:limits.get(key, 2)]:
                compact[key].append({
                    "title": item.get("title", ""),
                    "member": item.get("member", ""),
                    "date": item.get("date", ""),
                    "evidence_role": item.get("evidence_role", ""),
                    "topic": item.get("topic", ""),
                    "summary": self.clean_narrative_snippet(item.get("summary", ""), max_chars=180),
                    "explanation_context": self.clean_narrative_snippet(item.get("explanation_context", ""), max_chars=260),
                    "rlp_relevance": item.get("rlp_relevance", ""),
                    "zirp_relevance": item.get("zirp_relevance", ""),
                    "recommended_use": item.get("recommended_use", ""),
                })
        return compact

    def classified_signal_groups(self, events: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
        groups = {
            "main_briefing_signals": [],
            "convening_input_signals": [],
            "watchlist_signals": [],
            "appendix_signals": [],
            "dropped_signals": [],
        }
        for event in self.select_events_for_narrative(events):
            item = self.classify_signal_for_prompt(event)
            use = item.get("recommended_use", "watchlist")
            key = f"{use}_signals" if use in {"main_briefing", "convening_input", "watchlist", "appendix"} else "dropped_signals"
            groups[key].append(item)
        return groups

    def classify_signal_for_prompt(self, event: dict[str, Any]) -> dict[str, Any]:
        rlp_score = int(event.get("regional_score", 0) or 0)
        decision_score = int(event.get("decision_score", 0) or 0)
        direct_score = int(event.get("direct_member_score", 0) or 0)
        section = event.get("hauptsektion", "")
        themes = event.get("themenfelder", []) or []
        is_low = self.is_low_value_narrative_event(event)

        rlp_relevance = "hoch" if rlp_score >= 2 else ("mittel" if rlp_score >= 1 or direct_score >= 4 else "niedrig")
        zirp_relevance = "hoch" if decision_score >= 18 and not is_low else ("mittel" if decision_score >= 10 and not is_low else "niedrig")

        if is_low or (rlp_relevance == "niedrig" and zirp_relevance == "niedrig"):
            recommended_use = "drop"
        elif section == "kooperationen" or any(t in {"technologie", "wissen", "gesundheit", "wirtschaft", "nachhaltigkeit"} for t in themes):
            recommended_use = "convening_input" if zirp_relevance in {"hoch", "mittel"} else "watchlist"
        elif zirp_relevance == "hoch":
            recommended_use = "main_briefing"
        elif zirp_relevance == "mittel":
            recommended_use = "watchlist"
        else:
            recommended_use = "appendix"

        if recommended_use == "convening_input" and decision_score >= 18 and rlp_relevance in {"hoch", "mittel"}:
            recommended_use = "main_briefing"

        return {
            "title": event.get("titel", ""),
            "member": event.get("mitglied", ""),
            "date": event.get("datum", ""),
            "evidence_role": event.get("evidence_role", self.evidence_role_for_event(event)),
            "topic": REPORT_SECTIONS.get(section, {}).get("title", section),
            "summary": self.clean_narrative_snippet(event.get("snippet", ""), max_chars=240),
            "explanation_context": self.acronym_explanation_context(
                str(event.get("titel", "")),
                [str(event.get("snippet", ""))]
            ),
            "rlp_relevance": rlp_relevance,
            "zirp_relevance": zirp_relevance,
            "relevance_reason": self.signal_relevance_reason(event, rlp_relevance, zirp_relevance),
            "recommended_use": recommended_use,
        }

    def signal_relevance_reason(self, event: dict[str, Any], rlp_relevance: str, zirp_relevance: str) -> str:
        section = REPORT_SECTIONS.get(event.get("hauptsektion", ""), {}).get("title", event.get("hauptsektion", ""))
        return (
            f"RLP-Relevanz {rlp_relevance} durch Regional-/Mitgliedssignal; "
            f"ZIRP-Relevanz {zirp_relevance} wegen Bezug zu {section}, "
            f"Decision Score {event.get('decision_score', 0)}."
        )


    def generic_weekly_filler_score(self, text: str) -> int:
        lowered = str(text or "").lower()
        filler_phrases = [
            "signale und informationen",
            "qualitÃ¤t und relevanz der informationen",
            "qualitaet und relevanz der informationen",
            "eine vielzahl von themen und projekten",
            "konkrete projekte und initiativen",
            "sorgfÃ¤ltig auszuwerten und zu bewerten",
            "sorgfaeltig auszuwerten und zu bewerten",
            "die richtigen schlÃ¼sse",
            "die richtigen schluesse",
            "die richtigen entscheidungen",
            "ziele und interessen der zirp",
            "chancen maximieren",
            "risiken minimieren",
            "immer auf dem neuesten stand",
            "eigene projekte und initiativen",
            "von interesse sind",
            "nutzen, um konkrete",
            "fÃ¼r die zirp relevant sind",
            "fÃ¼r die zirp relevant sind",
            "sorgfÃ¤ltige auswertung",
            "sorgfaeltige auswertung",
        ]
        return sum(1 for phrase in filler_phrases if phrase in lowered)

    def is_bad_weekly_narrative(self, text: str) -> bool:
        value = str(text or "").strip()
        if not value:
            return True
        words = re.findall(r"\b\w+\b", value, flags=re.UNICODE)
        if len(words) < 220:
            return True
        if value[-1] not in ".!?":
            return True
        tail = value[-80:].strip()
        if re.search(r"\b(und|oder|mit|fÃ¼r|fÃ¼r|zur|zum|Sch|sch|die|der|das|eine|einen)$", tail):
            return True
        lowered = value.lower()
        if self.generic_weekly_filler_score(value) >= 2:
            return True
        blocked_phrases = [
            "prioritized action plan",
            "below is",
            "high-priority",
            "mid-priority",
            "lower-priority",
            "key takeaways",
            "next best action",
            "outcome goal",
            "why this matters",
            "members involved",
            "ordered by urgency",
            "do:",
            "watch:",
            "maybe:",
            "wozu:",
            "zusammenfassung",
            "signalien",
            "qualitÃ¤t und relevanz der informationen",
            "qualitaet und relevanz der informationen",
            "eine vielzahl von themen und projekten",
            "sorgfÃ¤ltig auszuwerten und zu bewerten",
            "sorgfaeltig auszuwerten und zu bewerten",
            "die richtigen schlÃ¼sse",
            "die richtigen schluesse",
            "die richtigen entscheidungen",
            "ziele und interessen der zirp",
            "eigene projekte und initiativen zu entwickeln",
            "konkrete handlungs- oder transferchancen zu identifizieren und umzusetzen",
            "basierend darauf kÃ¶nnten wir",
            "basierend darauf kÃ¶nnten wir",
            "folgende schritte unternehmen",
            "die analyse beendet sich hiermit",
            "chancen maximieren",
            "risiken minimieren",
            "richtige schritte unternehmen",
            "immer auf dem neuesten stand",
            "viele mÃ¶glichkeiten gibt",
            "viele mÃ¶glichkeiten gibt",
            "wirtschaft und die gesellschaft stÃ¤rken",
            "wirtschaft und die gesellschaft staerken",
            "umfassendes bild der aktuellen entwicklungen",
            "kombination von scip und crawler-funden",
            "do:",
            "watch:",
            "maybe:",
            "wozu:",
            "zusammenfassung",
            "signalien",
            "qualitÃ¤t und relevanz der informationen",
            "qualitaet und relevanz der informationen",
            "eine vielzahl von themen und projekten",
            "sorgfÃ¤ltig auszuwerten und zu bewerten",
            "sorgfaeltig auszuwerten und zu bewerten",
            "die richtigen schlÃ¼sse",
            "die richtigen schluesse",
            "die richtigen entscheidungen",
            "ziele und interessen der zirp",
            "eigene projekte und initiativen zu entwickeln",
            "konkrete handlungs- oder transferchancen zu identifizieren und umzusetzen",
            "basierend darauf kÃ¶nnten wir",
            "basierend darauf kÃ¶nnten wir",
            "folgende schritte unternehmen",
            "die analyse beendet sich hiermit",
            "chancen maximieren",
            "risiken minimieren",
            "richtige schritte unternehmen",
            "immer auf dem neuesten stand",
            "viele mÃ¶glichkeiten gibt",
            "viele mÃ¶glichkeiten gibt",
            "wirtschaft und die gesellschaft stÃ¤rken",
            "wirtschaft und die gesellschaft staerken",
            "umfassendes bild der aktuellen entwicklungen",
            "kombination von scip und crawler-funden",
            "zu den stÃ¤rksten akteuren zÃ¤hlen",
            "zu den stÃ¤rksten akteuren zÃ¤hlen",
        ]
        if any(phrase in lowered for phrase in blocked_phrases):
            return True
        if len(re.findall(r"(?m)^\s*\d+[\.)]\s+", value)) >= 2:
            return True
        english_markers = sum(1 for word in ["the", "and", "with", "based", "opportunities", "priority", "recommendations"] if f" {word} " in f" {lowered} ")
        german_markers = sum(1 for word in ["die", "und", "fÃ¼r", "fÃ¼r", "zirp", "woche", "anschlussfrage", "sondierung"] if f" {word} " in f" {lowered} ")
        return english_markers >= 3 and english_markers > german_markers

    def normalize_editorial_report_text(self, text: str) -> str:
        value = self.repair_mojibake(str(text or "")).strip()
        value = re.sub(r"^```(?:\w+)?\s*", "", value)
        value = re.sub(r"\s*```$", "", value).strip()
        value = re.sub(r"\n{3,}", "\n\n", value)
        value = re.sub(r"[ \t]+", " ", value)
        if not value:
            return ""

        dangling_tail = re.search(
            r"\b(und|oder|mit|für|fuer|zur|zum|die|der|das|eine|einen|ein|im|in|auf|durch)$",
            value[-100:].strip(),
            flags=re.IGNORECASE,
        )
        if value[-1] not in ".!?" or dangling_tail:
            last_stop = max(value.rfind("."), value.rfind("!"), value.rfind("?"))
            if last_stop >= max(80, int(len(value) * 0.62)):
                value = value[: last_stop + 1].strip()
            elif value[-1] not in ".!?":
                value = value.rstrip(" ,;:-") + "."
        return value

    def report_completion_issue(self, text: str) -> str:
        value = self.normalize_editorial_report_text(text)
        if not value:
            return "leer"
        if value[-1] not in ".!?":
            return "endet unvollstaendig"
        tail = value[-100:].strip()
        if re.search(
            r"\b(und|oder|mit|für|fuer|zur|zum|die|der|das|eine|einen|ein|im|in|auf|durch)$",
            tail,
            flags=re.IGNORECASE,
        ):
            return "unvollstaendiger Satz am Ende"
        return ""

    def groq_report_rejection_reason(self, text: str) -> str:
        value = self.normalize_editorial_report_text(text)
        if not value:
            return "leer"
        words = re.findall(r"\b\w+\b", value, flags=re.UNICODE)
        if len(words) < 280:
            return f"zu kurz ({len(words)} WÃ¶rter)"
        if len(words) > 850:
            return f"zu lang ({len(words)} WÃ¶rter)"
        completion_issue = self.report_completion_issue(value)
        if completion_issue:
            return completion_issue

        lowered = value.lower()
        hard_blocked = [
            "do:", "watch:", "maybe:", "wozu:", "signalien",
            "opportunity line", "actor anchors", "warum für zirp", "warum fuer zirp",
            "nächster schritt:", "naechster schritt:",
            "die zentrale herausforderung liegt weniger in",
            "prioritized action plan", "below is", "key takeaways",
            "high-priority", "mid-priority", "lower-priority",
            "die analyse beendet sich hiermit",
        ]
        for phrase in hard_blocked:
            if phrase in lowered:
                return f"verbotener Marker: {phrase}"

        # Groq is the preferred editorial writer. Do not reject it merely because one
        # safe-but-generic phrase appears. Only reject if the whole text is clearly filler.
        filler_score = self.generic_weekly_filler_score(value)
        if filler_score >= 4:
            return f"zu generisch (Filler-Score {filler_score})"

        english_markers = sum(
            1 for word in ["the", "and", "with", "based", "opportunities", "priority", "recommendations"]
            if f" {word} " in f" {lowered} "
        )
        german_markers = sum(
            1 for word in ["die", "und", "für", "fuer", "zirp", "woche", "anschlussfrage", "sondierung"]
            if f" {word} " in f" {lowered} "
        )
        if english_markers >= 3 and english_markers > german_markers:
            return "zu viel Englisch"
        return ""

    def is_usable_groq_report(self, text: str) -> bool:
        return self.groq_report_rejection_reason(text) == ""

    def is_usable_report_250_550(self, text: str) -> bool:
        value = self.normalize_editorial_report_text(text)
        if not value:
            return False
        words = re.findall(r"\b\w+\b", value, flags=re.UNICODE)
        if len(words) < 280 or len(words) > 780:
            return False
        if self.report_completion_issue(value):
            return False
        lowered = value.lower()
        if self.generic_weekly_filler_score(value) >= 4:
            return False
        blocked_phrases = [
            "prioritized action plan",
            "below is",
            "high-priority",
            "mid-priority",
            "lower-priority",
            "key takeaways",
            "next best action",
            "outcome goal",
            "why this matters",
            "members involved",
            "ordered by urgency",
            "do:",
            "watch:",
            "maybe:",
            "wozu:",
            "opportunity line",
            "actor anchors",
            "warum für zirp",
            "warum fuer zirp",
            "nächster schritt:",
            "naechster schritt:",
            "die zentrale herausforderung liegt weniger in",
            "zusammenfassung",
            "signalien",
            "qualitÃ¤t und relevanz der informationen",
            "qualitaet und relevanz der informationen",
            "eine vielzahl von themen und projekten",
            "sorgfÃ¤ltig auszuwerten und zu bewerten",
            "sorgfaeltig auszuwerten und zu bewerten",
            "die richtigen schlÃ¼sse",
            "die richtigen schluesse",
            "die richtigen entscheidungen",
            "ziele und interessen der zirp",
            "eigene projekte und initiativen zu entwickeln",
            "konkrete handlungs- oder transferchancen zu identifizieren und umzusetzen",
            "basierend darauf kÃ¶nnten wir",
            "basierend darauf kÃ¶nnten wir",
            "folgende schritte unternehmen",
            "die analyse beendet sich hiermit",
            "chancen maximieren",
            "risiken minimieren",
            "richtige schritte unternehmen",
            "immer auf dem neuesten stand",
            "viele mÃ¶glichkeiten gibt",
            "viele mÃ¶glichkeiten gibt",
            "wirtschaft und die gesellschaft stÃ¤rken",
            "wirtschaft und die gesellschaft staerken",
            "umfassendes bild der aktuellen entwicklungen",
            "kombination von scip und crawler-funden",
        ]
        if any(phrase in lowered for phrase in blocked_phrases):
            return False
        if len(re.findall(r"(?m)^\s*\d+[\.)]\s+", value)) >= 2:
            return False
        english_markers = sum(1 for word in ["the", "and", "with", "based", "opportunities", "priority", "recommendations"] if f" {word} " in f" {lowered} ")
        german_markers = sum(1 for word in ["die", "und", "für", "fuer", "zirp", "woche", "anschlussfrage", "sondierung"] if f" {word} " in f" {lowered} ")
        return not (english_markers >= 3 and english_markers > german_markers)

    def is_usable_qwen_fallback_report(self, text: str) -> bool:
        value = self.normalize_editorial_report_text(text)
        if not value:
            return False
        words = re.findall(r"\b\w+\b", value, flags=re.UNICODE)
        if len(words) < 260 or len(words) > 780:
            return False
        if self.report_completion_issue(value):
            return False
        lowered = value.lower()
        if self.generic_weekly_filler_score(value) >= 4:
            return False
        blocked_phrases = [
            "prioritized action plan",
            "below is",
            "high-priority",
            "mid-priority",
            "lower-priority",
            "key takeaways",
            "members involved",
            "ordered by urgency",
            "do:",
            "watch:",
            "maybe:",
            "wozu:",
            "opportunity line",
            "actor anchors",
            "warum für zirp",
            "warum fuer zirp",
            "nächster schritt:",
            "naechster schritt:",
            "die zentrale herausforderung liegt weniger in",
            "zusammenfassung",
            "signalien",
            "qualitÃ¤t und relevanz der informationen",
            "qualitaet und relevanz der informationen",
            "eine vielzahl von themen und projekten",
            "sorgfÃ¤ltig auszuwerten und zu bewerten",
            "sorgfaeltig auszuwerten und zu bewerten",
            "die richtigen schlÃ¼sse",
            "die richtigen schluesse",
            "die richtigen entscheidungen",
            "ziele und interessen der zirp",
            "eigene projekte und initiativen zu entwickeln",
            "konkrete handlungs- oder transferchancen zu identifizieren und umzusetzen",
            "basierend darauf kÃ¶nnten wir",
            "basierend darauf kÃ¶nnten wir",
            "folgende schritte unternehmen",
            "die analyse beendet sich hiermit",
            "chancen maximieren",
            "risiken minimieren",
            "richtige schritte unternehmen",
            "immer auf dem neuesten stand",
            "viele mÃ¶glichkeiten gibt",
            "viele mÃ¶glichkeiten gibt",
            "wirtschaft und die gesellschaft stÃ¤rken",
            "wirtschaft und die gesellschaft staerken",
            "umfassendes bild der aktuellen entwicklungen",
            "kombination von scip und crawler-funden",
        ]
        if any(phrase in lowered for phrase in blocked_phrases):
            return False
        english_markers = sum(1 for word in ["the", "and", "with", "based", "opportunities", "priority", "recommendations"] if f" {word} " in f" {lowered} ")
        german_markers = sum(1 for word in ["die", "und", "für", "fuer", "zirp", "woche", "bericht", "sondierung"] if f" {word} " in f" {lowered} ")
        return not (english_markers >= 3 and english_markers > german_markers)

    def enrich_weekly_meta(self, weekly_events: list[dict[str, Any]]) -> dict[str, Any]:
        period_start = self._report_period_start()
        period_end = self.today
        searched_member_count = self.checked_member_count()
        members_with_signals = len({event["mitglied"] for event in weekly_events})
        signal_count = len(weekly_events)
        fallback = {
            "headline": "Wochenanalyse",
            "meta_line": (
                f"{searched_member_count} Mitglieder geprÃ¼ft · "
                f"{members_with_signals} Mitglieder mit priorisierten Funden · "
                f"{signal_count} belastbare Signale · "
                f"Zeitraum {period_start.strftime('%d.%m.%Y')} bis {period_end.strftime('%d.%m.%Y')} · "
                f"erstellt {datetime.now().strftime('%d.%m.%Y, %H:%M:%S')}"
            ),
            "narrative_analysis": self.build_rule_based_weekly_analysis(weekly_events),
            "weekly_thesis": self.build_weekly_thesis(weekly_events[:3]) if weekly_events else "",
            "one_line_conclusion": "Die Auswertung konzentriert sich auf die sichtbarsten Muster im RÃ¼ckblickszeitraum.",
            "strategic_summary": "Die Analyse wurde aus den priorisierten Funden des RÃ¼ckblickszeitraums gebildet.",
        }

        if not weekly_events:
            return fallback

        # 1) Groq / Report-Client hat Vorrang fÃ¼r die oeffentliche Wochenanalyse.
        if self.report_openai_client is not None:
            narrative = self.request_report_plain_text(
                system_prompt=(
                    "Du schreibst einen klaren, nuechternen und strategischen Projekttext fuer eine ZIRP-Projektleitung. "
                    "Benenne Relevanz, konkreten Umsetzungsengpass, vorhandene Potenziale, strategischen Hebel, ZIRP-Rolle und naechsten begrenzten Schritt. "
                    "Titel als Frage, danach Fliesstext. Keine Bulletpoints, keine Floskeln, keine alten Labels."
                ),
                user_prompt=self.build_draft_rewrite_prompt(
                    draft=fallback["narrative_analysis"],
                    events=weekly_events,
                ),
                call_name="weekly_cloud_rewrite",
                max_tokens=WEEKLY_REWRITE_MAX_TOKENS,
            )
            if narrative and self.is_usable_groq_report(narrative):
                fallback["narrative_analysis"] = self.normalize_editorial_report_text(narrative)
                print("Report editorial: Groq-Bericht akzeptiert")
                return fallback
            if narrative:
                reason = self.groq_report_rejection_reason(narrative)
                print(f"Report editorial: Groq-Bericht verworfen ({reason}); pruefe lokalen Fallback")
            else:
                print("Report editorial: Groq-Bericht fehlgeschlagen; pruefe lokalen Fallback")

        # 2) Qwen bleibt Backup, darf aber keinen schlechten oeffentlichen Text durchreichen.
        if self.report_local_cascade and self.openai_client is not None and self.openai_base_url:
            print(
                "Report editorial: lokale Qwen-Kaskade aktiv "
                f"(draft={self.report_draft_model}, polish={self.report_polish_model})"
            )
            draft = self.request_local_plain_text(
                system_prompt=(
                    "Du schreibst einen kompakten deutschen Projekttext fuer eine ZIRP-Projektleitung. "
                    "Priorisiere eine klare Leitthese, den konkreten Engpass, die vorhandene Ausgangsbasis, den strategischen Hebel, die Rolle von ZIRP und einen pragmatischen naechsten Schritt. "
                    "Titel als Frage, danach klare Absaetze. Keine Newsletter-Sprache, keine Bulletpoints, keine alten Labels."
                ),
                user_prompt=self.build_draft_rewrite_prompt(
                    draft=fallback["narrative_analysis"],
                    events=weekly_events,
                ),
                call_name="weekly_qwen_draft",
                max_tokens=REPORT_DRAFT_MAX_TOKENS,
                model=self.report_draft_model,
            )
            if draft and self.is_usable_report_250_550(draft):
                fallback["narrative_analysis"] = self.normalize_editorial_report_text(draft)
                polish = self.request_local_plain_text(
                    system_prompt=(
                        "Du bist Schlussredakteur fuer einen nuechternen ZIRP-Projekttext. "
                        "Poliere nur Stil, Gewichtung und Lesefluss. Keine neuen Fakten, keine sichtbaren Denkprozesse. "
                        "Bewahre das Format: Titel als Frage, danach Fliesstext mit klaren Absaetzen."
                    ),
                    user_prompt=self.build_polish_report_prompt(
                        draft=draft,
                        events=weekly_events,
                    ),
                    call_name="weekly_qwen14_polish",
                    max_tokens=REPORT_POLISH_MAX_TOKENS,
                    model=self.report_polish_model,
                )
                if polish and self.is_usable_report_250_550(polish):
                    fallback["narrative_analysis"] = self.normalize_editorial_report_text(polish)
                    return fallback
                if polish:
                    print("Report editorial: Qwen14-Polish verworfen; nutze schnellen Qwen-Entwurf")
                else:
                    print("Report editorial: Qwen14-Polish fehlgeschlagen; nutze schnellen Qwen-Entwurf")
                return fallback
            if draft:
                print("Report editorial: schneller Qwen-Entwurf zu kurz oder unpassend; versuche Qwen-Erweiterung")
                expanded = self.request_local_plain_text(
                    system_prompt=(
                        "Du schreibst einen einseitigen strategischen Projekttext fuer ZIRP. "
                        "Erweitere den vorhandenen Entwurf entlang von Relevanz, Engpass, Potenzialen, ZIRP-Rolle und naechstem begrenztem Schritt. "
                        "Keine neuen Fakten, keine sichtbaren Denkprozesse, keine alten Abschnittslabels."
                    ),
                    user_prompt=self.build_polish_report_prompt(
                        draft=draft,
                        events=weekly_events,
                    ),
                    call_name="weekly_qwen_draft_expand",
                    max_tokens=REPORT_DRAFT_MAX_TOKENS,
                    model=self.report_draft_model,
                )
                if expanded and self.is_usable_qwen_fallback_report(expanded):
                    fallback["narrative_analysis"] = self.normalize_editorial_report_text(expanded)
                    return fallback
                if expanded:
                    print("Report editorial: Qwen-Erweiterung verworfen; nutze sauberen regelbasierten Fallback")
                    fallback["narrative_analysis"] = self.build_rule_based_weekly_analysis(weekly_events)
                    return fallback
                print("Report editorial: Qwen-Erweiterung fehlgeschlagen; nutze sauberen regelbasierten Fallback")
                fallback["narrative_analysis"] = self.build_rule_based_weekly_analysis(weekly_events)
                return fallback
            else:
                print("Report editorial: schneller Qwen-Entwurf fehlgeschlagen; nutze sauberen regelbasierten Fallback")
                fallback["narrative_analysis"] = self.build_rule_based_weekly_analysis(weekly_events)
                return fallback

        # 3) Sonstige lokale OpenAI-kompatible Modelle nur als nachrangige Reserve.
        if self.openai_client is not None and self.openai_base_url:
            narrative = self.request_local_plain_text(
                system_prompt=(
                    "Du schreibst einen deutschen strategischen Projekttext fuer ZIRP. "
                    "Bewahre Fakten, korrigiere Gewichtung und mache Relevanz, Engpass, Potenziale, ZIRP-Rolle und naechsten Umsetzungsschritt sichtbar. "
                    "Keine Newsletter-Sprache, keine Bulletpoints, keine alten Labels."
                ),
                user_prompt=self.build_draft_rewrite_prompt(
                    draft=fallback["narrative_analysis"],
                    events=weekly_events,
                ),
                call_name="weekly_draft_rewrite",
                max_tokens=WEEKLY_REWRITE_MAX_TOKENS,
            )
            if narrative and not self.is_bad_weekly_narrative(narrative):
                fallback["narrative_analysis"] = self.normalize_editorial_report_text(narrative)
                return fallback
            if narrative:
                print("OpenAI editorial: Entwurfs-Ueberarbeitung verworfen (zu kurz, abgebrochen oder falsches Format)")
            else:
                print("OpenAI editorial: Entwurfs-Ueberarbeitung fehlgeschlagen, versuche sehr kompakten Direktbericht")

            narrative = self.request_local_plain_text(
                system_prompt=(
                    "Du machst aus Rohsignalen einen praezisen deutschen Projekttext fuer eine ZIRP-Projektleitung. "
                    "Schreibe mit normalem Deutsch, ohne Modell- oder Verwaltungssprache. "
                    "Titel als Frage, danach Fliesstext. Keine Newsletter-Sprache, keine Bulletpoints, keine alten Labels."
                ),
                user_prompt=self.build_fast_direct_report_prompt(weekly_events),
                call_name="weekly_fast_direct",
                max_tokens=WEEKLY_DIRECT_MAX_TOKENS,
            )
            if narrative and not self.is_bad_weekly_narrative(narrative):
                fallback["narrative_analysis"] = self.normalize_editorial_report_text(narrative)
                return fallback
            if narrative:
                print("OpenAI editorial: lokale Wochenanalyse verworfen (zu kurz, abgebrochen oder falsches Format)")
            else:
                print("OpenAI editorial: Modellbericht fehlgeschlagen; nutze regelbasierten Notfalltext")
            return fallback

        result = self.request_structured_json(
            schema_name="zirp_weekly_meta",
            schema=self.weekly_meta_schema(),
            system_prompt=(
                "Du schreibst eine analytische Wochenanalyse fÃ¼r ein ZIRP-Mitgliederbriefing. "
                "Du arbeitest ausschliesslich mit den Ã¼bergebenen Daten, priorisierst echte Querschnittsmuster und vermeidest Kalender- oder Evidenzdump-Stil. "
                "Gib ausschliesslich gueltiges JSON nach Schema zurueck."
            ),
            user_prompt=self.build_weekly_meta_prompt(weekly_events),
        )
        if not result:
            return fallback

        merged = dict(fallback)
        merged.update(result)
        if self.is_bad_weekly_narrative(str(merged.get("narrative_analysis", ""))):
            print("OpenAI editorial: strukturierte Wochenanalyse verworfen (Listen-/Englischformat)")
            merged["narrative_analysis"] = fallback["narrative_analysis"]
        return merged

    def build_rule_based_weekly_analysis(self, weekly_events: list[dict[str, Any]]) -> str:
        period_start = self._report_period_start()
        searched_member_count = self.checked_member_count()
        if not weekly_events:
            return (
                f"FÃ¼r den Zeitraum vom {period_start.strftime('%d.%m.%Y')} bis "
                f"{self.today.strftime('%d.%m.%Y')} wurden {searched_member_count} Mitglieder geprÃ¼ft. "
                "Dabei wurden keine priorisierten Funde identifiziert."
            )
        return self.build_opportunity_line_weekly_analysis(weekly_events)

        ranked = self.rank_events_for_narrative(weekly_events)
        substantial = [
            e for e in ranked
            if self.is_report_worthy_event(e)
            and not self.is_low_value_narrative_event(e)
        ]
        if not substantial:
            return self.build_thin_week_report(ranked[:8])

        selected = substantial[:5]
        member_count = len({e["mitglied"] for e in selected})
        section_titles = self.join_german_list([
            REPORT_SECTIONS.get(section, {}).get("title", section)
            for section in sorted({e.get("hauptsektion", "beobachtung") for e in selected})
        ])
        top = selected[0]
        second = selected[1] if len(selected) > 1 else None
        third = selected[2] if len(selected) > 2 else None

        intro = (
            f"Die Auswertung der letzten {PUBLIC_REPORT_DAYS} Tage hat {searched_member_count} Mitglieder geprÃ¼ft "
            f"und daraus {len(selected)} belastbare Signale bei {member_count} Mitgliedern identifiziert. "
            f"Der Kern liegt nicht in der Menge der Meldungen, sondern in der Frage, "
            f"wo aus {section_titles} konkrete Handlungs- oder Transferchancen fÃ¼r ZIRP entstehen."
        )
        p1 = (
            f"Besonders relevant ist {self.compact_event_reference(top)}. "
            f"Das Signal weist Ã¼ber reine Kommunikation hinaus auf {self.event_implication_phrase(top)}."
        )
        if second:
            p1 += (
                f" Ein weiterer wichtiger Bezugspunkt ist {self.compact_event_reference(second)}. "
                f"Damit wird eine zweite Linie sichtbar, die an {self.event_implication_phrase(second)} anschlieÃŸt."
            )

        p2 = (
            "FÃ¼r den Radar zÃ¤hlen vor allem Hinweise, die institutionelle Entscheidungen, umsetzbare Programme, "
            "Partnerschaften oder regionalen Kompetenzaufbau erkennen lassen. SchwÃ¤chere PR- oder Nachwuchssignale "
            "werden deshalb nicht als eigener Trend gelesen, sondern nur als Umfeldrauschen."
        )
        if third:
            p2 += (
                f" In diese Kategorie fÃ¤llt auch {self.compact_event_reference(third)}, "
                f"weil es auf {self.event_implication_phrase(third)} verweist."
            )

        action_target = self.join_german_list(sorted({e["mitglied"] for e in selected[:3]}))
        actions = (
            "Was ZIRP daraus machen kÃ¶nnte:\n\n"
            f"ZIRP kÃ¶nnte mit {action_target} prÃ¼fen, ob aus den stÃ¤rksten Signalen ein kleines Sondierungsformat entsteht. "
            "Im Mittelpunkt sollte nicht eine allgemeine Trenddiskussion stehen, sondern die Frage, wo sich aus den aktuellen Signalen "
            "konkrete Transfer-, Umsetzungs- oder GesprÃ¤chschancen ergeben. Weiter beobachtet werden sollten vor allem Linien, "
            "bei denen neue Partner, Finanzierung, Programme oder praktische Umsetzungsschritte hinzukommen."
        )
        return "\n\n".join([intro, p1, p2, actions])

    def build_opportunity_line_weekly_analysis(self, weekly_events: list[dict[str, Any]]) -> str:
        searched_member_count = self.checked_member_count()
        ranked = self.rank_events_for_narrative(weekly_events)
        substantial = [
            event for event in ranked
            if self.is_report_worthy_event(event)
            and not self.is_low_value_narrative_event(event)
        ]
        if not substantial:
            return self.build_opportunity_line_thin_week_report(ranked[:8])

        selected = substantial[:5]
        anchors = selected[:4]
        member_count = len({event["mitglied"] for event in selected})
        section_titles = self.join_german_list([
            REPORT_SECTIONS.get(section, {}).get("title", section)
            for section in sorted({event.get("hauptsektion", "beobachtung") for event in selected})
        ])
        scip_cards = self.scip_cards_for_briefing(limit=3)
        primary_scip = scip_cards[0] if scip_cards else {}
        primary_theme = self.scip_card_theme_text(primary_scip) if primary_scip else section_titles
        primary_members = self.scip_card_members_text(primary_scip) if primary_scip else self.join_german_list(sorted({event["mitglied"] for event in anchors[:3]}))

        title = f"Kann {primary_theme} in eine konkrete Umsetzungslinie übersetzt werden?"
        opportunity_line = (
            f"# {title}\n\n"
            f"Die aktuellen Funde bei {member_count} Mitgliedern zeigen den Kontext: {section_titles} sind in dieser Woche keine isolierten Einzelmeldungen. "
            f"Die SCIP-Kuratierung verdichtet diese Lage zu konkreten Akteurslinien, vor allem rund um {primary_members}. "
            "Der Nutzen entsteht, wenn Kontextsignale und kuratierte Rollenlogik gemeinsam in eine prüfbare Umsetzungskonstellation führen."
        )

        anchor_sentences = []
        for idx, event in enumerate(anchors, start=1):
            prefix = ["Erster Anker", "Zweiter Anker", "Dritter Anker", "Vierter Anker"][idx - 1]
            anchor_sentences.append(
                f"{prefix}: {self.compact_event_reference(event)}. "
                f"Relevant ist der Fund, weil er auf {self.event_implication_phrase(event)} verweist."
            )
        context_anchors = " ".join(anchor_sentences)
        scip_anchors = " ".join(self.scip_anchor_sentence(card) for card in scip_cards[:3])
        actor_anchors = " ".join(part for part in [context_anchors, scip_anchors] if part)

        zirp_relevance = (
            "Für ZIRP liegt der Wert genau in der Verbindung beider Ebenen. "
            "Die Crawler-Funde zeigen, warum das Thema jetzt relevant ist; die SCIP-Karten zeigen, welche Akteure die Rollen für Bedarf, Kompetenz, Legitimation und Umsetzung übernehmen könnten. "
            "Der Engpass liegt deshalb weniger in zusätzlicher Aufmerksamkeit als in der Rollenklärung: Welche der kuratierten Linien hat einen konkreten Bedarf, einen belastbaren Ansprechpartner und einen realistischen nächsten Schritt?"
        )

        action_target = primary_members or self.join_german_list(sorted({event["mitglied"] for event in anchors[:3]}))
        scip_next = self.scip_card_next_step_text(primary_scip) if primary_scip else "eine konkrete Transfer-, Umsetzungs- oder Gesprächschance ableiten"
        next_step = (
            f"Ein pragmatischer nächster Schritt wäre ein kuratiertes Sondierungsformat mit {action_target}. "
            "Dieses Format sollte in einem begrenzten 90-Tage-Rahmen genau eine Frage klären: "
            f"welcher konkrete Bedarf hinter der Linie steht und wie sich als nächster Schritt {scip_next} lässt. "
            "Damit wird aus Beobachtung eine belastbare Projektprüfung."
        )

        return "\n\n".join([opportunity_line, actor_anchors, zirp_relevance, next_step])

    def build_opportunity_line_thin_week_report(self, events: list[dict[str, Any]]) -> str:
        period_start = self._report_period_start()
        searched_member_count = self.checked_member_count()
        candidate_events = [
            event for event in events
            if not self.is_macro_context_without_member_action(event)
            and not self.is_low_value_narrative_event(event)
        ][:4]

        opportunity_line = (
            "# Welche Funde rechtfertigen bereits eine Projektprüfung?\n\n"
            f"Die Auswertung vom {period_start.strftime('%d.%m.%Y')} bis {self.today.strftime('%d.%m.%Y')} "
            f"hat {searched_member_count} Mitglieder geprüft und liefert nur wenige robuste Leitungsimpulse. "
            "Die vorhandenen Hinweise reichen für eine gezielte Nachprüfung, aber noch nicht für eine breite Convening-Initiative."
        )

        if candidate_events:
            actor_anchors = " ".join(self.event_evidence_sentence(event) for event in candidate_events[:4])
        else:
            actor_anchors = (
                "Es gibt noch keine ausreichend belastbaren Anker für eine öffentliche Projektlinie."
            )

        zirp_relevance = (
            "Für ZIRP ist ein schwaches Wochenbild trotzdem nützlich, weil es die Grenze zwischen Beobachtung und Projektanlass markiert. "
            "Relevanz entsteht erst, wenn ein Thema wiederkehrt oder durch weitere Akteure, Programme oder Umsetzungsschritte konkreter wird. "
            "Der Engpass liegt deshalb nicht in der Aufmerksamkeit, sondern in der belastbaren Validierung."
        )

        next_step = (
            "Der nächste sinnvolle Schritt ist ein begrenzter Validierungsprozess: ZIRP sollte die belastbaren Einzelfunde nachrecherchieren "
            "und erst bei einem weiteren starken Signal ein kleines Sondierungsgespräch prüfen. "
            "So bleibt der Radar handlungsfähig, ohne schwache Evidenz zu überschätzen."
        )

        return "\n\n".join([opportunity_line, actor_anchors, zirp_relevance, next_step])

    def build_thin_week_report(self, events: list[dict[str, Any]]) -> str:
        period_start = self._report_period_start()
        searched_member_count = self.checked_member_count()
        candidate_events = [
            event for event in events
            if not self.is_macro_context_without_member_action(event)
            and not self.is_low_value_narrative_event(event)
        ][:5]
        lines = [
            (
                f"Die Auswertung vom {period_start.strftime('%d.%m.%Y')} bis "
                f"{self.today.strftime('%d.%m.%Y')} hat {searched_member_count} Mitglieder geprÃ¼ft, "
                "liefert aber nur wenige robuste Leitungsimpulse. Mehrere Treffer sind eher Monitoring-Material "
                "als belastbare Grundlage fÃ¼r eine strategische Wochenklammer."
            )
        ]
        if candidate_events:
            lines.append(
                "Als beobachtbare Signale bleiben vor allem einzelne Hinweise stehen: "
                + " ".join(self.event_evidence_sentence(event) for event in candidate_events[:3])
            )
            lines.append(
                "FÃ¼r ZIRP ist diese Woche deshalb eher als Watchlist zu lesen. Relevanz entsteht erst, wenn aus solchen Hinweisen "
                "konkrete Transferprojekte, regionale Umsetzungsschritte, neue Partnerschaften oder belastbare institutionelle Entscheidungen werden."
            )
        else:
            lines.append(
                "Die verbleibenden Treffer enthalten keine hinreichend klare Verbindung aus Mitgliedsbezug, regionaler Relevanz und konkreter Handlungsebene. "
                "Ein kÃ¼nstlicher Themencluster wÃ¤re in dieser Lage irrefÃ¼hrend."
            )
        lines.append(
            "Was ZIRP daraus machen kÃ¶nnte:\n\n"
            "ZIRP sollte aus dieser Woche keine grÃ¶ÃŸere Convening-Initiative ableiten, sondern die belastbaren Einzelfunde nachrecherchieren. "
            "Weiter beobachtet werden sollte, ob aus den Hochschul-, Transfer- oder Qualifizierungssignalen in den nÃ¤chsten Wochen konkrete Partner, Programme oder Umsetzungsformate hervorgehen. "
            "Bei einem weiteren starken Signal kann ein kleines SondierungsgesprÃ¤ch geprÃ¼ft werden."
        )
        return "\n\n".join(lines)

    def compact_event_reference(self, event: dict[str, Any]) -> str:
        member = str(event.get("mitglied", "")).strip()
        title = str(event.get("titel", "")).strip()
        snippet = self.clean_narrative_snippet(event.get("snippet", ""), max_chars=130)
        if snippet:
            return f"{member} mit '{title}' ({snippet})"
        return f"{member} mit '{title}'"

    def event_implication_phrase(self, event: dict[str, Any]) -> str:
        section = str(event.get("hauptsektion", ""))
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        if "resilienz" in text or "hochwasser" in text or "klima" in text:
            return "kommunale Resilienz, Wissenstransfer und Umsetzungsfragen"
        if section == "versorgung_gesundheit" or "pflege" in text:
            return "FachkrÃ¤ftesicherung und VersorgungskapazitÃ¤t"
        if section == "kooperationen" or "transfer" in text or "kooperation" in text:
            return "Kooperation, Transfer und anwendungsnahe Vernetzung"
        if section == "wirtschaftsentwicklung" or "emissionshandel" in text or "investition" in text:
            return "Standortbedingungen, Investitionsanreize und Transformationsdruck"
        return "einen mÃ¶glichen Ansatzpunkt fÃ¼r regionale HandlungsfÃ¤higkeit"

    def rank_events_for_narrative(self, events: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return sorted(
            events,
            key=lambda e: (
                -self.narrative_event_score(e),
                -e.get("decision_score", 0),
                -e.get("score", 0),
                -e["datum_obj"].timestamp(),
            ),
        )

    def narrative_event_score(self, event: dict[str, Any]) -> int:
        text = f"{event.get('titel', '')} {event.get('snippet', '')}".lower()
        frame = self.editorial_frame_for(event)
        role_bonus = {"top_signal": 8, "condensation": 5, "watchlist": 2, "ignore": -8}.get(
            frame.get("role_in_briefing", "watchlist"),
            0,
        )
        score = role_bonus
        score += min(event.get("direct_member_score", 0), 6)
        score += min(event.get("regional_score", 0), 4)
        score += min(len(event.get("themenfelder", [])), 4)
        score += 4 if any(t in text for t in STRONG_STRATEGIC_TERMS) else 0
        score += 3 if any(t in text for t in INSTITUTIONAL_ACTION_TERMS) else 0
        score += 2 if event.get("hauptsektion") in {"versorgung_gesundheit", "kooperationen", "wirtschaftsentwicklung"} else 0
        score -= 5 if self.is_low_value_narrative_event(event) else 0
        return score

    def select_events_for_narrative(self, events: list[dict[str, Any]]) -> list[dict[str, Any]]:
        ranked = self.rank_events_for_narrative(events)
        narrative_limit = self.dynamic_narrative_event_limit()
        selected = []
        used_titles = set()

        for event in ranked:
            if not self.is_report_worthy_event(event):
                continue
            title_key = re.sub(r"\s+", " ", str(event.get("titel", "")).lower()).strip()
            if title_key in used_titles:
                continue
            selected.append(event)
            used_titles.add(title_key)
            if len(selected) >= narrative_limit:
                break

        if len(selected) < min(6, narrative_limit):
            for event in ranked:
                if event in selected:
                    continue
                if self.is_macro_context_without_member_action(event):
                    continue
                if self.is_low_value_narrative_event(event):
                    continue
                selected.append(event)
                if len(selected) >= narrative_limit:
                    break

        return selected

    def is_low_value_narrative_event(self, event: dict[str, Any]) -> bool:
        text = f"{event.get('titel', '')} {event.get('snippet', '')} {event.get('url', '')}".lower()
        low_value_terms = [
            "alumni", "zurÃ¼ck am campus", "wiedersehen", "fischdetektiv",
            "messeauftritt", "erfolgreich auf der hannover messe", "sponsoring",
            "brustring", "business club", "gewinnspiel", "campus entdecken",
            "ladendiebstahl", "e-scooter", "so tanken sie", "heute am gÃ¼nstigsten",
            "waffenruhe wackelt", "usa und iran", "Ã¶lpreis", "bremsspuren des krieges",
            "infoabend", "nacht der wissenschaft", "girls'day", "girlsday",
            "girlsâ€™ day", "girls' day", "zukunft mint", "33 mÃ¤dchen", "33 maedchen",
            "das bestÃ¤ndige haus", "dissertation erfolgreich verteidigt",
            "flora: wenn aus wurzeln wunder wachsen", "ausstellung flora",
        ]
        low_value_terms.extend([
            "che-ranking", "bestnoten", "trump droht iran", "kapitalmÃ¤rkte daily",
            "kapitalmÃ¤rkte daily",
        ])
        return any(term in text for term in low_value_terms)

    def build_theme_paragraph(self, section: str, events: list[dict[str, Any]]) -> str:
        lead = self.theme_lead_sentence(section, events)
        evidence_parts = [self.event_evidence_sentence(event) for event in events[:2]]
        implication = self.theme_implication_sentence(section, events)
        return " ".join([lead, *evidence_parts, implication]).strip()

    def theme_lead_sentence(self, section: str, events: list[dict[str, Any]]) -> str:
        members = self.join_german_list(sorted({e["mitglied"] for e in events})[:4])
        if section == "versorgung_gesundheit":
            return f"Im Bereich Versorgung und Gesundheit entsteht ein verdichtetes Bild aus Ausbildungs-, Weiterbildungs- und Strukturfragen; sichtbar wird dies vor allem bei {members}."
        if section == "kooperationen":
            return f"Der zweite Strang liegt in Kooperation und Transfer: Mehrere Meldungen zeigen, dass Vernetzung zunehmend als Umsetzungsinstrument eingesetzt wird, unter anderem bei {members}."
        if section == "wirtschaftsentwicklung":
            return f"Wirtschaftlich relevant sind vor allem Signale, in denen Standortentwicklung, Investitionen, ProduktivitÃ¤t oder neue Werkzeuge konkrete Anpassungsprozesse anzeigen; Beispiele liefern {members}."
        if section == "soziales_engagement":
            return f"Im gesellschaftlichen Bereich treten Verantwortung, Weiterbildung und regionale Bindung als wiederkehrende Motive hervor; besonders sichtbar wird dies bei {members}."
        if section == "fuehrungswechsel":
            return f"Bei Führungs- und Personalthemen geht es weniger um Einzelpersonalien als um mögliche Verschiebungen institutioneller Prioritäten; einschlägig sind hier {members}."
        return f"Ein weiterer Beobachtungsstrang ergibt sich aus den Meldungen von {members}."

    def event_evidence_sentence(self, event: dict[str, Any]) -> str:
        title = str(event.get("titel", "")).strip()
        member = str(event.get("mitglied", "")).strip()
        snippet = self.clean_narrative_snippet(event.get("snippet", ""), max_chars=150)
        if snippet:
            return f"Bei {member} ist vor allem '{title}' relevant: {snippet}."
        return f"Bei {member} ist vor allem '{title}' relevant."

    def theme_implication_sentence(self, section: str, events: list[dict[str, Any]]) -> str:
        if section == "versorgung_gesundheit":
            return "FÃ¼r die ZIRP ist das relevant, weil sich daran FachkrÃ¤ftesicherung, VersorgungskapazitÃ¤t und institutionelle Verantwortung bÃ¼ndeln."
        if section == "kooperationen":
            return "Der gemeinsame Nenner ist nicht die Kooperation als Schlagwort, sondern die Frage, welche Partnerschaften tatsÃ¤chlich neue Lern-, Transfer- oder UmsetzungsmÃ¶glichkeiten erÃ¶ffnen."
        if section == "wirtschaftsentwicklung":
            return "Diese Funde sind vor allem als Hinweise auf WettbewerbsfÃ¤higkeit, regulatorische Anpassung, Investitionsdruck oder praktische Modernisierung lesbar."
        if section == "soziales_engagement":
            return "Der strategische Wert liegt dort, wo gesellschaftliche Verantwortung mit regionaler HandlungsfÃ¤higkeit oder Qualifizierung verbunden wird."
        if section == "fuehrungswechsel":
            return "BeobachtungswÃ¼rdig sind solche Signale vor allem dann, wenn sie neue Ansprechpartner, PrioritÃ¤ten oder strategische Akzente erwarten lassen."
        return "Als Muster ist dieser Strang vor allem dann relevant, wenn er sich in weiteren Meldungen bestÃ¤tigt."

    def clean_narrative_snippet(self, snippet: str, max_chars: int = 220) -> str:
        text = re.sub(r"\s+", " ", str(snippet)).strip()
        text = re.sub(r"\s*\.\.\.$", "", text)
        if len(text) > max_chars:
            text = text[:max_chars].rsplit(" ", 1)[0].rstrip() + "..."
        return text

    def normalize_report_text(self, text: str) -> str:
        text = self.repair_mojibake(str(text or ""))
        text = self.normalize_german_ascii_umlauts(text)

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
        text = re.sub(r"[ \t]+", " ", text)

        headings = [
            "Executive Brief",
            "Wochenanalyse",
            "Convening Radar",
            "Gesundheit / gesellschaftliche Verantwortung",
            "Transformation / Standort / Umsetzungspraxis",
        ]

        inline_evidence_headings = [
            "ZIRP-Relevanz",
            "Auffaellige Mitglieder",
            "Auff?llige Mitglieder",
            "Was ZIRP daraus machen kÃ¶nnte",
            "Was ZIRP daraus machen k?nnte",
        ]
        for heading in inline_evidence_headings:
            text = re.sub(
                rf"\s*(?:^|\n)\s*{re.escape(heading)}:?\s*(?:\n|\s)*",
                " ",
                text,
                flags=re.I,
            )

        for heading in headings:
            text = re.sub(
                rf"\s*({re.escape(heading)}:?)\s*",
                r"\n\n\1\n\n",
                text,
                flags=re.I,
            )

        text = re.sub(r"\s+(DO:)", r"\n\n\1", text)
        text = re.sub(r"\s+(WATCH:)", r"\n\n\1", text)
        text = re.sub(r"\s+(MAYBE:)", r"\n\n\1", text)
        text = re.sub(r"\s+(DROP:)", r"\n\n\1", text)
        text = re.sub(r"(?m)^\s+(\d+\.\s+)", r"\n\n\1", text)
        text = re.sub(r"(?m)^\s*---+\s*$", "", text)
        text = re.sub(r"(?im)^\s*Die Analyse beendet sich hiermit\.?\s*$", "", text)
        text = re.sub(r"Signalsignaten", "Signalen", text)
        text = re.sub(r"\s+auszukosten,", " zu nutzen,", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    @staticmethod
    def normalize_german_ascii_umlauts(text: str) -> str:
        replacements = {
            "Waehrend": "WÃ¤hrend",
            "waehrend": "wÃ¤hrend",
            "Waechst": "WÃ¤chst",
            "waechst": "wÃ¤chst",
            "FÃ¼r": "FÃ¼r",
            "fÃ¼r": "fÃ¼r",
            "Koennte": "KÃ¶nnte",
            "kÃ¶nnte": "kÃ¶nnte",
            "KÃ¶nnen": "KÃ¶nnen",
            "kÃ¶nnen": "kÃ¶nnen",
            "Oeffentliche": "Ã–ffentliche",
            "oeffentliche": "Ã¶ffentliche",
            "Oeffentlich": "Ã–ffentlich",
            "oeffentlich": "Ã¶ffentlich",
            "Moeglich": "MÃ¶glich",
            "mÃ¶glich": "mÃ¶glich",
            "mÃ¶gliche": "mÃ¶gliche",
            "Loesung": "LÃ¶sung",
            "loesung": "lÃ¶sung",
            "LÃ¶sungen": "LÃ¶sungen",
            "loesungen": "lÃ¶sungen",
            "schlieÃŸen": "schlieÃŸen",
            "Schliessen": "SchlieÃŸen",
        }
        for bad, good in replacements.items():
            text = text.replace(bad, good)
        return text

    @staticmethod
    def join_german_list(items: list[str]) -> str:
        clean = [str(item).strip() for item in items if str(item).strip()]
        if not clean:
            return "keine klaren Schwerpunkte"
        if len(clean) == 1:
            return clean[0]
        if len(clean) == 2:
            return f"{clean[0]} und {clean[1]}"
        return f"{', '.join(clean[:-1])} und {clean[-1]}"

    def build_narrative_report_text(self, weekly_events: list[dict[str, Any]]) -> str:
        analysis_events = self.weekly_events_for_main_analysis()
        if not analysis_events and weekly_events:
            analysis_events = [
                event for event in weekly_events
                if str(event.get("evidence_role", "")) != "exclude"
                and not self.is_macro_context_without_member_action(event)
            ][:MAX_NARRATIVE_EVENTS]
        if not analysis_events:
            if self.openai_client is not None:
                print("OpenAI editorial: keine Ereignisse fÃ¼r Wochenanalyse; Modellbericht Ã¼bersprungen")
            meta = self.enrich_weekly_meta([])
            narrative = self.normalize_report_text(meta.get("narrative_analysis", "").strip())
            return self.normalize_report_text(
                "\n".join(
                    line for line in [
                        meta.get("headline", "Wochenanalyse"),
                        meta.get("meta_line", ""),
                        "",
                        narrative,
                    ]
                    if line is not None
                )
            )
        if self.openai_client is not None:
            print(f"OpenAI editorial: erstelle finale Wochenanalyse mit {min(len(analysis_events), MAX_NARRATIVE_EVENTS)} Ereignissen")

        meta = self.enrich_weekly_meta(analysis_events)
        narrative = self.normalize_report_text(meta.get("narrative_analysis", "").strip())

        lines = [
            meta.get("headline", "Wochenanalyse"),
            meta.get("meta_line", ""),
            "",
            narrative,
        ]

        if INCLUDE_DAILY_CALENDAR:
            lines.extend(["", self.build_week_calendar_text(weekly_events)])

        return self.normalize_report_text(
            "\n".join(line for line in lines if line is not None)
        )

    def render_top_signal_from_frame(self, event: dict[str, Any], frame: dict[str, Any]) -> dict[str, str]:
        title = str(event.get("titel", "")).strip()
        snippet = re.sub(r"\s+", " ", str(event.get("snippet", "")).strip())
        if len(snippet) > 240:
            snippet = snippet[:237].rstrip() + "..."

        happened = f"Bei **{event['mitglied']}** fÃ¤llt die Meldung **â€ž{title}â€œ** auf. {snippet}".strip()

        why = frame.get("world_relevance", "").strip() or self.build_why_it_matters(
            event, self.classify_signal_strength(event, self.weekly_events())
        )
        zirp = frame.get("zirp_relevance", "").strip() or self.build_zirp_relevance(event)
        next_step = frame.get("recommended_next_step", "").strip() or self.suggest_next_step(
            event,
            self.classify_signal_strength(event, self.weekly_events()),
            self.classify_signal_scope(event),
        )

        return {
            "signaltyp": self.classify_signal_scope(event),
            "signalstaerke": self.classify_signal_strength(event, self.weekly_events()),
            "was_ist_passiert": happened,
            "warum_relevant": why,
            "zirp_relevanz": zirp,
            "nÃ¤chster_schritt": next_step,
            "source_line": self.build_source_line(event),
        }

    def render_watchlist_line_from_frame(self, event: dict[str, Any], frame: dict[str, Any]) -> str:
        strength = self.classify_signal_strength(event, self.weekly_events())
        zirp_relevance = frame.get("zirp_relevance", "").strip() or self.build_zirp_relevance(event)
        next_step = frame.get("recommended_next_step", "").strip() or self.suggest_next_step(
            event, strength, self.classify_signal_scope(event)
        )
        return (
            f"- **{event['titel']}** ({event['mitglied']})  **{strength}**. "
            f"{zirp_relevance} Daraus folgt als begrenzter Prüfschritt: {next_step}."
        )

    def render_condensation_line_from_frame(
            self,
            event: dict[str, Any],
            frame: dict[str, Any],
            weekly_events: list[dict[str, Any]]
    ) -> str:
        pattern_score = self.strategic_pattern_score(event, weekly_events)
        themes = ", ".join(event.get("themenfelder", [])) or "keine klare Zuordnung"
        angle = frame.get("dominant_angle", "").strip()
        if angle:
            angle = f" Einordnung: {angle}"
        return (
            f"- **{event['titel']}** ({event['mitglied']})  "
            f"Verdichtungswert **{pattern_score}**, Themenfelder: **{themes}**.{angle}"
        )


    def _report_period_start(self) -> datetime:
        today_start = self.today.replace(hour=0, minute=0, second=0, microsecond=0)
        return today_start - timedelta(days=PUBLIC_REPORT_DAYS - 1)

    def _weekday_label(self, dt: datetime) -> str:
        names = {
            0: "Montag",
            1: "Dienstag",
            2: "Mittwoch",
            3: "Donnerstag",
            4: "Freitag",
            5: "Samstag",
            6: "Sonntag",
        }
        return names[dt.weekday()]

    def thematic_label_for_event(self, event: dict[str, Any]) -> str:
        title = str(event.get("titel", "")).lower()
        snippet = str(event.get("snippet", "")).lower()
        section = str(event.get("hauptsektion", "")).lower()
        themes = set(event.get("themenfelder", []))

        if "international" in title or "international" in snippet or "internationalisierung" in title or "internationalisierung" in snippet:
            return "Internationalisierung"

        if "mint" in title or "mint" in snippet:
            return "Kooperation / MINT / Vernetzung"

        if "justiz" in title or "wissenschaft" in title or "transfer" in snippet:
            return "Wissenstransfer / institutionelle Ã–ffnung"

        if "kommune" in title or "kommunen" in title or "klima" in title or "biodivers" in snippet:
            return "Transformation / kommunale Praxis"

        if "iwf" in title or "wachstumsprognosen" in title or "konjunktur" in snippet:
            return "Wirtschaftlicher Kontext / Beobachtung"

        if section == "kooperationen":
            return "Kooperation / Vernetzung"
        if section == "versorgung_gesundheit":
            return "Versorgung / Gesundheit"
        if section == "fuehrungswechsel":
            return "Führung / Personal"
        if section == "soziales_engagement":
            return "Gesellschaft / Engagement"

        if "wissen" in themes and "wirtschaft" in themes:
            return "Wissen / Wirtschaft"
        if "wissen" in themes:
            return "Wissen / Transfer"
        if "technologie" in themes and "nachhaltigkeit" in themes:
            return "Technologie / Transformation"
        if "wirtschaft" in themes:
            return "Wirtschaft / Entwicklung"
        if "nachhaltigkeit" in themes:
            return "Transformation / Nachhaltigkeit"

        return "Beobachtung"

    def build_week_calendar_text(self, weekly_events: list[dict[str, Any]]) -> str:
        period_start = self._report_period_start()
        days = [period_start + timedelta(days=i) for i in range(PUBLIC_REPORT_DAYS)]

        by_day: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for event in weekly_events:
            by_day[event["datum"]].append(event)

        for key in by_day:
            by_day[key] = sorted(
                by_day[key],
                key=lambda e: (
                    -self.briefing_rank_score(e, weekly_events),
                    -e.get("decision_score", 0),
                    -e.get("score", 0),
                    e.get("mitglied", ""),
                ),
            )

        lines: list[str] = []
        lines.append(f"Kalender der letzten {PUBLIC_REPORT_DAYS} Tage")
        lines.append("")

        for day in days:
            date_str = day.strftime("%Y-%m-%d")
            header = f"{self._weekday_label(day)} · {day.strftime('%d.%m.%Y')}"
            lines.append(header)

            items = by_day.get(date_str, [])
            if not items:
                lines.append("Keine priorisierten Funde")
                lines.append("")
                continue

            for event in items:
                lines.append(event["mitglied"])
                lines.append(event["titel"])
                lines.append(f"Themenfeld: {self.thematic_label_for_event(event)}")
                lines.append("")

        return "\n".join(lines).strip()

    def build_leitungsfassung_text(self, weekly_events: list[dict[str, Any]]) -> str:
        ranked = sorted(
            weekly_events,
            key=lambda x: (
                -self.briefing_rank_score(x, weekly_events),
                -x["decision_score"],
                -x["score"],
                -x["datum_obj"].timestamp(),
            ),
        )

        period_start = self._report_period_start()
        period_end = self.today

        lines: list[str] = []
        lines.append("Leitungsfassung")
        lines.append("")

        if not ranked:
            lines.append(
                f"Im Zeitraum vom {period_start.strftime('%d.%m.%Y')} bis {period_end.strftime('%d.%m.%Y')} "
                "wurden keine priorisierten Funde identifiziert."
            )
            return "\n".join(lines).strip()

        member_count = len({event["mitglied"] for event in ranked})
        section_counts = Counter(event.get("hauptsektion", "") for event in ranked)
        leading_sections = [
            REPORT_SECTIONS.get(section, {}).get("title", section)
            for section, _ in section_counts.most_common(3)
            if section
        ]
        section_text = ", ".join(leading_sections) if leading_sections else "keine klare Schwerpunktsetzung"

        lines.append(
            f"Der Bericht blickt auf den Zeitraum vom {period_start.strftime('%d.%m.%Y')} bis "
            f"{period_end.strftime('%d.%m.%Y')} zurueck. In diesem Fenster wurden {len(ranked)} "
            f"priorisierte Funde bei {member_count} Mitgliedern identifiziert. Die sichtbaren Schwerpunkte liegen bei: "
            f"{section_text}."
        )
        lines.append("")

        lines.append("Priorisierte Signale:")
        for event in ranked[:5]:
            frame = self.editorial_frame_for(event) if self.editorial_frames else self.build_rule_based_editorial_frame(event)
            role = frame.get("role_in_briefing", "watchlist")
            why = frame.get("zirp_relevance", "").strip() or self.build_zirp_relevance(event)
            lines.append(
                f"- {event['datum']} | {event['mitglied']}: {event['titel']} "
                f"({self.thematic_label_for_event(event)}, Rolle: {role}). {why}"
            )
        lines.append("")

        context_items = [
            event for event in ranked
            if (self.editorial_frame_for(event) if self.editorial_frames else self.build_rule_based_editorial_frame(event)).get("is_context_only")
        ]
        if context_items:
            lines.append(
                f"Als Kontextsignale bleiben {len(context_items)} Funde beobachtungsrelevant; sie werden jedoch nicht als aktive Leitsignale gewertet."
            )
        else:
            lines.append(
                "Es wurden keine gesonderten Kontextsignale priorisiert; die Bewertung konzentriert sich auf die oben genannten Funde."
            )

        return "\n".join(lines).strip()

        top_signals = ranked[:3]
        condensation = ranked[3] if len(ranked) > 3 else None
        context_signal = None
        for event in ranked:
            frame = self.editorial_frame_for(event) if self.editorial_frames else self.build_rule_based_editorial_frame(event)
            if frame.get("event_type") == "macro_context" or "iwf" in event.get("titel", "").lower():
                context_signal = event
                break

        lines: list[str] = []
        lines.append("Leitungsfassung")
        lines.append("")

        lines.append(
            "Die Woche ist geprÃ¤gt von belastbaren Kooperations- und Innovationssignalen. "
            "Besonders auffÃ¤llig ist, dass Zukunftsthemen nicht abstrakt bleiben, sondern in konkrete Partnerschaften, "
            "Transferformate und institutionelle AnschlussbezÃ¼ge Ã¼bersetzt werden. "
            "FÃ¼r die ZIRP entstehen daraus Ansatzpunkte fÃ¼r GesprÃ¤chsfÃ¼hrung, Vernetzung und Themensteuerung."
        )
        lines.append("")

        if len(top_signals) >= 3:
            a, b, c = top_signals[:3]
            lines.append(
                f"Besonders sichtbar wird dies an drei Entwicklungen: "
                f"{a['mitglied']} setzt mit â€ž{a['titel']}â€œ ein Signal im Themenfeld {self.thematic_label_for_event(a).lower()} "
                f"und macht damit eine aktiv organisierte Kooperations- und Projektlogik sichtbar. "
                f"{b['mitglied']} verdichtet mit â€ž{b['titel']}â€œ die sektorÃ¼bergreifende Zusammenarbeit und steht damit fÃ¼r "
                f"{self.thematic_label_for_event(b).lower()}. "
                f"Hinzu kommt das Signal von {c['mitglied']} mit â€ž{c['titel']}â€œ, das zeigt, dass "
                f"{self.thematic_label_for_event(c).lower()} zunehmend Ã¼ber konkrete Umsetzungsleistung sichtbar wird."
            )
            lines.append("")
        elif top_signals:
            lines.append(
                "Im Vordergrund stehen mehrere belastbare Einzelsignale, die vor allem auf Kooperation, Transfer und regionale AnschlussfÃ¤higkeit verweisen."
            )
            lines.append("")

        if condensation:
            lines.append(
                f"In der Verdichtung fÃ¤llt zudem {condensation['mitglied']} mit â€ž{condensation['titel']}â€œ auf, "
                f"das vor allem mit Blick auf {self.thematic_label_for_event(condensation).lower()} relevant ist."
            )
            lines.append("")

        if context_signal:
            lines.append(
                f"Als Rahmensignal bleibt schlieÃŸlich â€ž{context_signal['titel']}â€œ von {context_signal['mitglied']} beobachtungsrelevant, "
                f"weil es auf ein fragileres wirtschaftliches Umfeld und mÃ¶gliche Belastungen fÃ¼r Investitionen, Planungssicherheit "
                f"und Transformation verweist, derzeit jedoch ohne Vorrang fÃ¼r die aktive ThemenfÃ¼hrung."
            )
            lines.append("")
        else:
            lines.append(
                "Kontextsignale bleiben beobachtungsrelevant, stehen derzeit aber hinter den konkreten Kooperations- und Umsetzungsimpulsen zurÃ¼ck."
            )
            lines.append("")

        lines.append(
            "Insgesamt zeigt die Woche mehrere anschlussfÃ¤hige Signale fÃ¼r Beobachtung, GesprÃ¤chsvorbereitung und strategische Formatentwicklung; "
            "besonders relevant sind Entwicklungen, in denen Kooperation, Wissenstransfer und regionale UmsetzungsfÃ¤higkeit zusammenfallen."
        )

        return "\n".join(lines).strip()


    def build_report_texts(self) -> tuple[str, str]:
        weekly_events = self.weekly_events()

        if weekly_events and MODEL_DECIDES_RELEVANCE and not self.editorial_frames:
            self.enrich_weekly_events_editorially(weekly_events)

        if REPORT_STYLE == "narrative":
            briefing_text = self.build_narrative_report_text(weekly_events)
            if not INCLUDE_EVIDENCE_SECTION:
                return briefing_text, ""

        if not weekly_events:
            calendar_text = self.build_week_calendar_text([])
            leitungsfassung_text = self.build_leitungsfassung_text([])
            briefing_text = f"{calendar_text}\n\n{leitungsfassung_text}"
            evidence_text = "# Evidenzteil\n\nKeine belastbaren Funde im Berichtszeitraum."
            return briefing_text, evidence_text

        calendar_text = self.build_week_calendar_text(weekly_events)
        leitungsfassung_text = self.build_leitungsfassung_text(weekly_events)
        briefing_text = f"{calendar_text}\n\n{leitungsfassung_text}"

        evidence = []
        evidence.append("# Evidenzteil")
        evidence.append("")
        evidence.append(
            "Die folgenden Funde bilden die Grundlage des Leiterinnen-Briefings. "
            "BerÃ¼cksichtigt wurden nur Artikel mit hinreichender Artikellogik, belastbarem Titel, verwertbarem Inhalt "
            f"und einer Relevanzentscheidung fÃ¼r den Zeitraum der letzten {REPORT_LOOKBACK_DAYS} Tage."
        )
        evidence.append("")

        evidence_profiles = self.build_evidence_profiles()
        for section_key, cfg in REPORT_SECTIONS.items():
            evidence.append(f"## {cfg['title']}")
            evidence.append("")
            items = evidence_profiles.get(section_key, [])
            if not items:
                evidence.append("Keine priorisierten Funde im Berichtszeitraum.")
                evidence.append("")
                continue

            for event in items:
                frame = self.editorial_frame_for(event)
                snippet = re.sub(r"\s+", " ", event.get("snippet", "")).strip()
                if len(snippet) > 260:
                    snippet = snippet[:257].rstrip() + "..."
                themes = ", ".join(event.get("themenfelder", [])) or "keine Zuordnung"
                evidence.append(
                    f"**{event['mitglied']}** | **{event['datum']}** | **{event['titel']}**\n"
                    f"- Score: {event['score']} | Decision Score: {event['decision_score']}\n"
                    f"- Themenfelder: {themes}\n"
                    f"- Direct Member Score: {event.get('direct_member_score', 0)} | Regional Score: {event.get('regional_score', 0)} | Static Fit: {event.get('static_member_fit_score', 0)} | Article Quality: {event.get('article_quality_score', 0)}\n"
                    f"- Editorial Frame: {json.dumps(frame, ensure_ascii=False)}\n"
                    f"- URL: {event['url']}\n"
                    f"- Snippet: {snippet}"
                )
                evidence.append("")

        evidence_text = "\n".join(evidence)
        return briefing_text, evidence_text

    def create_outputs(self) -> None:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        member_homepage_csv = self.output_dir / f"zirp_mitglied_homepages_{timestamp}.csv"
        hits_csv = self.output_dir / f"zirp_seitenfunde_{timestamp}.csv"
        keyword_csv = self.output_dir / f"zirp_begriffe_gewichtet_{timestamp}.csv"
        events_csv = self.output_dir / f"zirp_ereignisse_{timestamp}.csv"
        pipeline_csv = self.output_dir / f"zirp_raw_events_{timestamp}.csv"
        report_docx = self.output_dir / f"zirp_leiterinnen_briefing_{timestamp}.docx"
        dashboard_html = self.output_dir / "zirp_dashboard.html"

        pd.DataFrame(self.member_homepages).to_csv(member_homepage_csv, index=False, encoding="utf-8-sig")
        pd.DataFrame(self.page_hits).to_csv(hits_csv, index=False, encoding="utf-8-sig")

        weekly_events = self.weekly_events()
        if weekly_events and MODEL_DECIDES_RELEVANCE and not self.editorial_frames:
            self.enrich_weekly_events_editorially(weekly_events)

        weighted_df = self.build_weighted_keyword_dataframe()
        weighted_df.to_csv(keyword_csv, index=False, encoding="utf-8-sig")

        events_df = self.build_events_dataframe()
        events_df.to_csv(events_csv, index=False, encoding="utf-8-sig")

        pipeline_df = self.build_pipeline_events_dataframe()
        pipeline_df.to_csv(pipeline_csv, index=False, encoding="utf-8-sig")

        print("Pipeline: 1/4 Crawler-Daten gespeichert")
        optimizer_output = self.run_meeting_optimizer_after_report(enable_ai=False)
        print("Pipeline: 2/4 SCIP-Muster ohne KI-Radartext berechnet")

        briefing_text, evidence_text = self.build_report_texts()
        self.create_word_report(report_docx, briefing_text, evidence_text)
        print("Pipeline: 3/4 Bericht erstellt (KI oder schneller Fallback)")
        briefing_context_path = self.output_dir / "zirp_current_briefing_context.txt"
        briefing_context_path.write_text(briefing_text, encoding="utf-8")
        radar_ai_output = self.run_meeting_optimizer_after_report(
            enable_ai=True,
            report_context_path=briefing_context_path,
        )
        print("Pipeline: 4/4 Convening Radar mit Berichtskontext per KI erstellt")
        self.create_dashboard(
            dashboard_html,
            briefing_text=briefing_text,
            weekly_events=weekly_events,
            events_df=events_df,
            weighted_df=weighted_df,
            timestamp=timestamp,
        )
        print("Pipeline: Dashboard erstellt")
        if CREATE_WORDCLOUD:
            wordcloud_png = self.output_dir / f"zirp_wordcloud_{timestamp}.png"
            self.create_wordcloud(wordcloud_png)

        print("Ergebnisse gespeichert:")
        print(f"  Homepages: {member_homepage_csv.resolve()}")
        print(f"  Seitenfunde: {hits_csv.resolve()}")
        print(f"  Gewichtete Begriffe: {keyword_csv.resolve()}")
        print(f"  Ereignisse: {events_csv.resolve()}")
        print(f"  Pipeline Events: {pipeline_csv.resolve()}")
        print(f"  Bericht: {report_docx.resolve()}")
        print(f"  Dashboard: {dashboard_html.resolve()}")
        if optimizer_output:
            print(optimizer_output)
        if radar_ai_output:
            print(radar_ai_output)
        if CREATE_WORDCLOUD and WordCloud is not None:
            print(f"  WordCloud: {wordcloud_png.resolve()}")

    def run_meeting_optimizer_after_report(
            self,
            enable_ai: bool = True,
            report_context_path: Optional[Path] = None,
    ) -> str:
        if not RUN_MEETING_OPTIMIZER_AFTER_REPORT:
            return "Meeting Optimizer: Ã¼bersprungen (RUN_MEETING_OPTIMIZER_AFTER_REPORT=False)"

        # Prefer the cleaned Excel-linked optimizer if it exists; otherwise keep the
        # old filename as fallback so prototype_v2 remains compatible with v1 folders.
        optimizer_candidates = [
            Path(__file__).resolve().with_name("zirp_meeting_optimizer.py"),
            Path(__file__).resolve().with_name("scip_optimizer_v2.py"),
        ]
        optimizer_path = next((path for path in optimizer_candidates if path.exists()), optimizer_candidates[0])
        if not optimizer_path.exists():
            return f"Meeting Optimizer: Ã¼bersprungen ({optimizer_path.name} nicht gefunden)"

        self.meeting_optimizer_started_at = time.time()
        env = os.environ.copy()
        env["ZIRP_CONVENING_AI"] = "1" if enable_ai else "0"
        env["ZIRP_STATIC_CANDIDATE_PATH"] = str(_resolve_project_path(STATIC_ZIRP_EXCEL_PATH))
        env["STATIC_ZIRP_EXCEL_PATH"] = str(_resolve_project_path(STATIC_ZIRP_EXCEL_PATH))
        if report_context_path is not None:
            env["ZIRP_CONVENING_CONTEXT_FILE"] = str(report_context_path)
        try:
            proc = subprocess.run(
                [sys.executable, str(optimizer_path)],
                cwd=str(optimizer_path.parent),
                capture_output=True,
                text=True,
                timeout=MEETING_OPTIMIZER_TIMEOUT,
                env=env,
            )
        except subprocess.TimeoutExpired:
            return f"Meeting Optimizer: Timeout nach {MEETING_OPTIMIZER_TIMEOUT} Sekunden"
        except Exception as exc:
            return f"Meeting Optimizer: konnte nicht gestartet werden ({exc})"

        output = "\n".join(part.strip() for part in [proc.stdout, proc.stderr] if part and part.strip())
        if proc.returncode != 0:
            return f"Meeting Optimizer: Fehlercode {proc.returncode}\n{output}"
        return f"Meeting Optimizer: abgeschlossen\n{output}"

    def build_weighted_keyword_dataframe(self) -> pd.DataFrame:
        if not self.keyword_rows:
            return pd.DataFrame(columns=[
                "begriff", "raw_count", "weighted_score",
                "wirtschaftsentwicklung", "versorgung_gesundheit", "fuehrungswechsel",
                "soziales_engagement", "kooperationen"
            ])

        df = pd.DataFrame(self.keyword_rows)
        grouped = df.groupby("begriff", as_index=False).agg({
            "count": "sum",
            "wirtschaftsentwicklung": "sum",
            "versorgung_gesundheit": "sum",
            "fuehrungswechsel": "sum",
            "soziales_engagement": "sum",
            "kooperationen": "sum",
        }).rename(columns={"count": "raw_count"})

        grouped["weighted_score"] = (
                grouped["raw_count"]
                + grouped["wirtschaftsentwicklung"] * 2
                + grouped["versorgung_gesundheit"] * 2
                + grouped["fuehrungswechsel"] * 2
                + grouped["soziales_engagement"] * 2
                + grouped["kooperationen"] * 2
        )

        self.global_weighted_counter.clear()
        for _, row in grouped.iterrows():
            self.global_weighted_counter[str(row["begriff"])] = int(row["weighted_score"])

        return grouped.sort_values(
            ["weighted_score", "raw_count", "begriff"],
            ascending=[False, False, True]
        )

    def build_events_dataframe(self) -> pd.DataFrame:
        if not self.events:
            return pd.DataFrame(columns=[
                "member_id", "mitglied", "titel", "datum", "url", "hauptsektion",
                "themenfelder", "snippet", "score", "decision_score", "is_updated",
                "event_type", "dominant_angle", "world_relevance", "zirp_relevance",
                "recommended_next_step", "confidence", "is_context_only", "role_in_briefing",
                "one_sentence_context", "why_it_matters", "mechanism", "named_actors",
                "concrete_assets", "next_date_or_window", "implementation_angle",
                "scip_reasoning", "opportunity_strength", "recommended_zirp_use",
                "briefing_sentence", "opportunity_keywords",
            ])

        rows = []
        for e in self.events:
            frame = self.editorial_frame_for(e) if self.editorial_frames else self.build_rule_based_editorial_frame(
                e)
            opportunity = self.build_signal_opportunity_fields(e, frame)
            rows.append({
                "member_id": e.get("member_id"),
                "mitglied": e["mitglied"],
                "titel": e["titel"],
                "datum": e["datum"],
                "url": e["url"],
                "hauptsektion": e["hauptsektion"],
                "themenfelder": ", ".join(e["themenfelder"]),
                "snippet": e["snippet"],
                "score": e["score"],
                "decision_score": e["decision_score"],
                "recency_weight": e["recency_weight"],
                "direct_member_score": e.get("direct_member_score", 0),
                "regional_score": e.get("regional_score", 0),
                "static_member_fit_score": e.get("static_member_fit_score", 0),
                "article_quality_score": e.get("article_quality_score", 0),
                "is_updated": e["is_updated"],
                "event_type": frame.get("event_type", ""),
                "dominant_angle": frame.get("dominant_angle", ""),
                "world_relevance": frame.get("world_relevance", ""),
                "zirp_relevance": frame.get("zirp_relevance", ""),
                "recommended_next_step": frame.get("recommended_next_step", ""),
                "confidence": frame.get("confidence", ""),
                "is_context_only": frame.get("is_context_only", False),
                "role_in_briefing": frame.get("role_in_briefing", ""),
                "zirp_label": e.get("zirp_label", ""),
                "zirp_reason": e.get("zirp_reason", ""),
                "llm_judge_source": e.get("llm_judge_source", ""),
                "llm_strategic_score": e.get("llm_strategic_score", ""),
                "llm_concreteness_score": e.get("llm_concreteness_score", ""),
                "llm_zirp_relevance": e.get("llm_zirp_relevance", ""),
                "llm_recommended_use": e.get("llm_recommended_use", ""),
                "llm_judge_reason": e.get("llm_judge_reason", ""),
                "one_sentence_context": opportunity.get("one_sentence_context", ""),
                "why_it_matters": opportunity.get("why_it_matters", ""),
                "mechanism": opportunity.get("mechanism", ""),
                "named_actors": opportunity.get("named_actors", ""),
                "concrete_assets": opportunity.get("concrete_assets", ""),
                "next_date_or_window": opportunity.get("next_date_or_window", ""),
                "implementation_angle": opportunity.get("implementation_angle", ""),
                "scip_reasoning": opportunity.get("scip_reasoning", ""),
                "opportunity_strength": opportunity.get("opportunity_strength", ""),
                "recommended_zirp_use": opportunity.get("recommended_zirp_use", ""),
                "briefing_sentence": opportunity.get("briefing_sentence", ""),
                "opportunity_keywords": opportunity.get("opportunity_keywords", ""),
            })

        return pd.DataFrame(rows).sort_values(
            ["static_member_fit_score", "direct_member_score", "regional_score", "decision_score", "score", "datum"],
            ascending=[False, False, False, False, False, False]
        )

    def build_pipeline_events_dataframe(self) -> pd.DataFrame:
        columns = [
            "event_id",
            "member_id",
            "member_name",
            "event_date",
            "title",
            "snippet",
            "url",
            "source_trust_score",
            "snippet_length",
            "duplicate_overlap_score",
            "event_coherence_score",
            "date_validity_score",
            "article_quality_score",
            "theme_density",
            "actor_density",
            "member_density",
            "cross_theme_score",
            "persistence_score",
            "emergence_score",
            "bridge_signal_score",
            "direct_member_score",
            "regional_score",
            "static_member_fit_score",
            "theme_fit_score",
            "actor_fit_score",
            "forecast_fit_score",
            "timing_fit_score",
            "bridge_potential_score",
            "cluster_alignment_score",
            "person_name_detected",
            "person_names_raw",
            "person_count",
            "hauptsektion",
            "themenfelder",
            "score",
            "decision_score",
            "is_updated",
            "event_type",
            "dominant_angle",
            "world_relevance",
            "zirp_relevance",
            "recommended_next_step",
            "confidence",
            "is_context_only",
            "role_in_briefing",
            "one_sentence_context",
            "why_it_matters",
            "mechanism",
            "named_actors",
            "concrete_assets",
            "next_date_or_window",
            "implementation_angle",
            "scip_reasoning",
            "opportunity_strength",
            "recommended_zirp_use",
            "briefing_sentence",
            "opportunity_keywords",
        ]

        if not self.events:
            return pd.DataFrame(columns=columns)

        rows = []
        for idx, e in enumerate(self.events, start=1):
            member_meta = self.member_lookup.get(e["mitglied"], {})
            snippet = e.get("snippet", "") or ""
            themes = e.get("themenfelder", []) or []
            section = e.get("hauptsektion", "")
            frame = self.editorial_frame_for(e) if self.editorial_frames else self.build_rule_based_editorial_frame(
                e)
            opportunity = self.build_signal_opportunity_fields(e, frame)

            source_trust_score = min(1.0, max(0.0, e.get("article_quality_score", 0) / 30))
            snippet_length = min(1.0, len(snippet.split()) / 80)
            duplicate_overlap_score = 0.0
            event_coherence_score = min(1.0, 0.4 + (len(snippet.split()) / 100))
            date_validity_score = 1.0 if e.get("datum") else 0.0

            theme_density = min(1.0, len(themes) / 5)
            actor_density = min(1.0, e.get("direct_member_score", 0) / 6)
            member_density = min(1.0, e.get("direct_member_score", 0) / 6)
            cross_theme_score = min(1.0, max(0, len(themes) - 1) / 4)
            persistence_score = 0.7 if e.get("recency_weight", 0) >= 2 else 0.4
            emergence_score = min(1.0, e.get("decision_score", 0) / 30)
            bridge_signal_score = 1.0 if section == "kooperationen" else 0.4
            theme_fit_score = min(1.0, e.get("decision_score", 0) / 25)
            static_fit_norm = min(1.0, e.get("static_member_fit_score", 0) / 10)
            actor_fit_score = min(1.0, max(e.get("direct_member_score", 0) / 6, static_fit_norm))
            forecast_fit_score = min(1.0, e.get("score", 0) / 50)
            timing_fit_score = min(1.0, e.get("recency_weight", 0) / 4)
            bridge_potential_score = 1.0 if section in {"kooperationen", "versorgung_gesundheit"} else 0.5
            cluster_alignment_score = min(1.0, len(themes) / 4)

            rows.append({
                "event_id": idx,
                "member_id": member_meta.get("member_id"),
                "member_name": e["mitglied"],
                "event_date": e["datum"],
                "title": e["titel"],
                "snippet": snippet,
                "url": e["url"],
                "source_trust_score": source_trust_score,
                "snippet_length": snippet_length,
                "duplicate_overlap_score": duplicate_overlap_score,
                "event_coherence_score": event_coherence_score,
                "date_validity_score": date_validity_score,
                "article_quality_score": e.get("article_quality_score", 0),
                "theme_density": theme_density,
                "actor_density": actor_density,
                "member_density": member_density,
                "cross_theme_score": cross_theme_score,
                "persistence_score": persistence_score,
                "emergence_score": emergence_score,
                "bridge_signal_score": bridge_signal_score,
                "direct_member_score": e.get("direct_member_score", 0),
                "regional_score": e.get("regional_score", 0),
                "static_member_fit_score": e.get("static_member_fit_score", 0),
                "theme_fit_score": theme_fit_score,
                "actor_fit_score": actor_fit_score,
                "forecast_fit_score": forecast_fit_score,
                "timing_fit_score": timing_fit_score,
                "bridge_potential_score": bridge_potential_score,
                "cluster_alignment_score": cluster_alignment_score,
                "person_name_detected": 0,
                "person_names_raw": "",
                "person_count": 0,
                "hauptsektion": section,
                "themenfelder": ", ".join(themes),
                "score": e.get("score", 0),
                "decision_score": e.get("decision_score", 0),
                "is_updated": e.get("is_updated", False),
                "event_type": frame.get("event_type", ""),
                "dominant_angle": frame.get("dominant_angle", ""),
                "world_relevance": frame.get("world_relevance", ""),
                "zirp_relevance": frame.get("zirp_relevance", ""),
                "recommended_next_step": frame.get("recommended_next_step", ""),
                "confidence": frame.get("confidence", ""),
                "is_context_only": frame.get("is_context_only", False),
                "role_in_briefing": frame.get("role_in_briefing", ""),
                "one_sentence_context": opportunity.get("one_sentence_context", ""),
                "why_it_matters": opportunity.get("why_it_matters", ""),
                "mechanism": opportunity.get("mechanism", ""),
                "named_actors": opportunity.get("named_actors", ""),
                "concrete_assets": opportunity.get("concrete_assets", ""),
                "next_date_or_window": opportunity.get("next_date_or_window", ""),
                "implementation_angle": opportunity.get("implementation_angle", ""),
                "scip_reasoning": opportunity.get("scip_reasoning", ""),
                "opportunity_strength": opportunity.get("opportunity_strength", ""),
                "recommended_zirp_use": opportunity.get("recommended_zirp_use", ""),
                "briefing_sentence": opportunity.get("briefing_sentence", ""),
                "opportunity_keywords": opportunity.get("opportunity_keywords", ""),
            })

        return pd.DataFrame(rows)

    def create_dashboard(
            self,
            filepath: Path,
            *,
            briefing_text: str,
            weekly_events: list[dict[str, Any]],
            events_df: pd.DataFrame,
            weighted_df: pd.DataFrame,
            timestamp: str,
    ) -> None:
        history_file = filepath.with_name("zirp_dashboard_history.json")
        history = self.load_dashboard_history(history_file)
        entry = self.build_dashboard_entry(
            briefing_text=briefing_text,
            weekly_events=weekly_events,
            events_df=events_df,
            weighted_df=weighted_df,
            timestamp=timestamp,
        )

        history = [item for item in history if item.get("id") != entry["id"]]
        history.insert(0, entry)
        history = history[:40]

        history_file.write_text(
            json.dumps(history, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        self.ensure_dashboard_logo_asset()
        filepath.write_text(self.repair_mojibake(self.render_dashboard_html(history)), encoding="utf-8")
        filepath.with_name("archive.html").write_text(
            self.repair_mojibake(self.render_archive_html(history)),
            encoding="utf-8",
        )
        filepath.with_name("scip_archive.html").write_text(
            self.repair_mojibake(self.render_scip_archive_html(history)),
            encoding="utf-8",
        )
        filepath.with_name("signals.html").write_text(
            self.repair_mojibake(self.render_signals_html(history)),
            encoding="utf-8",
        )

        root_dashboard = self.script_dir / "zirp_dashboard.html"
        root_index = self.script_dir / "index.html"
        root_html = self.repair_mojibake(self.render_dashboard_html(history, asset_prefix="zirp_berichte/"))
        root_dashboard.write_text(root_html, encoding="utf-8")
        root_index.write_text(root_html, encoding="utf-8")

    def load_dashboard_history(self, history_file: Path) -> list[dict[str, Any]]:
        if not history_file.exists():
            return []
        try:
            raw = json.loads(history_file.read_text(encoding="utf-8"))
            return raw if isinstance(raw, list) else []
        except Exception:
            return []

    def build_dashboard_entry(
            self,
            *,
            briefing_text: str,
            weekly_events: list[dict[str, Any]],
            events_df: pd.DataFrame,
            weighted_df: pd.DataFrame,
            timestamp: str,
    ) -> dict[str, Any]:
        headline, meta_line, analysis = self.split_briefing_text(briefing_text)
        scip_radar = self.latest_meeting_recommendations()
        period_start = self._report_period_start().strftime("%d.%m.%Y")
        period_end = self.today.strftime("%d.%m.%Y")
        members = sorted({event.get("mitglied", "") for event in weekly_events if event.get("mitglied")})
        sections = Counter(event.get("hauptsektion", "") for event in weekly_events)
        top_sections = [
            {
                "key": key,
                "title": REPORT_SECTIONS.get(key, {}).get("title", key),
                "count": count,
            }
            for key, count in sections.most_common(5)
            if key
        ]
        top_events = []
        for event in self.rank_events_for_narrative(weekly_events)[:8]:
            frame = self.editorial_frame_for(event) if self.editorial_frames else self.build_rule_based_editorial_frame(event)
            opportunity = self.build_signal_opportunity_fields(event, frame)
            top_events.append({
                "date": event.get("datum", ""),
                "member": event.get("mitglied", ""),
                "title": event.get("titel", ""),
                "theme": self.thematic_label_for_event(event),
                "url": event.get("url", ""),
                "one_sentence_context": opportunity.get("one_sentence_context", ""),
                "why_it_matters": opportunity.get("why_it_matters", ""),
                "mechanism": opportunity.get("mechanism", ""),
                "named_actors": opportunity.get("named_actors", ""),
                "concrete_assets": opportunity.get("concrete_assets", ""),
                "next_date_or_window": opportunity.get("next_date_or_window", ""),
                "implementation_angle": opportunity.get("implementation_angle", ""),
                "scip_reasoning": opportunity.get("scip_reasoning", ""),
            })

        top_terms = []
        if weighted_df is not None and not weighted_df.empty:
            for _, row in weighted_df.head(12).iterrows():
                top_terms.append({
                    "term": str(row.get("begriff", "")),
                    "score": int(row.get("weighted_score", 0) or 0),
                })

        entry_id = hashlib.sha256(
            f"{period_start}|{period_end}|{analysis}|{timestamp}".encode("utf-8")
        ).hexdigest()[:16]

        return {
            "id": entry_id,
            "timestamp": timestamp,
            "created_at": datetime.now().strftime("%d.%m.%Y, %H:%M:%S"),
            "period_start": period_start,
            "period_end": period_end,
            "headline": headline or "Wochenanalyse",
            "meta_line": meta_line,
            "analysis": analysis,
            "event_count": len(weekly_events),
            "member_count": len(members),
            "members": members,
            "top_sections": top_sections,
            "top_events": top_events,
            "top_terms": top_terms,
            "scip_radar": scip_radar,
        }

    def split_briefing_text(self, briefing_text: str) -> tuple[str, str, str]:
        lines = [line.rstrip() for line in str(briefing_text).splitlines()]
        non_empty = [line for line in lines if line.strip()]
        if not non_empty:
            return "Wochenanalyse", "", ""
        headline = re.sub(r"^\s*#{1,6}\s*", "", non_empty[0]).strip()
        meta_line = non_empty[1] if len(non_empty) > 1 else ""
        analysis_start = 2 if len(non_empty) > 1 else 1
        analysis = "\n\n".join(line.strip() for line in non_empty[analysis_start:] if line.strip())
        return headline, meta_line, analysis

    def latest_meeting_recommendations(self) -> dict[str, Any]:
        pattern_files = sorted(
            self.output_dir.glob("zirp_meeting_patterns_*.csv"),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        if not pattern_files:
            return {}
        if self.meeting_optimizer_started_at is not None:
            pattern_files = [
                path for path in pattern_files
                if path.stat().st_mtime >= self.meeting_optimizer_started_at - 5
            ]
            if not pattern_files:
                return {}
        newest_mtime = pattern_files[0].stat().st_mtime
        pattern_files = [
            path for path in pattern_files
            if newest_mtime - path.stat().st_mtime <= 300
        ]
        pattern_files.sort(
            key=lambda path: (
                0 if "_ai_" in path.name else 1 if "_rulebased_" not in path.name else 2,
                -path.stat().st_mtime,
            )
        )
        html_files = sorted(
            self.output_dir.glob("zirp_meeting_recommendations_*.html"),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        selected = pd.DataFrame()
        pattern_file = None
        for candidate in pattern_files:
            try:
                df = pd.read_csv(candidate, encoding="utf-8-sig")
            except Exception:
                continue
            if df.empty or "selected" not in df.columns:
                continue
            raw_selected = df[df["selected"].astype(str) == "1"].copy()
            selected = raw_selected
            if not selected.empty:
                format_text = selected.get("suggested_format", "").astype(str).str.lower()
                empfehlung_text = selected.get("empfehlung", "").astype(str).str.lower()
                public_selected = selected[
                    ~format_text.str.contains("nicht weiterverfolgen", na=False)
                    & ~empfehlung_text.str.contains("drop", na=False)
                    & ~empfehlung_text.str.contains("nicht weiterverfolgen", na=False)
                ].copy()
                if not public_selected.empty:
                    selected = public_selected
            if not selected.empty:
                pattern_file = candidate
                break
        if selected.empty or pattern_file is None:
            return {}
        pattern_suffix = pattern_file.stem.replace("zirp_meeting_patterns_", "")
        run_suffix = re.sub(r"^(ai_|rulebased_)", "", pattern_suffix)
        matching_html = self.output_dir / f"zirp_meeting_recommendations_{pattern_suffix}.html"
        html_file_name = matching_html.name if matching_html.exists() else (html_files[0].name if html_files else "")
        leaderboard_file = self.output_dir / f"zirp_match_leaderboard_{run_suffix}.csv"
        model_summary_file = self.output_dir / f"zirp_match_model_summary_{run_suffix}.json"
        model_summary: dict[str, Any] = {}

        if leaderboard_file.exists():
            try:
                leaderboard_df = pd.read_csv(leaderboard_file, encoding="utf-8-sig")
                if not leaderboard_df.empty and "pattern_id" in leaderboard_df.columns:
                    wanted_cols = [
                        col for col in [
                            "pattern_id", "rank_all_pairs", "why_selected", "why_not_selected",
                            "blocked_by_member", "blocked_by_cluster", "blocked_by_opportunity",
                        ]
                        if col in leaderboard_df.columns
                    ]
                    selected = selected.merge(
                        leaderboard_df[wanted_cols],
                        on="pattern_id",
                        how="left",
                        suffixes=("", "_leaderboard"),
                    )
            except Exception:
                pass

        if model_summary_file.exists():
            try:
                loaded = json.loads(model_summary_file.read_text(encoding="utf-8"))
                if isinstance(loaded, dict):
                    model_summary = loaded
            except Exception:
                model_summary = {}

        if "score" in selected.columns:
            selected["score_num"] = pd.to_numeric(selected["score"], errors="coerce").fillna(0)
            selected = selected.sort_values("score_num", ascending=False)
            if "display_points" in selected.columns:
                selected["display_points"] = pd.to_numeric(selected["display_points"], errors="coerce").fillna(
                    selected["score_num"].round()
                ).round().astype(int)
            else:
                selected["display_points"] = selected["score_num"].round().astype(int)
            selected = selected[selected["display_points"] >= 8].copy()
            public_role_selected = selected[selected.apply(self.row_has_public_actor_roles, axis=1)].copy()
            if len(public_role_selected) >= max(2, min(self.dynamic_public_scip_card_limit(), 4)):
                selected = public_role_selected
            if selected.empty:
                return {}

        cards = []
        public_card_limit = self.dynamic_public_scip_card_limit()
        for _, row in selected.head(public_card_limit).iterrows():
            cards.append({
                "pattern_id": str(row.get("pattern_id", "")),
                "recommendation_id": str(row.get("recommendation_id", "")),
                "opportunity_id": str(row.get("opportunity_id", "")),
                "members": str(row.get("members", "")),
                "score": str(row.get("score", "")),
                "display_points": str(row.get("display_points", "")),
                "selected": str(row.get("selected", "1")),
                "public_status": str(row.get("public_status", "")),
                "card_title": str(row.get("card_title", "")),
                "decision_summary": str(row.get("decision_summary", "")),
                "decision_line": str(row.get("decision_line", "")),
                "risk_line": str(row.get("risk_line", "")),
                "next_line": str(row.get("next_line", "")),
                "raw_score": str(row.get("raw_score", "")),
                "rule_final_score": str(row.get("rule_final_score", "")),
                "nn_adjustment": str(row.get("nn_adjustment", "")),
                "feedback_count": str(row.get("feedback_count", "")),
                "word_memory_hits": str(row.get("word_memory_hits", "")),
                "feedback_evidence_count": str(row.get("feedback_evidence_count", "")),
                "critical_action_penalty": str(row.get("critical_action_penalty", "")),
                "academic_support_count": str(row.get("academic_support_count", "")),
                "profit_anchor_count": str(row.get("profit_anchor_count", "")),
                "implementation_anchor_count": str(row.get("implementation_anchor_count", "")),
                "context_actor_count": str(row.get("context_actor_count", "")),
                "empfehlung": str(row.get("empfehlung", "")),
                "zirp_fit": str(row.get("zirp_fit", "")),
                "aktionsreife": str(row.get("aktionsreife", "")),
                "signalstaerke": str(row.get("signalstaerke", "")),
                "theme": str(row.get("convening_theme", "")),
                "convening_typology": str(row.get("convening_typology", "")),
                "convening_stage": str(row.get("convening_stage", "")),
                "purpose_maturity": str(row.get("purpose_maturity", "")),
                "next_engagement_move": str(row.get("next_engagement_move", "")),
                "expected_outputs": str(row.get("expected_outputs", "")),
                "do_not_ask_for": str(row.get("do_not_ask_for", "")),
                "stage_gate_status": str(row.get("stage_gate_status", "")),
                "stage_gate_reason": str(row.get("stage_gate_reason", "")),
                "north_star_purpose": str(row.get("north_star_purpose", "")),
                "participants_logic": str(row.get("participants_logic", "")),
                "shared_topics": str(row.get("shared_topics", "")),
                "bridge_topics": str(row.get("bridge_topics", "")),
                "concrete_clusters": str(row.get("concrete_clusters", "")),
                "cluster_problem": str(row.get("cluster_problem", "")),
                "format": str(row.get("suggested_format", "") or row.get("format", "")),
                "agenda": str(row.get("possible_agenda", "")),
                "why": str(row.get("editorial_justification", "")),
                "why_now": str(row.get("why_now", "")),
                "shared_tension": str(row.get("shared_tension", "")),
                "ideal_guests": str(row.get("ideal_guests", "")),
                "possible_output": str(row.get("possible_output", "")),
                "recommendation_limit": str(row.get("recommendation_limit", "")),
                "next_best_action": str(row.get("next_best_action", "")),
                "maturity_level": str(row.get("maturity_level", "")),
                "why_this_actor": str(row.get("why_this_actor", "")),
                "why_this_counterpart": str(row.get("why_this_counterpart", "")),
                "concrete_joint_action": str(row.get("concrete_joint_action", "")),
                "evidence_quality": str(row.get("evidence_quality", "")),
                "sufficiency": str(row.get("sufficiency", "")),
                "consequence": str(row.get("consequence", "")),
                "uncertainty": str(row.get("uncertainty", "")),
                "action_decision": str(row.get("action_decision", "")),
                "critical_decision_reason": str(row.get("critical_decision_reason", "")),
                "rlp_relevance": str(row.get("rlp_relevance", "")),
                "recent_signals": str(row.get("recent_signals", "")),
                "rank_all_pairs": str(row.get("rank_all_pairs", "")),
                "why_selected": str(row.get("why_selected", "")),
                "why_not_selected": str(row.get("why_not_selected", "")),
                "blocked_by_member": str(row.get("blocked_by_member", "")),
                "blocked_by_cluster": str(row.get("blocked_by_cluster", "")),
                "blocked_by_opportunity": str(row.get("blocked_by_opportunity", "")),
            })

        return {
            "count": int(len(selected)),
            "pattern_file": pattern_file.name,
            "html_file": html_file_name,
            "leaderboard_file": leaderboard_file.name if leaderboard_file.exists() else "",
            "model_summary_file": model_summary_file.name if model_summary_file.exists() else "",
            "created_at": datetime.fromtimestamp(pattern_file.stat().st_mtime).strftime("%d.%m.%Y, %H:%M:%S"),
            "stats": self.latest_radar_funnel_stats(pattern_file),
            "model_summary": model_summary,
            "cards": cards,
        }

    def latest_radar_funnel_stats(self, pattern_file: Path) -> dict[str, Any]:
        stats: dict[str, Any] = {}
        try:
            homepage_files = sorted(
                self.output_dir.glob("zirp_mitglied_homepages_*.csv"),
                key=lambda path: path.stat().st_mtime,
                reverse=True,
            )
            if homepage_files:
                members_df = pd.read_csv(homepage_files[0], encoding="utf-8-sig")
                stats["observed_members"] = int(len(members_df))
        except Exception:
            pass
        try:
            event_files = sorted(
                self.output_dir.glob("zirp_ereignisse_*.csv"),
                key=lambda path: path.stat().st_mtime,
                reverse=True,
            )
            if event_files:
                events_df = pd.read_csv(event_files[0], encoding="utf-8-sig")
                stats["signal_count"] = int(len(events_df))
                if "mitglied" in events_df.columns:
                    stats["signal_members"] = int(events_df["mitglied"].nunique())
                if "role_in_briefing" in events_df.columns:
                    counts = events_df["role_in_briefing"].fillna("").astype(str).value_counts().to_dict()
                    stats["main_signals"] = int(counts.get("top_signal", 0))
                    stats["context_signals"] = int(counts.get("condensation", 0))
                    stats["watchlist_signals"] = int(counts.get("watchlist", 0))
        except Exception:
            pass
        try:
            patterns_df = pd.read_csv(pattern_file, encoding="utf-8-sig")
            stats["candidate_count"] = int(len(patterns_df))
        except Exception:
            pass
        return stats

    def render_meeting_panel(self, data: dict[str, Any]) -> str:
        if not data:
            return (
                '<section class="card meeting-panel">'
                '<div class="meeting-head"><div><h2>SCIP Match Selector</h2>'
                '<p>Noch keine SCIP Opportunity Matches gespeichert.</p></div></div>'
                '</section>'
            )

        cards = "\n".join(self.render_meeting_card(item, index) for index, item in enumerate(data.get("cards", [])))
        html_file = html.escape(str(data.get("html_file", "")))
        leaderboard_file = html.escape(str(data.get("leaderboard_file", "")))
        model_summary_file = html.escape(str(data.get("model_summary_file", "")))
        model_summary = data.get("model_summary", {}) if isinstance(data.get("model_summary", {}), dict) else {}
        stats = data.get("stats", {}) if isinstance(data.get("stats", {}), dict) else {}
        observed_members = stats.get("observed_members")
        signal_count = stats.get("signal_count")
        signal_members = stats.get("signal_members")
        candidate_count = stats.get("candidate_count")
        main_signals = stats.get("main_signals")
        context_signals = stats.get("context_signals")
        watchlist_signals = stats.get("watchlist_signals")
        count = html.escape(str(data.get("count", 0)))
        coverage_bits = [f"{count} Ã¶ffentliche Meetinglinien"]
        if observed_members:
            coverage_bits.append(f"{int(observed_members)} Mitglieder beobachtet")
        if signal_count and signal_members:
            coverage_bits.append(f"{int(signal_count)} Signale bei {int(signal_members)} Mitgliedern")
        if candidate_count:
            coverage_bits.append(f"{int(candidate_count)} SCIP-Kandidaten geprÃ¼ft")
        coverage_line = " · ".join(coverage_bits)
        candidate_total = model_summary.get("candidate_count") or candidate_count
        selected_total = model_summary.get("selected_count")
        coverage_line = f"{count} kuratierte Opportunity Matches"
        links = []
        if html_file:
            links.append(f'<a class="button-link" href="{html_file}" target="_blank" rel="noopener">Alle Pairings bewerten</a>')
        link = "".join(links)
        return f"""
    <section class="card meeting-panel">
      <div class="meeting-head">
        <div>
          <h2>SCIP Match Selector</h2>
          <p>{html.escape(coverage_line)} · aktualisiert {html.escape(str(data.get("created_at", "")))}</p>
        </div>
        {link}
      </div>
      <div class="meeting-grid">
        {cards}
      </div>
    </section>"""

    def critical_decision_grid_html(self, item: dict[str, Any]) -> str:
        rows = [
            ("Why this actor", item.get("why_this_actor", "")),
            ("Why this counterpart", item.get("why_this_counterpart", "")),
            ("Why now", item.get("why_now", "")),
            ("Concrete action", item.get("concrete_joint_action", "") or item.get("possible_output", "")),
        ]
        bits = []
        for label, value in rows:
            text = self.clean_dashboard_value(value)
            if text:
                bits.append(
                    '<div class="decision-line">'
                    f'<span>{html.escape(label)}</span>'
                    f'<p>{html.escape(text)}</p>'
                    '</div>'
                )
        action_decision = self.clean_dashboard_value(item.get("action_decision", ""))
        sufficiency = self.clean_dashboard_value(item.get("sufficiency", ""))
        evidence_quality = self.clean_dashboard_value(item.get("evidence_quality", ""))
        uncertainty = self.clean_dashboard_value(item.get("uncertainty", ""))
        footer_bits = []
        if action_decision:
            footer_bits.append(f"Decision: {action_decision}")
        if evidence_quality:
            footer_bits.append(f"Evidence: {evidence_quality}")
        if sufficiency:
            footer_bits.append(f"Sufficiency: {sufficiency}")
        if uncertainty:
            footer_bits.append(f"Uncertainty: {uncertainty}")
        footer = ""
        if footer_bits:
            footer = f'<div class="decision-footer">{html.escape(" · ".join(footer_bits))}</div>'
        if not bits and not footer:
            return ""
        return '<div class="critical-decision-grid">' + "".join(bits) + footer + '</div>'


    def render_meeting_card(self, item: dict[str, Any], index: int = 0) -> str:
        members = self.compact_scip_members(item.get("members", ""))
        why_compact = self.compact_convening_why(item)
        signals_compact = self.compact_signal_items(item.get("recent_signals", ""))
        format_compact = self.compact_convening_format(item)
        card_label = self.scip_card_label(item, index)
        points = self.scip_display_points(item)
        card_title = self.clean_dashboard_value(item.get("card_title", ""))
        decision_summary = self.clean_dashboard_value(item.get("decision_summary", ""))
        decision_line = self.clean_dashboard_value(item.get("decision_line", ""))
        risk_line = self.clean_dashboard_value(item.get("risk_line", ""))
        next_line = self.clean_dashboard_value(item.get("next_line", ""))
        hypothesis = decision_summary or self.compact_scip_hypothesis(item)
        timeline = self.opportunity_timeline_html(item)
        status_badges = self.opportunity_status_badges_html(item)
        topic = self.clean_dashboard_value(item.get("theme", ""))
        actors = [part.strip() for part in self.clean_dashboard_value(item.get("members", "")).split("|") if part.strip()]
        actor_types = {actor: self.actor_role_type(actor) for actor in actors}
        role_logic = self.actor_role_logic_html(item, actors, actor_types)
        public_status = self.scip_public_status(item)
        signal_count = len(signals_compact)
        score_explanation = self.scip_score_explanation_html(item, actor_types, signal_count)
        if decision_line or risk_line or next_line:
            decision_bits = []
            if decision_line:
                decision_bits.append(f"<p><strong>Decision:</strong> {html.escape(decision_line)}</p>")
            if risk_line:
                decision_bits.append(f"<p><strong>Risk:</strong> {html.escape(risk_line)}</p>")
            if next_line:
                decision_bits.append(f"<p><strong>Next:</strong> {html.escape(next_line)}</p>")
            decision_grid = '<div class="critical-decision-grid compact-decision-lines">' + "".join(decision_bits) + "</div>"
        else:
            decision_grid = self.critical_decision_grid_html(item)
        words = re.findall(
            r"[A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ][A-Za-zÃ„Ã–ÃœÃ¤Ã¶Ã¼ÃŸ\-]{2,}",
            " ".join([topic, why_compact, " ".join(title for _, title in signals_compact)]),
        )[:30]
        feedback_payload = {
            "recommendation_id": self.clean_dashboard_value(item.get("recommendation_id", "")),
            "opportunity_id": self.clean_dashboard_value(item.get("opportunity_id", "")),
            "pattern_id": self.clean_dashboard_value(item.get("pattern_id", "")),
            "rank": self.clean_dashboard_value(item.get("rank_all_pairs", "")),
            "selected": str(item.get("selected", "1")).strip().lower() not in {"0", "false", "no"},
            "score": self.safe_float(item.get("score", 0)),
            "display_points": points,
            "pair": " x ".join(actors),
            "actors": actors,
            "actor_types": actor_types,
            "topic": topic,
            "topics": {
                "theme": topic,
                "shared": self.dashboard_cell_list(item.get("shared_topics", ""), 4),
                "bridge": self.dashboard_cell_list(item.get("bridge_topics", ""), 4),
                "concrete": self.dashboard_cell_list(item.get("concrete_clusters", ""), 4),
            },
            "reason": why_compact,
            "selection_reason": self.clean_dashboard_value(item.get("why_selected", "")),
            "non_selection_reason": self.clean_dashboard_value(item.get("why_not_selected", "")),
            "feedback_layer": "dashboard_selected_card",
            "words": words,
            "public_status": public_status,
            "features": {
                "scip_score": self.safe_float(item.get("score", 0)),
                "display_points": points,
                "role_coverage_score": self.role_coverage_score(actor_types),
                "number_of_signals": signal_count,
                "signal_strength": self.safe_float(item.get("signalstaerke", 0)),
                "source_quality": min(1.0, signal_count / 4.0),
                "recency": 1.0 if self.clean_dashboard_value(item.get("why_now", "")) else 0.0,
                "llm_judge_score": self.safe_float(item.get("nn_adjustment", 0)),
                "public_status_score": {"hauptlinie": 3, "sondierungshypothese": 2, "watchlist": 1, "hidden": 0}.get(public_status, 0),
                "score": self.safe_float(item.get("score", 0)),
                "raw_score": self.safe_float(item.get("raw_score", 0)),
                "rule_final_score": self.safe_float(item.get("rule_final_score", 0)),
                "nn_adjustment": self.safe_float(item.get("nn_adjustment", 0)),
                "word_memory_hits": self.safe_float(item.get("word_memory_hits", 0)),
                "feedback_evidence_count": self.safe_float(item.get("feedback_evidence_count", 0)),
                "critical_action_penalty": self.safe_float(item.get("critical_action_penalty", 0)),
                "academic_support_count": self.safe_float(item.get("academic_support_count", 0)),
                "profit_anchor_count": self.safe_float(item.get("profit_anchor_count", 0)),
                "implementation_anchor_count": self.safe_float(item.get("implementation_anchor_count", 0)),
                "context_actor_count": self.safe_float(item.get("context_actor_count", 0)),
            },
            "critical_judgment": {
                "why_this_actor": self.clean_dashboard_value(item.get("why_this_actor", "")),
                "why_this_counterpart": self.clean_dashboard_value(item.get("why_this_counterpart", "")),
                "why_now": self.clean_dashboard_value(item.get("why_now", "")),
                "concrete_joint_action": self.clean_dashboard_value(item.get("concrete_joint_action", "")),
                "evidence_quality": self.clean_dashboard_value(item.get("evidence_quality", "")),
                "sufficiency": self.clean_dashboard_value(item.get("sufficiency", "")),
                "consequence": self.clean_dashboard_value(item.get("consequence", "")),
                "uncertainty": self.clean_dashboard_value(item.get("uncertainty", "")),
                "action_decision": self.clean_dashboard_value(item.get("action_decision", "")),
                "critical_decision_reason": self.clean_dashboard_value(item.get("critical_decision_reason", "")),
            },
            "source": "dashboard_scip_card",
        }
        payload_attr = html.escape(json.dumps(feedback_payload, ensure_ascii=False), quote=True)
        selected_reason = self.clean_dashboard_value(item.get("why_selected", ""))
        selected_reason_html = (
            f'<p><strong>Warum ausgewählt?</strong><br>{html.escape(selected_reason)}</p>'
            if selected_reason else ""
        )

        return (
            f'<article class="meeting-card compact-meeting-card" data-feedback-payload="{payload_attr}">'
            f'<div class="scip-label">{html.escape(card_label)}</div>'
            f'<h3><span class="meeting-title">{html.escape(card_title or members)}</span><span class="points-chip">{points}/100</span></h3>'
            f'<p class="scip-hypothesis">{html.escape(hypothesis)}</p>'
            f'{decision_grid}'
            '<div class="feedback-box" aria-label="SCIP Feedback">'
            '<span>Feedback für SCIP:</span>'
            '<button type="button" onclick="prepareScipFeedback(this, \'useful\')">Nützlich</button>'
            '<button type="button" onclick="prepareScipFeedback(this, \'interesting_but_weak\')">Schwach</button>'
            '<button type="button" onclick="prepareScipFeedback(this, \'wrong_connection\')">Verbindung falsch</button>'
            '<button type="button" onclick="prepareScipFeedback(this, \'good_topic_wrong_actors\')">Akteure falsch</button>'
            '<button type="button" onclick="prepareScipFeedback(this, \'not_relevant\')">Nicht relevant</button>'
            '<div class="feedback-comment-box" hidden>'
            '<label>Was bewertest du?</label>'
            '<select class="feedback-target-select" aria-label="Feedback-Ziel">'
            '<option value="">Ziel auswählen</option>'
            '<option value="actor_match">Akteursmatch</option>'
            '<option value="evidence_signal">Evidenzsignal</option>'
            '<option value="rlp_relevance">RLP-Relevanz</option>'
            '<option value="next_action">Nächster Schritt</option>'
            '<option value="typology_stage">Typologie/Reifegrad</option>'
            '<option value="public_framing">Öffentliche Darstellung</option>'
            '</select>'
            '<label>Warum?</label>'
            '<select class="feedback-reason-select" aria-label="Feedback-Grund">'
            '<option value="">Grund auswählen</option>'
            '<option value="good_match">Guter Match</option>'
            '<option value="good_match_weak_evidence">Guter Match, aber schwache Evidenz</option>'
            '<option value="good_match_weak_rlp">Guter Match, aber schwacher RLP-Bezug</option>'
            '<option value="indirect_transfer_legitimacy">Indirekte Transfer-Legitimation, keine direkte RLP-Anwendung</option>'
            '<option value="missing_rlp_application_anchor">RLP-Anwendungsanker fehlt</option>'
            '<option value="source_signal_not_application">Signal zeigt Kompetenz, aber keinen Anwendungskontext</option>'
            '<option value="source_only_actor">Akteur ist nur Quelle/Kontext</option>'
            '<option value="wrong_counterpart">Gegenpart passt nicht</option>'
            '<option value="insufficient_evidence">Zu wenig Evidenz</option>'
            '<option value="weak_why_now">Warum jetzt unklar</option>'
            '<option value="too_broad_action">Vorschlag zu breit</option>'
            '<option value="wrong_pilot_lane">Falsches Pilotfeld</option>'
            '<option value="not_public">Nicht öffentlich zeigen</option>'
            '<option value="downgrade_not_drop">Nicht löschen, sondern herunterstufen</option>'
            '</select>'
            '<textarea rows="3" placeholder="z.B. Akteursrolle oder Verbindung passt nicht."></textarea>'
            '<button type="button" onclick="submitScipFeedback(this)">Feedback speichern</button>'
            '<button type="button" onclick="cancelScipFeedback(this)">Abbrechen</button>'
            '</div>'
            '<small class="feedback-status"></small>'
            '</div>'
            '</article>'
        )

    def safe_float(self, value: Any) -> float:
        try:
            if value is None or pd.isna(value):
                return 0.0
        except Exception:
            pass
        try:
            return float(str(value).replace(",", "."))
        except Exception:
            return 0.0


    def compact_member_name(self, member: str) -> str:
        value = self.clean_dashboard_value(member)
        lower = value.lower()
        if "iqib" in lower:
            return "IQIB"
        if "universit" in lower and "trier" in lower:
            return "Uni Trier"
        if "technische hochschule" in lower and "bingen" in lower:
            return "TH Bingen"
        if "hochschule mainz" in lower:
            return "Hochschule Mainz"
        if "1. fsv mainz" in lower:
            return "Mainz 05"
        if "landesbank baden" in lower or "lbbw" in lower:
            return "LBBW"
        if len(value) > 46:
            return value[:43].rstrip() + "..."
        return value


    def actor_role_rank(self, member: str) -> int:
        role = self.actor_role_type(member)
        if role == "profit_anchor":
            return 0
        if role == "implementation_anchor":
            return 1
        if role == "academic_support":
            return 3
        return 2


    def actor_role_type(self, member: str) -> str:
        lower = f" {self.clean_dashboard_value(member).lower()} "
        is_academic = any(marker in lower for marker in ["hochschule", "universit", "uni ", "whu"])
        is_business = any(marker in lower for marker in [" gmbh", " ag", " se", " kg", " bank", "landesbank", "sparkasse", "versicherung", "werke", "industrie", "wirtschaft"])
        is_implementation = any(marker in lower for marker in ["kammer", "verband", "klinikum", "krankenhaus", "stadt ", "kreis ", "ministerium", "agentur", "institut"])
        if is_business and not is_academic:
            return "profit_anchor"
        if is_implementation and not is_academic:
            return "implementation_anchor"
        if is_academic:
            return "academic_support"
        return "context_actor"


    def row_has_public_actor_roles(self, row: pd.Series) -> bool:
        members = [part.strip() for part in self.clean_dashboard_value(row.get("members", "")).split("|") if part.strip()]
        if not members:
            return False
        item = {
            "members": self.clean_dashboard_value(row.get("members", "")),
            "theme": self.clean_dashboard_value(row.get("convening_theme", "")),
            "shared_topics": self.clean_dashboard_value(row.get("shared_topics", "")),
            "bridge_topics": self.clean_dashboard_value(row.get("bridge_topics", "")),
            "concrete_clusters": self.clean_dashboard_value(row.get("concrete_clusters", "")),
            "recent_signals": self.clean_dashboard_value(row.get("recent_signals", "")),
            "why": self.clean_dashboard_value(row.get("editorial_justification", "")),
        }
        return all(self.actor_has_public_role(actor, item) for actor in members)


    def actor_has_public_role(self, actor: str, item: dict[str, Any]) -> bool:
        role = self.actor_role_type(actor)
        if role in {"profit_anchor", "implementation_anchor", "academic_support"}:
            return True
        actor_lower = self.clean_dashboard_value(actor).lower()
        text = " ".join(
            self.clean_dashboard_value(item.get(key, ""))
            for key in ["theme", "shared_topics", "bridge_topics", "concrete_clusters", "recent_signals", "why"]
        ).lower()
        if "swr" in actor_lower:
            return any(term in text for term in ["versorgung", "Ã¶ffentlich", "oeffentlich", "kommunen", "problem", "studie", "alarm", "debatte", "sichtbarkeit"])
        if "zdf" in actor_lower:
            return any(term in text for term in ["Ã¶ffentlich", "oeffentlich", "debatte", "gesellschaft", "sichtbarkeit", "demokratie"])
        if "mainz 05" in actor_lower or "fsv mainz" in actor_lower:
            return any(term in text for term in ["jugend", "prÃ¤vention", "praevention", "gesund", "community", "bildung", "spende", "stadion", "fans"])
        return False


    def actor_role_sentence(self, actor: str, item: dict[str, Any]) -> str:
        actor_display = self.compact_member_name(actor)
        role = self.actor_role_type(actor)
        actor_lower = self.clean_dashboard_value(actor).lower()
        text = " ".join(
            self.clean_dashboard_value(item.get(key, ""))
            for key in ["theme", "shared_topics", "bridge_topics", "concrete_clusters", "recent_signals", "why"]
        ).lower()
        if "iqib" in actor_lower:
            role_text = "kommunale Resilienz, Wissenstransfer und angewandte Innovationsforschung."
        elif "universit" in actor_lower and "trier" in actor_lower:
            role_text = "Ethik, Gesundheit/Pflege, Digitalisierung und Nachhaltigkeit."
        elif "swr" in actor_lower:
            role_text = "regionale Ã–ffentlichkeit, Problemwahrnehmung und Sichtbarkeit von VersorgungslÃ¼cken."
        elif "zahnen" in actor_lower:
            role_text = "mittelstÃ¤ndische Umsetzungspraxis und digitale Werkzeuge."
        elif "handwerkskammer" in actor_lower:
            role_text = "Zugang zum Mittelstand, Betrieben und praktischer Umsetzung."
        elif "hochschule mainz" in actor_lower:
            role_text = "Wissenstransfer, angewandte Forschung und Projektzugang."
        elif "technische hochschule" in actor_lower or "th bingen" in actor_lower:
            role_text = "technische Umsetzungskompetenz, Transfer und Nachwuchsperspektive."
        elif role == "profit_anchor":
            role_text = "Praxisanker mit Umsetzungs-, Markt- oder Anwendungsperspektive."
        elif role == "implementation_anchor":
            role_text = "institutioneller Zugang zu Umsetzung, Fachpraxis oder Mitgliedern."
        elif role == "academic_support":
            role_text = "Wissens-, Forschungs- und Transferperspektive."
        elif self.actor_has_public_role(actor, item):
            role_text = "Ã¶ffentliche Problemwahrnehmung oder gesellschaftliche Kontextperspektive."
        else:
            role_text = "keine ausreichend klare Ã¶ffentliche Rolle fÃ¼r diese Konstellation."
        if "versorgung" in text and "swr" in actor_lower:
            role_text = "regionale Ã–ffentlichkeit und Sichtbarkeit von VersorgungslÃ¼cken."
        return f"{actor_display}: {role_text}"


    def actor_role_logic_html(self, item: dict[str, Any], actors: list[str], actor_types: dict[str, str]) -> str:
        rows = []
        for actor in actors:
            sentence = self.actor_role_sentence(actor, item)
            if ":" in sentence:
                actor_name, role_text = sentence.split(":", 1)
                rows.append(f"<li><strong>{html.escape(actor_name.strip())}:</strong>{html.escape(role_text)}</li>")
            else:
                rows.append(f"<li>{html.escape(sentence)}</li>")
        return "".join(rows) or "<li>Rollenlogik noch nicht ausreichend belegt.</li>"


    def role_coverage_score(self, actor_types: dict[str, str]) -> float:
        roles = set(actor_types.values())
        score = 0.0
        if roles & {"profit_anchor", "implementation_anchor"}:
            score += 0.45
        if "academic_support" in roles:
            score += 0.20
        if "context_actor" in roles:
            score += 0.10
        if len(roles) >= 3:
            score += 0.15
        if "context_actor" in roles and not (roles & {"profit_anchor", "implementation_anchor"}):
            score -= 0.20
        return max(0.0, min(1.0, round(score, 3)))


    def compact_scip_members(self, members: Any) -> str:
        parts = [part.strip() for part in self.clean_dashboard_value(members).split("|") if part.strip()]
        parts.sort(key=lambda name: (self.actor_role_rank(name), name.lower()))
        return " x ".join(self.compact_member_name(part) for part in parts)


    def scip_display_points(self, item: dict[str, Any]) -> int:
        explicit = self.clean_dashboard_value(item.get("display_points", ""))
        if explicit:
            return max(0, min(100, int(round(self.safe_float(explicit)))))
        return max(0, min(100, int(round(self.safe_float(item.get("score", 0))))))


    def scip_card_label(self, item: dict[str, Any], index: int = 0) -> str:
        points = self.scip_display_points(item)
        rank = self.clean_dashboard_value(item.get("rank_all_pairs", ""))
        prefix = f"#{rank} · " if rank else ""
        if index == 0:
            return prefix + "Top-Match"
        if index <= 2:
            return prefix + "Starker Match"
        if points >= 15:
            return prefix + "Sondierungsmatch"
        if points >= 10:
            return prefix + "Beobachtungsmatch"
        return "Nicht öffentlich anzeigen"


    def scip_public_status(self, item: dict[str, Any]) -> str:
        points = self.scip_display_points(item)
        if points >= 20:
            return "hauptlinie"
        if points >= 15:
            return "sondierungshypothese"
        if points >= 10:
            return "watchlist"
        return "hidden"


    def dashboard_cell_list(self, value: Any, limit: int = 3) -> list[str]:
        text = self.clean_dashboard_value(value)
        if not text:
            return []
        pieces = re.split(r"\s*[,|]\s*", text)
        return [piece.strip() for piece in pieces if piece.strip()][:limit]


    def dashboard_match_actor_names(self, item: dict[str, Any]) -> list[str]:
        raw = self.clean_dashboard_value(item.get("members", ""))
        return [part.strip() for part in raw.split("|") if part.strip()]


    def pair_specific_scip_hypothesis(self, item: dict[str, Any]) -> str:
        actors = " | ".join(self.dashboard_match_actor_names(item)).lower()
        text = " ".join(
            self.clean_dashboard_value(item.get(key, ""))
            for key in [
                "theme",
                "shared_topics",
                "bridge_topics",
                "concrete_clusters",
                "cluster_problem",
                "recent_signals",
                "why",
            ]
        ).lower()
        if "zahnen" in actors and "hochschule mainz" in actors:
            return "Digitale Anlagen- und Prozesskompetenz trifft auf anwendungsnahen Transfer."
        if "iqib" in actors and ("universit" in actors and "trier" in actors):
            return "Kommunale Resilienz trifft auf ethische und gesellschaftliche Legitimation digitaler LÃ¶sungen."
        if "handwerkskammer" in actors and ("technische hochschule" in actors or "th bingen" in actors):
            return "Betriebliche Anwendungspraxis trifft auf technische Transferkompetenz."
        if "zahnen" in actors and "handwerkskammer" in actors:
            return "Digitale Prozesspraxis trifft auf Zugang zu mittelstÃ¤ndischer Anwendung."
        if "swr" in actors:
            if "zahnen" in actors:
                return "MittelstÃ¤ndische Umsetzungspraxis trifft auf Ã¶ffentliche Problemwahrnehmung."
            if "handwerkskammer" in actors:
                return "Betriebliche UmsetzungslÃ¼cken treffen auf regionale Sichtbarkeit."
            return "Ã–ffentliche Problemwahrnehmung trifft auf einen mÃ¶glichen regionalen Umsetzungskontext."
        if "resilienz" in text and any(term in text for term in ["ethik", "gesund", "digital", "nachhalt"]):
            return "Kommunale Resilienz trifft auf wissenschaftliche und gesellschaftliche Orientierung."
        if "industrie" in text or "emissionshandel" in text or "klima" in text:
            return "Regulatorischer Druck trifft auf konkrete PilotfÃ¤higkeit in Unternehmen."
        return ""


    def compact_scip_hypothesis(self, item: dict[str, Any]) -> str:
        specific = self.pair_specific_scip_hypothesis(item)
        if specific:
            return specific
        text = " ".join(
            self.clean_dashboard_value(item.get(key, ""))
            for key in [
                "members",
                "shared_topics",
                "bridge_topics",
                "concrete_clusters",
                "cluster_problem",
                "theme",
                "recent_signals",
            ]
        ).lower()
        if "resilienz" in text and any(term in text for term in ["ethik", "gesund", "digital", "nachhalt"]):
            return "Kommunale Resilienz trifft auf wissenschaftliche und gesellschaftliche Orientierung."
        if "resilienz" in text and any(term in text for term in ["wissen", "transfer", "technik", "hochschule"]):
            return "Resilienzpraxis trifft auf Wissenstransfer und technische Umsetzung."
        if "pflege" in text or "versorgung" in text or "gesundheit" in text:
            return "Versorgungsdruck trifft auf Qualifizierung und regionale Umsetzungspraxis."
        if "emissionshandel" in text or "klima" in text or "industrie" in text:
            return "Regulatorischer Druck trifft auf konkrete PilotfÃ¤higkeit in Unternehmen."
        shared = self.dashboard_cell_list(item.get("shared_topics", ""), 2)
        bridge = self.dashboard_cell_list(item.get("bridge_topics", ""), 2)
        if shared and bridge:
            return f"{' / '.join(shared)} trifft auf {' / '.join(bridge)}."
        if shared:
            return f"{' / '.join(shared)} liefert einen prÃ¼fbaren Anlass fÃ¼r ein kleines ZIRP-SondierungsgesprÃ¤ch."
        tension = self.clean_dashboard_value(item.get("shared_tension", ""))
        if tension:
            return tension
        return "Noch nicht stark genug fÃ¼r eine Empfehlung, aber beobachtenswert."


    def opportunity_history_path(self) -> Path:
        return self.script_dir / "data" / "feedback" / "opportunity_history.json"


    def load_opportunity_history(self) -> list[dict[str, Any]]:
        path = self.opportunity_history_path()
        if not path.exists():
            return []
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
            return value if isinstance(value, list) else []
        except Exception:
            return []


    def opportunity_timeline_items(self, opportunity_id: str, limit: int = 4) -> list[dict[str, Any]]:
        opportunity_id = self.clean_dashboard_value(opportunity_id)
        if not opportunity_id:
            return []
        rows = [
            item for item in self.load_opportunity_history()
            if self.clean_dashboard_value(item.get("opportunity_id", "")) == opportunity_id
        ]
        rows.sort(key=lambda item: self.clean_dashboard_value(item.get("created_at", "")))
        return rows[-limit:]


    def opportunity_timeline_html(self, item: dict[str, Any]) -> str:
        rows = self.opportunity_timeline_items(item.get("opportunity_id", ""))
        if not rows:
            return (
                '<div class="opportunity-line">'
                '<strong>Entwicklungslinie</strong>'
                '<span>Neue Linie in Beobachtung.</span>'
                '</div>'
            )
        first = rows[0]
        current = rows[-1]
        trend = self.clean_dashboard_value(current.get("trend", "")) or "neu"
        trend_class = "rising" if trend == "steigend" else "falling" if trend == "fallend" else "stable"
        previous_status = self.clean_dashboard_value(current.get("previous_status", "")) or "neu"
        current_status = self.clean_dashboard_value(current.get("status", "")) or self.scip_card_label(item)
        new_signals = self.clean_dashboard_value(current.get("new_signal_count", "0"))
        feedback = self.clean_dashboard_value(current.get("feedback_summary", "")) or "kein Feedback"
        first_seen = self.clean_dashboard_value(current.get("first_seen", "")) or self.clean_dashboard_value(first.get("date", ""))
        chips = []
        for row in rows:
            date = self.clean_dashboard_value(row.get("date", ""))
            status = self.clean_dashboard_value(row.get("status", ""))
            points = self.clean_dashboard_value(row.get("score_points", ""))
            chips.append(f"<span>{html.escape(date)} · {html.escape(status)} · {html.escape(points)} Punkte</span>")
        return (
            f'<div class="opportunity-line {trend_class}">'
            '<strong>Entwicklungslinie</strong>'
            f'<span>{html.escape(first_seen)} gestartet · jetzt: {html.escape(current_status)} · Entwicklung: {html.escape(trend)} · +{html.escape(new_signals)} neue Signale · Feedback: {html.escape(feedback)}</span>'
            f'<div class="timeline-chips">{"".join(chips)}</div>'
            '</div>'
        )


    def opportunity_status_badges_html(self, item: dict[str, Any]) -> str:
        rows = self.opportunity_timeline_items(item.get("opportunity_id", ""))
        current = rows[-1] if rows else {}
        status = self.clean_dashboard_value(current.get("status", "")) or self.scip_card_label(item)
        trend = self.clean_dashboard_value(current.get("trend", "")) or "neu"
        first_seen = self.clean_dashboard_value(current.get("first_seen", "")) or self.clean_dashboard_value(current.get("date", "")) or "neu"
        new_signals = self.clean_dashboard_value(current.get("new_signal_count", "")) or "0"
        window = self.opportunity_window_label(status, trend, new_signals)
        return (
            '<div class="scip-status-row">'
            f'<span>Status: {html.escape(status)}</span>'
            f'<span>Trend: {html.escape(trend)}</span>'
            f'<span>Opportunity Window: {html.escape(window)}</span>'
            f'<span>Erstmals gesehen: {html.escape(first_seen)}</span>'
            f'<span>Neue Signale: {html.escape(new_signals)}</span>'
            '</div>'
        )

    def opportunity_window_label(self, status: str, trend: str, new_signals: str) -> str:
        trend_l = self.clean_dashboard_value(trend).lower()
        status_l = self.clean_dashboard_value(status).lower()
        try:
            signal_count = int(float(str(new_signals).replace(",", ".")))
        except Exception:
            signal_count = 0
        if trend_l in {"steigend", "rising", "verstÃ¤rkt", "verstaerkt"} or signal_count >= 2:
            return "Ã¶ffnet sich"
        if trend_l in {"fallend", "falling", "schwÃ¤cher", "schwaecher"}:
            return "schwÃ¤cht sich ab"
        if "hauptlinie" in status_l:
            return "offen"
        if "watch" in status_l:
            return "beobachten"
        return "neu"

    def scip_score_explanation_html(self, item: dict[str, Any], actor_types: dict[str, str], signal_count: int) -> str:
        points = self.scip_display_points(item)
        role_score = self.role_coverage_score(actor_types)
        signal_strength = self.safe_float(item.get("signalstaerke", 0))
        memory_hits = self.safe_float(item.get("word_memory_hits", 0))
        feedback_count = self.safe_float(item.get("feedback_evidence_count", 0))
        parts = [
            f"{points} Punkte GesamtprioritÃ¤t",
            f"Rollenabdeckung {role_score:.2f}",
            f"{signal_count} aktuelle Signale",
        ]
        if signal_strength:
            parts.append(f"SignalstÃ¤rke {signal_strength:.1f}")
        if memory_hits:
            parts.append(f"Memory-Hinweise {memory_hits:.0f}")
        if feedback_count:
            parts.append(f"Feedback-Evidenz {feedback_count:.0f}")
        return (
            "<p><strong>Warum diese Punktzahl?</strong><br>"
            + html.escape(" · ".join(parts))
            + "<br><span class='members'>Die Punktzahl ist kein Wahrheitswert, sondern ein Priorisierungssignal aus Rollenlogik, SignalstÃ¤rke, AktualitÃ¤t, Memory und Feedback.</span></p>"
        )

    def compact_convening_why(self, item: dict[str, Any]) -> str:
        why = self.clean_dashboard_value(item.get("why", ""))
        theme = self.clean_dashboard_value(item.get("theme", ""))
        problem = self.clean_dashboard_value(item.get("cluster_problem", ""))
        if problem:
            primary_problem = [part.strip() for part in problem.split("|") if part.strip()][0]
            return f"PrÃ¼fpunkt im Themenfeld {theme or 'dieser Linie'}: {primary_problem} KlÃ¤ren, ob daraus ein konkreter gemeinsamer Arbeitsauftrag entsteht."
        if why:
            if len(why) > 260:
                why = why[:257].rstrip() + "..."
            return why
        if theme:
            return f"Pruefbarer Convening-Anlass im Themenfeld {theme}."
        return "Pruefbarer Convening-Anlass; redaktionell noch zu schaerfen."


    def compact_signal_items(self, raw_signals: Any) -> list[tuple[str, str]]:
        signals = []
        for raw in str(raw_signals or "").split("||"):
            raw = raw.strip()
            if not raw:
                continue
            if ":" in raw:
                member, rest = raw.split(":", 1)
            else:
                member, rest = "", raw
            member = self.clean_dashboard_value(member)
            rest = self.clean_dashboard_value(rest)
            pieces = re.split(r"\s{2,}", rest, maxsplit=1)
            title = pieces[0].strip()
            snippet = pieces[1].strip() if len(pieces) > 1 else ""
            title_lower = title.lower()
            if "einbaukarte" in title_lower or "ersatzbaustoff" in title_lower:
                title = "EBV-Einbaukarte"
            elif "hannover messe" in title_lower:
                title = "Hannover Messe 2026"
            elif "mainzed" in title_lower:
                title = "mainzed"
            evidence = f"{title} â€” {snippet}" if snippet else title
            if len(evidence) > 170:
                evidence = evidence[:167].rstrip() + "..."
            signals.append((member.strip(), evidence))
        return signals

    def compact_convening_format(self, item: dict[str, Any]) -> str:
        fmt = self.clean_dashboard_value(item.get("format", ""))
        if fmt:
            return fmt.replace("kleiner", "Kleiner")
        return "Kleiner Expertenkreis"

    def clean_dashboard_value(self, value: Any) -> str:
        try:
            if value is None or pd.isna(value):
                return ""
        except Exception:
            if value is None:
                return ""
        text = re.sub(r"\s+", " ", str(value)).strip()
        if text.lower() in {"nan", "none", "null", "nat"}:
            return ""
        text = self.repair_mojibake(text)
        member_match = re.search(r"(\d+)\s+Mitglieder\s+gepr(?:ü|Ã¼|ue)ft", text, re.IGNORECASE)
        created_match = re.search(r"erstellt\s+(\d{2}\.\d{2}\.\d{4},\s*\d{1,2}:\d{2}(?::\d{2})?)", text, re.IGNORECASE)
        if member_match and created_match:
            return f"{member_match.group(1)} Mitglieder geprüft · erstellt {created_match.group(1)}"
        return text

    @staticmethod
    def repair_mojibake(text: str) -> str:
        if not text:
            return ""
        repaired = str(text)
        replacements = {
            chr(0x00c3) + chr(0x0192) + chr(0x00a4): "\u00e4",
            chr(0x00c3) + chr(0x0192) + chr(0x00b6): "\u00f6",
            chr(0x00c3) + chr(0x0192) + chr(0x00bc): "\u00fc",
            chr(0x00c3) + chr(0x0192) + chr(0x009f): "\u00df",
            chr(0x00c3) + chr(0x201a) + chr(0x00b7): "\u00b7",
            chr(0x00c3) + chr(0x201a): "",
            "\u00c3\u00a4": "\u00e4",
            "\u00c3\u00b6": "\u00f6",
            "\u00c3\u00bc": "\u00fc",
            "\u00c3\u009f": "\u00df",
            "\u00c3\u201e": "\u00c4",
            "\u00c3\u2013": "\u00d6",
            "\u00c3\u2013ffentlich": "\u00d6ffentlich",
            "\u00c3\u0178": "\u00df",
            "\u00c3\u0153": "\u00dc",
            "\u00c2\u00b7": "\u00b7",
            "\u00c2": "",
            "\u00e2\u20ac\u201c": "\u2013",
            "\u00e2\u20ac\u201d": "\u2014",
            "\u00e2\u20ac\u017e": "\u201e",
            "\u00e2\u20ac\u0153": "\u201c",
            "\u00e2\u20ac\u009d": "\u201d",
            "\u00e2\u20ac\u2122": "\u2019",
            "â€“": "\u2013",
            "â€“": "\u2013",
            "â€”": "\u2014",
            "\u00c3\u00a2\u00e2\u201a\u00ac\u00e2\u20ac\u0153": "\u2013",
            "\u00c3\u00a2\u00e2\u201a\u00ac\u00c2\u009d": "\u201d",
            "Ã¤": "ä",
            "Ã¶": "ö",
            "Ã¼": "ü",
            "Ãœ": "Ü",
            "Ã„": "Ä",
            "Ã–": "Ö",
            "ÃŸ": "ß",
            "Ã¤": "ä",
            "Ã¶": "ö",
            "Ã¼": "ü",
            "Ãœ": "Ü",
            "ÃƒÆ’Ã†â€™â€“": "Ö",
            "ÃƒÆ’Ã†â€™â€ž": "Ä",
            "LÃ–WEN": "LÖWEN",
            "LÃƒÂ–WEN": "LÖWEN",
            "ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬â€œ": "â€“",
            "ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â": "â€”",
            "ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦Â¾": "â€ž",
            "ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€¦â€œ": "â€œ",
            "ÃƒÆ’Ã‚Â¢ÃƒÂ¢Ã¢â‚¬Å¡Ã‚Â¬Ãƒâ€šÃ‚Â": "â€",
        }
        double_separator = (
            chr(0x00c3)
            + chr(0x0192)
            + chr(0x00c6)
            + chr(0x2019)
            + chr(0x00c3)
            + chr(0x00a2)
            + chr(0x00e2)
            + chr(0x201a)
            + chr(0x00ac)
            + chr(0x00c5)
            + chr(0x00a1)
            + chr(0x00c3)
            + chr(0x0192)
            + chr(0x00e2)
            + chr(0x20ac)
            + chr(0x0161)
            + " \u00b7 "
        )
        replacements[double_separator] = " \u00b7 "
        for bad, good in replacements.items():
            repaired = repaired.replace(bad, good)
        for _ in range(2):
            if not any(marker in repaired for marker in ("\u00c3", "\u00c2", "\u00e2", "\u0192")):
                break
            try:
                fixed = repaired.encode("cp1252").decode("utf-8")
            except UnicodeError:
                break
            if fixed == repaired:
                break
            repaired = fixed
        return repaired

    def ensure_dashboard_logo_asset(self) -> None:
        preferred = self.script_dir / "5d549ab9-bdc2-4449-b585-c6ac8568b465.png"
        source = preferred if preferred.exists() else self.script_dir / "sor-logo.png"
        if not source.exists():
            return
        target = self.output_dir / "sor-logo-full.png"
        try:
            shutil.copy2(source, target)
            legacy_target = self.output_dir / "sor-logo.png"
            shutil.copy2(source, legacy_target)
            tab_source = self.script_dir / "Screenshot O.png"
            tab_target = self.output_dir / "sor-tab-o.png"
            if tab_source.exists():
                shutil.copy2(tab_source, tab_target)
                shutil.copy2(tab_source, self.output_dir / "sor-tab-logo.png")
        except Exception as exc:
            print(f"Dashboard logo: konnte nicht kopiert werden ({exc})")

    def dashboard_nav(self, active: str, asset_prefix: str = "") -> str:
        items = [
            ("home", "Wochenanalyse", f"{asset_prefix}zirp_dashboard.html" if asset_prefix else "zirp_dashboard.html"),
            ("scip", "SCIP Matchings", f"{asset_prefix}scip_archive.html"),
            ("archive", "Archiv", f"{asset_prefix}archive.html"),
        ]
        links = []
        for key, label, href in items:
            cls = "active" if key == active else ""
            links.append(f'<a class="{cls}" href="{html.escape(href)}">{html.escape(label)}</a>')
        return "\n".join(links)

    def dashboard_shell(
            self,
            *,
            title: str,
            subtitle: str,
            active: str,
            body: str,
            asset_prefix: str = "",
    ) -> str:
        logo_src = f"{asset_prefix}sor-logo-full.png"
        tab_logo_src = f"{asset_prefix}sor-tab-o.png"
        nav = self.dashboard_nav(active, asset_prefix=asset_prefix)
        return f"""<!DOCTYPE html>
<html lang="de">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{html.escape(title)} · Strategic Opportunity Radar</title>
  <link rel="icon" type="image/png" href="{html.escape(tab_logo_src)}" />
  <link rel="apple-touch-icon" href="{html.escape(tab_logo_src)}" />
  <style>
    :root {{
      --bg: #05070A;
      --panel: #07111F;
      --panel-2: #0A1626;
      --card: #F5F7FA;
      --card-2: #FFFFFF;
      --ink: #05070A;
      --muted: #647183;
      --line: #D8E3ED;
      --navy: #07111F;
      --navy-2: #0D1B2D;
      --blue: #00AEEF;
      --blue-soft: #EAF8FF;
      --blue-line: rgba(0,174,239,.34);
      --gold: #F5B82E;
      --gold-soft: #FFF6D9;
      --white: #F5F7FA;
      --shadow: 0 18px 44px rgba(0,0,0,.24);
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background:
        linear-gradient(180deg, #05070A 0%, #07111F 42%, #091320 100%);
      color: var(--white);
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    }}
    body::before {{
      content: "";
      position: fixed;
      inset: 0;
      pointer-events: none;
      opacity: .18;
      background:
        repeating-linear-gradient(90deg, rgba(0,174,239,.06) 0 1px, transparent 1px 96px),
        repeating-linear-gradient(0deg, rgba(0,174,239,.045) 0 1px, transparent 1px 96px);
      mask-image: linear-gradient(180deg, black, transparent 78%);
    }}
    .site-header {{
      position: relative;
      background: #05070A;
      color: var(--white);
      border-bottom: 1px solid var(--blue-line);
      box-shadow: 0 18px 60px rgba(0,0,0,.32);
    }}
    .site-header::after {{
      content: "";
      position: absolute;
      left: 0;
      right: 0;
      bottom: -1px;
      height: 1px;
      background: linear-gradient(90deg, transparent, var(--blue), transparent);
    }}
    .header-inner {{
      max-width: 1180px;
      margin: 0 auto;
      padding: 14px 18px;
      display: flex;
      gap: 22px;
      align-items: center;
      justify-content: space-between;
    }}
    .brand {{
      display: flex;
      align-items: center;
      gap: 15px;
      min-width: 0;
      color: inherit;
      text-decoration: none;
    }}
    .brand img {{
      width: min(260px, 30vw);
      height: auto;
      aspect-ratio: 16 / 9;
      object-fit: contain;
      object-position: center;
      border-radius: 6px;
      border: 1px solid rgba(0,174,239,.14);
      background: #000;
      box-shadow: 0 0 18px rgba(0,174,239,.10);
      flex: 0 0 auto;
    }}
    .brand-title {{ display: none; }}
    nav {{
      display: flex;
      gap: 8px;
      flex-wrap: wrap;
      justify-content: flex-end;
    }}
    nav a {{
      color: #D9E7F2;
      text-decoration: none;
      border: 1px solid rgba(170,179,192,.22);
      border-radius: 999px;
      padding: 8px 12px;
      font-size: .78rem;
      font-weight: 820;
      text-transform: uppercase;
      letter-spacing: .06em;
      background: rgba(255,255,255,.035);
    }}
    nav a.active, nav a:hover {{
      color: #05070A;
      background: var(--blue);
      border-color: var(--blue);
      box-shadow: 0 0 18px rgba(0,174,239,.32);
    }}
    .wrap {{
      position: relative;
      max-width: 1180px;
      margin: 0 auto;
      padding: 34px 18px 76px;
    }}
    .page-intro {{
      position: relative;
      min-height: 118px;
      margin-bottom: 22px;
      border: 1px solid rgba(0,174,239,.22);
      border-radius: 12px;
      padding: 24px 28px;
      overflow: hidden;
      background:
        linear-gradient(135deg, rgba(7,17,31,.98), rgba(5,7,10,.92));
      box-shadow: var(--shadow);
    }}
    h1 {{
      position: relative;
      margin: 0;
      max-width: 760px;
      font-size: clamp(1.52rem, 2.7vw, 2.35rem);
      line-height: 1.14;
      font-weight: 840;
      letter-spacing: .02em;
      color: var(--white);
    }}
    .subtitle {{
      position: relative;
      margin: 12px 0 0;
      color: #AAB3C0;
      font-weight: 650;
      line-height: 1.48;
      max-width: 850px;
    }}
    .card {{
      background: var(--card);
      border: 1px solid rgba(216,227,237,.9);
      border-top: 2px solid rgba(0,174,239,.55);
      border-radius: 8px;
      box-shadow: var(--shadow);
      padding: 26px;
      color: var(--ink);
    }}
    .analysis {{
      position: relative;
      max-width: 940px;
      padding: 34px 38px 38px;
      overflow: hidden;
    }}
    .analysis::before {{
      content: "";
      position: absolute;
      inset: 0 auto 0 0;
      width: 4px;
      background: var(--blue);
      box-shadow: 0 0 18px rgba(0,174,239,.45);
    }}
    .analysis p {{
      max-width: 820px;
      font-size: 1.02rem;
      line-height: 1.72;
      margin: 0 0 17px;
      color: #172033;
    }}
    .analysis .briefing-title {{
      max-width: 860px;
      font-size: clamp(1.32rem, 2vw, 1.76rem);
      line-height: 1.25;
      font-weight: 900;
      margin: 0 0 16px;
      padding-bottom: 16px;
      border-bottom: 1px solid var(--line);
      color: var(--navy);
    }}
    .analysis .executive-summary {{
      max-width: 840px;
      margin: 0 0 24px;
      padding: 14px 16px;
      border-left: 4px solid var(--blue);
      border-radius: 7px;
      background: #EAF8FF;
      font-size: 1.04rem;
      line-height: 1.62;
      font-weight: 850;
      color: var(--navy);
    }}
    .analysis h2, .analysis h3 {{ margin: 24px 0 10px; color: var(--navy); }}
    .two-col {{ display: grid; grid-template-columns: minmax(0, 1fr) 340px; gap: 18px; align-items: start; }}
    .home-grid {{ display: grid; grid-template-columns: minmax(0, 1fr) minmax(280px, 390px); gap: 18px; align-items: start; }}
    .home-insights {{ display: grid; grid-template-columns: minmax(0, 1fr) minmax(220px, .75fr) minmax(240px, .9fr); gap: 18px; margin-top: 18px; }}
    .source-card {{ padding: 22px; }}
    .source-list {{ display: grid; gap: 10px; }}
    .source-item {{ border-top: 1px solid var(--line); padding-top: 10px; }}
    .source-item:first-child {{ border-top: 0; padding-top: 0; }}
    .source-item a {{ color: var(--ink); text-decoration: none; font-weight: 840; line-height: 1.35; }}
    .source-item a:hover {{ color: #006C96; }}
    .stat-grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 10px; margin-bottom: 18px; }}
    .stat {{ background: #07111F; border: 1px solid rgba(0,174,239,.38); border-radius: 8px; padding: 12px; }}
    .stat b {{ display: block; font-size: 1.55rem; color: var(--blue); }}
    .stat span {{ color: #AAB3C0; font-size: .78rem; font-weight: 850; text-transform: uppercase; letter-spacing: .08em; }}
    h2 {{ margin: 0 0 14px; font-size: 1.12rem; color: var(--navy); text-transform: uppercase; letter-spacing: .05em; }}
    h3 {{ color: var(--navy); }}
    .section-row {{ margin: 10px 0; }}
    .section-top {{ display: flex; justify-content: space-between; gap: 12px; font-size: .9rem; color: var(--muted); }}
    .bar {{ height: 8px; background: #E6EDF4; border-radius: 999px; overflow: hidden; margin-top: 5px; }}
    .bar span {{ display: block; height: 100%; background: var(--blue); box-shadow: 0 0 14px rgba(0,174,239,.42); }}
    .event-chip {{ border-top: 1px solid var(--line); padding: 12px 0; }}
    .event-chip:first-child {{ border-top: 0; }}
    .event-chip a {{ color: var(--ink); text-decoration: none; font-weight: 820; }}
    .event-chip small {{ display: block; color: var(--muted); margin-top: 4px; line-height: 1.4; }}
    .term-row {{ display: flex; flex-wrap: wrap; gap: 8px; }}
    .term {{ background: var(--blue-soft); border: 1px solid rgba(0,174,239,.22); border-radius: 999px; padding: 7px 10px; font-size: .84rem; color: var(--navy); font-weight: 720; }}
    .members {{ color: var(--muted); line-height: 1.58; font-size: .94rem; }}
    .history-card, .scip-card {{ border-top: 1px solid var(--line); padding: 18px 0; }}
    .history-card:first-child, .scip-card:first-child {{ border-top: 0; }}
    .history-card h3, .scip-card h3 {{ margin: 0 0 5px; font-size: 1.02rem; }}
    .history-card p, .scip-card p {{ color: var(--muted); margin: 8px 0 0; line-height: 1.58; }}
    .artifact-row {{ display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; }}
    .artifact-row a, .button-link {{
      border: 1px solid rgba(0,174,239,.48);
      background: #07111F;
      border-radius: 999px;
      padding: 8px 11px;
      text-decoration: none;
      color: var(--white);
      font-weight: 850;
      font-size: .78rem;
      text-transform: uppercase;
      letter-spacing: .05em;
      box-shadow: 0 0 16px rgba(0,174,239,.14);
    }}
    .artifact-row a:hover, .button-link:hover {{ background: var(--blue); color: #05070A; }}
    details {{ margin-top: 10px; }}
    summary {{ cursor: pointer; color: var(--navy); font-weight: 850; }}
    .meeting-head {{ display: flex; justify-content: space-between; align-items: flex-start; gap: 14px; margin-bottom: 14px; }}
    .meeting-head > div {{ min-width: 0; }}
    .meeting-head p {{ margin: 6px 0 0; color: var(--muted); line-height: 1.45; }}
    .radar-filter-note {{ font-size: .9rem; }}
    .meeting-grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 12px; }}
    .meeting-card {{ border: 1px solid var(--line); border-left: 3px solid var(--blue); border-radius: 8px; padding: 15px; background: var(--card-2); }}
    .meeting-card h3 {{ margin: 0 0 6px; font-size: 1rem; display: flex; align-items: baseline; justify-content: space-between; gap: 10px; flex-wrap: wrap; }}
    .meeting-card h3 .meeting-title {{ min-width: 0; }}
    .meeting-card h3 .points-chip {{ color: var(--navy); font-size: .82rem; font-weight: 900; white-space: nowrap; background: var(--blue-soft); border: 1px solid rgba(0,174,239,.35); border-radius: 999px; padding: 3px 7px; }}
    .meeting-card p, .meeting-card li {{ color: var(--muted); line-height: 1.5; font-size: .92rem; }}
    .meeting-card li strong {{ color: var(--navy); }}
    .meeting-card ul {{ margin: 6px 0 10px 18px; padding: 0; }}
    .role-list {{ list-style: none; margin-left: 0 !important; display: grid; gap: 6px; }}
    .role-list li {{ margin: 0; }}
    .scip-label {{ color: var(--blue); font-weight: 900; text-transform: uppercase; letter-spacing: .09em; font-size: .72rem; margin-bottom: 5px; }}
    .scip-hypothesis {{ color: var(--navy) !important; font-weight: 760; }}
    .critical-decision-grid {{
      border: 1px solid #dfe7ef;
      background: white;
      border-radius: 8px;
      padding: 10px;
      margin: 10px 0 12px;
      display: grid;
      gap: 8px;
    }}
    .decision-line span {{ display: block; color: var(--blue); font-size: .7rem; font-weight: 900; text-transform: uppercase; letter-spacing: .08em; margin-bottom: 2px; }}
    .decision-line p {{ color: var(--navy) !important; margin: 0; font-size: .86rem; line-height: 1.42; }}
    .decision-footer {{ border-top: 1px solid var(--line); padding-top: 7px; color: var(--muted); font-size: .78rem; line-height: 1.4; }}
    .scip-status-row {{ display: flex; flex-wrap: wrap; gap: 6px; margin: 8px 0 4px; }}
    .scip-status-row span {{ background: var(--blue-soft); border: 1px solid rgba(0,174,239,.26); border-radius: 999px; padding: 4px 8px; color: var(--navy); font-size: .74rem; font-weight: 850; text-transform: uppercase; letter-spacing: .05em; }}
    .opportunity-line {{ border: 1px solid #dfe7ef; background: white; border-radius: 8px; padding: 9px 10px; margin: 10px 0; display: grid; gap: 5px; }}
    .opportunity-line strong {{ color: var(--navy); font-size: .84rem; }}
    .opportunity-line > span {{ color: var(--muted); font-size: .84rem; line-height: 1.45; }}
    .opportunity-line.rising {{ border-left: 4px solid var(--blue); }}
    .opportunity-line.falling {{ border-left: 4px solid #d9480f; }}
    .opportunity-line.stable {{ border-left: 4px solid var(--gold); }}
    .timeline-chips {{ display: flex; flex-wrap: wrap; gap: 5px; }}
    .timeline-chips span {{ background: var(--blue-soft); border: 1px solid rgba(0,174,239,.22); border-radius: 999px; padding: 4px 7px; color: var(--muted); font-size: .74rem; font-weight: 780; }}
    .feedback-box {{ border-top: 1px solid var(--line); margin-top: 12px; padding-top: 10px; display: flex; flex-wrap: wrap; gap: 7px; align-items: center; }}
    .feedback-box span {{ color: var(--muted); font-size: .82rem; font-weight: 850; margin-right: 3px; }}
    .feedback-box button {{ border: 1px solid rgba(0,174,239,.28); background: white; color: var(--navy); border-radius: 999px; padding: 6px 9px; font-size: .78rem; font-weight: 780; cursor: pointer; }}
    .feedback-box button:hover {{ border-color: var(--blue); background: var(--blue-soft); }}
    .feedback-box select {{ border: 1px solid #d6e0ea; background: white; color: var(--navy); border-radius: 8px; padding: 6px 9px; font-size: .78rem; font-weight: 750; max-width: 100%; }}
    .feedback-comment-box {{ flex-basis: 100%; display: grid; gap: 7px; margin-top: 6px; }}
    .feedback-comment-box[hidden] {{ display: none; }}
    .feedback-comment-box label {{ color: var(--muted); font-size: .8rem; font-weight: 850; }}
    .feedback-comment-box textarea {{ width: 100%; box-sizing: border-box; border: 1px solid #d6e0ea; border-radius: 8px; padding: 8px 9px; font: inherit; resize: vertical; color: var(--ink); background: white; }}
    .feedback-status {{ color: var(--muted); font-size: .78rem; margin-left: 2px; }}
    .site-footer {{ background: #05070A; color: rgba(245,247,250,.68); border-top: 1px solid var(--blue-line); }}
    .footer-inner {{ max-width: 1180px; margin: 0 auto; padding: 18px; display: flex; justify-content: space-between; gap: 16px; flex-wrap: wrap; font-size: .88rem; }}
    .footer-inner strong {{ color: white; }}
    @media (max-width: 900px) {{
      .header-inner {{ display: block; }}
      .brand img {{ width: min(260px, 72vw); }}
      nav {{ justify-content: flex-start; margin-top: 14px; }}
      .page-intro {{ padding: 24px 20px; }}
      .two-col {{ grid-template-columns: 1fr; }}
      .home-grid, .home-insights {{ grid-template-columns: 1fr; }}
      .meeting-head {{ display: block; }}
      .meeting-head .button-link {{ display: inline-block; margin-top: 10px; }}
      .meeting-grid {{ grid-template-columns: 1fr; }}
    }}  </style>
</head>
<body>
  <header class="site-header">
    <div class="header-inner">
      <a class="brand" href="{html.escape(f'{asset_prefix}zirp_dashboard.html' if asset_prefix else 'zirp_dashboard.html')}">
        <img src="{html.escape(logo_src)}" alt="SOR Strategic Opportunity Radar Logo" />
      </a>
      <nav>{nav}</nav>
    </div>
  </header>
  <main class="wrap">
    <section class="page-intro">
      <h1>{html.escape(title)}</h1>
      <p class="subtitle">{html.escape(subtitle)}</p>
    </section>
    {body}
  </main>
  <footer class="site-footer">
    <div class="footer-inner">
      <span><strong>SOR · Strategic Opportunity Radar</strong></span>
      <span>Â© {datetime.now().year} SOR / SCIP method, scoring logic and convening radar concept. All rights reserved.</span>
    </div>
  </footer>
  <script>
    function defaultFeedbackTarget(label) {{
      if (label === 'useful') return 'actor_match';
      if (label === 'good_topic_wrong_actors' || label === 'wrong_connection') return 'actor_match';
      if (label === 'not_relevant') return 'rlp_relevance';
      return '';
    }}

    function defaultFeedbackReason(label) {{
      if (label === 'useful') return 'good_match';
      if (label === 'not_relevant') return 'missing_rlp_application_anchor';
      return '';
    }}

    function prepareScipFeedback(button, label) {{
      const card = button.closest('[data-feedback-payload]');
      const status = card.querySelector('.feedback-status');
      const box = card.querySelector('.feedback-comment-box');
      const textarea = box.querySelector('textarea');
      const targetSelect = box.querySelector('.feedback-target-select');
      const reasonSelect = box.querySelector('.feedback-reason-select');
      card.dataset.pendingFeedbackLabel = label;
      box.hidden = false;
      if (targetSelect && !targetSelect.value) targetSelect.value = defaultFeedbackTarget(label);
      if (reasonSelect && !reasonSelect.value) reasonSelect.value = defaultFeedbackReason(label);
      textarea.focus();
      status.textContent = 'Feedback kann mehrfach pro Karte gespeichert werden';
    }}

    function cancelScipFeedback(button) {{
      const card = button.closest('[data-feedback-payload]');
      const status = card.querySelector('.feedback-status');
      const box = card.querySelector('.feedback-comment-box');
      const textarea = box.querySelector('textarea');
      const selects = box.querySelectorAll('select');
      card.dataset.pendingFeedbackLabel = '';
      textarea.value = '';
      selects.forEach((select) => select.value = '');
      box.hidden = true;
      status.textContent = '';
    }}

    async function submitScipFeedback(button) {{
      const card = button.closest('[data-feedback-payload]');
      const status = card.querySelector('.feedback-status');
      const box = card.querySelector('.feedback-comment-box');
      const textarea = box.querySelector('textarea');
      const targetSelect = box.querySelector('.feedback-target-select');
      const reasonSelect = box.querySelector('.feedback-reason-select');
      const label = card.dataset.pendingFeedbackLabel || 'interesting_but_weak';
      let payload = {{}};
      try {{
        payload = JSON.parse(card.dataset.feedbackPayload || '{{}}');
      }} catch (error) {{
        payload = {{}};
      }}
      payload.label = label;
      payload.human_feedback_label = label;
      payload.feedback_target = targetSelect ? targetSelect.value : '';
      payload.feedback_dimension = payload.feedback_target;
      payload.reason_category = reasonSelect ? reasonSelect.value : '';
      payload.feedback_reason_category = payload.reason_category;
      payload.human_comment = textarea.value.trim();
      payload.timestamp = new Date().toISOString();
      payload.page_url = window.location.href;
      payload.public_status = payload.public_status || 'unknown';
      status.textContent = 'saving...';
      try {{
        const response = await fetch('http://127.0.0.1:8766/feedback/scip', {{
          method: 'POST',
          headers: {{ 'Content-Type': 'application/json' }},
          body: JSON.stringify(payload)
        }});
        if (!response.ok) throw new Error('feedback server unavailable');
        status.textContent = 'saved';
      }} catch (error) {{
        const key = 'sor_scip_feedback_queue';
        const queue = JSON.parse(localStorage.getItem(key) || '[]');
        queue.push(payload);
        localStorage.setItem(key, JSON.stringify(queue));
        status.textContent = 'saved locally';
      }}
      textarea.value = '';
      if (targetSelect) targetSelect.value = '';
      if (reasonSelect) reasonSelect.value = '';
      box.hidden = true;
      card.dataset.pendingFeedbackLabel = '';
    }}
  </script>
</body>
</html>"""

    def render_dashboard_html(self, history: list[dict[str, Any]], asset_prefix: str = "") -> str:
        current = history[0] if history else {}
        current_analysis = self.paragraphs_to_html(
            self.sanitize_dashboard_analysis(
                self.strip_duplicate_analysis_heading(
                self.normalize_report_text(str(current.get("analysis", "")))
                )
            )
        )
        subtitle = self.public_meta_line(current.get("meta_line", "")) or "Aktuelle Wochenanalyse aus Crawler-Signalen und redaktioneller Verdichtung."
        sources = "\n".join(self.render_source_item(item) for item in current.get("top_events", [])[:6])
        top_terms = "\n".join(self.render_term_pill(item) for item in current.get("top_terms", [])[:18])
        top_sections = "\n".join(self.render_section_bar(item) for item in current.get("top_sections", []))
        members = self.repair_mojibake(", ".join(current.get("members", [])[:32]))
        body = f"""
        <section class="home-grid">
          <article class="card analysis">{current_analysis or "<p>Noch keine Wochenanalyse gespeichert.</p>"}</article>
          <aside class="card source-card">
            <h2>Quellen</h2>
            <div class="source-list">{sources or "<p class='members'>Keine Quellen im Zeitraum.</p>"}</div>
          </aside>
        </section>
        <section class="home-insights">
          <article class="card">
            <h2>Begriffe</h2>
            <div class="term-row">{top_terms or "<p class='members'>Keine Begriffe ermittelt.</p>"}</div>
          </article>
          <article class="card">
            <h2>Schwerpunkte</h2>
            {top_sections or "<p class='members'>Keine Schwerpunkte ermittelt.</p>"}
          </article>
          <article class="card">
            <h2>Mitglieder</h2>
            <p class="members">{html.escape(members)}</p>
          </article>
        </section>"""
        return self.dashboard_shell(
            title="SOR Wochenradar",
            subtitle=(subtitle + " · Detect signals. Connect actors. Act in time."),
            active="home",
            body=body,
            asset_prefix=asset_prefix,
        )

    def strip_duplicate_analysis_heading(self, text: str) -> str:
        lines = str(text or "").splitlines()
        while lines and not lines[0].strip():
            lines.pop(0)
        if lines and lines[0].strip().lower() == "wochenanalyse":
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
        return "\n".join(lines).strip()

    def sanitize_dashboard_analysis(self, text: str) -> str:
        paragraphs = re.split(r"\n{2,}", self.repair_mojibake(str(text or "")).strip())
        cleaned: list[str] = []
        seen_generic_next_step = False
        for paragraph in paragraphs:
            p = re.sub(r"\s+", " ", paragraph).strip()
            if not p:
                continue
            lower = p.lower()
            if "zentral fÃ¼r die sind" in lower or "zentral fuer die sind" in lower:
                continue
            if "zu testbar machen" in lower:
                continue
            is_generic_next_step = (
                "ein erster schritt kÃ¶nnte ein kuratiertes sondierungsgesprÃ¤ch sein" in lower
                and ("bau, industrie oder verwaltung" in lower or "im bau, industrie oder verwaltung" in lower)
            )
            if is_generic_next_step:
                if seen_generic_next_step:
                    continue
                seen_generic_next_step = True
                p = (
                    "Ein erster Schritt kÃ¶nnte ein kuratiertes SondierungsgesprÃ¤ch zur Auswahl eines "
                    "90-Tage-Testfelds fÃ¼r digitale Resilienz oder mittelstÃ¤ndische Prozessdigitalisierung sein."
                )
            cleaned.append(p)
        return "\n\n".join(cleaned)

    def public_meta_line(self, value: Any) -> str:
        text = self.clean_dashboard_value(value)
        text = re.sub(r"^\*\*(.*?)\*\*$", r"\1", text)
        text = re.sub(r"\s*Â·\s*Modell\s+[^Â·]+(?=\s*Â·)", "", text)
        text = re.sub(r"\s*Â·\s*Modell\s+[^Â·]+(?=\s*Â·)", "", text)
        text = re.sub(r"\s*[Â·Â·]\s*", " · ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return self.repair_mojibake(text)

    def render_archive_html(self, history: list[dict[str, Any]]) -> str:
        previous_cards = "\n".join(self.render_previous_entry(item) for item in history)
        scip_cards = self.render_scip_archive_cards(history[1:])
        body = (
            f'<section class="card"><h2>Archivierte Wochenanalysen</h2>{previous_cards or "<p class=\"members\">Noch keine Analysen gespeichert.</p>"}</section>'
            f'<section class="card" style="margin-top:18px;"><h2>Alte SCIP-Matchings mit Wochenzusammenfassung</h2>{scip_cards}</section>'
        )
        return self.dashboard_shell(
            title="Signal Archiv",
            subtitle="Gespeicherte Radar-Lagen, SCIP-Matchings, Berichte und Rohdaten.",
            active="archive",
            body=body,
        )

    def render_signals_html(self, history: list[dict[str, Any]]) -> str:
        current = history[0] if history else {}
        top_events = "\n".join(self.render_event_chip(item) for item in current.get("top_events", []))
        top_sections = "\n".join(self.render_section_bar(item) for item in current.get("top_sections", []))
        top_terms = "\n".join(self.render_term_pill(item) for item in current.get("top_terms", []))
        members = self.repair_mojibake(", ".join(current.get("members", [])[:24]))
        body = f"""
        <section class="two-col">
          <div>
            <article class="card">
              <h2>Priorisierte News-Signale</h2>
              {top_events or "<p class='members'>Keine Funde im Zeitraum.</p>"}
            </article>
            <article class="card" style="margin-top:18px;">
              <h2>Gewichtete Begriffe</h2>
              <div class="term-row">{top_terms or "<p class='members'>Keine Begriffe ermittelt.</p>"}</div>
            </article>
          </div>
          <aside>
            <div class="card">
              <div class="stat-grid">
                <div class="stat"><b>{current.get("member_count", 0)}</b><span>Mitglieder</span></div>
                <div class="stat"><b>{current.get("event_count", 0)}</b><span>Signale</span></div>
              </div>
              <h2>Schwerpunkte</h2>
              {top_sections or "<p class='members'>Keine Schwerpunkte ermittelt.</p>"}
            </div>
            <div class="card" style="margin-top:18px;">
              <h2>Mitglieder</h2>
              <p class="members">{html.escape(members)}</p>
            </div>
          </aside>
        </section>"""
        return self.dashboard_shell(
            title="Signal Board",
            subtitle="Detected signals, weighted terms and current opportunity indicators.",
            active="signals",
            body=body,
        )

    def render_scip_archive_html(self, history: list[dict[str, Any]]) -> str:
        meeting_panel = self.render_meeting_panel(self.latest_meeting_recommendations())
        body = meeting_panel
        return self.dashboard_shell(
            title="SCIP Match Radar",
            subtitle="Detected actor bridges, evidence status and next validation steps.",
            active="scip",
            body=body,
        )

    def render_scip_archive_cards(self, history: Optional[list[dict[str, Any]]] = None) -> str:
        if history:
            grouped_cards = []
            for item in history:
                radar = item.get("scip_radar") if isinstance(item, dict) else None
                if not isinstance(radar, dict) or not radar.get("cards"):
                    continue
                period = (
                    f"{self.clean_dashboard_value(item.get('period_start', ''))} bis "
                    f"{self.clean_dashboard_value(item.get('period_end', ''))}"
                )
                summary = self.scip_week_summary(item)
                cards_html = "\n".join(
                    self.render_meeting_card(card, index)
                    for index, card in enumerate(radar.get("cards", [])[:self.dynamic_public_scip_card_limit()])
                )
                links = self.render_dashboard_artifact_links(item)
                grouped_cards.append(
                    '<div class="scip-card scip-week-group">'
                    f'<h3>{html.escape(period)}</h3>'
                    f'<small>{html.escape(self.public_meta_line(item.get("meta_line", "")))}</small>'
                    f'<p>{html.escape(summary)}</p>'
                    f'{links}'
                    '<details class="scip-week-details">'
                    '<summary>SCIP-Matchings dieser Woche anzeigen</summary>'
                    f'<div class="meeting-grid archive-meeting-grid">{cards_html}</div>'
                    '</details>'
                    '</div>'
                )
            legacy_cards = self.render_legacy_scip_file_cards()
            if grouped_cards:
                extra = (
                    '<div class="scip-card"><h3>Fruehere SCIP-Dateien ohne Wochenzusammenfassung</h3>'
                    '<p>Diese Dateien stammen aus Laeufen, bevor SCIP-Matchings direkt in der Wochenhistorie gespeichert wurden.</p>'
                    f'{legacy_cards}</div>'
                    if legacy_cards else ""
                )
                return "\n".join(grouped_cards) + extra

        return self.render_legacy_scip_file_cards() or "<p class='members'>Noch keine SCIP-Archive gespeichert.</p>"

    def render_legacy_scip_file_cards(self) -> str:
        files = sorted(
            self.output_dir.glob("zirp_meeting_recommendations*.html"),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        cards = []
        seen = set()
        for path in files:
            if path.name in seen or path.name.endswith("_latest.html"):
                continue
            seen.add(path.name)
            label = path.stem.replace("zirp_meeting_recommendations_", "").replace("_ai_", " KI ")
            created = datetime.fromtimestamp(path.stat().st_mtime).strftime("%d.%m.%Y, %H:%M")
            cards.append(
                '<div class="scip-card">'
                f'<h3>{html.escape(label)}</h3>'
                f'<p>Erstellt {html.escape(created)} · enthält ausgewählte SCIP-Matchings und Begründungen aus diesem Lauf.</p>'
                f'<div class="artifact-row"><a href="{html.escape(path.name)}" target="_blank" rel="noopener">SCIP-Datei öffnen</a></div>'
                '</div>'
            )
            if len(cards) >= 24:
                break
        return "\n".join(cards)

    def scip_week_summary(self, item: dict[str, Any]) -> str:
        text = re.sub(r"\s+", " ", self.repair_mojibake(str(item.get("analysis", "")))).strip()
        if text.startswith("Wochenanalyse "):
            text = text[len("Wochenanalyse "):].strip()
        markers = [
            "Actor Anchors:",
            "Warum das fÃ¼r ZIRP relevant ist:",
            "Warum fÃ¼r ZIRP:",
            "Warum das fÃ¼r ZIRP relevant ist:",
            "NÃ¤chster Schritt:",
            "NÃ¤chster Schritt:",
        ]
        end_positions = [text.find(marker) for marker in markers if text.find(marker) > 0]
        if end_positions:
            text = text[:min(end_positions)].strip()
        if len(text) > 360:
            text = text[:357].rstrip() + "..."
        return text or "Wochenzusammenfassung fÃ¼r diese SCIP-Matchings."

    def paragraphs_to_html(self, text: str) -> str:
        blocks: list[str] = []
        list_items: list[str] = []
        paragraph_index = 0

        def flush_list() -> None:
            if list_items:
                blocks.append("<ul>" + "".join(list_items) + "</ul>")
                list_items.clear()

        for raw in self.repair_mojibake(str(text)).splitlines():
            line = raw.strip()
            if not line:
                flush_list()
                continue
            if line.startswith("## "):
                flush_list()
                blocks.append(f"<h3>{html.escape(line[3:].strip())}</h3>")
                paragraph_index += 1
                continue
            if line.startswith("# "):
                flush_list()
                blocks.append(f'<h2 class="briefing-title">{html.escape(line[2:].strip())}</h2>')
                paragraph_index += 1
                continue
            if line.startswith("- "):
                list_items.append(f"<li>{html.escape(line[2:].strip())}</li>")
                continue
            flush_list()
            bold_match = re.match(r"^\*\*(.*?)\*\*$", line)
            if paragraph_index == 0 and line.endswith("?") and len(line) <= 140:
                blocks.append(f'<h2 class="briefing-title">{html.escape(line)}</h2>')
            elif paragraph_index == 1 and (bold_match or len(line) <= 260):
                summary = bold_match.group(1).strip() if bold_match else line
                blocks.append(f'<p class="executive-summary"><strong>{html.escape(summary)}</strong></p>')
            else:
                blocks.append(f"<p>{html.escape(line)}</p>")
            paragraph_index += 1

        flush_list()
        return "\n".join(blocks)

    def render_previous_entry(self, item: dict[str, Any]) -> str:
        preview = re.sub(r"\s+", " ", self.repair_mojibake(str(item.get("analysis", "")))).strip()
        if len(preview) > 420:
            preview = preview[:417].rstrip() + "..."
        artifact_links = self.render_dashboard_artifact_links(item)
        full_text = self.paragraphs_to_html(self.normalize_report_text(str(item.get("analysis", ""))))
        return (
            '<div class="history-card">'
            f"<h3>{html.escape(self.clean_dashboard_value(item.get('period_start', '')))} bis {html.escape(self.clean_dashboard_value(item.get('period_end', '')))}</h3>"
            f"<small>{html.escape(self.public_meta_line(item.get('meta_line', '')))}</small>"
            f"<p>{html.escape(preview)}</p>"
            f"{artifact_links}"
            "<details>"
            "<summary>Briefing im Dashboard öffnen</summary>"
            f"<div class=\"history-full\">{full_text}</div>"
            "</details>"
            "</div>"
        )

    def render_dashboard_artifact_links(self, item: dict[str, Any]) -> str:
        timestamp = str(item.get("timestamp", "")).strip()
        if not timestamp:
            return ""

        candidates = [
            ("Word-Bericht", f"zirp_leiterinnen_briefing_{timestamp}.docx"),
            ("Ereignisse CSV", f"zirp_ereignisse_{timestamp}.csv"),
            ("Begriffe CSV", f"zirp_begriffe_gewichtet_{timestamp}.csv"),
            ("Raw Events CSV", f"zirp_raw_events_{timestamp}.csv"),
            ("Match Leaderboard", f"zirp_match_leaderboard_{timestamp}.csv"),
            ("Modell Summary", f"zirp_match_model_summary_{timestamp}.json"),
        ]
        links = []
        for label, filename in candidates:
            if (self.output_dir / filename).exists():
                links.append(
                    f'<a href="{html.escape(filename)}" target="_blank" rel="noopener">{html.escape(label)}</a>'
                )
        if not links:
            return ""
        return f'<div class="artifact-row">{"".join(links)}</div>'

    def render_event_chip(self, item: dict[str, Any]) -> str:
        url = html.escape(str(item.get("url", "")))
        title = html.escape(self.clean_dashboard_value(item.get("title", "")))
        member = html.escape(self.clean_dashboard_value(item.get("member", "")))
        date = html.escape(self.clean_dashboard_value(item.get("date", "")))
        theme = html.escape(self.clean_dashboard_value(item.get("theme", "")))
        title_html = f'<a href="{url}" target="_blank" rel="noopener">{title}</a>' if url else title
        return f'<div class="event-chip">{title_html}<small>{date} · {member} · {theme}</small></div>'

    def render_source_item(self, item: dict[str, Any]) -> str:
        url = html.escape(str(item.get("url", "")))
        title = html.escape(self.clean_dashboard_value(item.get("title", "")))
        member = html.escape(self.clean_dashboard_value(item.get("member", "")))
        date = html.escape(self.clean_dashboard_value(item.get("date", "")))
        theme = html.escape(self.clean_dashboard_value(item.get("theme", "")))
        context = self.clean_dashboard_value(item.get("one_sentence_context", "")) or self.clean_dashboard_value(item.get("why_it_matters", ""))
        context_html = ""
        title_html = f'<a href="{url}" target="_blank" rel="noopener">{title}</a>' if url else title
        return (
            '<div class="source-item">'
            f'{title_html}'
            f'{context_html}'
            '</div>'
        )

    def render_section_bar(self, item: dict[str, Any]) -> str:
        count = int(item.get("count", 0) or 0)
        width = min(100, max(8, count * 18))
        title = html.escape(str(item.get("title", "")))
        return (
            '<div class="section-row">'
            f'<div class="section-top"><span>{title}</span><strong>{count}</strong></div>'
            f'<div class="bar"><span style="width:{width}%"></span></div>'
            '</div>'
        )

    def render_term_pill(self, item: dict[str, Any]) -> str:
        term = html.escape(str(item.get("term", "")))
        score = html.escape(str(item.get("score", "")))
        return f'<span class="term">{term} · {score}</span>'


    def create_word_report(self, filepath: Path, briefing_text: str, evidence_text: str) -> None:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

        if REPORT_STYLE != "narrative":
            title = doc.add_paragraph()
            run = title.add_run(REPORT_TITLE)
            run.bold = True
            run.font.size = Pt(16)

            subtitle = doc.add_paragraph()
            run = subtitle.add_run(REPORT_SUBTITLE)
            run.italic = True
            run.font.size = Pt(10.5)

            doc.add_paragraph("")
        self.add_briefing_text_to_doc(doc, briefing_text)

        if evidence_text.strip():
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

            evidence_title = doc.add_paragraph()
            run = evidence_title.add_run("Evidenzteil")
            run.bold = True
            run.font.size = Pt(14)

            doc.add_paragraph("")
            self.add_evidence_text_to_doc(doc, evidence_text)

        doc.save(filepath)


    def add_briefing_text_to_doc(self, doc: Document, text: str) -> None:
        text = self.normalize_report_text(text)
        lines = text.splitlines()

        for line in lines:
            line = line.rstrip()

            if not line.strip():
                doc.add_paragraph("")
                continue

            if line in {"Wochenanalyse", "Kalender der Woche"} or line.startswith("Kalender der letzten "):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.bold = True
                r.font.size = Pt(15)

            elif re.search(r"Mitglieder.*Modell.*erstellt", line):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.italic = True
                r.font.size = Pt(9)

            elif line == "Leitungsfassung":
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.bold = True
                r.font.size = Pt(15)

            elif re.match(r"^(Montag|Dienstag|Mittwoch|Donnerstag|Freitag|Samstag|Sonntag)\s+Â·\s+\d{2}\.\d{2}\.\d{4}$", line):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.bold = True
                r.font.size = Pt(11.5)

            elif line.startswith("Themenfeld:"):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.italic = True

            elif line == "Keine priorisierten Funde":
                p = doc.add_paragraph()
                p.add_run(line)

            else:
                p = doc.add_paragraph()
                p.add_run(line)

    def add_evidence_text_to_doc(self, doc: Document, text: str) -> None:
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue

            if block.startswith("# "):
                p = doc.add_paragraph()
                r = p.add_run(block.replace("# ", ""))
                r.bold = True
                r.font.size = Pt(14)

            elif block.startswith("## "):
                p = doc.add_paragraph()
                r = p.add_run(block.replace("## ", ""))
                r.bold = True
                r.font.size = Pt(11.5)

            elif block.startswith("**") and "| **" in block:
                lines = block.split("\n")
                first = lines[0]

                p = doc.add_paragraph()
                self.add_inline_bold_text(p, first)

                for line in lines[1:]:
                    lp = doc.add_paragraph(style="List Bullet")
                    self.add_inline_bold_text(lp, line)

            else:
                p = doc.add_paragraph()
                self.add_inline_bold_text(p, block)

    def add_inline_bold_text(self, paragraph, text: str) -> None:
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def create_wordcloud(self, filepath: Path) -> None:
        if WordCloud is None or not self.global_weighted_counter:
            return

        wc = WordCloud(
            width=1800,
            height=1000,
            background_color="white",
            collocations=False
        ).generate_from_frequencies(dict(self.global_weighted_counter))

        plt.figure(figsize=(15, 8))
        plt.imshow(wc, interpolation="bilinear")
        plt.axis("off")
        plt.tight_layout()
        plt.savefig(filepath, dpi=200, bbox_inches="tight")
        plt.close()

    @staticmethod
    def normalize_url(url: str) -> str:
        parsed = urlparse(url)
        return parsed._replace(fragment="").geturl().rstrip("/")

    @staticmethod
    def is_valid_http_url(url: str) -> bool:
        parsed = urlparse(url)
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)

    @staticmethod
    def same_domain(url1: str, url2: str) -> bool:
        d1 = urlparse(url1).netloc.replace("www.", "").lower()
        d2 = urlparse(url2).netloc.replace("www.", "").lower()
        return d1 == d2

    @staticmethod
    def deduplicate_list(items: Iterable[str]) -> list[str]:
        seen = set()
        result = []
        for item in items:
            if item not in seen:
                seen.add(item)
                result.append(item)
        return result

    @staticmethod
    def make_hash(text: str) -> str:
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    @staticmethod
    def text_len(text: str) -> int:
        return len(re.sub(r"\s+", " ", text).strip())


def main() -> None:
    analyzer = ZIRPWeeklyAnalyzer()
    analyzer.run()


if __name__ == "__main__":
    main()









